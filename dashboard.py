# dashboard.py
import tkinter as tk
from tkinter import Toplevel, Label, Button, messagebox, Menu, ttk, END
import sqlite3
import shutil
import os
import sys
from tkinter import filedialog
from PIL import Image, ImageTk, ImageOps
from datetime import datetime
from tkcalendar import DateEntry
from pathlib import Path
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.colors import HexColor
import tempfile
import time
import calendar
from reportlab.lib import colors
import json
from fpdf import FPDF
from tkinter import Canvas, Scrollbar, Frame, BOTH, LEFT, RIGHT, Y, NW, END
import platform
import subprocess
import reportlab.pdfgen.canvas as pdfcanvas
from reportlab.lib.pagesizes import A5
from reportlab.lib.utils import ImageReader
import math
from PIL import Image
from PIL import ImageTk
from PIL import Image as PILImage
from PIL.Image import Resampling    
from reportlab.lib.pagesizes import A4, landscape
# Matplotlib embedding
import matplotlib
matplotlib.use("Agg")   # use Agg for drawing to PNG (we'll embed image in Tk with PhotoImage via PIL)
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from PIL import Image, ImageTk
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from tempfile import mkdtemp

from db_helper import get_connection, set_window_icon
from styles import (COLORS, FONTS, apply_theme, center_window, setup_modal_window,
                    setup_fullscreen_window, create_window_header, create_section_header,
                    styled_entry, styled_label, create_button_bar, ORG_SUFFIX)

# PDF Generation
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from reportlab.lib.units import inch
import win32api

# DOCX Generation  
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from docx.enum.table        import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import copy, os, win32api

from PIL import Image, ImageDraw, ImageFont
import arabic_reshaper
from bidi.algorithm import get_display
import tempfile



# Add this near the top of your file (outside functions), once:
if getattr(sys, 'frozen', False):
    # Running as compiled .exe
    APP_DIR = os.path.dirname(sys.executable)
else:
    # Running as .py script
    APP_DIR = os.path.dirname(os.path.abspath(__file__))

PHOTO_DIR = os.path.join(APP_DIR, "Photo")

#///////////////////////////
def open_file_cross_platform(filepath):
    try:
        if platform.system() == 'Windows':
            os.startfile(filepath)
        elif platform.system() == 'Darwin':
            subprocess.Popen(['open', filepath])
        else:
            subprocess.Popen(['xdg-open', filepath])
    except Exception as e:
        print(f"Could not open file: {e}")


def format_db_date(value, fmt="%d/%m/%Y"):
    if not value:
        return ""
    if isinstance(value, datetime):
        return value.strftime(fmt)
    if isinstance(value, str):
        for parse_fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return datetime.strptime(value, parse_fmt).strftime(fmt)
            except ValueError:
                continue
        return value
    return str(value)


def parse_db_date(value):
    if not value:
        return None
    if isinstance(value, datetime):
        return value
    if isinstance(value, str):
        for parse_fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%d/%m/%Y", "%d-%m-%Y"):
            try:
                return datetime.strptime(value, parse_fmt)
            except ValueError:
                continue
    return None


SERVER_PHOTO_DIR = r"\\server\AghoshCareSys\Photo"
# ─────────────────────────────────────────────────────────────────────────────
 
def _resolve_photo_path(photo_path_raw: str) -> str:
    """
    Return the best absolute path for a child photo, or "" if not found.
    photo_path_raw comes from tblChildren.PhotoPath (could be a bare filename,
    a relative path, or a full absolute/UNC path).
    """
    if not photo_path_raw:
        return ""
 
    photo_path_raw = photo_path_raw.strip()
 
    # ── 1. Already absolute and exists? ──────────────────────────────────────
    if os.path.isabs(photo_path_raw) or photo_path_raw.startswith("\\\\"):
        if os.path.exists(photo_path_raw):
            return photo_path_raw
 
    # ── 2. Relative to EXE/script directory ──────────────────────────────────
    #    When launched via \\server\AghoshCareSys\AghoshCareSys.exe,
    #    sys.executable is the UNC path so this resolves correctly over LAN.
    try:
        if getattr(sys, "frozen", False):
            base = os.path.dirname(sys.executable)        # PyInstaller EXE
        else:
            base = os.path.dirname(os.path.abspath(__file__))  # plain .py
 
        candidate = os.path.normpath(os.path.join(base, photo_path_raw))
        if os.path.exists(candidate):
            return candidate
    except Exception:
        pass
 
    # ── 3. Bare filename → try directly inside the known Photo folder ─────────
    filename = os.path.basename(photo_path_raw)
    candidate = os.path.join(SERVER_PHOTO_DIR, filename)
    if os.path.exists(candidate):
        return candidate
 
    # ── 4. Not found anywhere ─────────────────────────────────────────────────
    return ""


def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


global edit_name, photo_path, photo_label
report_rows = []

# Determine the base directory where the app is running
if getattr(sys, 'frozen', False):
    # If the app is bundled using PyInstaller
    BASE_DIR = sys._MEIPASS
else:
    # When running from source (development)
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# Full path to the placeholder image inside "Photo" subfolder
placeholder_path = os.path.join(BASE_DIR, "Photo", "placeholder.png")
# Dynamically get path to the Photo folder
PHOTO_DIR = os.path.join(BASE_DIR, "Photo")



#Base directory of the application (works for both .py and .exe)
#BASE_DIR = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(sys.argv[0])))
# ============================================================
# Checking Commitments to be expired
# ============================================================


# ============================================================
# Update sponsorship if not active now
# ============================================================
#from datetime import datetime

def deactivate_expired_sponsorships():
    try:
        today = datetime.today().strftime('%Y-%m-%d')

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            UPDATE tblSponsorships
            SET IsActive = 0
            WHERE IsActive = 1
              AND EndDate IS NOT NULL
              AND EndDate < ?
        """, (today,))

        conn.commit()
        conn.close()

    except Exception as e:
        messagebox.showerror(
            "Startup Error",
            f"Failed to deactivate expired sponsorships:\n{e}"
        )
        
deactivate_expired_sponsorships()


def open_placeholder(title):
    """Placeholder function to open a module window."""
    win = tk.Toplevel(root)
    win.title(title + ORG_SUFFIX)
    win.geometry("600x400")
    win.grab_set()
    tk.Label(win, text=title, font=("Arial", 18, "bold")).pack(pady=20)
    tk.Label(win, text="This module is under development.", font=("Arial", 14)).pack()
    tk.Button(win, text="Close", command=win.destroy).pack(pady=10)
    #messagebox.showinfo("","Kashif")
# ============================================================
# Date format helpers
# ============================================================



def validate_and_convert_date(date_str):
    try:
        # Try parsing user date input (dd/mm/yyyy)
        dt = datetime.strptime(date_str, "%d/%m/%Y")
        # Return SQL-friendly format (yyyy-mm-dd)
        #return dt.strftime("%Y-%m-%d")
        return dt.strftime("%Y-%m-%d")
    except ValueError:
        messagebox.showerror("Invalid Date", f"Date '{date_str}' is not in DD/MM/YYYY format.")
        return None


# Use this helper function to show placeholder text
def set_placeholder(entry_widget, placeholder="DD/MM/YYYY"):
    entry_widget.insert(0, placeholder)
    entry_widget.bind("<FocusIn>", lambda e: entry_widget.delete(0, tk.END) if entry_widget.get() == placeholder else None)
    entry_widget.bind("<FocusOut>", lambda e: entry_widget.insert(0, placeholder) if not entry_widget.get() else None)

# ============================================================
# EDIT/UPDATE COLLECTORS
# ============================================================

def open_collector_edit_window():
    """Open the collector edit/update window."""
    edit_win = tk.Toplevel(root)
    edit_win.title("Edit Collector" + ORG_SUFFIX)
    edit_win.configure(bg=COLORS["bg"])
    center_window(edit_win, 700, 500)
    edit_win.grab_set()

    create_window_header(edit_win, "Search and Edit Collector")

    search_frame = tk.Frame(edit_win, bg=COLORS["bg"])
    search_frame.pack(padx=20, pady=10)

    tk.Label(search_frame, text="Search Collector:", font=FONTS["body"], bg=COLORS["bg"]).pack(side="left")
    search_input = tk.Entry(search_frame, font=FONTS["entry"], width=30)
    search_input.pack(side="left", padx=5)
    search_input.focus_set()

    # Listbox to show search results
    result_listbox = tk.Listbox(edit_win, height=8, font=("Arial", 12))
    result_listbox.pack(padx=20, pady=(0, 10), fill="x")

    scrollbar = tk.Scrollbar(edit_win, orient="vertical", command=result_listbox.yview)
    scrollbar.pack(side="right", fill="y")
    result_listbox.config(yscrollcommand=scrollbar.set)

    form_frame = tk.Frame(edit_win, bg=COLORS["bg_card"], bd=1, relief="solid", highlightbackground=COLORS["border"], highlightthickness=1)
    form_frame.pack(padx=30, pady=10)

    tk.Label(form_frame, text="Full Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=0, column=0, sticky="e", pady=5)
    edit_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    edit_name.grid(row=0, column=1)

    tk.Label(form_frame, text="Contact Number:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=1, column=0, sticky="e", pady=5)
    edit_contact = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    edit_contact.grid(row=1, column=1)

    tk.Label(form_frame, text="Address:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=2, column=0, sticky="e", pady=5)
    edit_address = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    edit_address.grid(row=2, column=1)

    tk.Label(form_frame, text="Center:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=3, column=0, sticky="e", pady=5)
    center_var = tk.StringVar()
    center_dropdown = ttk.Combobox(
        form_frame,
        textvariable=center_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    center_dropdown.grid(row=3, column=1)

    selected_collector_id = [None]  # To store current collector ID

    def load_centers():
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT CenterID, CenterName FROM tblCenters ORDER BY CenterName")
                rows = cursor.fetchall()
                conn.close()

                center_options = [row.CenterName for row in rows]
                center_dropdown["values"] = center_options
                global center_id_map
                center_id_map = {row.CenterName: row.CenterID for row in rows}
                if center_options:
                    center_dropdown.set(center_options[0])  # Default selection
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load centers: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_search(event=None):
        query = search_input.get().strip()
        if not query:
            return

        result_listbox.delete(0, tk.END)

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT c.CollectorID, c.FullName, c.ContactNumber, c.Address, ce.CenterName
                    FROM tblCollectors c
                    JOIN tblCenters ce ON c.CenterID = ce.CenterID
                    WHERE c.FullName LIKE ?
                    ORDER BY c.FullName
                """, ('%' + query + '%',))
                rows = cursor.fetchall()
                conn.close()

                global collector_data_map
                collector_data_map = {}

                for row in rows:
                    result_listbox.insert(tk.END, row.FullName)
                    collector_data_map[row.FullName] = {
                        'CollectorID': row.CollectorID,
                        'ContactNumber': row.ContactNumber,
                        'Address': row.Address,
                        'CenterName': row.CenterName
                    }

            except Exception as e:
                messagebox.showerror("Database Error", f"Search failed: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_select(event=None):
        try:
            selected = result_listbox.get(result_listbox.curselection())
            data = collector_data_map[selected]

            edit_name.delete(0, tk.END)
            edit_name.insert(0, selected)

            edit_contact.delete(0, tk.END)
            edit_contact.insert(0, data['ContactNumber'])

            edit_address.delete(0, tk.END)
            edit_address.insert(0, data['Address'])

            center_dropdown.set(data['CenterName'])
            selected_collector_id[0] = data['CollectorID']

        except Exception as e:
            pass

    result_listbox.bind("<<ListboxSelect>>", on_select)

    def on_update():
        name = edit_name.get().strip()
        contact = edit_contact.get().strip()
        address = edit_address.get().strip()
        selected_center = center_var.get()

        if not name:
            messagebox.showwarning("Input Error", "Full Name is required.")
            return

        if not selected_center or selected_center not in center_id_map:
            messagebox.showwarning("Input Error", "Please select a valid center.")
            return

        if selected_collector_id[0] is None:
            messagebox.showerror("Selection Error", "No collector selected for update.")
            return

        center_id = center_id_map[selected_center]

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE tblCollectors SET
                        FullName = ?,
                        ContactNumber = ?,
                        Address = ?,
                        CenterID = ?
                    WHERE CollectorID = ?
                """, (name, contact, address, center_id, selected_collector_id[0]))
                conn.commit()
                conn.close()
                messagebox.showinfo("Success", "Collector updated successfully!")
                #edit_win.destroy()
            except Exception as e:
                conn.rollback()
                messagebox.showerror("Error", f"Update failed: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_close():
        edit_win.grab_release()
        edit_win.destroy()

    btn_frame = tk.Frame(edit_win, bg=COLORS["bg"])
    btn_frame.pack(pady=10)

    ttk.Button(btn_frame, text="Update", command=on_update, style="Success.TButton", width=12).grid(row=0, column=0, padx=5)
    ttk.Button(btn_frame, text="Close", command=on_close, style="Danger.TButton", width=10).grid(row=0, column=1, padx=5)

    load_centers()

    # Load all collectors initially
    def load_all_collectors():
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT c.CollectorID, c.FullName, c.ContactNumber, c.Address, ce.CenterName
                    FROM tblCollectors c
                    JOIN tblCenters ce ON c.CenterID = ce.CenterID
                    ORDER BY c.FullName
                """)
                rows = cursor.fetchall()
                conn.close()

                result_listbox.delete(0, tk.END)
                global collector_data_map
                collector_data_map = {}

                for row in rows:
                    result_listbox.insert(tk.END, row.FullName)
                    collector_data_map[row.FullName] = {
                        'CollectorID': row.CollectorID,
                        'ContactNumber': row.ContactNumber,
                        'Address': row.Address,
                        'CenterName': row.CenterName
                    }

            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load collectors: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    # Populate listbox with all collectors initially
    load_all_collectors()

    # Bind search input
    search_input.bind('<KeyRelease>', lambda e: on_search())

    # === KEYBOARD SHORTCUTS ===
    edit_win.bind('<Return>', lambda e: on_update())
    edit_win.bind('<Escape>', lambda e: on_close())    
    
    
    
# ============================================================
# Collector Management
# ============================================================
def open_collector_mgmt():
    global logged_in_role  # Ensure role is accessible

    if logged_in_role.lower() != "admin":
        messagebox.showerror("Access Denied", "You are not authorized to manage collectors.")
        return

    collector_win = tk.Toplevel(root)
    collector_win.title("Add New Collector" + ORG_SUFFIX)
    collector_win.configure(bg=COLORS["bg"])

    global add_collector_window
    add_collector_window = collector_win

    center_window(collector_win, 700, 620)
    collector_win.grab_set()

    create_window_header(collector_win, "Add New Collector")

    form_frame = tk.Frame(collector_win, bg=COLORS["bg_card"], bd=1, relief="solid", highlightbackground=COLORS["border"], highlightthickness=1)
    form_frame.pack(padx=30, pady=10)

    tk.Label(form_frame, text="Full Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=0, column=0, sticky="e", pady=5)
    collector_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    collector_name.grid(row=0, column=1)

    tk.Label(form_frame, text="Contact Number:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=1, column=0, sticky="e", pady=5)
    collector_contact = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    collector_contact.grid(row=1, column=1)

    tk.Label(form_frame, text="Address:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=2, column=0, sticky="e", pady=5)
    collector_address = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    collector_address.grid(row=2, column=1)

    tk.Label(form_frame, text="Center:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=3, column=0, sticky="e", pady=5)
    center_var = tk.StringVar()
    center_dropdown = ttk.Combobox(
        form_frame,
        textvariable=center_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    center_dropdown.grid(row=3, column=1)

    def load_centers():
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT CenterID, CenterName FROM tblCenters ORDER BY CenterName")
                rows = cursor.fetchall()
                conn.close()

                center_options = [row.CenterName for row in rows]
                center_dropdown["values"] = center_options
                global center_id_map
                center_id_map = {row.CenterName: row.CenterID for row in rows}
                if center_options:
                    center_dropdown.set(center_options[0])  # Default selection
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load centers: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to database.")

    # === BUTTON PANEL ===
    btn_frame = tk.Frame(collector_win)
    btn_frame.pack(pady=10)
    collector_name.focus_set()  # <-- Set focus to collector full name
    
    def open_edit_window():
        collector_win.destroy()  # <<< This destroys Add window before opening Edit
        open_collector_edit_window()  # Call edit window without args now
    
    def on_add():
        name = collector_name.get().strip()
        contact = collector_contact.get().strip()
        address = collector_address.get().strip()
        selected_center = center_var.get()

        if not name:
            messagebox.showwarning("Input Error", "Collector Full Name is required.")
            return

        if not selected_center or selected_center not in center_id_map:
            messagebox.showwarning("Input Error", "Please select a valid center.")
            return

        center_id = center_id_map[selected_center]

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO tblCollectors (FullName, ContactNumber, Address, CenterID)
                    VALUES (?, ?, ?, ?)
                """, (name, contact, address, center_id))
                conn.commit()
                conn.close()
                messagebox.showinfo("Success", "Collector added successfully!")
                clear_form_collector()
                display_collectors()
                collector_name.focus_set()  # <-- Set focus to collector full name
            except Exception as e:
                conn.rollback()
                messagebox.showerror("Error", f"Failed to add collector: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def clear_form_collector():
        collector_name.delete(0, tk.END)
        collector_contact.delete(0, tk.END)
        collector_address.delete(0, tk.END)
        #center_dropdown.set("")  # Reset dropdown
        collector_name.focus_set()
        
    # === SCROLLABLE FRAME FOR COLLECTORS LIST ===
    container = tk.Frame(collector_win)
    container.pack(fill="both", padx=30, pady=10, expand=True)

    canvas = tk.Canvas(container, height=200)
    scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
    collector_data_frame = tk.Frame(canvas)

    collector_data_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=collector_data_frame, anchor="nw", width=660)
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Store frame so we can access it later
    global collector_list_container
    collector_list_container = {
        'canvas': canvas,
        'data_frame': collector_data_frame
    }

    

    # === TABLE HEADER ===
    header_frame = tk.Frame(collector_win)
    header_frame.pack(padx=30, pady=(0, 5))

    tk.Label(header_frame, text="Collector Name", font=("Arial", 12, "bold"), width=25, anchor="w").pack(side="left")
    tk.Label(header_frame, text="Center", font=("Arial", 12, "bold"), width=25, anchor="w").pack(side="left")

    # === BUTTONS ===
    ttk.Button(btn_frame, text="Add Collector", command=on_add, style="Success.TButton", width=15).grid(row=0, column=0, padx=5)
    ttk.Button(btn_frame, text="Clear", command=clear_form_collector, style="Outline.TButton", width=10).grid(row=0, column=1, padx=5)
    ttk.Button(btn_frame, text="Edit/Update", command=open_edit_window, style="Primary.TButton", width=12).grid(row=0, column=2, padx=5)
    ttk.Button(btn_frame, text="Close", command=collector_win.destroy, style="Danger.TButton", width=10).grid(row=0, column=3, padx=5)

    # === DISPLAY EXISTING COLLECTORS ===
    def display_collectors():
        # Clear existing data
        for widget in collector_list_container['data_frame'].winfo_children():
            widget.destroy()

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT c.FullName, ce.CenterName 
                    FROM tblCollectors c
                    JOIN tblCenters ce ON c.CenterID = ce.CenterID
                    ORDER BY ce.CenterName, c.FullName
                """)
                rows = cursor.fetchall()
                conn.close()

                for row in rows:
                    item_frame = tk.Frame(collector_list_container['data_frame'])
                    item_frame.pack(fill="x")

                    tk.Label(item_frame, text=row.FullName, font=("Arial", 11), width=25, anchor="w").pack(side="left")
                    tk.Label(item_frame, text=row.CenterName, font=("Arial", 11), width=25, anchor="w").pack(side="left")

                # Update scroll region
                collector_list_container['canvas'].update_idletasks()
                collector_list_container['canvas'].configure(scrollregion=collector_list_container['canvas'].bbox("all"))

            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load collectors: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    # Load centers and collectors
    load_centers()
    display_collectors()
    
    
    def on_close():
        collector_win.grab_release()
        collector_win.destroy()

    ttk.Button(btn_frame, text="Close", command=on_close, style="Danger.TButton", width=10).grid(row=0, column=3, padx=5)

    collector_win.bind('<Return>', lambda e: on_add())
    collector_win.bind('<Escape>', lambda e: on_close())
    
    
# ============================================================
# Edit Child Information
# ============================================================
#Best Practice (Important)

#Database datetime fields ke liye hamesha ye pattern use karein:

#date_entry.set_date(parse_db_date(db_value))

#Ye aapko 3 problems se bachata hai:

#✔ datetime with time
#✔ string date
#✔ NULL values
def parse_db_date(value):
    if not value:
        return None
    try:
        if isinstance(value, str):
            value = value.split(" ")[0]
            return datetime.strptime(value, "%Y-%m-%d").date()
        return value.date()
    except:
        return None    

    
def open_edit_child_form(child_id):
    edit_win = tk.Toplevel(root)
    edit_win.title("Edit Child Details" + ORG_SUFFIX)
    edit_win.state('zoomed')
    edit_win.grab_set()
    wcm_entries = []

    # ── Colour & font constants ──────────────────────────────────────────────
    HEADER_BG   = "#1a3c5e"
    HEADER_FG   = "#ffffff"
    SECTION_BG  = "#f0f4f8"
    FRAME_BG    = "#ffffff"
    LABEL_FG    = "#374151"
    ACCENT      = "#2563eb"
    BTN_SAVE_BG = "#16a34a"
    BTN_CLOSE_BG= "#dc2626"
    FONT_LABEL  = ("Segoe UI", 9)
    FONT_ENTRY  = ("Segoe UI", 9)
    FONT_TITLE  = ("Segoe UI", 10, "bold")
    FONT_HEADER = ("Segoe UI", 13, "bold")
    FONT_BTN    = ("Segoe UI", 10, "bold")
    PAD         = {"padx": 6, "pady": 4}


    def safe_value(val):
        if val is None:
            return " "
        val = str(val).strip()
        return val if val != "" else " "

    # ── Helper: styled LabelFrame section ────────────────────────────────────
    def make_section(parent, title, row, col=0, colspan=1, rowspan=1):
        lf = tk.LabelFrame(
            parent, text=f"  {title}  ",
            font=FONT_TITLE, fg=ACCENT, bg=FRAME_BG,
            bd=1, relief="groove",
            labelanchor="nw"
        )
        lf.grid(row=row, column=col, columnspan=colspan, rowspan=rowspan,
                sticky="nsew", padx=8, pady=6)
        lf.columnconfigure((0, 2, 4), weight=0)
        lf.columnconfigure((1, 3, 5), weight=1)
        return lf

    # ── Helper: label + entry pair ────────────────────────────────────────────
    def lbl_entry(parent, text, row, col, width=22, default=""):
        tk.Label(parent, text=text, font=FONT_LABEL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        e = tk.Entry(parent, font=FONT_ENTRY, width=width,
                     relief="solid", bd=1)
        e.grid(row=row, column=col + 1, sticky="ew", **PAD)
        if default:
            e.insert(0, default)
        return e

    # ── Helper: label + combobox pair ─────────────────────────────────────────
    def lbl_combo(parent, text, row, col, values, default="", width=20):
        tk.Label(parent, text=text, font=FONT_LABEL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        var = tk.StringVar(value=default)
        cb = ttk.Combobox(parent, textvariable=var, values=values,
                          state="readonly", font=FONT_ENTRY, width=width)
        cb.grid(row=row, column=col + 1, sticky="ew", **PAD)
        return var, cb

    # ── Helper: label + DateEntry pair ────────────────────────────────────────
    def lbl_date(parent, text, row, col, default_date=None, width=14):
        tk.Label(parent, text=text, font=FONT_LABEL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        de = DateEntry(parent, date_pattern="dd/MM/yyyy",
                       font=FONT_ENTRY, width=width)
        de.grid(row=row, column=col + 1, sticky="ew", **PAD)
        if default_date:
            parsed = parse_db_date(default_date)
            if parsed:
                de.set_date(parsed)
        return de

    # ═════════════════════════════════════════════════════════════════════════
    # FETCH DATA
    # ═════════════════════════════════════════════════════════════════════════
    conn = get_connection()
    if not conn:
        messagebox.showerror("Error", "Database connection failed.")
        edit_win.destroy()
        return

    cursor = conn.cursor()
    cursor.execute("SELECT * FROM tblChildren WHERE ChildID=?", (child_id,))
    data = cursor.fetchone()
    conn.close()

    if not data:
        messagebox.showerror("Error", "Child record not found.")
        edit_win.destroy()
        return

    columns = [col[0] for col in cursor.description]
    row_dict = dict(zip(columns, data))

    # ═════════════════════════════════════════════════════════════════════════
    # OUTER LAYOUT: header bar + scrollable body + footer buttons
    # ═════════════════════════════════════════════════════════════════════════

    # ── Header bar ────────────────────────────────────────────────────────────
    header = tk.Frame(edit_win, bg=HEADER_BG, height=50)
    header.pack(fill="x", side="top")
    tk.Label(header, text="✎  Edit Child Record",
             font=FONT_HEADER, bg=HEADER_BG, fg=HEADER_FG,
             pady=10).pack(side="left", padx=18)

    # ── Scrollable body ───────────────────────────────────────────────────────
    body_frame = tk.Frame(edit_win, bg=SECTION_BG)
    body_frame.pack(fill="both", expand=True)

    canvas = tk.Canvas(body_frame, bg=SECTION_BG, highlightthickness=0)
    v_scroll = tk.Scrollbar(body_frame, orient="vertical",   command=canvas.yview)
    h_scroll = tk.Scrollbar(body_frame, orient="horizontal", command=canvas.xview)

    h_scroll.pack(side="bottom", fill="x")
    v_scroll.pack(side="right",  fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    canvas.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)

    scroll_frame = tk.Frame(canvas, bg=SECTION_BG)
    win_id = canvas.create_window((0, 0), window=scroll_frame, anchor="nw")

    def _on_frame_configure(e):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def _on_canvas_configure(e):
        canvas.itemconfig(win_id, width=max(e.width, scroll_frame.winfo_reqwidth()))

    scroll_frame.bind("<Configure>", _on_frame_configure)
    canvas.bind("<Configure>", _on_canvas_configure)

    def _mousewheel(e):
        canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
    canvas.bind_all("<MouseWheel>", _mousewheel)

    # ── Footer button bar ─────────────────────────────────────────────────────
    footer = tk.Frame(edit_win, bg="#e5e7eb", pady=10)
    footer.pack(fill="x", side="bottom")

    # ═════════════════════════════════════════════════════════════════════════
    # CONTENT GRID inside scroll_frame
    # (3 columns of sections)
    # ═════════════════════════════════════════════════════════════════════════
    scroll_frame.columnconfigure((0, 1, 2), weight=1)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 0 — PHOTO (col 2, rows 0-1) + BASIC INFO (col 0-1, row 0)
    # ─────────────────────────────────────────────────────────────────────────

    # ── PHOTO SECTION ─────────────────────────────────────────────────────────
    photo_lf = make_section(scroll_frame, "Photo", row=0, col=2, rowspan=2)

    photo_container = tk.Frame(photo_lf, width=140, height=140,
                               bg="#d1d5db", relief="solid", bd=1)
    photo_container.grid(row=0, column=0, columnspan=2, padx=10, pady=8)
    photo_container.pack_propagate(False)

    photo_label = tk.Label(photo_container, text="No Image",
                           bg="#d1d5db", font=FONT_LABEL)
    photo_label.pack(expand=True)

    photo_path = tk.Entry(photo_lf, font=FONT_ENTRY, width=24,
                          relief="solid", bd=1)
    photo_path.grid(row=1, column=0, columnspan=2, padx=10, pady=2, sticky="ew")
    photo_path.insert(0, data.PhotoPath if data.PhotoPath else "")

    def load_photo(path):
        candidates = [
            path,
            os.path.join("Photo", os.path.basename(path)),
            os.path.join("photo", os.path.basename(path)),
        ]
        for p in candidates:
            if p and os.path.exists(p):
                try:
                    img = Image.open(p)
                    img.thumbnail((140, 140))
                    img_tk = ImageTk.PhotoImage(img)
                    photo_label.configure(image=img_tk, text="", bg="white")
                    photo_label.image = img_tk
                    photo_container.configure(bg="white")
                    return
                except Exception:
                    pass
        photo_label.configure(text="No Image", image="", bg="#d1d5db")
        photo_container.configure(bg="#d1d5db")

    load_photo(photo_path.get())

    def browse_photo():
        file_path = filedialog.askopenfilename(
            filetypes=[("Image Files", "*.jpg *.jpeg *.png")])
        if file_path:
            try:
                img = Image.open(file_path)
                img.thumbnail((140, 140))
                img_tk = ImageTk.PhotoImage(img)
                photo_label.configure(image=img_tk, text="", bg="white")
                photo_label.image = img_tk
                photo_container.configure(bg="white")
                filename = os.path.basename(file_path)
                photo_path.delete(0, tk.END)
                photo_path.insert(0, f"Photo\\{filename}")
            except Exception as e:
                messagebox.showerror("Photo Error", f"Failed to load photo: {e}")

    tk.Button(photo_lf, text="📁  Browse Photo",
              command=browse_photo, font=FONT_LABEL,
              bg=ACCENT, fg="white", relief="flat",
              cursor="hand2", padx=8, pady=4).grid(
        row=2, column=0, columnspan=2, pady=(6, 4))

    # Status + Required Amount inside photo frame
    is_active_child = tk.IntVar(value=1 if data.Status else 0)
    tk.Checkbutton(photo_lf, text="Active Record",
                   variable=is_active_child,
                   font=FONT_LABEL, bg=FRAME_BG, fg=LABEL_FG,
                   activebackground=FRAME_BG).grid(
        row=3, column=0, columnspan=2, pady=3)

    tk.Label(photo_lf, text="Required Amount (PKR):",
             font=FONT_LABEL, fg=LABEL_FG, bg=FRAME_BG).grid(
        row=4, column=0, sticky="e", **PAD)
    required_amount_child = tk.StringVar(
        value=data.ChildRequiredAmount if data.ChildRequiredAmount else "")
    tk.Entry(photo_lf, textvariable=required_amount_child,
             font=FONT_ENTRY, width=12,
             relief="solid", bd=1).grid(
        row=4, column=1, sticky="ew", **PAD)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 0 — BASIC INFORMATION (col 0-1)
    # ─────────────────────────────────────────────────────────────────────────
    basic_lf = make_section(scroll_frame, "Basic Information", row=0, col=0, colspan=2)

    full_name  = lbl_entry(basic_lf, "Full Name",          0, 0, default=data.FullName)
    father_name= lbl_entry(basic_lf, "Father Name",        0, 2, default=data.FatherName)
    gender_var, _ = lbl_combo(basic_lf, "Gender", 0, 4,
                              ["Male", "Female"], default=data.Gender)

    dob_entry  = lbl_date(basic_lf,  "Date of Birth",      1, 0, data.DateOfBirth)
    adm_date_entry = lbl_date(basic_lf, "Admission Date",  1, 2, data.AdmissionDate)
    reg_number = lbl_entry(basic_lf, "Registration No.",   1, 4, default=data.RegistrationNumber)

    school_name = lbl_entry(basic_lf, "School Name",       2, 0, default=data.SchoolName)
    class_entry = lbl_entry(basic_lf, "Class",             2, 2, default=data.Class)
    intelligence_var, _ = lbl_combo(
        basic_lf, "Child Category", 2, 4,
        ["Single Orphan", "Double Orphan"], default=data.Intelligence)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 1 — CENTER SELECTION (col 0-1)
    # ─────────────────────────────────────────────────────────────────────────
    center_lf = make_section(scroll_frame, "Center Assignment", row=1, col=0, colspan=2)

    current_center_id   = tk.IntVar(value=data.CenterID)
    current_center_name = tk.StringVar()

    center_lookup = {}
    conn2 = get_connection()
    if conn2:
        try:
            cur2 = conn2.cursor()
            cur2.execute("SELECT CenterID, CenterName FROM tblCenters ORDER BY CenterName")
            for r in cur2.fetchall():
                center_lookup[r.CenterName] = r.CenterID
            for name, cid in center_lookup.items():
                if cid == data.CenterID:
                    current_center_name.set(name)
                    break
            if not current_center_name.get() and center_lookup:
                first = next(iter(center_lookup))
                current_center_name.set(first)
                current_center_id.set(center_lookup[first])
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load centers: {e}")
        finally:
            conn2.close()

    tk.Label(center_lf, text="Assigned Center:", font=FONT_LABEL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=0, column=0, sticky="e", **PAD)

    center_combo = ttk.Combobox(center_lf, textvariable=current_center_name,
                                values=list(center_lookup.keys()),
                                font=FONT_ENTRY, width=28, state="readonly")
    center_combo.grid(row=0, column=1, sticky="ew", **PAD)
    center_combo.set(current_center_name.get())

    def handle_center_selection(event=None):
        current_center_id.set(center_lookup.get(center_combo.get(), 0))

    center_combo.bind("<<ComboboxSelected>>", handle_center_selection)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 2 — HEALTH & FATHER INFO (col 0-1) | MOTHER INFO (col 2)
    # ─────────────────────────────────────────────────────────────────────────
    health_lf = make_section(scroll_frame, "Health & Father Details", row=2, col=0, colspan=2)

    disability     = lbl_entry(health_lf, "Disability",          0, 0, default=data.Disability)
    health_condition_var, _ = lbl_combo(
        health_lf, "Health Condition", 0, 2,
        ["Very Good", "Good", "Normal", "Bad"], default=data.HealthCondition)
    reason_father_death = lbl_entry(health_lf, "Reason Father Death", 0, 4,
                                    default=data.ReasonFatherDeath)

    father_death_date  = lbl_date(health_lf,  "Father Death Date", 1, 0,
                                  data.FatherDeathDate)
    father_occupation  = lbl_entry(health_lf, "Father Occupation", 1, 2,
                                   default=data.FatherOccupation)
    father_designation = lbl_entry(health_lf, "Father Designation",1, 4,
                                   default=data.FatherDesignation)

    mother_lf = make_section(scroll_frame, "Mother Details", row=2, col=2)

    mother_name       = lbl_entry(mother_lf, "Mother Name",       0, 0, default=data.MotherName)
    mother_status_var, _ = lbl_combo(
        mother_lf, "Mother Status", 1, 0,
        ["Alive", "Late"], default=data.MotherStatus)
    mother_death_date = lbl_date(mother_lf,  "Mother Death Date", 2, 0,
                                 data.MotherDeathDate)
    mother_cnic       = lbl_entry(mother_lf, "Mother CNIC",       3, 0, default=data.MotherCNIC)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 3 — ADDRESS (col 0-1) | GUARDIAN (col 2)
    # ─────────────────────────────────────────────────────────────────────────
    addr_lf = make_section(scroll_frame, "Address Information", row=3, col=0, colspan=2)

    permanent_address  = lbl_entry(addr_lf, "Permanent Address",  0, 0, width=30,
                                   default=data.PermanentAddress)
    temporary_address  = lbl_entry(addr_lf, "Temporary Address",  0, 2, width=30,
                                   default=data.TemporaryAddress)

    guard_lf = make_section(scroll_frame, "Guardian Details", row=3, col=2)

    guardian_name     = lbl_entry(guard_lf, "Guardian Name",     0, 0, default=data.GuardianName)
    guardian_relation = lbl_entry(guard_lf, "Relation",          1, 0, default=data.GuardianRelation)
    guardian_cnic     = lbl_entry(guard_lf, "Guardian CNIC",     2, 0, default=data.GuardianCNIC)
    guardian_contact  = lbl_entry(guard_lf, "Contact No.",       3, 0, default=data.GuardianContact)
    guardian_address  = lbl_entry(guard_lf, "Guardian Address",  4, 0, width=22,
                                  default=data.Guardianaddress)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 4 — SIBLINGS (full width)
    # ─────────────────────────────────────────────────────────────────────────
    sib_lf = make_section(scroll_frame, "Siblings", row=4, col=0, colspan=3)
    sib_lf.columnconfigure((1, 3, 5), weight=1)

    # Heading row
    for col_idx, heading in enumerate(["#", "Name", "", "Gender", "", "Date of Birth"]):
        tk.Label(sib_lf, text=heading, font=("Segoe UI", 9, "bold"),
                 fg=ACCENT, bg=FRAME_BG).grid(
            row=0, column=col_idx, padx=6, pady=(4, 2))

    name_idx   = [23, 26, 29, 32, 35]
    gender_idx = [24, 27, 30, 33, 36]
    dob_idx    = [25, 28, 31, 34, 37]
    sibling_entries = []

    for i in range(5):
        r = i + 1
        tk.Label(sib_lf, text=f"{i+1}.", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=0, padx=(8, 2))

        s_name = tk.Entry(sib_lf, font=FONT_ENTRY, width=24, relief="solid", bd=1)
        s_name.grid(row=r, column=1, padx=6, pady=3, sticky="ew")
        s_name.insert(0, data[name_idx[i]] if data[name_idx[i]] else "")

        tk.Label(sib_lf, text="Gender:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=2, sticky="e", padx=(12,2))

        s_gender_var = tk.StringVar()
        gender_val = data[gender_idx[i]]
        gender_val = gender_val.strip().capitalize() if gender_val else "Male"
        if gender_val not in ("Male", "Female"):
            gender_val = "Male"
        s_gender_var.set(gender_val)

        s_gender_dd = ttk.Combobox(sib_lf, textvariable=s_gender_var,
                                   values=["Male", "Female"],
                                   state="readonly", font=FONT_ENTRY, width=10)
        s_gender_dd.grid(row=r, column=3, padx=6, pady=3, sticky="ew")

        tk.Label(sib_lf, text="DOB:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=4, sticky="e", padx=(12,2))

        s_dob = DateEntry(sib_lf, width=13, date_pattern="dd/MM/yyyy", font=FONT_ENTRY)
        s_dob.grid(row=r, column=5, padx=6, pady=3, sticky="ew")
        if data[dob_idx[i]]:
            parsed = parse_db_date(data[dob_idx[i]])
            if parsed:
                s_dob.set_date(parsed)

        sibling_entries.append((s_name, s_gender_var, s_dob))

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 5 — WHO CAN MEET (full width)
    # ─────────────────────────────────────────────────────────────────────────
    meet_lf = make_section(scroll_frame, "Who Can Meet", row=5, col=0, colspan=3)
    meet_lf.columnconfigure((1, 3, 5), weight=1)

    for col_idx, heading in enumerate(["#", "Name", "", "CNIC", "", "Contact No."]):
        tk.Label(meet_lf, text=heading, font=("Segoe UI", 9, "bold"),
                 fg=ACCENT, bg=FRAME_BG).grid(
            row=0, column=col_idx, padx=6, pady=(4, 2))

    who_idx  = [38, 41, 44, 47, 50]
    cnic_idx = [39, 42, 45, 48, 51]
    ph_idx   = [40, 43, 46, 49, 52]
    whomeet_entries = []

    for i in range(5):
        r = i + 1
        tk.Label(meet_lf, text=f"{i+1}.", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=0, padx=(8, 2))

        w_name = tk.Entry(meet_lf, font=FONT_ENTRY, width=24, relief="solid", bd=1)
        w_name.grid(row=r, column=1, padx=6, pady=3, sticky="ew")
        w_name.insert(0, data[who_idx[i]] if data[who_idx[i]] else "")

        tk.Label(meet_lf, text="CNIC:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=2, sticky="e", padx=(12,2))

        w_cnic = tk.Entry(meet_lf, font=FONT_ENTRY, width=18, relief="solid", bd=1)
        w_cnic.grid(row=r, column=3, padx=6, pady=3, sticky="ew")
        w_cnic.insert(0, data[cnic_idx[i]] if data[cnic_idx[i]] else "")

        tk.Label(meet_lf, text="Contact:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=4, sticky="e", padx=(12,2))

        w_contact = tk.Entry(meet_lf, font=FONT_ENTRY, width=18, relief="solid", bd=1)
        w_contact.grid(row=r, column=5, padx=6, pady=3, sticky="ew")
        w_contact.insert(0, data[ph_idx[i]] if data[ph_idx[i]] else "")

        whomeet_entries.append((w_name, w_cnic, w_contact))

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 6 — INTRODUCER (full width)
    # ─────────────────────────────────────────────────────────────────────────
    intro_lf = make_section(scroll_frame, "Introducer", row=6, col=0, colspan=3)
    intro_lf.columnconfigure((1, 3, 5), weight=1)

    introducer_name    = lbl_entry(intro_lf, "Name",    0, 0, width=24, default=data.IntroducerName)
    introducer_cnic    = lbl_entry(intro_lf, "CNIC",    0, 2, width=20, default=data.IntroducerCNIC)
    introducer_contact = lbl_entry(intro_lf, "Contact", 0, 4, width=20, default=data.IntroducerContact)

    tk.Label(intro_lf, text="Address:", font=FONT_LABEL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=1, column=0, sticky="e", **PAD)
    introducer_address = tk.Entry(intro_lf, font=FONT_ENTRY, relief="solid", bd=1)
    introducer_address.grid(row=1, column=1, columnspan=5, sticky="ew", **PAD)
    introducer_address.insert(0, data.IntroducerAddress)

    # ─────────────────────────────────────────────────────────────────────────
    # ROW 7 — DOCUMENTS (full width)
    # ─────────────────────────────────────────────────────────────────────────
    def browse_doc(entry_widget):
        filename = filedialog.askopenfilename(
            filetypes=[("Image / PDF files", "*.jpg *.jpeg *.png *.bmp *.gif *.pdf")])
        if filename:
            basename = os.path.basename(filename)
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, f"Photo\\{basename}")

    doc_lf = make_section(scroll_frame, "Documents", row=7, col=0, colspan=3)
    doc_lf.columnconfigure(1, weight=1)
    doc_lf.columnconfigure(4, weight=1)

    doc_fields = [
        ("School Certificate",    data.DocSchoolCertificate),
        ("B-Form",                data.DocBForm),
        ("Father CNIC",           data.DocFatherCNIC),
        ("Mother CNIC",           data.DocMotherCNIC),
        ("Father Death Certificate", data.DocFatherDeathCert),
        ("Other Document 1",      data.OtherDoc1),
        ("Other Document 2",      data.OtherDoc2),
        ("Other Document 3",      data.OtherDoc3),
    ]

    doc_entries = {}
    for idx, (label_text, db_val) in enumerate(doc_fields):
        base_row = (idx // 2) * 1
        base_col = (idx % 2) * 3

        tk.Label(doc_lf, text=label_text + ":", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
            row=base_row, column=base_col, sticky="e", **PAD)

        e = tk.Entry(doc_lf, font=FONT_ENTRY, relief="solid", bd=1)
        e.grid(row=base_row, column=base_col + 1, sticky="ew", **PAD)
        e.insert(0, db_val if db_val else "")

        tk.Button(doc_lf, text="Browse",
                  command=lambda ew=e: browse_doc(ew),
                  font=FONT_LABEL, bg="#e5e7eb", relief="flat",
                  cursor="hand2", padx=6, pady=2).grid(
            row=base_row, column=base_col + 2, padx=(0, 10), pady=4)

        doc_entries[label_text] = e

    # configure doc columns
    for c in [1, 4]:
        doc_lf.columnconfigure(c, weight=1)

    # ═════════════════════════════════════════════════════════════════════════
    # SAVE / UPDATE LOGIC
    # ═════════════════════════════════════════════════════════════════════════
    def format_date(value):
        if value:
            try:
                day, month, year = value.split('/')
                return f"{year}-{month}-{day}"
            except Exception:
                return None
        return None

    def on_update():
        try:
            # Validate required fields
            child_required_amount_val = required_amount_child.get().strip()
            if not child_required_amount_val:
                messagebox.showwarning("Input Error", "Required Amount is mandatory.")
                return
            
            if not full_name.get().strip():
                messagebox.showwarning("Input Error", "Full Name is mandatory.")
                return

            # FIXED: Conditional logic for Mother Death Date
            mother_status = safe_value(mother_status_var.get())
            if mother_status and mother_status.lower() == "alive":
                mother_death_date_value = " "  # Insert space if mother is alive
            else:
                mother_death_date_value = format_date(mother_death_date.get())

            # FIXED: Conditional logic for Father Death Date
            father_death_date_raw = father_death_date.get()
            if not father_death_date_raw or str(father_death_date_raw).strip() == "":
                father_death_date_value = " "  # Insert space if father is alive (no death date)
            else:
                father_death_date_value = format_date(father_death_date_raw)

            # FIXED: Collect siblings data - sibling_entries is defined in outer scope
            siblings_data = []
            for s_name, s_gender_var, s_dob in sibling_entries:
                siblings_data.extend([
                    safe_value(s_name.get()),
                    safe_value(s_gender_var.get()),
                    format_date(s_dob.get()),
                ])

            # FIXED: Collect who can meet data - using whomeet_entries (not wcm_entries)
            meet_persons_data = []
            for wcm_name, wcm_cnic, wcm_contact in whomeet_entries:
                meet_persons_data.extend([
                    safe_value(wcm_name.get()),
                    safe_value(wcm_cnic.get()),
                    safe_value(wcm_contact.get()),
                ])

            # FIXED: Collect document entries - using correct keys
            doc_school_cert = safe_value(doc_entries["School Certificate"].get())
            doc_bform = safe_value(doc_entries["B-Form"].get())
            doc_father_cnic = safe_value(doc_entries["Father CNIC"].get())
            doc_mother_cnic = safe_value(doc_entries["Mother CNIC"].get())
            doc_father_death = safe_value(doc_entries["Father Death Certificate"].get())
            doc_other1 = safe_value(doc_entries["Other Document 1"].get())
            doc_other2 = safe_value(doc_entries["Other Document 2"].get())
            doc_other3 = safe_value(doc_entries["Other Document 3"].get())

            # FIXED: Collect introducer data - all variables are defined in outer scope
            intro_name = safe_value(introducer_name.get())
            intro_cnic = safe_value(introducer_cnic.get())
            intro_contact = safe_value(introducer_contact.get())
            intro_address = safe_value(introducer_address.get())

            # FIXED: Session ID - use data.SessionID from database
            session_id = data.SessionID if data.SessionID else None

            conn = get_connection()
            if not conn:
                messagebox.showerror("Error", "Database connection failed.")
                return

            cursor = conn.cursor()
            
            # UPDATE query with all fields
            cursor.execute("""
                UPDATE tblChildren SET
                    CenterID = ?,
                    FullName = ?,
                    FatherName = ?,
                    Gender = ?,
                    DateOfBirth = ?,
                    AdmissionDate = ?,
                    RegistrationNumber = ?,
                    SchoolName = ?,
                    Class = ?,
                    Intelligence = ?,
                    Disability = ?,
                    HealthCondition = ?,
                    ReasonFatherDeath = ?,
                    FatherDeathDate = ?,
                    FatherOccupation = ?,
                    FatherDesignation = ?,
                    MotherName = ?,
                    MotherStatus = ?,
                    MotherDeathDate = ?,
                    MotherCNIC = ?,
                    PermanentAddress = ?,
                    TemporaryAddress = ?,
                    GuardianName = ?,
                    GuardianRelation = ?,
                    GuardianCNIC = ?,
                    GuardianContact = ?,
                    Guardianaddress = ?,
                    Sibling1Name = ?,
                    Sibling1Gender = ?,
                    Sibling1DOB = ?,
                    Sibling2Name = ?,
                    Sibling2Gender = ?,
                    Sibling2DOB = ?,
                    Sibling3Name = ?,
                    Sibling3Gender = ?,
                    Sibling3DOB = ?,
                    Sibling4Name = ?,
                    Sibling4Gender = ?,
                    Sibling4DOB = ?,
                    Sibling5Name = ?,
                    Sibling5Gender = ?,
                    Sibling5DOB = ?,
                    MeetPerson1Name = ?,
                    MeetPerson1CNIC = ?,
                    MeetPerson1Contact = ?,
                    MeetPerson2Name = ?,
                    MeetPerson2CNIC = ?,
                    MeetPerson2Contact = ?,
                    MeetPerson3Name = ?,
                    MeetPerson3CNIC = ?,
                    MeetPerson3Contact = ?,
                    MeetPerson4Name = ?,
                    MeetPerson4CNIC = ?,
                    MeetPerson4Contact = ?,
                    MeetPerson5Name = ?,
                    MeetPerson5CNIC = ?,
                    MeetPerson5Contact = ?,
                    IntroducerName = ?,
                    IntroducerCNIC = ?,
                    IntroducerContact = ?,
                    IntroducerAddress = ?,
                    DocSchoolCertificate = ?,
                    DocBForm = ?,
                    DocFatherCNIC = ?,
                    DocMotherCNIC = ?,
                    DocFatherDeathCert = ?,
                    OtherDoc1 = ?,
                    OtherDoc2 = ?,
                    OtherDoc3 = ?,
                    Status = ?,
                    PhotoPath = ?,
                    ChildRequiredAmount = ?,
                    SessionID = ?
                WHERE ChildID = ?
            """, (
                current_center_id.get(),
                safe_value(full_name.get()),
                safe_value(father_name.get()),
                safe_value(gender_var.get()),
                format_date(dob_entry.get()),
                format_date(adm_date_entry.get()),
                safe_value(reg_number.get()),
                safe_value(school_name.get()),
                safe_value(class_entry.get()),
                safe_value(intelligence_var.get()),
                safe_value(disability.get()),
                safe_value(health_condition_var.get()),
                safe_value(reason_father_death.get()),
                father_death_date_value,  # FIXED: Conditional value
                safe_value(father_occupation.get()),
                safe_value(father_designation.get()),
                safe_value(mother_name.get()),
                mother_status,
                mother_death_date_value,  # FIXED: Conditional value
                safe_value(mother_cnic.get()),
                safe_value(permanent_address.get()),
                safe_value(temporary_address.get()),
                safe_value(guardian_name.get()),
                safe_value(guardian_relation.get()),
                safe_value(guardian_cnic.get()),
                safe_value(guardian_contact.get()),
                safe_value(guardian_address.get()),
                *siblings_data,       # 15 values (5 siblings × 3 fields)
                *meet_persons_data,   # 15 values (5 persons × 3 fields)
                intro_name,
                intro_cnic,
                intro_contact,
                intro_address,
                doc_school_cert,
                doc_bform,
                doc_father_cnic,
                doc_mother_cnic,
                doc_father_death,
                doc_other1,
                doc_other2,
                doc_other3,
                is_active_child.get(),
                photo_path.get().strip(),
                child_required_amount_val,
                session_id,
                child_id  # WHERE clause - the ID of the child being edited
            ))

            conn.commit()
            conn.close()
            
            messagebox.showinfo("Success", "Child record updated successfully!")
            edit_win.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to update child:\n{e}")
            import traceback
            traceback.print_exc()  # Print full error for debugging
    def delete_child(child_id):
        """Delete child record with confirmation"""
        result = messagebox.askyesno(
            "Confirm Deletion",
            f"Are you sure you want to delete this child record?\n\n"
            f"Child ID: {child_id}\n"
            f"Name: {full_name.get()}\n\n"
            f"This action cannot be undone!"
        )
        
        if not result:
            return
        
        try:
            conn = get_connection()
            if not conn:
                messagebox.showerror("Error", "Database connection failed.")
                return
            
            cursor = conn.cursor()
            
            # Check if child has active sponsorships
            cursor.execute("""
                SELECT COUNT(*) FROM tblSponsorships 
                WHERE ChildID = ? AND IsActive = 1
            """, (child_id,))
            
            active_sponsorships = cursor.fetchone()[0]
            
            if active_sponsorships > 0:
                messagebox.showwarning(
                    "Cannot Delete",
                    f"This child has {active_sponsorships} active sponsorship(s).\n\n"
                    f"Please deactivate or remove sponsorships first."
                )
                conn.close()
                return
            
            # Delete the child
            cursor.execute("DELETE FROM tblChildren WHERE ChildID = ?", (child_id,))
            conn.commit()
            conn.close()
            
            messagebox.showinfo("Deleted", "Child record deleted successfully!")
            edit_win.destroy()
            
            # Refresh parent window
            #if 'load_treeview' in globals():
            #    load_treeview()
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete child:\n{e}")
    # ═════════════════════════════════════════════════════════════════════════
    # FOOTER BUTTONS
    # ═════════════════════════════════════════════════════════════════════════
    
    tk.Button(
        footer, text="✔  Save Changes",
        command=on_update,
        font=FONT_BTN, bg=BTN_SAVE_BG, fg="white",
        relief="flat", cursor="hand2",
        padx=24, pady=8
    ).pack(side="left", padx=20)

    tk.Button(
        footer, 
        text="Cancel", 
        command=edit_win.destroy, 
        bg="gray", 
        fg="white", 
        width=20,
        font=("Arial", 11, "bold")
    ).pack(side="left", padx=10)

    tk.Button(
        footer, 
        text="Delete Child", 
        command=lambda: delete_child(child_id), 
        bg="red", 
        fg="white", 
        width=20,
        font=("Arial", 11, "bold")
    ).pack(side="left", padx=10)

    tk.Button(
        footer, text="✖  Close",
        command=edit_win.destroy,
        font=FONT_BTN, bg=BTN_CLOSE_BG, fg="white",
        relief="flat", cursor="hand2",
        padx=24, pady=8
    ).pack(side="right", padx=20)

    edit_win.bind('<Return>', lambda e: on_update())
    edit_win.bind('<Escape>', lambda e: edit_win.destroy())

    # Scroll to top after rendering
    edit_win.update_idletasks()
    canvas.yview_moveto(0)

#\\\\\\\\\\\Open Child List to edit child information\\\\\\\\\\\\\\\\\\\\\\\\\
def open_child_list():
    list_win = tk.Toplevel(root)
    list_win.title("Edit Child" + ORG_SUFFIX)
    #list_win.geometry("900x1000")
    list_win.state("zoomed")   # Opens window maximized
    list_win.grab_set()

    tk.Label(list_win, text="Search Child:", font=("Arial", 12)).pack(pady=5)

    search_var = tk.StringVar()
    search_entry = tk.Entry(list_win, textvariable=search_var, font=("Arial", 12), width=50)
    search_entry.pack(pady=5)

    # Treeview
    columns = ("ChildID", "FullName", "FatherName")
    tree = ttk.Treeview(list_win, columns=columns, show="headings", height=20)

    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=280)

    tree.pack(fill="both", expand=True, padx=10, pady=10)

    # ---------------- LOAD CHILDREN ----------------
    def load_children(filter_text=""):
        tree.delete(*tree.get_children())

        conn = get_connection()
        if conn:
            cursor = conn.cursor()
            query = "SELECT ChildID, FullName, FatherName FROM tblChildren"

            params = []
            if filter_text:
                query += " WHERE FullName LIKE ? OR FatherName LIKE ?"
                params = [f"%{filter_text}%", f"%{filter_text}%"]

            query += " ORDER BY FullName"
            cursor.execute(query, params)

            rows = cursor.fetchall()
            for r in rows:
                tree.insert("", "end", values=(r.ChildID, r.FullName, r.FatherName))

            conn.close()

        # 🔹 AUTO-SELECT FIRST ROW
        children = tree.get_children()
        if children:
            tree.selection_set(children[0])
            tree.focus(children[0])

    # Initial load
    load_children()
    search_entry.focus_set()

    # ---------------- SEARCH HANDLER ----------------
    def on_search(*args):
        load_children(search_var.get().strip())

    search_var.trace_add("write", on_search)

    # ---------------- EDIT HANDLER ----------------
    def edit_selected(event=None):
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select a child to edit.")
            return

        child_id = tree.item(selected[0], "values")[0]
        list_win.destroy()
        open_edit_child_form(child_id)

    # Buttons
    btn_frame = tk.Frame(list_win)
    btn_frame.pack(pady=10)

    edit_btn = tk.Button(
        btn_frame,
        text="Edit",
        command=edit_selected,
        font=("Arial", 12),
        width=15
    )
    edit_btn.grid(row=0, column=0, padx=5)

    tk.Button(
        btn_frame,
        text="Close",
        command=list_win.destroy,
        font=("Arial", 12),
        width=10
    ).grid(row=0, column=1, padx=5)

    # ---------------- KEYBOARD SHORTCUTS ----------------
    list_win.bind("<Escape>", lambda e: list_win.destroy())
    list_win.bind("<Return>", edit_selected)   # Enter = Edit
    tree.bind("<Double-1>", edit_selected)     # Double-click = Edit




# Define all placeholder functions first
#def open_child_mgmt():
 #   open_placeholder("Child Management")

#def open_donor_mgmt():
  #  open_placeholder("Donor Management")


    
# ============================================================
# Child Management
# ============================================================

def open_child_mgmt():

    child_win = tk.Toplevel(root)
    child_win.withdraw()
    child_win.title("Add New Child" + ORG_SUFFIX)

    screen_w = child_win.winfo_screenwidth()
    screen_h = child_win.winfo_screenheight()
    win_w    = int(screen_w * 0.99)
    win_h    = int(screen_h * 0.85)
    pos_x    = (screen_w - win_w) // 2
    pos_y    = (screen_h - win_h) // 2
    child_win.geometry(f"{win_w}x{win_h}+{pos_x}+{pos_y}")
    child_win.grab_set()
    child_win.transient(root)

    conn = get_connection()
    cursor = conn.cursor()

    # Get all table names
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name")
    tables = [row[0] for row in cursor.fetchall()]

    for table in tables:
        print(f"\n{'='*50}")
        print(f"TABLE: {table}")
        print(f"{'='*50}")
        cursor.execute(f"PRAGMA table_info({table})")
        for col in cursor.fetchall():
            pk = " [PK]" if col['pk'] else ""
            null = " NOT NULL" if col['notnull'] else ""
            default = f" DEFAULT={col['dflt_value']}" if col['dflt_value'] is not None else ""
            print(f"  {col['cid']:>3}  {col['name']:<35} {col['type']:<15}{pk}{null}{default}")

    conn.close()


    # ── Colour & font constants ──────────────────────────────────────────────
    HEADER_BG    = "#1a3c5e"
    HEADER_FG    = "#ffffff"
    SECTION_BG   = "#f0f4f8"
    FRAME_BG     = "#ffffff"
    LABEL_FG     = "#374151"
    ACCENT       = "#2563eb"
    BTN_ADD_BG   = "#16a34a"
    BTN_EDIT_BG  = "#2563eb"
    BTN_CLOSE_BG = "#dc2626"
    FONT_LABEL   = ("Segoe UI", 9)
    FONT_ENTRY   = ("Segoe UI", 9)
    FONT_TITLE   = ("Segoe UI", 10, "bold")
    FONT_HEADER  = ("Segoe UI", 13, "bold")
    FONT_BTN     = ("Segoe UI", 10, "bold")
    PAD          = {"padx": 6, "pady": 4}

    # ── Helper: styled LabelFrame section ────────────────────────────────────
    def make_section(parent, title, row, col=0, colspan=1, rowspan=1):
        lf = tk.LabelFrame(
            parent, text=f"  {title}  ",
            font=FONT_TITLE, fg=ACCENT, bg=FRAME_BG,
            bd=1, relief="groove", labelanchor="nw"
        )
        lf.grid(row=row, column=col, columnspan=colspan, rowspan=rowspan,
                sticky="nsew", padx=8, pady=6)
        lf.columnconfigure((0, 2, 4), weight=0)
        lf.columnconfigure((1, 3, 5), weight=1)
        return lf

    # ── Helper: label + entry pair ────────────────────────────────────────────
    def lbl_entry(parent, text, row, col, width=22, default=""):
        tk.Label(parent, text=text, font=FONT_LABEL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        e = tk.Entry(parent, font=FONT_ENTRY, width=width,
                     relief="solid", bd=1)
        e.grid(row=row, column=col + 1, sticky="ew", **PAD)
        if default:
            e.insert(0, default)
        return e

    # ── Helper: label + combobox pair ─────────────────────────────────────────
    def lbl_combo(parent, text, row, col, values, default="", width=20):
        tk.Label(parent, text=text, font=FONT_LABEL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        var = tk.StringVar(value=default)
        cb = ttk.Combobox(parent, textvariable=var, values=values,
                          state="readonly", font=FONT_ENTRY, width=width)
        cb.grid(row=row, column=col + 1, sticky="ew", **PAD)
        return var, cb

    # ── Helper: label + DateEntry pair ────────────────────────────────────────
    def lbl_date(parent, text, row, col, width=14):
        tk.Label(parent, text=text, font=FONT_LABEL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        de = DateEntry(parent, date_pattern="dd/MM/yyyy",
                       font=FONT_ENTRY, width=width)
        de.grid(row=row, column=col + 1, sticky="ew", **PAD)
        return de

    # ═════════════════════════════════════════════════════════════════════════
    # OUTER LAYOUT  — header / scrollable body / footer
    # ═════════════════════════════════════════════════════════════════════════

    # ── Header bar ────────────────────────────────────────────────────────────
    header = tk.Frame(child_win, bg=HEADER_BG, height=50)
    header.pack(fill="x", side="top")
    tk.Label(header, text="➕  Add New Child",
             font=FONT_HEADER, bg=HEADER_BG, fg=HEADER_FG,
             pady=10).pack(side="left", padx=18)

    # ── Footer bar (packed before body so it stays fixed at bottom) ───────────
    footer = tk.Frame(child_win, bg="#e5e7eb", pady=10)
    footer.pack(fill="x", side="bottom")

    # ── Scrollable body ───────────────────────────────────────────────────────
    body_frame = tk.Frame(child_win, bg=SECTION_BG)
    body_frame.pack(fill="both", expand=True)

    canvas   = tk.Canvas(body_frame, bg=SECTION_BG, highlightthickness=0)
    v_scroll = tk.Scrollbar(body_frame, orient="vertical",   command=canvas.yview)
    h_scroll = tk.Scrollbar(body_frame, orient="horizontal", command=canvas.xview)

    h_scroll.pack(side="bottom", fill="x")
    v_scroll.pack(side="right",  fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    canvas.configure(yscrollcommand=v_scroll.set, xscrollcommand=h_scroll.set)

    scroll_frame = tk.Frame(canvas, bg=SECTION_BG)
    win_id = canvas.create_window((0, 0), window=scroll_frame, anchor="nw")

    def _on_frame_configure(e):
        canvas.configure(scrollregion=canvas.bbox("all"))

    def _on_canvas_configure(e):
        canvas.itemconfig(win_id, width=max(e.width, scroll_frame.winfo_reqwidth()))

    scroll_frame.bind("<Configure>", _on_frame_configure)
    canvas.bind("<Configure>",       _on_canvas_configure)

    def _mousewheel(e):
        try:
            canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        except tk.TclError:
            pass

    canvas.bind_all("<MouseWheel>", _mousewheel)

    # ── Content grid columns ──────────────────────────────────────────────────
    scroll_frame.columnconfigure((0, 1, 2), weight=1)

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 0 — PHOTO (col 2, rowspan 2)  +  BASIC INFO (col 0-1)
    # ═════════════════════════════════════════════════════════════════════════

    # ── PHOTO SECTION ─────────────────────────────────────────────────────────
    photo_lf = make_section(scroll_frame, "Photo", row=0, col=2, rowspan=2)

    photo_container = tk.Frame(photo_lf, width=140, height=140,
                               bg="#d1d5db", relief="solid", bd=1)
    photo_container.grid(row=0, column=0, columnspan=2, padx=10, pady=8)
    photo_container.pack_propagate(False)

    photo_label = tk.Label(photo_container, text="No Image",
                           bg="#d1d5db", font=FONT_LABEL)
    photo_label.pack(expand=True)

    photo_path = tk.Entry(photo_lf, font=FONT_ENTRY, width=24,
                          relief="solid", bd=1)
    photo_path.grid(row=1, column=0, columnspan=2, padx=10, pady=2, sticky="ew")

    '''def load_photo(path):
        candidates = [
            path,
            os.path.join("Photo", os.path.basename(path)),
            os.path.join("photo", os.path.basename(path)),
        ]
        for p in candidates:
            if p and os.path.exists(p):
                try:
                    img = Image.open(p)
                    img.thumbnail((140, 140))
                    img_tk = ImageTk.PhotoImage(img)
                    photo_label.configure(image=img_tk, text="", bg="white")
                    photo_label.image = img_tk
                    photo_container.configure(bg="white")
                    return
                except Exception:
                    pass
        photo_label.configure(text="No Image", image="", bg="#d1d5db")
        photo_container.configure(bg="#d1d5db")'''
    
    def load_photo(path):
        candidates = []

        if path and path.strip():
            # 1. Absolute path as-is
            candidates.append(path.strip())
            # 2. Just the filename inside the server's Photo folder
            candidates.append(os.path.join(PHOTO_DIR, os.path.basename(path.strip())))

        # 3. Fallback placeholder
        candidates.append(os.path.join(PHOTO_DIR, "placeholder.png"))

        for p in candidates:
            if p and os.path.exists(p):
                try:
                    img = Image.open(p)
                    img.thumbnail((140, 140))
                    img_tk = ImageTk.PhotoImage(img)
                    photo_label.configure(image=img_tk, text="", bg="white")
                    photo_label.image = img_tk
                    photo_container.configure(bg="white")
                    return
                except Exception:
                    pass

        photo_label.configure(text="No Image", image="", bg="#d1d5db")
        photo_container.configure(bg="#d1d5db")


    # Load placeholder on startup
    load_photo(os.path.join(PHOTO_DIR, "placeholder.png"))

    # Load placeholder on startup
    #load_photo("photo/placeholder.png")

    '''def browse_photo():
        file_path = filedialog.askopenfilename(
            filetypes=[("Image Files", "*.jpg *.jpeg *.png")])
        if file_path:
            try:
                img = Image.open(file_path)
                img.thumbnail((140, 140))
                img_tk = ImageTk.PhotoImage(img)
                photo_label.configure(image=img_tk, text="", bg="white")
                photo_label.image = img_tk
                photo_container.configure(bg="white")
                filename = os.path.basename(file_path)
                photo_path.delete(0, tk.END)
                photo_path.insert(0, f"Photo\\{filename}")
            except Exception as e:
                messagebox.showerror("Photo Error", f"Failed to load photo: {e}")'''
    
    def browse_photo():
        file_path = filedialog.askopenfilename(
            filetypes=[("Image Files", "*.jpg *.jpeg *.png")])
        if file_path:
            try:
                img = Image.open(file_path)
                img.thumbnail((140, 140))
                img_tk = ImageTk.PhotoImage(img)
                photo_label.configure(image=img_tk, text="", bg="white")
                photo_label.image = img_tk
                photo_container.configure(bg="white")

                filename = os.path.basename(file_path)
                dest_path = os.path.join(PHOTO_DIR, filename)

                # Copy photo to server's Photo folder if not already there
                if os.path.abspath(file_path) != os.path.abspath(dest_path):
                    import shutil
                    os.makedirs(PHOTO_DIR, exist_ok=True)
                    shutil.copy2(file_path, dest_path)

                # Save only the filename (portable across machines)
                photo_path.delete(0, tk.END)
                photo_path.insert(0, f"Photo\\{filename}")

            except Exception as e:
                messagebox.showerror("Photo Error", f"Failed to load photo: {e}")


    tk.Button(photo_lf, text="📁  Browse Photo",
              command=browse_photo, font=FONT_LABEL,
              bg=ACCENT, fg="white", relief="flat",
              cursor="hand2", padx=8, pady=4).grid(
        row=2, column=0, columnspan=2, pady=(6, 4))

    # Active checkbox + Required Amount inside photo frame
    is_active_child = tk.IntVar(value=1)
    tk.Checkbutton(photo_lf, text="Active Record",
                   variable=is_active_child,
                   font=FONT_LABEL, bg=FRAME_BG, fg=LABEL_FG,
                   activebackground=FRAME_BG).grid(
        row=3, column=0, columnspan=2, pady=3)

    tk.Label(photo_lf, text="Required Amount (PKR):",
             font=FONT_LABEL, fg=LABEL_FG, bg=FRAME_BG).grid(
        row=4, column=0, sticky="e", **PAD)
    required_amount_child = tk.StringVar(value="20000")
    tk.Entry(photo_lf, textvariable=required_amount_child,
             font=FONT_ENTRY, width=12,
             relief="solid", bd=1).grid(
        row=4, column=1, sticky="ew", **PAD)

    # ── BASIC INFORMATION ─────────────────────────────────────────────────────
    basic_lf = make_section(scroll_frame, "Basic Information", row=0, col=0, colspan=2)

    # Auto-generate registration number
    def get_next_registration_number():
        try:
            import re
            from tkinter import simpledialog

            conn = get_connection()
            cursor = conn.cursor()

            # Get last registration number
            cursor.execute("""
                SELECT RegistrationNumber 
                FROM tblChildren
                WHERE RegistrationNumber IS NOT NULL
                ORDER BY ChildID DESC LIMIT 1
            """)
            r = cursor.fetchone()

            # ───── CASE 1: RECORD EXISTS ─────
            if r and r.RegistrationNumber:
                reg = r.RegistrationNumber.strip()

                # Extract prefix (everything before last hyphen part)
                parts = reg.split("-")
                if len(parts) >= 2:
                    prefix = "-".join(parts[:-1]) + "-"
                    last_num = int(parts[-1])
                    conn.close()
                    return f"{prefix}{last_num + 1:03d}"

            conn.close()

            # ───── CASE 2: NO RECORD FOUND ─────
            messagebox.showinfo("Child's Aghosh ID Prefix Required (e.g. AGH-KHI-)", "No previous record found.\nPlease enter prefix (e.g. AGH-KHI-)")

            while True:
                user_prefix = simpledialog.askstring(
                    "Enter Prefix",
                    "Enter prefix (max 8 chars, e.g. AGH-KHI-):"
                )

                if user_prefix is None:
                    return ""  # user cancelled

                user_prefix = user_prefix.strip().upper()

                # Validate length
                if len(user_prefix) > 8:
                    messagebox.showerror("Invalid", "Prefix must be maximum 8 characters.")
                    continue

                # Validate format (must end with hyphen)
                if not user_prefix.endswith("-"):
                    messagebox.showerror("Invalid", "Prefix must end with '-' (e.g. AGH-KHI-)")
                    continue

                # Optional: basic pattern check (letters + hyphen)
                if not re.match(r'^[A-Z\-]+$', user_prefix):
                    messagebox.showerror("Invalid", "Only alphabets and '-' allowed.")
                    continue

                return f"{user_prefix}001"

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate registration number:\n{e}")
            return ""

    full_name   = lbl_entry(basic_lf, "Full Name",         0, 0)
    father_name = lbl_entry(basic_lf, "Father Name",       0, 2)
    gender_var, _ = lbl_combo(basic_lf, "Gender",          0, 4,
                              ["Male", "Female"], default="Female")

    dob_entry      = lbl_date(basic_lf,  "Date of Birth",  1, 0)
    adm_date_entry = lbl_date(basic_lf,  "Admission Date", 1, 2)
    reg_number     = lbl_entry(basic_lf, "Registration No.", 1, 4,
                               default=get_next_registration_number())

    school_name  = lbl_entry(basic_lf, "School Name",      2, 0)
    class_entry  = lbl_entry(basic_lf, "Class",            2, 2)
    intelligence_var, _ = lbl_combo(
        basic_lf, "Child Category", 2, 4,
        ["Single Orphan", "Double Orphan"], default="Single Orphan")

    full_name.focus_set()

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 1 — CENTER & SESSION (col 0-1)
    # ═════════════════════════════════════════════════════════════════════════
    center_lf = make_section(scroll_frame, "Center & Session", row=1, col=0, colspan=2)
    center_lf.columnconfigure((1, 3), weight=1)

    # Center
    current_center_id   = tk.IntVar()
    current_center_name = tk.StringVar()
    center_lookup = {}

    conn = get_connection()
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute("SELECT CenterID, CenterName FROM tblCenters ORDER BY CenterName")
            for r in cursor.fetchall():
                center_lookup[r.CenterName] = r.CenterID
            if center_lookup:
                first = next(iter(center_lookup))
                current_center_name.set(first)
                current_center_id.set(center_lookup[first])
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load centers: {e}")
        finally:
            conn.close()

    tk.Label(center_lf, text="Assigned Center:", font=FONT_LABEL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=0, column=0, sticky="e", **PAD)
    center_combo = ttk.Combobox(center_lf, textvariable=current_center_name,
                                values=list(center_lookup.keys()),
                                font=FONT_ENTRY, width=24, state="readonly")
    center_combo.grid(row=0, column=1, sticky="ew", **PAD)
    center_combo.set(current_center_name.get())

    def handle_center_selection(event=None):
        current_center_id.set(center_lookup.get(center_combo.get(), 0))

    center_combo.bind("<<ComboboxSelected>>", handle_center_selection)

    # Session
    current_session_id   = tk.IntVar(value=0)
    current_session_name = tk.StringVar()
    session_lookup = {}

    tk.Label(center_lf, text="Session:", font=FONT_LABEL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=0, column=2, sticky="e", **PAD)
    session_combo = ttk.Combobox(center_lf, textvariable=current_session_name,
                                 font=FONT_ENTRY, width=24, state="readonly")
    session_combo.grid(row=0, column=3, sticky="ew", **PAD)

    def handle_session_selection(event=None):
        name = current_session_name.get()
        current_session_id.set(session_lookup.get(name, 0))

    session_combo.bind("<<ComboboxSelected>>", handle_session_selection)

    def load_sessions():
        try:
            conn = get_connection()
            cur = conn.cursor()
            cur.execute("""
                SELECT SessionID, SessionName, IsCurrentSession
                FROM tblSessions ORDER BY StartDate DESC
            """)
            rows = cur.fetchall()
            conn.close()
            session_lookup.clear()
            names, default_name = [], None
            for r in rows:
                session_lookup[r.SessionName] = r.SessionID
                names.append(r.SessionName)
                if r.IsCurrentSession:
                    default_name = r.SessionName
            session_combo['values'] = names
            if names:
                chosen = default_name or names[0]
                session_combo.set(chosen)
                current_session_name.set(chosen)
                current_session_id.set(session_lookup[chosen])
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load sessions:\n{e}")

    load_sessions()

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 2 — HEALTH & FATHER (col 0-1)  +  MOTHER (col 2)
    # ═════════════════════════════════════════════════════════════════════════
    health_lf = make_section(scroll_frame, "Health & Father Details", row=2, col=0, colspan=2)

    disability         = lbl_entry(health_lf, "Disability",           0, 0)
    health_condition_var, _ = lbl_combo(
        health_lf, "Health Condition", 0, 2,
        ["Very Good", "Good", "Normal", "Bad"], default=" ")
    reason_father_death = lbl_entry(health_lf, "Reason Father Death", 0, 4,
                                    default="Medical")

    father_death_date  = lbl_date(health_lf,  "Father Death Date",    1, 0)
    father_occupation  = lbl_entry(health_lf, "Father Occupation",    1, 2)
    father_designation = lbl_entry(health_lf, "Father Designation",   1, 4)

    mother_lf = make_section(scroll_frame, "Mother Details", row=2, col=2)

    mother_name = lbl_entry(mother_lf, "Mother Name",   0, 0)

    mother_status_var, mother_status_cb = lbl_combo(
        mother_lf, "Mother Status", 1, 0,
        ["Alive", "Late"], default="Alive")

    tk.Label(mother_lf, text="Mother Death Date:", font=FONT_LABEL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=2, column=0, sticky="e", **PAD)
    mother_death_date = DateEntry(mother_lf, date_pattern="dd/MM/yyyy",
                                  font=FONT_ENTRY, width=14)
    mother_death_date.grid(row=2, column=1, sticky="ew", **PAD)

    mother_cnic = lbl_entry(mother_lf, "Mother CNIC", 3, 0)

    def on_mother_status_change(event=None):
        if mother_status_var.get() == "Alive":
            mother_death_date.config(state="disabled")
        else:
            mother_death_date.config(state="normal")

    mother_status_cb.bind("<<ComboboxSelected>>", on_mother_status_change)
    on_mother_status_change()

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 3 — ADDRESS (col 0-1)  +  GUARDIAN (col 2)
    # ═════════════════════════════════════════════════════════════════════════
    addr_lf = make_section(scroll_frame, "Address Information", row=3, col=0, colspan=2)

    permanent_address = lbl_entry(addr_lf, "Permanent Address", 0, 0, width=30)
    temporary_address = lbl_entry(addr_lf, "Temporary Address", 0, 2, width=30)

    guard_lf = make_section(scroll_frame, "Guardian Details", row=3, col=2)

    guardian_name     = lbl_entry(guard_lf, "Guardian Name",    0, 0)
    guardian_relation = lbl_entry(guard_lf, "Relation",         1, 0)
    guardian_cnic     = lbl_entry(guard_lf, "Guardian CNIC",    2, 0)
    guardian_contact  = lbl_entry(guard_lf, "Contact No.",      3, 0)
    guardian_address  = lbl_entry(guard_lf, "Guardian Address", 4, 0, width=22)

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 4 — SIBLINGS
    # ═════════════════════════════════════════════════════════════════════════
    sib_lf = make_section(scroll_frame, "Siblings", row=4, col=0, colspan=3)
    sib_lf.columnconfigure((1, 3, 5), weight=1)

    for col_idx, heading in enumerate(["#", "Name", "", "Gender", "", "Date of Birth"]):
        tk.Label(sib_lf, text=heading, font=("Segoe UI", 9, "bold"),
                 fg=ACCENT, bg=FRAME_BG).grid(
            row=0, column=col_idx, padx=6, pady=(4, 2))

    sibling_entries = []
    for i in range(5):
        r = i + 1
        tk.Label(sib_lf, text=f"{i+1}.", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=0, padx=(8, 2))

        s_name = tk.Entry(sib_lf, font=FONT_ENTRY, width=24, relief="solid", bd=1)
        s_name.grid(row=r, column=1, padx=6, pady=3, sticky="ew")

        tk.Label(sib_lf, text="Gender:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=2, sticky="e", padx=(12, 2))

        s_gender_var = tk.StringVar(value="Female")
        s_gender_dd = ttk.Combobox(sib_lf, textvariable=s_gender_var,
                                   values=["Male", "Female"],
                                   state="readonly", font=FONT_ENTRY, width=10)
        s_gender_dd.grid(row=r, column=3, padx=6, pady=3, sticky="ew")

        tk.Label(sib_lf, text="DOB:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=4, sticky="e", padx=(12, 2))

        s_dob = DateEntry(sib_lf, width=13, date_pattern="dd/MM/yyyy", font=FONT_ENTRY)
        s_dob.grid(row=r, column=5, padx=6, pady=3, sticky="ew")

        sibling_entries.append((s_name, s_gender_var, s_dob))

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 5 — WHO CAN MEET
    # ═════════════════════════════════════════════════════════════════════════
    meet_lf = make_section(scroll_frame, "Who Can Meet", row=5, col=0, colspan=3)
    meet_lf.columnconfigure((1, 3, 5), weight=1)

    for col_idx, heading in enumerate(["#", "Name", "", "CNIC", "", "Contact No."]):
        tk.Label(meet_lf, text=heading, font=("Segoe UI", 9, "bold"),
                 fg=ACCENT, bg=FRAME_BG).grid(
            row=0, column=col_idx, padx=6, pady=(4, 2))

    wcm_entries = []
    for i in range(5):
        r = i + 1
        tk.Label(meet_lf, text=f"{i+1}.", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=0, padx=(8, 2))

        wcm_name = tk.Entry(meet_lf, font=FONT_ENTRY, width=24, relief="solid", bd=1)
        wcm_name.grid(row=r, column=1, padx=6, pady=3, sticky="ew")

        tk.Label(meet_lf, text="CNIC:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=2, sticky="e", padx=(12, 2))

        wcm_cnic = tk.Entry(meet_lf, font=FONT_ENTRY, width=18, relief="solid", bd=1)
        wcm_cnic.grid(row=r, column=3, padx=6, pady=3, sticky="ew")

        tk.Label(meet_lf, text="Contact:", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG).grid(row=r, column=4, sticky="e", padx=(12, 2))

        wcm_contact = tk.Entry(meet_lf, font=FONT_ENTRY, width=18, relief="solid", bd=1)
        wcm_contact.grid(row=r, column=5, padx=6, pady=3, sticky="ew")

        wcm_entries.append((wcm_name, wcm_cnic, wcm_contact))

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 6 — INTRODUCER
    # ═════════════════════════════════════════════════════════════════════════
    intro_lf = make_section(scroll_frame, "Introducer", row=6, col=0, colspan=3)
    intro_lf.columnconfigure((1, 3, 5), weight=1)

    introducer_name    = lbl_entry(intro_lf, "Name",    0, 0, width=24)
    introducer_cnic    = lbl_entry(intro_lf, "CNIC",    0, 2, width=20)
    introducer_contact = lbl_entry(intro_lf, "Contact", 0, 4, width=20)

    tk.Label(intro_lf, text="Address:", font=FONT_LABEL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=1, column=0, sticky="e", **PAD)
    introducer_address = tk.Entry(intro_lf, font=FONT_ENTRY, relief="solid", bd=1)
    introducer_address.grid(row=1, column=1, columnspan=5, sticky="ew", **PAD)

    # ═════════════════════════════════════════════════════════════════════════
    # ROW 7 — DOCUMENTS
    # ═════════════════════════════════════════════════════════════════════════
    def browse_doc(entry_widget):
        filename = filedialog.askopenfilename(
            filetypes=[("Image / PDF files", "*.jpg *.jpeg *.png *.bmp *.gif *.pdf")])
        if filename:
            basename = os.path.basename(filename)
            entry_widget.delete(0, tk.END)
            entry_widget.insert(0, f"Photo\\{basename}")

    doc_lf = make_section(scroll_frame, "Documents", row=7, col=0, colspan=3)
    for c in [1, 4]:
        doc_lf.columnconfigure(c, weight=1)

    doc_field_names = [
        "Admission Form",
        "B-Form",
        "Father CNIC",
        "Mother CNIC",
        "Father Death Certificate",
        "Other Document 1",
        "Other Document 2",
        "Other Document 3",
    ]

    doc_entries = {}
    for idx, label_text in enumerate(doc_field_names):
        base_row = idx // 2
        base_col = (idx % 2) * 3

        tk.Label(doc_lf, text=label_text + ":", font=FONT_LABEL,
                 fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
            row=base_row, column=base_col, sticky="e", **PAD)

        e = tk.Entry(doc_lf, font=FONT_ENTRY, relief="solid", bd=1)
        e.grid(row=base_row, column=base_col + 1, sticky="ew", **PAD)

        tk.Button(doc_lf, text="Browse",
                  command=lambda ew=e: browse_doc(ew),
                  font=FONT_LABEL, bg="#e5e7eb", relief="flat",
                  cursor="hand2", padx=6, pady=2).grid(
            row=base_row, column=base_col + 2, padx=(0, 10), pady=4)

        doc_entries[label_text] = e

    # ═════════════════════════════════════════════════════════════════════════
    # ON ADD — save to DB
    # ═════════════════════════════════════════════════════════════════════════
    def format_date(value):
        if value:
            try:
                day, month, year = value.split('/')
                return f"{year}-{month}-{day}"
            except Exception:
                return None
        return None

    def safe_value(val):
        if val is None:
            return " "
        val = str(val).strip()
        return val if val != "" else " "

    def on_add():
        try:
            # Validate required field
            child_required_amount_val = required_amount_child.get().strip()
            if not child_required_amount_val:
                messagebox.showwarning("Input Error", "Required Amount is mandatory.")
                return

            # Collect siblings
            siblings_data = []
            for s_name, s_gender_var, s_dob in sibling_entries:
                siblings_data.extend([
                    safe_value(s_name.get()),
                    safe_value(s_gender_var.get()),
                    format_date(s_dob.get()),
                ])

            # Collect who can meet
            meet_persons_data = []
            for wcm_name, wcm_cnic, wcm_contact in wcm_entries:
                meet_persons_data.extend([
                    safe_value(wcm_name.get()),
                    safe_value(wcm_cnic.get()),
                    safe_value(wcm_contact.get()),
                ])

            # FIXED: Conditional logic for Mother Death Date
            mother_status = safe_value(mother_status_var.get())
            if mother_status and mother_status.lower() == "alive":
                mother_death_date_value = " "  # Insert space if mother is alive
            else:
                mother_death_date_value = format_date(mother_death_date.get())

            # FIXED: Conditional logic for Father Death Date
            # Assuming you have a father_status variable or check based on death date
            # If you don't have father status, we can check if death date is empty
            father_death_date_raw = father_death_date.get()
            
            # Option 1: If you want to insert space when father is alive (no death date entered)
            if not father_death_date_raw or father_death_date_raw.strip() == "":
                father_death_date_value = " "  # Insert space if father is alive (no death date)
            else:
                father_death_date_value = format_date(father_death_date_raw)

            conn = get_connection()
            if not conn:
                messagebox.showerror("Error", "Database connection failed.")
                return

            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO tblChildren (
                    CenterID, FullName, FatherName, Gender,
                    DateOfBirth, AdmissionDate, RegistrationNumber,
                    SchoolName, Class, Intelligence,
                    Disability, HealthCondition,
                    ReasonFatherDeath, FatherDeathDate,
                    FatherOccupation, FatherDesignation,
                    MotherName, MotherStatus, MotherDeathDate, MotherCNIC,
                    PermanentAddress, TemporaryAddress,
                    GuardianName, GuardianRelation, GuardianCNIC,
                    GuardianContact, Guardianaddress,
                    Sibling1Name, Sibling1Gender, Sibling1DOB,
                    Sibling2Name, Sibling2Gender, Sibling2DOB,
                    Sibling3Name, Sibling3Gender, Sibling3DOB,
                    Sibling4Name, Sibling4Gender, Sibling4DOB,
                    Sibling5Name, Sibling5Gender, Sibling5DOB,
                    MeetPerson1Name, MeetPerson1CNIC, MeetPerson1Contact,
                    MeetPerson2Name, MeetPerson2CNIC, MeetPerson2Contact,
                    MeetPerson3Name, MeetPerson3CNIC, MeetPerson3Contact,
                    MeetPerson4Name, MeetPerson4CNIC, MeetPerson4Contact,
                    MeetPerson5Name, MeetPerson5CNIC, MeetPerson5Contact,
                    IntroducerName, IntroducerCNIC, IntroducerContact, IntroducerAddress,
                    DocSchoolCertificate, DocBForm, DocFatherCNIC, DocMotherCNIC,
                    DocFatherDeathCert, OtherDoc1, OtherDoc2, OtherDoc3,
                    Status, PhotoPath, ChildRequiredAmount, SessionID
                ) VALUES (
                    ?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,
                    ?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?
                )
            """, (
                current_center_id.get(),
                safe_value(full_name.get()),
                safe_value(father_name.get()),
                safe_value(gender_var.get()),
                format_date(dob_entry.get()),
                format_date(adm_date_entry.get()),
                safe_value(reg_number.get()),
                safe_value(school_name.get()),
                safe_value(class_entry.get()),
                safe_value(intelligence_var.get()),
                safe_value(disability.get()),
                safe_value(health_condition_var.get()),
                safe_value(reason_father_death.get()),
                father_death_date_value,  # FIXED: Conditional value
                safe_value(father_occupation.get()),
                safe_value(father_designation.get()),
                safe_value(mother_name.get()),
                mother_status,
                mother_death_date_value,  # FIXED: Conditional value
                safe_value(mother_cnic.get()),
                safe_value(permanent_address.get()),
                safe_value(temporary_address.get()),
                safe_value(guardian_name.get()),
                safe_value(guardian_relation.get()),
                safe_value(guardian_cnic.get()),
                safe_value(guardian_contact.get()),
                safe_value(guardian_address.get()),
                *siblings_data,       # 15 values
                *meet_persons_data,   # 15 values
                safe_value(introducer_name.get()),
                safe_value(introducer_cnic.get()),
                safe_value(introducer_contact.get()),
                safe_value(introducer_address.get()),
                safe_value(doc_entries["Admission Form"].get()),
                safe_value(doc_entries["B-Form"].get()),
                safe_value(doc_entries["Father CNIC"].get()),
                safe_value(doc_entries["Mother CNIC"].get()),
                safe_value(doc_entries["Father Death Certificate"].get()),
                safe_value(doc_entries["Other Document 1"].get()),
                safe_value(doc_entries["Other Document 2"].get()),
                safe_value(doc_entries["Other Document 3"].get()),
                is_active_child.get(),
                photo_path.get().strip(),
                child_required_amount_val,
                current_session_id.get(),
            ))

            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Child added successfully!")
            child_win.destroy()
            open_child_mgmt()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to add child:\n{e}")

    # ═════════════════════════════════════════════════════════════════════════
    # FOOTER BUTTONS
    # ═════════════════════════════════════════════════════════════════════════
    tk.Button(
        footer, text="✔  Add Child",
        command=on_add,
        font=FONT_BTN, bg=BTN_ADD_BG, fg="white",
        relief="flat", cursor="hand2",
        padx=24, pady=8
    ).pack(side="left", padx=20)

    tk.Button(
        footer, text="📋  Edit / Update",
        command=open_child_list,
        font=FONT_BTN, bg=BTN_EDIT_BG, fg="white",
        relief="flat", cursor="hand2",
        padx=24, pady=8
    ).pack(side="left", padx=6)

    tk.Button(
        footer, text="✖  Close",
        command=child_win.destroy,
        font=FONT_BTN, bg=BTN_CLOSE_BG, fg="white",
        relief="flat", cursor="hand2",
        padx=24, pady=8
    ).pack(side="right", padx=20)

    # ── Keyboard shortcuts ────────────────────────────────────────────────────
    child_win.bind('<Return>', lambda e: on_add())
    child_win.bind('<Escape>', lambda e: child_win.destroy())

    # ── Show window ───────────────────────────────────────────────────────────
    child_win.update_idletasks()
    child_win.deiconify()
    child_win.focus_force()
    canvas.yview_moveto(0)


    
    # End of on_add()
    


# ============================================================
# Child Edit/Update
# ============================================================

def open_child_edit_window():
    #edit_win = tk.Toplevel(root)
    #edit_win.title("Edit Child" + ORG_SUFFIX)
    #edit_win.geometry("1150x800")

    # Center the window
    #window_width = 1150
    #window_height = 1050
    #screen_width = root.winfo_screenwidth()
    #screen_height = root.winfo_screenheight()
    #x = (screen_width // 2) - (window_width // 2)
    #y = (screen_height // 2) - (window_height // 2)
    #edit_win.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    
    edit_win = tk.Toplevel()
    edit_win.withdraw()
    edit_win.title("Edit Child" + ORG_SUFFIX)
    edit_win.configure(bg=COLORS["bg"])
    edit_win.state('zoomed')
    edit_win.grab_set()

    create_window_header(edit_win, "Search and Edit Child")

    search_frame = tk.Frame(edit_win, bg=COLORS["bg"])
    search_frame.pack(padx=20, pady=10)

    tk.Label(search_frame, text="Search Child:", font=FONTS["body"], bg=COLORS["bg"]).pack(side="left")
    search_input = tk.Entry(search_frame, font=FONTS["entry"], width=30)
    search_input.pack(side="left", padx=5)
    search_input.focus_set()

    list_container = tk.Frame(edit_win, bg=COLORS["bg"])
    list_container.pack(fill="both", expand=True, padx=20, pady=10)

    result_listbox = tk.Listbox(list_container, font=FONTS["body"])
    result_listbox.pack(side="left", fill="both", expand=True)

    scrollbar = tk.Scrollbar(list_container, orient="vertical", command=result_listbox.yview)
    scrollbar.pack(side="right", fill="y")
    result_listbox.config(yscrollcommand=scrollbar.set)

    selected_child_id = [None]  # To store donor ID from listbox

    form_frame = tk.Frame(edit_win, bg=COLORS["bg_card"], bd=1, relief="solid", highlightbackground=COLORS["border"], highlightthickness=1)
    form_frame.pack(padx=30, pady=10)

    tk.Label(form_frame, text="Full Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=0, column=0, sticky="e", pady=5)
    edit_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    edit_name.grid(row=0, column=1)

    tk.Label(form_frame, text="Gender:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=0, column=2, sticky="e", padx=10, pady=5)
    gender_var = tk.StringVar()
    gender_dropdown = ttk.Combobox(
        form_frame,
        textvariable=gender_var,
        values=["Male", "Female", "Other"],
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    gender_dropdown.grid(row=0, column=3)
    gender_dropdown.set("Female")

    tk.Label(form_frame, text="Date of Birth:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=2, column=0, sticky="e", pady=5)
    dob_entry = DateEntry(form_frame, font=FONTS["entry"], width=28, date_pattern='dd/MM/yyyy')
    dob_entry.grid(row=2, column=1)

    tk.Label(form_frame, text="Admission Date:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=2, column=2, sticky="e", padx=10, pady=5)
    adm_entry = DateEntry(form_frame, font=FONTS["entry"], width=28, date_pattern='dd/MM/yyyy')
    adm_entry.grid(row=2, column=3)

    tk.Label(form_frame, text="Health Info:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=4, column=0, sticky="e", pady=5)
    health_info = tk.Text(form_frame, font=FONTS["entry"], height=3, width=30)
    health_info.grid(row=4, column=1)

    tk.Label(form_frame, text="Educational Background:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=4, column=2, sticky="e", padx=10, pady=5)
    edu_background = tk.Text(form_frame, font=FONTS["entry"], height=3, width=30)
    edu_background.grid(row=4, column=3)

    tk.Label(form_frame, text="Status:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=6, column=0, sticky="e", pady=5)
    status_var = tk.StringVar()
    status_dropdown = ttk.Combobox(
        form_frame,
        textvariable=status_var,
        values=["Active", "Inactive", "Graduated", "Deceased"],
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    status_dropdown.grid(row=6, column=1)
    status_dropdown.set("Active")

    tk.Label(form_frame, text="Photo Path:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=6, column=2, sticky="e", padx=10, pady=5)
    photo_path = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    photo_path.grid(row=6, column=3)

    def browse_photo():
        path = filedialog.askopenfilename(filetypes=[("Image Files", "*.jpg *.png *.jpeg")])
        if path:
            basename = os.path.basename(path)
            simple_path = f"Photo\\{basename}"
            photo_path.delete(0, tk.END)
            photo_path.insert(0, simple_path)
            update_child_photo_live(simple_path)

    ttk.Button(form_frame, text="Browse", command=browse_photo, style="Outline.TButton").grid(row=6, column=4, padx=5)

    tk.Label(form_frame, text="ERP ID:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=7, column=0, sticky="e", pady=5)
    erp_id = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    erp_id.grid(row=7, column=1)

    tk.Label(form_frame, text="Old ID:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=7, column=2, sticky="e", padx=10, pady=5)
    old_id = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    old_id.grid(row=7, column=3)

    tk.Label(form_frame, text="Center:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=9, column=0, sticky="e", pady=5)
    center_var = tk.StringVar()
    center_dropdown = ttk.Combobox(
        form_frame,
        textvariable=center_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    center_dropdown.grid(row=9, column=1)

    tk.Label(form_frame, text="Category:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=9, column=2, sticky="e", padx=10, pady=5)
    category_var = tk.StringVar()
    category_dropdown = ttk.Combobox(
        form_frame,
        textvariable=category_var,
        values=["Orphan", "Dependent", "Special Needs"],
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    category_dropdown.grid(row=9, column=3)
    category_dropdown.set("Orphan")

    tk.Label(form_frame, text="Father Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=11, column=0, sticky="e", pady=5)
    father_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    father_name.grid(row=11, column=1)

    tk.Label(form_frame, text="Father Death Date:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=11, column=2, sticky="e", padx=10, pady=5)
    father_death_date = DateEntry(form_frame, font=FONTS["entry"], width=28, date_pattern='dd/MM/yyyy')
    father_death_date.grid(row=11, column=3)

    tk.Label(form_frame, text="Mother Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=13, column=0, sticky="e", pady=5)
    mother_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    mother_name.grid(row=13, column=1)

    tk.Label(form_frame, text="Guardian Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=13, column=2, sticky="e", padx=10, pady=5)
    guardian_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    guardian_name.grid(row=13, column=3)

    tk.Label(form_frame, text="Guardian Relation:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=15, column=0, sticky="e", pady=5)
    relation_var = tk.StringVar()
    relation_dropdown = ttk.Combobox(
        form_frame,
        textvariable=relation_var,
        values=["Parent", "Relative", "Other"],
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    relation_dropdown.grid(row=15, column=1)
    relation_dropdown.set("Parent")

    tk.Label(form_frame, text="Guardian CNIC:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=15, column=2, sticky="e", padx=10, pady=5)
    guardian_cnic = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    guardian_cnic.grid(row=15, column=3)

    tk.Label(form_frame, text="Guardian Contact:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=17, column=0, sticky="e", pady=5)
    guardian_contact = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    guardian_contact.grid(row=17, column=1)

    tk.Label(form_frame, text="Has Profile:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=17, column=2, sticky="e", padx=10, pady=5)
    has_profile_var = tk.BooleanVar()
    has_profile_check = tk.Checkbutton(form_frame, variable=has_profile_var, font=FONTS["body"], bg=COLORS["bg_card"])
    has_profile_check.grid(row=17, column=3, sticky="w")

    tk.Label(form_frame, text="Has Sponsorship Cert:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=19, column=0, sticky="e", pady=5)
    has_sponsorship_cert_var = tk.BooleanVar()
    has_sponsorship_cert_check = tk.Checkbutton(form_frame, variable=has_sponsorship_cert_var, font=FONTS["body"], bg=COLORS["bg_card"])
    has_sponsorship_cert_check.grid(row=19, column=1, sticky="w")

    # Photo Label
    photo_label = tk.Label(form_frame)
    photo_label.grid(row=20, column=0, columnspan=4, pady=20)
    

    def load_centers():
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT CenterID, CenterName FROM tblCenters ORDER BY CenterName")
                rows = cursor.fetchall()
                conn.close()

                center_options = [row.CenterName for row in rows]
                global center_id_map
                center_id_map = {row.CenterName: row.CenterID for row in rows}
                if center_options:
                    center_dropdown["values"] = center_options
                    center_dropdown.set(center_options[0])
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load centers: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def display_children(query=None):
        result_listbox.delete(0, tk.END)

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                if query:
                    cursor.execute("""
                        SELECT c.ChildID, c.FullName, c.DateOfBirth, c.AdmissionDate, 
                               c.Gender, c.Status, ce.CenterName, c.FatherName, c.MotherName,
                               c.FatherDeathDate, c.Category, c.PhotoPath, c.ERPId, c.OldID,
                               c.EducationalBackground, c.HealthInfo, c.GuardianName, 
                               c.GuardianRelation, c.GuardianCNIC, c.GuardianContact,
                               c.HasProfile, c.HasSponsorshipCertificate
                        FROM tblChildren c
                        JOIN tblCenters ce ON c.CenterID = ce.CenterID
                        WHERE c.FullName LIKE ?
                        ORDER BY c.FullName
                    """, ('%' + query + '%',))
                else:
                    cursor.execute("""
                        SELECT c.ChildID, c.FullName, c.DateOfBirth, c.AdmissionDate, 
                               c.Gender, c.Status, ce.CenterName, c.FatherName, c.MotherName,
                               c.FatherDeathDate, c.Category, c.PhotoPath, c.ERPId, c.OldID,
                               c.EducationalBackground, c.HealthInfo, c.GuardianName, 
                               c.GuardianRelation, c.GuardianCNIC, c.GuardianContact,
                               c.HasProfile, c.HasSponsorshipCertificate
                        FROM tblChildren c
                        JOIN tblCenters ce ON c.CenterID = ce.CenterID
                        ORDER BY c.FullName
                    """)
                rows = cursor.fetchall()
                conn.close()

                global child_data_map
                child_data_map = {}

                for row in rows:
                    dob = format_db_date(row.DateOfBirth, "%d/%m/%Y")
                    adm = format_db_date(row.AdmissionDate, "%d/%m/%Y")
                    death = format_db_date(row.FatherDeathDate, "%d/%m/%Y")

                    result_listbox.insert(tk.END, row.FullName)
                    child_data_map[row.FullName] = {
                        'ChildID': row.ChildID,
                        'DateOfBirth': dob,
                        'AdmissionDate': adm,
                        'FatherDeathDate': death,
                        'Gender': row.Gender,
                        'Status': row.Status,
                        'CenterName': row.CenterName,
                        'FatherName': row.FatherName,
                        'MotherName': row.MotherName,
                        'Category': row.Category,
                        'PhotoPath': row.PhotoPath,
                        'ERPId': row.ERPId,
                        'OldID': row.OldID,
                        'EducationalBackground': row.EducationalBackground,
                        'HealthInfo': row.HealthInfo,
                        'GuardianName': row.GuardianName,
                        'GuardianRelation': row.GuardianRelation,
                        'GuardianCNIC': row.GuardianCNIC,
                        'GuardianContact': row.GuardianContact,
                        'HasProfile': bool(row.HasProfile),
                        'HasSponsorshipCertificate': bool(row.HasSponsorshipCertificate)
                    }

            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load children: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_search(event=None):
        query = search_input.get().strip()
        display_children(query)

    def on_select_child(event=None):
        #clear_form_update_child()
        try:
            selected = result_listbox.get(result_listbox.curselection())
            data = child_data_map[selected]

            # Populate form fields
            edit_name.delete(0, tk.END)
            edit_name.insert(0, selected)

            gender_dropdown.set(data['Gender'])
            dob_entry.delete(0, tk.END)
            dob_entry.insert(0, data['DateOfBirth'])

            adm_entry.delete(0, tk.END)
            adm_entry.insert(0, data['AdmissionDate'])

            status_dropdown.set(data['Status'])
            center_dropdown.set(data['CenterName'])
            category_dropdown.set(data['Category'])

            father_name.delete(0, tk.END)
            father_name.insert(0, data['FatherName'])

            mother_name.delete(0, tk.END)
            mother_name.insert(0, data['MotherName'])

            father_death_date.delete(0, tk.END)
            father_death_date.insert(0, data['FatherDeathDate'])

            guardian_name.delete(0, tk.END)
            guardian_name.insert(0, data['GuardianName'])

            relation_dropdown.set(data['GuardianRelation'])
            guardian_cnic.delete(0, tk.END)
            guardian_cnic.insert(0, data['GuardianCNIC'])

            guardian_contact.delete(0, tk.END)
            guardian_contact.insert(0, data['GuardianContact'])

            has_profile_var.set(data['HasProfile'])
            has_sponsorship_cert_var.set(data['HasSponsorshipCertificate'])

            photo_path.delete(0, tk.END)
            photo_path.insert(0, data['PhotoPath'])

            selected_child_id[0] = data['ChildID']

            # Display child photo
            update_child_photo(data['PhotoPath'])

        except Exception as e:
            pass

    result_listbox.bind("<<ListboxSelect>>", on_select_child)

    def update_child_photo(photo_file_name):
        """Load and display child's photo from /photo folder"""
        photo_folder = os.path.join(os.path.dirname(__file__), 'photo')
        full_photo_path = os.path.join(photo_folder, str(photo_file_name).strip())

        if not photo_file_name or not os.path.exists(full_photo_path):
            photo_label.config(image='', text='No Image Found')
            photo_label.image = None
            return

        try:
            img = Image.open(full_photo_path)
            img.thumbnail((150, 150))  # Resize for preview
            img_tk = ImageTk.PhotoImage(img)
            photo_label.config(image=img_tk, text="")
            photo_label.image = img_tk
        except Exception as e:
            photo_label.config(image='', text='Error Loading Image')
            photo_label.image = None
            print(f"Error loading image: {e}")
            
    #photo path
    '''def update_child_photo_live(photo_file_path):
        """Load and display a selected photo file directly"""
        if not photo_file_path or not os.path.exists(photo_file_path):
            photo_label.config(image='', text='No Image Found')
            photo_label.image = None
            return

        try:
            img = Image.open(photo_file_path)
            img.thumbnail((150, 150))
            img_tk = ImageTk.PhotoImage(img)
            photo_label.config(image=img_tk, text="")
            photo_label.image = img_tk
        except Exception as e:
            photo_label.config(image='', text='Error Loading Image')
            photo_label.image = None
            print(f"Error loading image: {e}")'''
            
    def update_child_photo_live(photo_file_path):
        """Load and display a selected photo file directly"""
        # Convert Photo\filename.ext to full path
        if photo_file_path.startswith("Photo\\"):
            base_dir = os.path.dirname(os.path.abspath(__file__))
            full_path = os.path.join(base_dir, photo_file_path)
        else:
            full_path = photo_file_path
        
        if not os.path.exists(full_path):
            photo_label.config(image='', text='No Image Found')
            photo_label.image = None
            return

        try:
            img = Image.open(full_path)
            img.thumbnail((150, 150))
            img_tk = ImageTk.PhotoImage(img)
            photo_label.config(image=img_tk, text="")
            photo_label.image = img_tk
        except Exception as e:
            photo_label.config(image='', text='Error Loading Image')
            photo_label.image = None
            print(f"Error loading image: {e}")  
    
    
    
    

    def clear_form_update_child():
        edit_name.delete(0, tk.END)
        gender_dropdown.set("")
        dob_entry.delete(0, tk.END)
        set_placeholder(dob_entry, "DD/MM/YYYY")
        adm_entry.delete(0, tk.END)
        set_placeholder(adm_entry, "DD/MM/YYYY")
        father_name.delete(0, tk.END)
        father_death_date.delete(0, tk.END)
        set_placeholder(father_death_date, "DD/MM/YYYY")
        mother_name.delete(0, tk.END)
        guardian_name.delete(0, tk.END)
        relation_dropdown.set("")
        guardian_cnic.delete(0, tk.END)
        guardian_contact.delete(0, tk.END)
        status_dropdown.set("")
        center_dropdown.set("")
        category_dropdown.set("")
        has_profile_var.set(False)
        has_sponsorship_cert_var.set(False)
        photo_path.delete(0, tk.END)
        photo_label.config(image='', text='No Image')
        photo_label.image = None

    def on_update():
        name = edit_name.get().strip()
        gender = gender_var.get()
        dob = dob_entry.get().strip()
        adm = adm_entry.get().strip()
        status = status_var.get()
        center = center_var.get()
        category = category_var.get()
        father = father_name.get().strip()
        death_date = father_death_date.get().strip()
        mother = mother_name.get().strip()
        guardian = guardian_name.get().strip()
        relation = relation_var.get()
        cnic = guardian_cnic.get().strip()
        contact = guardian_contact.get().strip()
        has_profile = has_profile_var.get()
        has_certificate = has_sponsorship_cert_var.get()
        #photo = photo_path.get().strip()
        photo = os.path.basename(photo_path.get().strip())
        erp = erp_id.get().strip()
        old = old_id.get().strip()
        education = edu_background.get("1.0", tk.END).strip()
        health = health_info.get("1.0", tk.END).strip()

        if not name:
            messagebox.showwarning("Input Error", "Child Full Name is required.")
            return

        if not gender:
            messagebox.showwarning("Input Error", "Gender is required.")
            return

        if not status:
            messagebox.showwarning("Input Error", "Status is required.")
            return

        if not center or center not in center_id_map:
            messagebox.showwarning("Input Error", "Please select a valid center.")
            return

        converted_dob = validate_and_convert_date(dob)
        converted_adm = validate_and_convert_date(adm)
        converted_death_date = validate_and_convert_date(death_date)

        if None in (converted_dob, converted_adm, converted_death_date):
            messagebox.showerror("Date Format Error", "Invalid date format. Please use DD/MM/YYYY.")
            return

        center_id = center_id_map[center]

        if selected_child_id[0] is None:
            messagebox.showerror("Selection Error", "No child selected for update.")
            return

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE tblChildren SET
                        FullName = ?,
                        Gender = ?,
                        DateOfBirth = ?,
                        AdmissionDate = ?,
                        Status = ?,
                        CenterID = ?,
                        FatherName = ?,
                        FatherDeathDate = ?,
                        MotherName = ?,
                        GuardianName = ?,
                        GuardianRelation = ?,
                        GuardianCNIC = ?,
                        GuardianContact = ?,
                        HasProfile = ?,
                        HasSponsorshipCertificate = ?,
                        PhotoPath = ?,
                        ERPId = ?,
                        OldID = ?,
                        EducationalBackground = ?,
                        HealthInfo = ?,
                        Category = ?
                    WHERE ChildID = ?
                """, (
                    name, gender, converted_dob, converted_adm, status, center_id,
                    father, converted_death_date, mother, guardian, relation, cnic,
                    contact, has_profile, has_certificate, photo, erp, old, education,
                    health, category, selected_child_id[0]
                ))
                conn.commit()
                conn.close()
                messagebox.showinfo("Success", "Child updated successfully!")
                display_children()
                clear_form_update_child()
            except Exception as e:
                conn.rollback()
                messagebox.showerror("Error", f"Update failed: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_close():
        edit_win.grab_release()
        edit_win.destroy()

    def on_clear():
        clear_form_update_child()
        edit_name.focus_set()

    # Load centers into dropdown
    load_centers()
    display_children()  # Load all children initially

    # Bind search input
    search_input.bind('<KeyRelease>', lambda e: on_search())

    # === BUTTON PANEL ===
    btn_frame = tk.Frame(edit_win)
    btn_frame.pack(pady=10)

    tk.Button(btn_frame, text="Update", command=on_update, font=("Arial", 12), width=12).grid(row=0, column=0, padx=5)
    tk.Button(btn_frame, text="Clear", command=on_clear, font=("Arial", 12), width=10).grid(row=0, column=1, padx=5)
    tk.Button(btn_frame, text="Close", command=on_close, font=("Arial", 12), width=10).grid(row=0, column=2, padx=5)

    # === KEYBOARD SHORTCUTS ===
    edit_win.bind('<Return>', lambda e: on_update())
    edit_win.bind('<Escape>', lambda e: on_close())

    edit_win.update_idletasks()
    edit_win.deiconify()
    edit_win.focus_force()


# ============================================================
# Center Management
# ============================================================

def open_center_mgmt():
    global logged_in_role  # Ensure role is accessible

    if logged_in_role.lower() != "admin":
        messagebox.showerror("Access Denied", "You are not authorized to manage centers.")
        return

    center_win = tk.Toplevel(root)
    center_win.title("Add New Center" + ORG_SUFFIX)
    center_win.configure(bg=COLORS["bg"])
    center_window(center_win, 700, 620)
    center_win.grab_set()

    create_window_header(center_win, "Add New Center")

    form_frame = tk.Frame(center_win, bg=COLORS["bg_card"], bd=1, relief="solid", highlightbackground=COLORS["border"], highlightthickness=1)
    form_frame.pack(padx=30, pady=10)

    tk.Label(form_frame, text="Center Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=0, column=0, sticky="e", pady=5)
    center_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    center_name.grid(row=0, column=1)

    tk.Label(form_frame, text="City:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=1, column=0, sticky="e", pady=5)
    center_city = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    center_city.grid(row=1, column=1)

    tk.Label(form_frame, text="Address:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=2, column=0, sticky="e", pady=5)
    center_address = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    center_address.grid(row=2, column=1)

    tk.Label(form_frame, text="Capacity:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=3, column=0, sticky="e", pady=5)
    center_capacity = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    center_capacity.grid(row=3, column=1)

    tk.Label(form_frame, text="Operational Status:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=4, column=0, sticky="e", pady=5)
    status_var = tk.StringVar()
    status_dropdown = ttk.Combobox(
        form_frame,
        textvariable=status_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    status_dropdown.grid(row=4, column=1)
    status_dropdown["values"] = ["Active", "Inactive", "Under Construction"]
    status_dropdown.set("Active")

    tk.Label(form_frame, text="Region:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=5, column=0, sticky="e", pady=5)
    center_region = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    center_region.grid(row=5, column=1)

    tk.Label(form_frame, text="Administrator Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=6, column=0, sticky="e", pady=5)
    admin_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    admin_name.grid(row=6, column=1)

    tk.Label(form_frame, text="Admin Contact Number:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=7, column=0, sticky="e", pady=5)
    admin_contact = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    admin_contact.grid(row=7, column=1)

    def clear_form_center():
        center_name.delete(0, tk.END)
        center_city.delete(0, tk.END)
        center_address.delete(0, tk.END)
        center_capacity.delete(0, tk.END)
        status_dropdown.set("Active")
        center_region.delete(0, tk.END)
        admin_name.delete(0, tk.END)
        admin_contact.delete(0, tk.END)

    def on_add():
        name = center_name.get().strip()
        city = center_city.get().strip()
        address = center_address.get().strip()
        capacity = center_capacity.get().strip()
        status = status_var.get()
        region = center_region.get().strip()
        admin = admin_name.get().strip()
        contact = admin_contact.get().strip()

        if not name:
            messagebox.showwarning("Input Error", "Center Name is required.")
            return
        if not city:
            messagebox.showwarning("Input Error", "City is required.")
            return
        if not address:
            messagebox.showwarning("Input Error", "Address is required.")
            return
        if not capacity:
            messagebox.showwarning("Input Error", "Capacity is required.")
            return
        if not status:
            messagebox.showwarning("Input Error", "Operational Status is required.")
            return
        if not region:
            messagebox.showwarning("Input Error", "Region is required.")
            return
        if not admin:
            messagebox.showwarning("Input Error", "Administrator Name is required.")
            return
        if not contact:
            messagebox.showwarning("Input Error", "Admin Contact Number is required.")
            return

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO tblCenters (
                        CenterName, City, Address, Capacity, OperationalStatus, Region,
                        AdministratorName, AdminContactNumber
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                """, (name, city, address, int(capacity), status, region, admin, contact))
                conn.commit()
                conn.close()
                messagebox.showinfo("Success", "Center added successfully!")
                clear_form_center()
                display_centers()  # Refresh list
                
            except Exception as e:
                conn.rollback()
                messagebox.showerror("Error", f"Failed to add center: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_close():
        center_win.grab_release()
        center_win.destroy()

    # === BUTTON PANEL ===
    btn_frame = tk.Frame(center_win)
    btn_frame.pack(pady=10)
    center_name.focus_set()

    add_center_btn = ttk.Button(btn_frame, text="Add Center", command=on_add, style="Success.TButton", width=15)
    add_center_btn.grid(row=0, column=0, padx=5)

    ttk.Button(btn_frame, text="Clear", command=clear_form_center, style="Outline.TButton", width=10).grid(row=0, column=1, padx=5)
    ttk.Button(btn_frame, text="Close", command=on_close, style="Danger.TButton", width=10).grid(row=0, column=2, padx=5)

    # === SCROLLABLE FRAME FOR CENTER LIST ===
    container = tk.Frame(center_win)
    container.pack(fill="both", padx=30, pady=10, expand=True)

    canvas = tk.Canvas(container, height=150)
    scrollbar = tk.Scrollbar(container, orient="vertical", command=canvas.yview)
    center_data_frame = tk.Frame(canvas)

    center_data_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=center_data_frame, anchor="nw", width=640)
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")

    # Store reference for refresh
    global center_list_container
    center_list_container = {
        'canvas': canvas,
        'data_frame': center_data_frame
    }

    def display_centers():
        # Clear existing data
        for widget in center_list_container['data_frame'].winfo_children():
            widget.destroy()

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT CenterName, City, Region FROM tblCenters ORDER BY CenterName")
                rows = cursor.fetchall()
                conn.close()

                # Header
                header_frame = tk.Frame(center_list_container['data_frame'])
                header_frame.pack(fill="x")
                tk.Label(header_frame, text="Center Name", font=("Arial", 12, "bold"), width=25, anchor="w").pack(side="left")
                tk.Label(header_frame, text="City", font=("Arial", 12, "bold"), width=15, anchor="w").pack(side="left")
                tk.Label(header_frame, text="Region", font=("Arial", 12, "bold"), width=15, anchor="w").pack(side="left")

                divider = tk.Frame(center_list_container['data_frame'], height=2, bg="gray")
                divider.pack(fill="x", pady=2)

                # Data Rows
                for row in rows:
                    item_frame = tk.Frame(center_list_container['data_frame'])
                    item_frame.pack(fill="x")

                    tk.Label(item_frame, text=row.CenterName, font=("Arial", 11), width=25, anchor="w").pack(side="left")
                    tk.Label(item_frame, text=row.City, font=("Arial", 11), width=15, anchor="w").pack(side="left")
                    tk.Label(item_frame, text=row.Region, font=("Arial", 11), width=15, anchor="w").pack(side="left")

                # Update scroll region
                center_list_container['canvas'].update_idletasks()
                center_list_container['canvas'].configure(scrollregion=center_list_container['canvas'].bbox("all"))

                if len(rows) >= 1:
                    add_center_btn.config(state="disabled")
                else:
                    add_center_btn.config(state="normal")



            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to reload center list: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    # Call display_centers() to load list
    display_centers()

    # === KEYBOARD SHORTCUTS ===
    center_win.bind('<Return>', lambda e: on_add())
    center_win.bind('<Escape>', lambda e: on_close())



# ============================================================

# ============================================================
# Donor Management
# ============================================================


def open_donor_mgmt():
    global logged_in_role  # Ensure role is accessible

    if logged_in_role.lower() != "admin":
        messagebox.showerror("Access Denied", "You are not authorized to manage donors.")
        return

    donor_win = tk.Toplevel(root)
    donor_win.title("Manage Donors" + ORG_SUFFIX)
    donor_win.configure(bg=COLORS["bg"])
    center_window(donor_win, 700, 650)
    donor_win.grab_set()

    create_window_header(donor_win, "Add New Donor")

    form_frame = tk.Frame(donor_win, bg=COLORS["bg_card"], bd=1, relief="solid", highlightbackground=COLORS["border"], highlightthickness=1)
    form_frame.pack(padx=30, pady=10)

    tk.Label(form_frame, text="Full Name:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=0, column=0, sticky="e", pady=5)
    donor_name = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    donor_name.grid(row=0, column=1)

    '''tk.Label(form_frame, text="CNIC:", font=("Arial", 12)).grid(row=1, column=0, sticky="e", pady=5)
    donor_cnic = tk.Entry(form_frame, font=("Arial", 12), width=30)
    donor_cnic.grid(row=1, column=1)

    tk.Label(form_frame, text="Address:", font=("Arial", 12)).grid(row=2, column=0, sticky="e", pady=5)
    donor_address = tk.Entry(form_frame, font=("Arial", 12), width=30)
    donor_address.grid(row=2, column=1)

    tk.Label(form_frame, text="Contact Number:", font=("Arial", 12)).grid(row=3, column=0, sticky="e", pady=5)
    donor_contact = tk.Entry(form_frame, font=("Arial", 12), width=30)
    donor_contact.grid(row=3, column=1)'''
    
    # Mask template: fixed parts + placeholders
    TEMPLATE = list("Donor-___-___-___")  # We'll work with a list so we can replace characters

    def on_focus_in(event):
        """Place cursor at first underscore on focus."""
        widget = event.widget
        text = widget.get()
        first_underscore = text.find("_")
        if first_underscore != -1:
            widget.icursor(first_underscore)

    def on_key_press(event):
        allowed_positions = [6, 7, 8, 10, 11, 12, 14, 15, 16]
        pos = event.widget.index(tk.INSERT)

        # Handle left/right arrows normally
        if event.keysym in ("Left", "Right", "Tab", "Shift_L", "Shift_R"):
            return

        # Handle backspace
        if event.keysym == "BackSpace":
            new_pos = pos - 1
            while new_pos >= 0 and new_pos not in allowed_positions:
                new_pos -= 1
            if new_pos >= 0:
                event.widget.delete(new_pos, new_pos+1)
                event.widget.insert(new_pos, "_")
                event.widget.icursor(new_pos)
            return "break"

        # Handle delete key
        if event.keysym == "Delete":
            while pos < len(event.widget.get()) and pos not in allowed_positions:
                pos += 1
            if pos < len(event.widget.get()):
                event.widget.delete(pos, pos+1)
                event.widget.insert(pos, "_")
                event.widget.icursor(pos)
            return "break"

        # Only allow valid input in valid positions
        if pos in [6, 7, 8]:  # city letters
            if not event.char.isalpha():
                return "break"
            event.widget.delete(pos, pos+1)
            event.widget.insert(pos, event.char.upper())
            move_cursor(event.widget, pos+1)
            return "break"

        if pos in [10, 11, 12, 14, 15, 16]:  # digits
            if not event.char.isdigit():
                return "break"
            event.widget.delete(pos, pos+1)
            event.widget.insert(pos, event.char)
            move_cursor(event.widget, pos+1)
            return "break"

        return "break"  # Block typing on fixed chars

    def move_cursor(widget, pos):
        """Move cursor to next editable position, skipping hyphens."""
        while pos < len(widget.get()) and widget.get()[pos] == "-":
            pos += 1
        widget.icursor(pos)
    
    
        
    
    
    def check_office_donor_id(event=None):
        office_id = donor_ID.get().strip()
        conn = get_connection()
        cursor = conn.cursor()
        if not office_id:
            return  # Skip if empty

        cursor.execute("""
            SELECT COUNT(*) 
            FROM tblDonors 
            WHERE OfficeDonorID = ?
        """, (office_id,))
        
        exists = cursor.fetchone()[0]
        if exists > 0:
            messagebox.showwarning("Duplicate Office Donor ID", f"Office Donor ID '{office_id}' already exists.")
            donor_ID.delete(0, tk.END)
            donor_ID.insert(0, "Donor-___-___-___")
            donor_ID.focus_set()
            on_focus_in()
        conn.close()


    
    tk.Label(form_frame, text="Office Donor ID:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=4, column=0, sticky="e", pady=5)
    donor_ID = tk.Entry(form_frame, font=FONTS["entry"], width=30)
    donor_ID.grid(row=4, column=1)
    donor_ID.insert(0, "Donor-___-___-___")
    donor_ID.bind("<FocusOut>", check_office_donor_id)
    donor_ID.bind("<Key>", on_key_press)
    #donor_ID.bind("<FocusOut>", check_duplicate)
    donor_ID.bind("<FocusIn>", on_focus_in)

    # Initialize
    #update_entry()
    donor_ID.icursor(6)  # Start after "Donor-"
    

    tk.Label(form_frame, text="Donation Type:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=5, column=0, sticky="e", pady=5)
    donation_type_var = tk.StringVar()
    donation_type_dropdown = ttk.Combobox(
        form_frame,
        textvariable=donation_type_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    donation_type_dropdown.grid(row=5, column=1)
    donation_type_dropdown["values"] = ["Cash", "In Kind", "Other"]
    donation_type_dropdown.set("Cash")

    tk.Label(form_frame, text="Collector:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=7, column=0, sticky="e", pady=5)
    collector_var = tk.StringVar()
    collector_dropdown = ttk.Combobox(
        form_frame,
        textvariable=collector_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    collector_dropdown.grid(row=7, column=1)

    tk.Label(form_frame, text="Payment Method:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=8, column=0, sticky="e", pady=5)
    payment_method_var = tk.StringVar()
    payment_method_dropdown = ttk.Combobox(
        form_frame,
        textvariable=payment_method_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    payment_method_dropdown.grid(row=8, column=1)
    payment_method_dropdown["values"] = ["Bank Transfer", "Cash", "Cheque", "Online Payment"]
    payment_method_dropdown.set("Bank Transfer")

    tk.Label(form_frame, text="Frequency:", font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=9, column=0, sticky="e", pady=5)
    frequency_var = tk.StringVar()
    frequency_dropdown = ttk.Combobox(
        form_frame,
        textvariable=frequency_var,
        state="readonly",
        font=FONTS["entry"],
        width=28
    )
    frequency_dropdown.grid(row=9, column=1)
    frequency_dropdown["values"] = ["One-Time", "Monthly", "Quarterly", "Yearly"]
    frequency_dropdown.set("One-Time")
    
        # Commitment Start Date
    #tk.Label(form_frame, text="Commitment Start Date:", font=("Arial", 12)).grid(row=10, column=0, sticky="e", pady=5)
    #start_date_entry = DateEntry(form_frame, font=("Arial", 12), width=28, date_pattern='yyyy-mm-dd')
    #start_date_entry.grid(row=10, column=1)

    # Commitment End Date
    #tk.Label(form_frame, text="Commitment End Date:", font=("Arial", 12)).grid(row=11, column=0, sticky="e", pady=5)
    #end_date_entry = DateEntry(form_frame, font=("Arial", 12), width=28, date_pattern='yyyy-mm-dd')
    #end_date_entry.grid(row=11, column=1)

    # Is Active Checkbox
    is_active_var = tk.IntVar(value=1)
    tk.Checkbutton(form_frame, text="Is Active", variable=is_active_var, font=FONTS["body"], bg=COLORS["bg_card"]).grid(row=12, column=1, sticky="w", pady=5)


    def load_collectors():
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT CollectorID, FullName FROM tblCollectors ORDER BY FullName")
                rows = cursor.fetchall()
                conn.close()

                collector_options = [row.FullName for row in rows]
                collector_dropdown["values"] = collector_options
                global collector_id_map
                collector_id_map = {row.FullName: row.CollectorID for row in rows}
                if collector_options:
                    collector_dropdown.set(collector_options[0])  # Default selection
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load collectors: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to database.")

    def clear_form():
        donor_name.delete(0, tk.END)
       # donor_cnic.delete(0, tk.END)
       # donor_address.delete(0, tk.END)
       # donor_contact.delete(0, tk.END)
        donor_ID.delete(0, tk.END)
        donation_type_var.set("Cash")
        #monthly_commitment.delete(0, tk.END)
        collector_dropdown.set("")
        payment_method_var.set("Bank Transfer")
        frequency_var.set("One-Time")
        #added one line below
        donor_win.grab_set()  # Keep focus on this window
        #start_date_entry.set_date(datetime.today())
        #end_date_entry.set_date(datetime.today())
        is_active_var.set(1)
        

    def on_add():
        name = donor_name.get().strip()
        #cnic = donor_cnic.get().strip()
        cnic = ""
        address = ""
        contact = ""
        donor_ID_var = donor_ID.get().strip()
        donation_type = donation_type_var.get()
        #commitment = monthly_commitment.get().strip()
        collector = collector_var.get()
        payment_method = payment_method_var.get()
        frequency = frequency_var.get()
        #commitment_start = start_date_entry.get_date().strftime('%Y-%m-%d')
        #commitment_end = end_date_entry.get_date().strftime('%Y-%m-%d')
        is_active = is_active_var.get()
        is_active = 1 if is_active_var.get() else 0



        #if not name:
        #    messagebox.showwarning("Input Error", "Donor Full Name is required.")
        #    donor_win.grab_set()  # Keep focus on this window
        #    return
        
        if not name:
            messagebox.showwarning("Input Error", "Donor Full Name is required.")
            donor_win.grab_set()  # Keep focus on this window
            return

        if not collector or collector not in collector_id_map:
            messagebox.showwarning("Input Error", "Please select a valid collector.")
            donor_win.grab_set()  # Keep focus on this window
            return

        collector_id = collector_id_map[collector]

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    INSERT INTO tblDonors (
                        FullName, CNIC, Address, ContactNumber, OfficeDonorID, DonationType, 
                        CollectorID, PaymentMethod, Frequency, IsActive
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (name, cnic, address, contact, donor_ID_var, donation_type, collector_id,
                    payment_method, frequency, is_active))
                
                conn.commit()
                conn.close()
                messagebox.showinfo("Success", "Donor added successfully!")
                clear_form()
                display_donors()
                donor_name.focus_set()
            except Exception as e:
                conn.rollback()
                messagebox.showerror("Error", f"Failed to add donor: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def display_donors():
        result_listbox.delete(0, tk.END)

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT FullName, CNIC, Address, ContactNumber, OfficeDonorID, DonationType, CollectorID, PaymentMethod, Frequency FROM tblDonors ORDER BY FullName")
                rows = cursor.fetchall()
                conn.close()

                for row in rows:
                    result_listbox.insert(tk.END, f"{row.FullName} | {row.OfficeDonorID}")
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load donors: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    # === BUTTON PANEL ===
    btn_frame = tk.Frame(donor_win)
    btn_frame.pack(pady=10)

    ttk.Button(btn_frame, text="Add Donor", command=on_add, style="Success.TButton", width=15).grid(row=0, column=0, padx=5)
    ttk.Button(btn_frame, text="Clear", command=clear_form, style="Outline.TButton", width=10).grid(row=0, column=1, padx=5)
    ttk.Button(btn_frame, text="Edit/Update", command=open_donor_edit_window, style="Primary.TButton", width=12).grid(row=0, column=2, padx=5)
    ttk.Button(btn_frame, text="Close", command=lambda: donor_win.destroy(), style="Danger.TButton", width=10).grid(row=0, column=3, padx=5)

    # === DONOR LISTBOX ===
    tk.Label(donor_win, text="Available Donors", font=("Arial", 14, "bold")).pack(pady=(20, 0))

    result_listbox = tk.Listbox(donor_win, height=5, font=("Arial", 12))
    result_listbox.pack(padx=20, pady=(0, 10), fill="both")

    scrollbar = tk.Scrollbar(donor_win, orient="vertical", command=result_listbox.yview)
    scrollbar.pack(side="right", fill="y")
    result_listbox.config(yscrollcommand=scrollbar.set)

    # Load donors initially
    load_collectors()
    display_donors()

    # === KEYBOARD SHORTCUTS ===
    donor_win.bind('<Return>', lambda e: on_add())
    donor_win.bind('<Escape>', lambda e: donor_win.destroy())

    # Focus on Donor Name field initially
    donor_name.focus_set()
    
# ============================================================
# Edit Donor Window
# ============================================================

def open_donor_edit_window():
    edit_win = tk.Toplevel(root)
    edit_win.title("Edit Donor" + ORG_SUFFIX)
    edit_win.geometry("700x500")

    # Center the window
    window_width = 700
    window_height = 650
    screen_width = root.winfo_screenwidth()
    screen_height = root.winfo_screenheight()
    x = (screen_width // 2) - (window_width // 2)
    y = (screen_height // 2) - (window_height // 2)
    edit_win.geometry(f"{window_width}x{window_height}+{x}+{y}")
    
    edit_win.grab_set()  # Keep focus on this window

    tk.Label(edit_win, text="Search and Edit Donor", font=("Arial", 16, "bold")).pack(pady=10)

    # === SEARCH BOX ===
    search_frame = tk.Frame(edit_win)
    search_frame.pack(padx=20, pady=10)

    tk.Label(search_frame, text="Search Donor:", font=("Arial", 12)).pack(side="left")
    search_input = tk.Entry(search_frame, font=("Arial", 12), width=30)
    search_input.pack(side="left", padx=5)
    search_input.focus_set()

    # Listbox to show search results
    result_listbox = tk.Listbox(edit_win, height=8, font=("Arial", 12))
    result_listbox.pack(padx=20, pady=(0, 10), fill="x")

    scrollbar = tk.Scrollbar(edit_win, orient="vertical", command=result_listbox.yview)
    scrollbar.pack(side="right", fill="y")
    result_listbox.config(yscrollcommand=scrollbar.set)

    selected_donor_id = [None]  # To store donor ID from listbox

    # Form fields
    form_frame = tk.Frame(edit_win)
    form_frame.pack(padx=30, pady=10)

    tk.Label(form_frame, text="Full Name:", font=("Arial", 12)).grid(row=0, column=0, sticky="e", pady=5)
    edit_name = tk.Entry(form_frame, font=("Arial", 12), width=30)
    edit_name.grid(row=0, column=1)

    '''tk.Label(form_frame, text="CNIC:", font=("Arial", 12)).grid(row=1, column=0, sticky="e", pady=5)
    edit_cnic = tk.Entry(form_frame, font=("Arial", 12), width=30)
    edit_cnic.grid(row=1, column=1)

    tk.Label(form_frame, text="Address:", font=("Arial", 12)).grid(row=2, column=0, sticky="e", pady=5)
    edit_address = tk.Entry(form_frame, font=("Arial", 12), width=30)
    edit_address.grid(row=2, column=1)

    tk.Label(form_frame, text="Contact Number:", font=("Arial", 12)).grid(row=3, column=0, sticky="e", pady=5)
    edit_contact = tk.Entry(form_frame, font=("Arial", 12), width=30)
    edit_contact.grid(row=3, column=1)'''
    
    # Mask template: fixed parts + placeholders
    TEMPLATE2 = list("Donor-___-___-___")  # We'll work with a list so we can replace characters

    def on_focus_in(event):
        """Place cursor at first underscore on focus."""
        widget = event.widget
        text = widget.get()
        first_underscore = text.find("_")
        if first_underscore != -1:
            widget.icursor(first_underscore)

    def on_key_press(event):
        allowed_positions = [6, 7, 8, 10, 11, 12, 14, 15, 16]
        pos = event.widget.index(tk.INSERT)

        # Handle left/right arrows normally
        if event.keysym in ("Left", "Right", "Tab", "Shift_L", "Shift_R"):
            return

        # Handle backspace
        if event.keysym == "BackSpace":
            new_pos = pos - 1
            while new_pos >= 0 and new_pos not in allowed_positions:
                new_pos -= 1
            if new_pos >= 0:
                event.widget.delete(new_pos, new_pos+1)
                event.widget.insert(new_pos, "_")
                event.widget.icursor(new_pos)
            return "break"

        # Handle delete key
        if event.keysym == "Delete":
            while pos < len(event.widget.get()) and pos not in allowed_positions:
                pos += 1
            if pos < len(event.widget.get()):
                event.widget.delete(pos, pos+1)
                event.widget.insert(pos, "_")
                event.widget.icursor(pos)
            return "break"

        # Only allow valid input in valid positions
        if pos in [6, 7, 8]:  # city letters
            if not event.char.isalpha():
                return "break"
            event.widget.delete(pos, pos+1)
            event.widget.insert(pos, event.char.upper())
            move_cursor(event.widget, pos+1)
            return "break"

        if pos in [10, 11, 12, 14, 15, 16]:  # digits
            if not event.char.isdigit():
                return "break"
            event.widget.delete(pos, pos+1)
            event.widget.insert(pos, event.char)
            move_cursor(event.widget, pos+1)
            return "break"

        return "break"  # Block typing on fixed chars

    def move_cursor(widget, pos):
        """Move cursor to next editable position, skipping hyphens."""
        while pos < len(widget.get()) and widget.get()[pos] == "-":
            pos += 1
        widget.icursor(pos)
    
    
        
    
    
    def check_office_donor_id(event=None):
        global current_donor_id

        office_id = edit_donorid.get().strip()
        if not office_id or not current_donor_id:
            return  # Skip if no donor loaded

        conn = get_connection()
        cursor = conn.cursor()

        # Get current OfficeDonorID from DB
        cursor.execute("""
            SELECT OfficeDonorID
            FROM tblDonors
            WHERE DonorID = ?
        """, (current_donor_id,))
        row = cursor.fetchone()
        current_office_id = row[0] if row else None

        if office_id == current_office_id:
            conn.close()
            return

        # Check duplication in other donors
        cursor.execute("""
            SELECT COUNT(*)
            FROM tblDonors
            WHERE OfficeDonorID = ? AND DonorID <> ?
        """, (office_id, current_donor_id))
        exists = cursor.fetchone()[0]
        conn.close()

        if exists > 0:
            messagebox.showwarning(
                "Duplicate Office Donor ID",
                f"Office Donor ID '{office_id}' already exists for another donor."
            )
            edit_donorid.focus_set()
            return

        confirm = messagebox.askyesno(
            "Confirm Change",
            f"The Office Donor ID is being changed for this donor.\n\n"
            f"Previous: {current_office_id}\nNew: {office_id}\n\nAre you sure?"
        )

        if confirm:
            donation_type_dropdown.focus_set()
        else:
            edit_donorid.delete(0, tk.END)
            edit_donorid.insert(0, current_office_id)
            edit_donorid.focus_set()


    

    tk.Label(form_frame, text="Office Donor ID:", font=("Arial", 12)).grid(row=4, column=0, sticky="e", pady=5)
    edit_donorid = tk.Entry(form_frame, font=("Arial", 12), width=30)
    edit_donorid.grid(row=4, column=1)
    edit_donorid.insert(0, "Donor-___-___-___")
    edit_donorid.bind("<FocusOut>", check_office_donor_id)
    edit_donorid.bind("<Key>", on_key_press)
    #donor_ID.bind("<FocusOut>", check_duplicate)
    edit_donorid.bind("<FocusIn>", on_focus_in)
    
    

    tk.Label(form_frame, text="Donation Type:", font=("Arial", 12)).grid(row=5, column=0, sticky="e", pady=5)
    donation_type_var = tk.StringVar()
    donation_type_dropdown = ttk.Combobox(
        form_frame,
        textvariable=donation_type_var,
        state="readonly",
        font=("Arial", 12),
        width=28
    )
    donation_type_dropdown.grid(row=5, column=1)
    donation_type_dropdown["values"] = ["Cash", "In Kind", "Other"]
    donation_type_dropdown.set("Cash")

    #tk.Label(form_frame, text="Monthly Commitment:", font=("Arial", 12)).grid(row=6, column=0, sticky="e", pady=5)
    #edit_commitment = tk.Entry(form_frame, font=("Arial", 12), width=30)
    #edit_commitment.grid(row=6, column=1)

    tk.Label(form_frame, text="Collector:", font=("Arial", 12)).grid(row=7, column=0, sticky="e", pady=5)
    collector_var = tk.StringVar()
    collector_dropdown = ttk.Combobox(
        form_frame,
        textvariable=collector_var,
        state="readonly",
        font=("Arial", 12),
        width=28
    )
    collector_dropdown.grid(row=7, column=1)

    tk.Label(form_frame, text="Payment Method:", font=("Arial", 12)).grid(row=8, column=0, sticky="e", pady=5)
    payment_method_var = tk.StringVar()
    payment_method_dropdown = ttk.Combobox(
        form_frame,
        textvariable=payment_method_var,
        state="readonly",
        font=("Arial", 12),
        width=28
    )
    payment_method_dropdown.grid(row=8, column=1)
    payment_method_dropdown["values"] = ["Bank Transfer", "Cash", "Cheque", "Online Payment"]
    payment_method_dropdown.set("Bank Transfer")

    tk.Label(form_frame, text="Frequency:", font=("Arial", 12)).grid(row=9, column=0, sticky="e", pady=5)
    frequency_var = tk.StringVar()
    frequency_dropdown = ttk.Combobox(
        form_frame,
        textvariable=frequency_var,
        state="readonly",
        font=("Arial", 12),
        width=28
    )
    frequency_dropdown.grid(row=9, column=1)
    frequency_dropdown["values"] = ["One-Time", "Monthly", "Quarterly", "Yearly"]
    frequency_dropdown.set("One-Time")
    
    # Commitment Start Date
    #tk.Label(form_frame, text="Commitment Start Date:", font=("Arial", 12)).grid(row=10, column=0, sticky="e", pady=5)
    #edit_start_date = DateEntry(form_frame, font=("Arial", 12), width=28, date_pattern='yyyy-mm-dd')
    #edit_start_date.grid(row=10, column=1)

    # Commitment End Date
    #tk.Label(form_frame, text="Commitment End Date:", font=("Arial", 12)).grid(row=11, column=0, sticky="e", pady=5)
    #edit_end_date = DateEntry(form_frame, font=("Arial", 12), width=28, date_pattern='yyyy-mm-dd')
    #edit_end_date.grid(row=11, column=1)

    # Is Active checkbox
    is_active_var = tk.IntVar(value=1)
    tk.Checkbutton(form_frame, text="Is Active", variable=is_active_var, font=("Arial", 12)).grid(row=12, column=1, sticky="w", pady=5)


    def load_collectors():
        """Load all collectors into dropdown"""
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("SELECT CollectorID, FullName FROM tblCollectors ORDER BY FullName")
                rows = cursor.fetchall()
                conn.close()

                collector_options = [row.FullName for row in rows]
                collector_dropdown["values"] = collector_options
                global collector_id_map
                collector_id_map = {row.FullName: row.CollectorID for row in rows}
                if collector_options:
                    collector_dropdown.set(collector_options[0])
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load collectors: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def display_donors(query=None):
        #Load donors from DB and populate listbox
        result_listbox.delete(0, tk.END)

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                if query:
                    cursor.execute("""
                        SELECT d.DonorID, d.FullName, d.CNIC, d.Address, d.ContactNumber, 
                        d.OfficeDonorID, d.DonationType, c.FullName AS CollectorName, 
                        d.PaymentMethod, d.Frequency, d.IsActive
                        FROM tblDonors d
                        JOIN tblCollectors c ON d.CollectorID = c.CollectorID
                        WHERE d.FullName LIKE ?
                        ORDER BY d.FullName
                    """, ('%' + query + '%',))
                else:
                    cursor.execute("""
                        SELECT d.DonorID, d.FullName, d.CNIC, d.Address, d.ContactNumber, 
                        d.OfficeDonorID, d.DonationType, c.FullName AS CollectorName, 
                        d.PaymentMethod, d.Frequency, d.IsActive
                        FROM tblDonors d
                        JOIN tblCollectors c ON d.CollectorID = c.CollectorID
                        ORDER BY d.FullName
                    """)
                rows = cursor.fetchall()
                conn.close()

                global donor_data_map
                donor_data_map = {}

                for row in rows:
                    result_listbox.insert(tk.END, row.FullName)
                    donor_data_map[row.FullName] = {
                        'DonorID': row.DonorID,
                        'CNIC': row.CNIC,
                        'Address': row.Address,
                        'ContactNumber': row.ContactNumber,
                        'OfficeDonorID': row.OfficeDonorID,
                        'DonationType': row.DonationType,
                        'CollectorName': row.CollectorName,
                        'PaymentMethod': row.PaymentMethod,
                        'Frequency': row.Frequency,
                        'IsActive': row.IsActive
                        
                        
                    }

            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to load donors: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_search(event=None):
        query = search_input.get().strip()
        display_donors(query)

    def on_select(event=None):
        global current_donor_id
        
        try:
            selected = result_listbox.get(result_listbox.curselection())
            data = donor_data_map[selected]
            current_donor_id = data['DonorID']
            #messagebox.showinfo("",current_donor_id)
            selected_donor_id[0] = current_donor_id
            # Populate form fields
            edit_name.delete(0, tk.END)
            edit_name.insert(0, selected)
            

            #edit_cnic.delete(0, tk.END)
            #edit_cnic.insert(0, data['CNIC'])

            #edit_address.delete(0, tk.END)
            #edit_address.insert(0, data['Address'])

            #edit_contact.delete(0, tk.END)
            #edit_contact.insert(0, data['ContactNumber'])

            edit_donorid.delete(0, tk.END)
            edit_donorid.insert(0, data['OfficeDonorID'])
            
            # Populate dropdowns
            donation_type_dropdown.set(data['DonationType'])
            #edit_commitment.delete(0, tk.END)
            #edit_commitment.insert(0, data['MonthlyCommitment'])
            
            collector_dropdown.set(data['CollectorName'])
            payment_method_dropdown.set(data['PaymentMethod'])
            frequency_dropdown.set(data['Frequency'])
            
            # Populate start and end dates
            #edit_start_date.set_date(data['CommitmentStartDate'])
            #edit_end_date.set_date(data['CommitmentEndDate'])
            is_active_var.set(data['IsActive'])

            
            
            # Store selected donor ID
            #selected_donor_id[0] = data['DonorID']
            #selected_donor_id[0] = current_donor_id
            
        except Exception as e:
            pass

    result_listbox.bind("<<ListboxSelect>>", on_select)

    def on_update():
        name = edit_name.get().strip()
        #cnic = edit_cnic.get().strip()
        cnic = ""
        #address = edit_address.get().strip()
        address = ""
        #contact = edit_contact.get().strip()
        contact = ""
        donoridvar = edit_donorid.get().strip()
        donation_type = donation_type_var.get()
        #commitment = edit_commitment.get().strip()
        collector = collector_var.get()
        payment_method = payment_method_var.get()
        frequency = frequency_var.get()
        #start_date = edit_start_date.get_date().strftime('%Y-%m-%d')
        #end_date = edit_end_date.get_date().strftime('%Y-%m-%d')
        is_active = is_active_var.get()


        if not name:
            messagebox.showwarning("Input Error", "Donor Full Name is required.")
            return

        #if not contact:
        #    messagebox.showwarning("Input Error", "Contact Number is required.")
        #    return

        if not collector or collector not in collector_id_map:
            messagebox.showwarning("Input Error", "Please select a valid collector.")
            return

        if not payment_method:
            messagebox.showwarning("Input Error", "Payment method is required.")
            return

        if not frequency:
            messagebox.showwarning("Input Error", "Donation frequency is required.")
            return

        collector_id = collector_id_map[collector]

        if selected_donor_id[0] is None:
            messagebox.showerror("Selection Error", "No donor selected for update.")
            return

        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE tblDonors SET
                        FullName = ?,
                        CNIC = ?,
                        Address = ?,
                        ContactNumber = ?,
                        OfficeDonorID = ?,
                        DonationType = ?,
                        CollectorID = ?,
                        PaymentMethod = ?,
                        Frequency = ?,
                        IsActive = ?
                    WHERE DonorID = ?
                """, (name, cnic, address, contact, donoridvar, donation_type, 
                      collector_id, payment_method, frequency, is_active, selected_donor_id[0]))
                conn.commit()
                conn.close()
                messagebox.showinfo("Success", "Donor updated successfully!")
                edit_win.destroy()
            except Exception as e:
                conn.rollback()
                messagebox.showerror("Error", f"Update failed: {e}")
        else:
            messagebox.showerror("Connection Failed", "Unable to connect to the database.")

    def on_close():
        edit_win.grab_release()
        edit_win.destroy()

    # === BUTTON PANEL ===
    btn_frame = tk.Frame(edit_win)
    btn_frame.pack(pady=10)

    tk.Button(btn_frame, text="Update", command=on_update, font=("Arial", 12), width=12).grid(row=0, column=0, padx=5)
    tk.Button(btn_frame, text="Close", command=on_close, font=("Arial", 12), width=10).grid(row=0, column=1, padx=5)

    # === LOAD DATA ===
    load_collectors()
    display_donors()  # Load all donors initially

    # Bind search input
    search_input.bind('<KeyRelease>', lambda e: on_search())

    # === KEYBOARD SHORTCUTS ===
    edit_win.bind('<Return>', lambda e: on_update())
    edit_win.bind('<Escape>', lambda e: on_close())

# ============================================================


#from tkinter import ttk
class AutocompleteCombobox(ttk.Combobox):
    def set_completion_list(self, completion_list):
        self._completion_list = sorted(completion_list, key=str.lower)
        self.configure(values=self._completion_list)
        self.bind('<KeyRelease>', self._handle_keyrelease)

    def _handle_keyrelease(self, event):
        if event.keysym in ("Left", "Right", "Down", "Up", "Return", "Escape"):
            return

        value = self.get().strip()
        if not value:
            data = self._completion_list
        else:
            data = [item for item in self._completion_list if item.lower().startswith(value.lower())]

        # Don’t auto-select after one letter — let user type
        self.configure(values=data)
        if len(value) >= 2 and data:
            self.event_generate('<Down>')

'''class AutocompleteCombobox(ttk.Combobox):
    def set_completion_list(self, completion_list):
        self._completion_list = sorted(completion_list, key=str.lower)
        self.configure(values=self._completion_list)
        self.bind('<KeyRelease>', self._handle_keyrelease)

    def _handle_keyrelease(self, event):
        if event.keysym in ("BackSpace", "Left", "Right", "Down", "Up", "Return"):
            return

        value = self.get()
        if value == '':
            data = self._completion_list
        else:
            data = [item for item in self._completion_list if item.lower().startswith(value.lower())]

        self.configure(values=data)
        if data:
            self.event_generate('<Down>')'''


# ============================================================
# Sponsorship Assignment
# ============================================================
#donor_dict = {}  # To store donor name → ID mapping




'''def open_sponsorship_mgmt():
    donor_summary_label = None  # Reference to label that shows summary
    global donor_dict
    donor_dict = {}
    
    win = tk.Toplevel()
    win.title("Sponsorship Assignment" + ORG_SUFFIX)
    win.configure(bg=COLORS["bg"])
    center_window(win, 1000, 500)
    win.transient()
    win.grab_set()
    win.focus_force()
    win.bind('<Escape>', lambda event: win.destroy())
    win.bind('<Return>', lambda e: assign_sponsorship())

    create_window_header(win, "Sponsorship Assignment")

    def get_remaining_amount_for_child(child_id):
        cursor.execute("""
            SELECT IFNULL(SUM(CAST(SponsorshipAmount AS REAL)), 0)
            FROM tblSponsorships
            WHERE ChildID=? AND IsActive=1
        """, (child_id,))
        sponsored = float(cursor.fetchone()[0] or 0)  # ✅ Convert to float

        cursor.execute("""
            SELECT CAST(ChildRequiredAmount AS REAL)
            FROM tblChildren
            WHERE ChildID=?
        """, (child_id,))
        required = float(cursor.fetchone()[0] or 0)  # ✅ Convert to float

        return required - sponsored



    def get_remaining_donor_capacity(donor_id):
        cursor.execute("SELECT IFNULL(SUM(CAST(SponsorshipAmount AS REAL)), 0) FROM tblSponsorships WHERE DonorID=? AND IsActive=1", (donor_id,))
        sponsored = cursor.fetchone()[0]
        cursor.execute("SELECT MonthlyCommitment FROM tblDonors WHERE DonorID=?", (donor_id,))
        capacity = cursor.fetchone()[0]
        return capacity - sponsored


    
    def load_children():
        today = datetime.today().strftime('%Y-%m-%d')

        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("""
            SELECT 
                c.ChildID,
                c.FullName,
                CAST(c.ChildRequiredAmount AS REAL) AS RequiredAmount,
                IFNULL(SUM(
                    CASE
                        WHEN s.IsActive = 1
                        AND (s.EndDate IS NULL OR s.EndDate >= ?)
                        THEN CAST(s.SponsorshipAmount AS REAL)
                        ELSE 0
                    END
                ), 0) AS SponsoredAmount
            FROM tblChildren c
            LEFT JOIN tblSponsorships s
                ON c.ChildID = s.ChildID
            GROUP BY
                c.ChildID,
                c.FullName,
                c.ChildRequiredAmount
            HAVING
                IFNULL(SUM(
                    CASE
                        WHEN s.IsActive = 1
                        AND (s.EndDate IS NULL OR s.EndDate >= ?)
                        THEN CAST(s.SponsorshipAmount AS REAL)
                        ELSE 0
                    END
                ), 0) < CAST(c.ChildRequiredAmount AS REAL)
            ORDER BY c.FullName
        """, (today, today))

        rows = cursor.fetchall()
        conn.close()

        return {
            f"{row.FullName} (ID:{row.ChildID})": row.ChildID
            for row in rows
        }

    

    donor_dict = {}  # Global dictionary to map display text -> DonorID


       
        
    def load_donors():
        global donor_dict

        cursor.execute("""
            SELECT d.DonorID, d.FullName
            FROM tblDonors d
            WHERE d.IsActive = 1 AND EXISTS (
                SELECT 1
                FROM (
                    SELECT c.DonorID,
                        IFNULL(SUM(c.MonthlyAmount), 0) AS TotalMonthlyCommitment,
                        IFNULL((
                            SELECT SUM(CAST(s.SponsorshipAmount AS REAL))
                            FROM tblSponsorships s
                            WHERE s.DonorID = c.DonorID AND s.IsActive = 1
                        ), 0) AS TotalSponsored
                    FROM tblCommitments c
                    WHERE c.IsActive = 1
                    GROUP BY c.DonorID
                ) t
                WHERE t.DonorID = d.DonorID
                AND (t.TotalMonthlyCommitment - t.TotalSponsored) > 0
            )
        """)
        
        donors = cursor.fetchall()
        donor_dict = {f"{row.FullName} (ID:{row.DonorID})": row.DonorID for row in donors}
        return list(donor_dict.keys())

        # Binding works as before
        #donor_combo.bind("<<ComboboxSelected>>", on_donor_selected)
      
    #donors_dict = donor_dict  # Alias used in some parts

    def on_child_select(event=None):
        selected = child_combo.get()
        if selected:
            child_id = children_dict.get(selected)
            cursor.execute("SELECT ChildRequiredAmount FROM tblChildren WHERE ChildID=?", (child_id,))
            result = cursor.fetchone()
            if result:
                required_amount_var.set(str(result.ChildRequiredAmount))
                update_percentage()
            remaining = get_remaining_amount_for_child(child_id)
            still_required_amount_var.set(f"{remaining:.2f}")  # ✅ Set into StringVar
            if remaining <= 0:
                amount_entry.config(state="disabled")
                assign_btn.config(state="disabled")
            else:
                amount_entry.config(state="normal")
                assign_btn.config(state="normal")
            
            #messagebox.showinfo("",still_required_amount_var)    
                
                
                
    def get_donor_financial_summary(donor_id):
        # Get donor info
        cursor.execute("SELECT FullName, MonthlyCommitment FROM tblDonors WHERE DonorID = ?", (donor_id,))
        donor = cursor.fetchone()
        donor_name = donor.FullName
        monthly_commitment = donor.MonthlyCommitment or 0

        # Get total assigned sponsorships
        cursor.execute("""
            SELECT IFNULL(SUM(CAST(SponsorshipAmount AS REAL)), 0) 
            FROM tblSponsorships 
            WHERE DonorID = ? AND IsActive = 1
        """, (donor_id,))
        assigned = cursor.fetchone()[0] or 0

        # Get total donations (optional)
        cursor.execute("""
            SELECT IFNULL(SUM(DonationAmount), 0)
            FROM tblDonations WHERE DonorID = ?
        """, (donor_id,))
        total_donated = cursor.fetchone()[0] or 0

        remaining = monthly_commitment - assigned

        summary = (
            f"Donor: {donor_name}\n"
            f"Monthly Commitment: {monthly_commitment}\n"
            f"Assigned Sponsorships: {assigned}\n"
            f"Remaining Capacity: {remaining}\n"
            f"Total Donations: {total_donated}"
        )

        return summary
    


                
                
                

    def update_percentage(event=None):
        try:
            required = float(required_amount_var.get())
            given = float(amount_entry.get())
            if required > 0:
                percent = round((given / required) * 100, 2)
                percentage_entry.config(state='normal')
                percentage_entry.delete(0, tk.END)
                percentage_entry.insert(0, str(percent))
                percentage_entry.config(state='readonly')
        except:
            percentage_entry.config(state='normal')
            percentage_entry.delete(0, tk.END)
            percentage_entry.config(state='readonly')

    def update_committed_amount(event=None):
        try:
            from dateutil.relativedelta import relativedelta
            start = start_date.get_date()
            end = end_date.get_date()
            months = (end.year - start.year) * 12 + (end.month - start.month) + 1
            amount = float(amount_entry.get())
            committed = round(amount * months, 2)
            #committed_entry.delete(0, tk.END)
            #committed_entry.insert(0, str(committed))
        except:
            pass

    def on_amount_change(event=None):
        update_percentage()
        update_committed_amount()

    
            
    def assign_sponsorship():
        try:
            if not child_combo.get().strip() or not donor_combo.get().strip():
                messagebox.showwarning("Validation Error", "Select both Child and Donor.")
                return

            amount = float(amount_entry.get())
            still_required = float(still_required_amount_var.get() or 0)

            if amount > still_required:
                messagebox.showwarning(
                    "Validation Error",
                    "Sponsorship amount cannot be greater than the still required amount."
                )
                return

            
            
            percentage = float(percentage_entry.get())
            #start = start_date.get_date()
            #end = end_date.get_date() if end_date.get() else None
            start = start_date.get_date().strftime('%Y-%m-%d')
            end = end_date.get_date().strftime('%Y-%m-%d') if end_date.get() else None

            notes = notes_entry.get("1.0", tk.END).strip()            
            child_id = children_dict[child_combo.get()]            
            donor_id = donor_dict[donor_combo.get()]
            
            
            if end and end < start:
                messagebox.showerror("Date Error", "End date cannot be earlier than start date.")
                return

            # Step 1: Find a commitment that matches date range
            # FIXED: Proper SQLite date comparison with NULL handling
            if end:
                # If end date provided, find commitment that covers the entire range
                cursor.execute("""
                    SELECT CommitmentID
                    FROM tblCommitments
                    WHERE DonorID = ?
                    AND IsActive = 1
                    AND DATE(CommitmentStartDate) <= DATE(?)
                    AND DATE(CommitmentEndDate) >= DATE(?)
                    ORDER BY CommitmentStartDate
                    LIMIT 1
                """, (donor_id, start, end))
            else:
                # If no end date, just check if start date falls within commitment period
                cursor.execute("""
                    SELECT CommitmentID
                    FROM tblCommitments
                    WHERE DonorID = ?
                    AND IsActive = 1
                    AND DATE(CommitmentStartDate) <= DATE(?)
                    AND (CommitmentEndDate IS NULL OR DATE(CommitmentEndDate) >= DATE(?))
                    ORDER BY CommitmentStartDate
                    LIMIT 1
                """, (donor_id, start, start))

            row = cursor.fetchone()
            if not row:
                # IMPROVED ERROR MESSAGE with debugging info
                cursor.execute("""
                    SELECT CommitmentStartDate, CommitmentEndDate, IsActive
                    FROM tblCommitments
                    WHERE DonorID = ?
                    ORDER BY CommitmentStartDate
                """, (donor_id,))
                
                all_commitments = cursor.fetchall()
                
                if not all_commitments:
                    messagebox.showerror(
                        "No Commitments Found",
                        f"This donor has no commitments in the system.\n\n"
                        f"Please create a commitment for this donor first."
                    )
                else:
                    commitment_details = "\n".join([
                        f"- {c[0]} to {c[1]} (Active: {c[2]})"
                        for c in all_commitments
                    ])
                    messagebox.showerror(
                        "Commitment Date Mismatch",
                        f"No commitment found covering sponsorship dates:\n"
                        f"Start: {start}\n"
                        f"End: {end if end else 'Not specified'}\n\n"
                        f"Available commitments for this donor:\n{commitment_details}\n\n"
                        f"Please adjust sponsorship dates to match an existing commitment."
                    )
                return

            commitment_id = row[0]

            # Step 2: Insert or update
            if selected_id.get():
                # Update existing
                cursor.execute("""
                    UPDATE tblSponsorships
                    SET DonorID=?, ChildID=?, SponsorshipAmount=?, Percentage=?,
                        StartDate=?, EndDate=?, Notes=?, IsActive=1, CommitmentID=?
                    WHERE SponsorshipID=?
                """, (donor_id, child_id, str(amount), percentage,
                    start, end if end else None,
                    notes, commitment_id, selected_id.get()))
            else:
                # Insert new
                cursor.execute("""
                    INSERT INTO tblSponsorships
                        (DonorID, ChildID, SponsorshipAmount, Percentage, StartDate, EndDate, Notes, IsActive, CommitmentID)
                    VALUES (?, ?, ?, ?, ?, ?, ?, 1, ?)
                """, (donor_id, child_id, str(amount), percentage,
                    start, end if end else None,
                    notes, commitment_id))

            conn.commit()
            load_treeview()
            
            win.grab_release()
            win.destroy()
            messagebox.showinfo("Success", "Sponsorship saved!")
            win.grab_release()
            open_sponsorship_mgmt()

        except Exception as e:
            messagebox.showerror("Error", str(e))



    def clear_form():
        child_combo.set('')
        donor_combo.set('')
        amount_entry.delete(0, tk.END)
        percentage_entry.config(state='normal')
        percentage_entry.delete(0, tk.END)
        percentage_entry.config(state='readonly')
        notes_entry.delete("1.0", tk.END)
        selected_id.set('')
        amount_entry.config(state='normal')
        amount_entry.delete(0, tk.END)
        #amount_entry.config(state='readonly')
        remaining_committed_entry.config(state='normal')
        remaining_committed_entry.delete(0, tk.END)
        remaining_committed_entry.config(state='readonly')
        donor_summary_label.config(text="")
        required_amount_var.set(str(''))
        
        

    def on_tree_select(event):
        selected = tree.selection()
        if selected:
            sid = selected[0]
            selected_id.set(sid)
            cursor.execute("SELECT * FROM tblSponsorships WHERE SponsorshipID=?", (sid,))
            row = cursor.fetchone()
            for k, v in children_dict.items():
                if v == row.ChildID: child_combo.set(k)
            for k, v in donors_dict.items():
                if v == row.DonorID: donor_combo.set(k)

            amount_entry.delete(0, tk.END)
            amount_entry.insert(0, row.SponsorshipAmount)

            percentage_entry.config(state='normal')
            percentage_entry.delete(0, tk.END)
            percentage_entry.insert(0, row.Percentage)
            percentage_entry.config(state='readonly')

            #committed_entry.delete(0, tk.END)
            #committed_entry.insert(0, row.AmountCommitted)

            start_date.set_date(parse_db_date(row.StartDate) or row.StartDate)
            if row.EndDate:
                end_date.set_date(parse_db_date(row.EndDate) or row.EndDate)
            notes_entry.delete("1.0", tk.END)
            notes_entry.insert(tk.END, row.Notes)

    def del_sponsorship():
        selected_item = tree.selection()
        if not selected_item:
            messagebox.showwarning("No Selection", "Please select a record to delete.")
            return
        sponsorship_id = tree.item(selected_item)['values'][0]
        if not messagebox.askyesno("Confirm Deletion", "Are you sure you want to delete this sponsorship?"):
            return
        cursor.execute("DELETE FROM tblSponsorships WHERE SponsorshipID=?", (sponsorship_id,))
        conn.commit()
        tree.delete(selected_item)
        messagebox.showinfo("Deleted", "Sponsorship deleted.")
        clear_form()

    def load_treeview():
        tree.delete(*tree.get_children())
        cursor.execute("""
            SELECT 
                s.SponsorshipID, 
                c.FullName AS ChildName, 
                d.FullName AS DonorName, 
                s.SponsorshipAmount,
                s.Percentage, 
                s.StartDate, 
                s.EndDate
            FROM tblSponsorships s
            JOIN tblChildren c ON s.ChildID = c.ChildID
            JOIN tblDonors d ON s.DonorID = d.DonorID
        """)
        for row in cursor.fetchall():
            tree.insert('', 'end', iid=row.SponsorshipID,
                        values=(
                            row.SponsorshipID,
                            row.ChildName,
                            row.DonorName,
                            row.SponsorshipAmount,
                            row.Percentage,
                            str(row.StartDate)[:10] if row.StartDate else '',
                            str(row.EndDate)[:10] if row.EndDate else ''
                        ))

    
            
    
        
    def on_donor_selected(event):
        selected_text = donor_combo.get()
        donor_id = donor_dict.get(selected_text)
        global remaining_monthly

        if donor_id:
            try:
                # Step 1: Get total monthly commitment from tblCommitments
                cursor.execute("""
                    SELECT IFNULL(SUM(MonthlyAmount), 0)
                    FROM tblCommitments
                    WHERE DonorID = ? AND IsActive = 1
                """, (donor_id,))
                total_monthly_commitment = cursor.fetchone()[0] or 0

                # Step 2: Get total monthly sponsorship allocations for this donor
                cursor.execute("""
                    SELECT IFNULL(SUM(CAST(SponsorshipAmount AS REAL)), 0)
                    FROM tblSponsorships
                    WHERE DonorID = ? AND IsActive = 1
                """, (donor_id,))
                total_allocated_monthly = cursor.fetchone()[0] or 0

                # Step 3: Calculate remaining monthly capacity
                remaining_monthly = total_monthly_commitment - total_allocated_monthly

                # Step 4: Display in Remaining Commitment Entry
                remaining_committed_entry.config(state='normal')
                remaining_committed_entry.delete(0, tk.END)
                remaining_committed_entry.insert(0, f"{remaining_monthly:.0f}")
                remaining_committed_entry.config(state='readonly')

                # Step 5: Fill Sponsorship Amount with remaining_monthly
                amount_entry.delete(0, tk.END)
                amount_entry.insert(0, f"{remaining_monthly:.0f}")

                # Step 6: Get active commitment duration
                # Step 6: Get active commitment duration and auto-fill dates
                cursor.execute("""
                    SELECT CommitmentStartDate, CommitmentEndDate
                    FROM tblCommitments
                    WHERE DonorID = ? AND IsActive = 1
                    ORDER BY CommitmentStartDate DESC
                    LIMIT 1
                """, (donor_id,))
                commitment = cursor.fetchone()
                if commitment:
                    start_dt, end_dt = commitment
                    # FIXED: Proper date parsing for SQLite strings
                    try:
                        if isinstance(start_dt, str):
                            parsed_start = datetime.strptime(start_dt[:10], "%Y-%m-%d")
                        else:
                            parsed_start = start_dt
                        start_date.set_date(parsed_start)
                    except Exception as e:
                        print(f"Start date parse error: {e}")
                    
                    try:
                        if end_dt and isinstance(end_dt, str):
                            parsed_end = datetime.strptime(end_dt[:10], "%Y-%m-%d")
                        elif end_dt:
                            parsed_end = end_dt
                        else:
                            parsed_end = None
                        
                        if parsed_end:
                            end_date.set_date(parsed_end)
                    except Exception as e:
                        print(f"End date parse error: {e}")

                # Step 7: Display detailed summary
                summary = (
                    f"Monthly Commitment: Rs. {total_monthly_commitment:.0f}\n"
                    f"Allocated (Sponsorship): Rs. {total_allocated_monthly:.0f}\n"
                    f"Remaining Monthly: Rs. {remaining_monthly:.0f}"
                )
                donor_summary_label.config(text=summary)
                update_percentage()

            except Exception as e:
                messagebox.showerror("Error", f"Failed to load donor commitments: {e}")


        


    # DB Setup
    conn = get_connection()
    cursor = conn.cursor()
    children_dict = load_children()
    donors_dict = donor_dict
    selected_id = tk.StringVar()

    form_frame = tk.Frame(win)
    form_frame.pack(pady=10)

    # --- Widgets setup below ---
    tk.Label(form_frame, text="Child:").grid(row=0, column=0)
    child_combo = AutocompleteCombobox(form_frame, width=30)
    child_combo.set_completion_list(list(children_dict.keys()))
    child_combo.grid(row=0, column=1)

    tk.Label(form_frame, text="Required Amount:").grid(row=0, column=4)
    required_amount_var = tk.StringVar()
    tk.Entry(form_frame, textvariable=required_amount_var, state='readonly').grid(row=0, column=5)
    
    tk.Label(form_frame, text="Still Required Amount:").grid(row=1, column=4)
    still_required_amount_var = tk.StringVar()
    tk.Entry(form_frame, textvariable=still_required_amount_var, state='readonly').grid(row=1, column=5)

    tk.Label(form_frame, text="Donor:").grid(row=1, column=0)

    donor_combo = AutocompleteCombobox(form_frame, width=30)
    donor_combo.grid(row=1, column=1)

    # Load donors and set combo list
    donor_names = load_donors()  # returns list of eligible donor display names
    donor_combo.set_completion_list(donor_names)

    donor_combo.bind("<<ComboboxSelected>>", on_donor_selected)

    # Add a label below or beside the combo to display summary
    donor_summary_label = tk.Label(form_frame, text="", anchor='w', justify='left', fg="blue")
    donor_summary_label.grid(row=2, column=1, columnspan=4, sticky='w', pady=5)


    tk.Label(form_frame, text="Sponsorship Amount:").grid(row=0, column=2)

    def validate_sponsorship_input(new_value):
        try:
            if new_value == "":
                return True  # Allow empty (user might be deleting)

            entered = float(new_value)

            # Ensure remaining_committed_entry has a valid number
            remaining_text = remaining_committed_entry.get()
            
            if not remaining_text:
                return True  # Nothing to compare against yet

            remaining = float(remaining_text)

            return entered <= remaining
        except ValueError:
            return False  # Reject invalid float
            
    vcmd = form_frame.register(validate_sponsorship_input)

    amount_entry = tk.Entry(
        form_frame,
        validate="key",
        validatecommand=(vcmd, '%P')
    )
    amount_entry.grid(row=0, column=3)


    tk.Label(form_frame, text="Percentage:").grid(row=1, column=2)
    percentage_entry = tk.Entry(form_frame, state='readonly')
    percentage_entry.grid(row=1, column=3)

 
    tk.Label(form_frame, text="Remaining Committed Amount:").grid(row=2, column=0)
    remaining_committed_entry = tk.Entry(form_frame)
    remaining_committed_entry.grid(row=3, column=0)

    tk.Label(form_frame, text="Start Date:").grid(row=2, column=2)
    start_date = DateEntry(form_frame, width=12)
    start_date.grid(row=2, column=3)

    tk.Label(form_frame, text="End Date:").grid(row=2, column=4)
    end_date = DateEntry(form_frame, width=12)
    end_date.grid(row=2, column=5)

    tk.Label(form_frame, text="Notes:").grid(row=4, column=0, pady=5)
    notes_entry = tk.Text(form_frame, height=3, width=70)
    notes_entry.grid(row=4, column=1, columnspan=3, pady=5)
    

    assign_btn = tk.Button(form_frame, text="Assign", command=assign_sponsorship, state="disabled" , bg="green", fg="white", width=20)
    assign_btn.grid(row=5, column=1, pady=5)

    
    tk.Button(form_frame, text="Clear", command=clear_form, bg="gray", fg="white", width=20).grid(row=5, column=2, pady=5)
    tk.Button(form_frame, text="Delete", command=del_sponsorship, bg="red", fg="white", width=20).grid(row=5, column=3, pady=5)

    tree = ttk.Treeview(win, columns=("ID", "Child", "Donor", "Amount", "Percent", "Start", "End"), show='headings')
    for col in tree["columns"]:
        tree.heading(col, text=col)
        tree.column(col, width=110)
    tree.pack(pady=10, fill='x')
    tree.bind('<<TreeviewSelect>>', on_tree_select)

    child_combo.bind("<<ComboboxSelected>>", on_child_select)
    amount_entry.bind("<KeyRelease>", on_amount_change)
    start_date.bind("<<DateEntrySelected>>", update_committed_amount)
    end_date.bind("<<DateEntrySelected>>", update_committed_amount)
    amount_entry.bind("<FocusOut>", on_amount_change)
    donor_combo.bind("<<ComboboxSelected>>", on_donor_selected)

    load_treeview()'''
def open_sponsorship_mgmt():

    # ─────────────────────────────────────────────────────────────────────────
    # WINDOW SETUP
    # ─────────────────────────────────────────────────────────────────────────
    win = tk.Toplevel()
    win.title("Sponsorship Assignment" + ORG_SUFFIX)
    win.state('zoomed')
    win.grab_set()
    win.focus_force()
    win.bind('<Escape>', lambda e: win.destroy())

    # ── Design tokens ─────────────────────────────────────────────────────────
    HEADER_BG    = "#0f2540"
    HEADER_FG    = "#f8fafc"
    SECTION_BG   = "#eef2f7"
    FRAME_BG     = "#ffffff"
    LABEL_FG     = "#374151"
    ACCENT       = "#1d6fd8"
    MUTED        = "#94a3b8"
    BTN_GREEN    = "#16a34a"
    BTN_GRAY     = "#64748b"
    BTN_RED      = "#dc2626"
    BTN_BLUE     = "#1d6fd8"
    BTN_ORANGE   = "#d97706"
    BTN_PURPLE   = "#7c3aed"
    FONT_LBL     = ("Segoe UI", 9)
    FONT_ENTRY   = ("Segoe UI", 9)
    FONT_HEADER  = ("Segoe UI", 13, "bold")
    FONT_SECTION = ("Segoe UI", 10, "bold")
    FONT_BTN     = ("Segoe UI", 9,  "bold")
    PAD          = {"padx": 6, "pady": 4}

    win.configure(bg=SECTION_BG)

    # ── DB connection ─────────────────────────────────────────────────────────
    conn   = get_connection()
    cursor = conn.cursor()

    # ─────────────────────────────────────────────────────────────────────────
    # STATE
    # ─────────────────────────────────────────────────────────────────────────
    donor_dict        = {}
    children_dict     = {}
    selected_id       = tk.StringVar()
    remaining_monthly = 0.0

    # ─────────────────────────────────────────────────────────────────────────
    # HEADER
    # ─────────────────────────────────────────────────────────────────────────
    header = tk.Frame(win, bg=HEADER_BG, height=56)
    header.pack(fill="x", side="top")
    header.pack_propagate(False)
    tk.Label(header, text="🤝  Sponsorship Assignment",
             font=FONT_HEADER, bg=HEADER_BG, fg=HEADER_FG,
             pady=10).pack(side="left", padx=20)
    tk.Frame(win, bg=ACCENT, height=3).pack(fill="x")

    # ─────────────────────────────────────────────────────────────────────────
    # FOOTER
    # ─────────────────────────────────────────────────────────────────────────
    footer = tk.Frame(win, bg="#e2e8f0", pady=8)
    footer.pack(fill="x", side="bottom")

    # ─────────────────────────────────────────────────────────────────────────
    # BODY
    # ─────────────────────────────────────────────────────────────────────────
    body = tk.Frame(win, bg=SECTION_BG)
    body.pack(fill="both", expand=True)

    # ── Scrollable form canvas ────────────────────────────────────────────────
    form_canvas   = tk.Canvas(body, bg=SECTION_BG, highlightthickness=0)
    form_v_scroll = tk.Scrollbar(body, orient="vertical",   command=form_canvas.yview)
    form_h_scroll = tk.Scrollbar(body, orient="horizontal", command=form_canvas.xview)
    form_h_scroll.pack(side="bottom", fill="x")
    form_v_scroll.pack(side="right",  fill="y")
    form_canvas.pack(side="top", fill="both", expand=False)
    form_canvas.configure(yscrollcommand=form_v_scroll.set,
                          xscrollcommand=form_h_scroll.set, height=320)

    inner    = tk.Frame(form_canvas, bg=SECTION_BG)
    inner_id = form_canvas.create_window((0, 0), window=inner, anchor="nw")

    def _on_inner_configure(e):
        form_canvas.configure(scrollregion=form_canvas.bbox("all"))
    def _on_canvas_configure(e):
        form_canvas.itemconfig(inner_id,
                               width=max(e.width, inner.winfo_reqwidth()))
    inner.bind("<Configure>",       _on_inner_configure)
    form_canvas.bind("<Configure>", _on_canvas_configure)

    def _mousewheel(e):
        try:
            form_canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
        except tk.TclError:
            pass
    form_canvas.bind_all("<MouseWheel>", _mousewheel)
    inner.columnconfigure((0, 1, 2), weight=1)

    # ── UI helpers ────────────────────────────────────────────────────────────
    def make_section(parent, title, row, col=0, colspan=1):
        lf = tk.LabelFrame(parent, text=f"  {title}  ",
                           font=FONT_SECTION, fg=ACCENT, bg=FRAME_BG,
                           bd=1, relief="groove", labelanchor="nw")
        lf.grid(row=row, column=col, columnspan=colspan,
                sticky="nsew", padx=8, pady=6)
        return lf

    def lbl(parent, text, row, col, **kwargs):
        tk.Label(parent, text=text, font=FONT_LBL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e", **kwargs).grid(
            row=row, column=col, sticky="e", **PAD)

    def ro_entry(parent, row, col, width=18, textvariable=None):
        e = tk.Entry(parent, font=FONT_ENTRY, width=width, relief="solid",
                     bd=1, state="readonly", readonlybackground="#f1f5f9",
                     textvariable=textvariable)
        e.grid(row=row, column=col, sticky="ew", **PAD)
        return e

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1 — SELECTION
    # ─────────────────────────────────────────────────────────────────────────
    sel_lf = make_section(inner, "Selection", row=0, col=0, colspan=3)
    sel_lf.columnconfigure((1, 3, 5), weight=1)

    lbl(sel_lf, "Collector:", 0, 0)
    collector_var   = tk.StringVar()
    collector_combo = ttk.Combobox(sel_lf, textvariable=collector_var,
                                   font=FONT_ENTRY, width=24, state="readonly")
    collector_combo.grid(row=0, column=1, sticky="ew", **PAD)

    lbl(sel_lf, "Child:", 0, 2)
    child_combo = AutocompleteCombobox(sel_lf, font=FONT_ENTRY, width=28)
    child_combo.grid(row=0, column=3, sticky="ew", **PAD)

    lbl(sel_lf, "Donor:", 0, 4)
    donor_combo = AutocompleteCombobox(sel_lf, font=FONT_ENTRY, width=28)
    donor_combo.grid(row=0, column=5, sticky="ew", **PAD)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2 — AMOUNTS
    # ─────────────────────────────────────────────────────────────────────────
    amt_lf = make_section(inner, "Amounts", row=1, col=0)
    amt_lf.columnconfigure(1, weight=1)

    lbl(amt_lf, "Required Amount:",      0, 0)
    required_amount_var = tk.StringVar()
    ro_entry(amt_lf, 0, 1, textvariable=required_amount_var)

    lbl(amt_lf, "Still Required:",       1, 0)
    still_required_amount_var = tk.StringVar()
    ro_entry(amt_lf, 1, 1, textvariable=still_required_amount_var)

    lbl(amt_lf, "Remaining Commitment:", 2, 0)
    remaining_committed_entry = tk.Entry(amt_lf, font=FONT_ENTRY, width=18,
                                         relief="solid", bd=1, state="readonly",
                                         readonlybackground="#f1f5f9")
    remaining_committed_entry.grid(row=2, column=1, sticky="ew", **PAD)

    lbl(amt_lf, "Sponsorship Amount:",   3, 0)

    def validate_sponsorship_input(new_value):
        try:
            if new_value == "":
                return True
            entered  = float(new_value)
            rem_text = remaining_committed_entry.get()
            if not rem_text:
                return True
            return entered <= float(rem_text)
        except ValueError:
            return False

    vcmd = amt_lf.register(validate_sponsorship_input)
    amount_entry = tk.Entry(amt_lf, font=FONT_ENTRY, width=18,
                            relief="solid", bd=1, validate="key",
                            validatecommand=(vcmd, '%P'))
    amount_entry.grid(row=3, column=1, sticky="ew", **PAD)

    lbl(amt_lf, "Percentage (%):", 4, 0)
    percentage_entry = tk.Entry(amt_lf, font=FONT_ENTRY, width=18,
                                relief="solid", bd=1, state="readonly",
                                readonlybackground="#f1f5f9")
    percentage_entry.grid(row=4, column=1, sticky="ew", **PAD)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3 — DONOR SUMMARY
    # ─────────────────────────────────────────────────────────────────────────
    sum_lf = make_section(inner, "Donor Financial Summary", row=1, col=1)
    sum_lf.columnconfigure(0, weight=1)
    sum_lf.rowconfigure(0, weight=1)

    donor_summary_label = tk.Label(sum_lf,
                                   text="Select a donor to view summary.",
                                   font=FONT_LBL, fg=MUTED, bg=FRAME_BG,
                                   anchor="nw", justify="left", wraplength=260)
    donor_summary_label.grid(row=0, column=0, sticky="nsew", padx=10, pady=8)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 4 — DATES & NOTES
    # ─────────────────────────────────────────────────────────────────────────
    date_lf = make_section(inner, "Dates & Notes", row=1, col=2)
    date_lf.columnconfigure(1, weight=1)

    lbl(date_lf, "Start Date:", 0, 0)
    start_date = DateEntry(date_lf, date_pattern="dd/MM/yyyy",
                           font=FONT_ENTRY, width=14)
    start_date.grid(row=0, column=1, sticky="ew", **PAD)

    lbl(date_lf, "End Date:", 1, 0)
    end_date = DateEntry(date_lf, date_pattern="dd/MM/yyyy",
                         font=FONT_ENTRY, width=14)
    end_date.grid(row=1, column=1, sticky="ew", **PAD)

    lbl(date_lf, "Notes:", 2, 0)
    notes_entry = tk.Text(date_lf, height=4, width=28,
                          font=FONT_ENTRY, relief="solid", bd=1)
    notes_entry.grid(row=2, column=1, sticky="ew", padx=6, pady=4)

    # ─────────────────────────────────────────────────────────────────────────
    # TREEVIEW
    # ─────────────────────────────────────────────────────────────────────────
    tree_lf = tk.LabelFrame(body, text="  Sponsorship Records  ",
                             font=FONT_SECTION, fg=ACCENT, bg=FRAME_BG,
                             bd=1, relief="groove", labelanchor="nw")
    tree_lf.pack(fill="both", expand=True, padx=8, pady=(0, 4))

    tv_v = tk.Scrollbar(tree_lf, orient="vertical")
    tv_h = tk.Scrollbar(tree_lf, orient="horizontal")

    TREE_COLS = ("S.No", "Child ID", "Child Name", "Age", "Class",
                 "School", "Donor/s", "Amount", "Start Date", "End Date", "Notes")

    tree = ttk.Treeview(tree_lf, columns=TREE_COLS, show="headings",
                        yscrollcommand=tv_v.set, xscrollcommand=tv_h.set)
    tv_v.config(command=tree.yview)
    tv_h.config(command=tree.xview)

    col_cfg = {
        "S.No": 45, "Child ID": 110, "Child Name": 175, "Age": 45,
        "Class": 75, "School": 175, "Donor/s": 185,
        "Amount": 90, "Start Date": 100, "End Date": 100, "Notes": 180,
    }
    for col in TREE_COLS:
        tree.heading(col, text=col)
        tree.column(col, width=col_cfg.get(col, 100), minwidth=40)

    tv_h.pack(side="bottom", fill="x")
    tv_v.pack(side="right",  fill="y")
    tree.pack(side="left", fill="both", expand=True)

    tree.tag_configure("odd",      background="#f0f4fa")
    tree.tag_configure("even",     background="#ffffff")
    tree.tag_configure("odd_sub",  background="#f0f4fa", foreground="#555555")
    tree.tag_configure("even_sub", background="#ffffff", foreground="#555555")

    # ─────────────────────────────────────────────────────────────────────────
    # SHARED HELPERS
    # ─────────────────────────────────────────────────────────────────────────
    def _calc_age(dob_str):
        import datetime as _dt
        if not dob_str:
            return ""
        try:
            dob   = _dt.datetime.strptime(str(dob_str)[:10], "%Y-%m-%d").date()
            today = _dt.date.today()
            return today.year - dob.year - (
                (today.month, today.day) < (dob.month, dob.day))
        except Exception:
            return ""

    def _fmt_date(raw):
        """YYYY-MM-DD for Excel."""
        if not raw:
            return ""
        try:
            return str(raw)[:10]
        except Exception:
            return str(raw)

    def _fmt_date_display(raw):
        """DD/MM/YYYY for Treeview."""
        if not raw:
            return ""
        try:
            y, m, d = str(raw)[:10].split("-")
            return f"{d}/{m}/{y}"
        except Exception:
            return str(raw)

    # ─────────────────────────────────────────────────────────────────────────
    # PIVOT QUERY  (sponsored children only — for treeview & export_sponsored)
    # ─────────────────────────────────────────────────────────────────────────
    def _build_pivot():
        from collections import OrderedDict
        cursor.execute("""
            SELECT
                c.ChildID,
                c.RegistrationNumber,
                c.FullName      AS ChildName,
                c.DateOfBirth,
                c.Class,
                c.SchoolName,
                d.FullName      AS DonorName,
                s.SponsorshipAmount,
                s.StartDate,
                s.EndDate,
                s.Notes,
                s.SponsorshipID
            FROM tblSponsorships s
            JOIN tblChildren c ON s.ChildID = c.ChildID
            JOIN tblDonors   d ON s.DonorID = d.DonorID
            ORDER BY c.FullName, s.SponsorshipID
        """)
        raw = cursor.fetchall()
        child_map = OrderedDict()
        for row in raw:
            cid = row.ChildID
            if cid not in child_map:
                child_map[cid] = {
                    "ChildID":            cid,
                    "RegistrationNumber": row.RegistrationNumber or "",
                    "ChildName":          row.ChildName,
                    "Age":                _calc_age(row.DateOfBirth),
                    "Class":              row.Class      or "",
                    "School":             row.SchoolName or "",
                    "sponsors": [],
                }
            child_map[cid]["sponsors"].append({
                "Donor":         row.DonorName,
                "Amount":        row.SponsorshipAmount or "",
                "StartDate":     _fmt_date(row.StartDate),
                "EndDate":       _fmt_date(row.EndDate),
                "Notes":         row.Notes or "",
                "SponsorshipID": row.SponsorshipID,
            })
        return list(child_map.values())

    # ─────────────────────────────────────────────────────────────────────────
    # ALL-CHILDREN QUERY  (for the full-list export)
    # Returns (sponsored_list, unsponsored_list)
    # sponsored   → same structure as _build_pivot()
    # unsponsored → same structure but sponsors=[]
    # ─────────────────────────────────────────────────────────────────────────
    def _build_all_children():
        from collections import OrderedDict
        import datetime as _dt
        today = _dt.date.today().strftime('%Y-%m-%d')

        # ── All active children ───────────────────────────────────────────────
        cursor.execute("""
            SELECT ChildID, RegistrationNumber, FullName,
                   DateOfBirth, Class, SchoolName
            FROM   tblChildren
            WHERE  Status != 'Inactive'
            ORDER  BY FullName
        """)
        all_children = {row.ChildID: {
            "ChildID":            row.ChildID,
            "RegistrationNumber": row.RegistrationNumber or "",
            "ChildName":          row.FullName,
            "Age":                _calc_age(row.DateOfBirth),
            "Class":              row.Class      or "",
            "School":             row.SchoolName or "",
            "sponsors":           [],
        } for row in cursor.fetchall()}

        # ── Fetch all active sponsorships ─────────────────────────────────────
        cursor.execute("""
            SELECT
                s.ChildID,
                d.FullName      AS DonorName,
                s.SponsorshipAmount,
                s.StartDate,
                s.EndDate,
                s.Notes,
                s.SponsorshipID
            FROM tblSponsorships s
            JOIN tblDonors d ON s.DonorID = d.DonorID
            WHERE s.IsActive = 1
            ORDER BY s.ChildID, s.SponsorshipID
        """)
        for row in cursor.fetchall():
            cid = row.ChildID
            if cid in all_children:
                all_children[cid]["sponsors"].append({
                    "Donor":         row.DonorName,
                    "Amount":        row.SponsorshipAmount or "",
                    "StartDate":     _fmt_date(row.StartDate),
                    "EndDate":       _fmt_date(row.EndDate),
                    "Notes":         row.Notes or "",
                    "SponsorshipID": row.SponsorshipID,
                })

        sponsored   = [c for c in all_children.values() if     c["sponsors"]]
        unsponsored = [c for c in all_children.values() if not c["sponsors"]]

        # Sort each group by name
        sponsored.sort(  key=lambda x: x["ChildName"])
        unsponsored.sort(key=lambda x: x["ChildName"])
        return sponsored, unsponsored

    # ─────────────────────────────────────────────────────────────────────────
    # DATA LOADERS
    # ─────────────────────────────────────────────────────────────────────────
    def load_children():
        nonlocal children_dict
        import datetime as _dt
        today = _dt.date.today().strftime('%Y-%m-%d')
        cursor.execute("""
            SELECT c.ChildID, c.FullName,
                   CAST(c.ChildRequiredAmount AS REAL) AS RequiredAmount,
                   IFNULL(SUM(
                       CASE WHEN s.IsActive=1
                             AND (s.EndDate IS NULL OR s.EndDate >= ?)
                            THEN CAST(s.SponsorshipAmount AS REAL) ELSE 0 END
                   ), 0) AS SponsoredAmount
            FROM   tblChildren c
            LEFT   JOIN tblSponsorships s ON c.ChildID = s.ChildID
            WHERE  c.Status != 'Inactive'
            GROUP  BY c.ChildID, c.FullName, c.ChildRequiredAmount
            HAVING IFNULL(SUM(
                       CASE WHEN s.IsActive=1
                             AND (s.EndDate IS NULL OR s.EndDate >= ?)
                            THEN CAST(s.SponsorshipAmount AS REAL) ELSE 0 END
                   ), 0) < CAST(c.ChildRequiredAmount AS REAL)
            ORDER  BY c.FullName
        """, (today, today))
        rows = cursor.fetchall()
        children_dict = {f"{r.FullName} (ID:{r.ChildID})": r.ChildID for r in rows}
        return children_dict

    collector_lookup = {"-- All Collectors --": None}

    def load_collectors():
        try:
            cursor.execute(
                "SELECT CollectorID, FullName FROM tblCollectors ORDER BY FullName")
            for r in cursor.fetchall():
                collector_lookup[r.FullName] = r.CollectorID
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load collectors:\n{e}")
        collector_combo["values"] = list(collector_lookup.keys())
        collector_combo.current(0)

    def load_donors_for_collector(event=None):
        nonlocal donor_dict
        cid = collector_lookup.get(collector_var.get())
        collector_clause = "AND d.CollectorID = ?" if cid is not None else ""
        params = [cid] if cid is not None else []
        cursor.execute(f"""
            SELECT d.DonorID, d.FullName
            FROM   tblDonors d
            WHERE  d.IsActive = 1 {collector_clause}
            AND EXISTS (
                SELECT 1 FROM (
                    SELECT c.DonorID,
                           IFNULL(SUM(c.MonthlyAmount), 0) AS TotalMonthly,
                           IFNULL((
                               SELECT SUM(CAST(s.SponsorshipAmount AS REAL))
                               FROM   tblSponsorships s
                               WHERE  s.DonorID = c.DonorID AND s.IsActive = 1
                           ), 0) AS TotalSponsored
                    FROM   tblCommitments c WHERE c.IsActive = 1
                    GROUP  BY c.DonorID
                ) t
                WHERE t.DonorID = d.DonorID
                  AND (t.TotalMonthly - t.TotalSponsored) > 0
            )
            ORDER BY d.FullName
        """, params)
        rows = cursor.fetchall()
        donor_dict = {f"{r.FullName} (ID:{r.DonorID})": r.DonorID for r in rows}
        donor_combo.set_completion_list(list(donor_dict.keys()))
        donor_var_str.set("")
        donor_summary_label.config(text="Select a donor to view summary.", fg=MUTED)
        remaining_committed_entry.config(state="normal")
        remaining_committed_entry.delete(0, tk.END)
        remaining_committed_entry.config(state="readonly")

    collector_combo.bind("<<ComboboxSelected>>", load_donors_for_collector)

    # ─────────────────────────────────────────────────────────────────────────
    # TREEVIEW LOADER
    # ─────────────────────────────────────────────────────────────────────────
    def load_treeview():
        tree.delete(*tree.get_children())
        pivot_rows = _build_pivot()
        sno = 0
        for grp_idx, child in enumerate(pivot_rows):
            sno += 1
            tag_main = "odd"     if grp_idx % 2 else "even"
            tag_sub  = "odd_sub" if grp_idx % 2 else "even_sub"
            for sp_idx, sp in enumerate(child["sponsors"]):
                if sp_idx == 0:
                    row_vals = (
                        sno,
                        child["RegistrationNumber"], child["ChildName"],
                        child["Age"], child["Class"], child["School"],
                        sp["Donor"], sp["Amount"],
                        _fmt_date_display(sp["StartDate"]),
                        _fmt_date_display(sp["EndDate"]),
                        sp["Notes"],
                    )
                    iid = f"child_{child['ChildID']}"
                    tree.insert("", "end", iid=iid, values=row_vals,
                                tags=(tag_main,))
                else:
                    row_vals = (
                        "", "", "", "", "", "",
                        sp["Donor"], sp["Amount"],
                        _fmt_date_display(sp["StartDate"]),
                        _fmt_date_display(sp["EndDate"]),
                        sp["Notes"],
                    )
                    iid = f"child_{child['ChildID']}_sp_{sp['SponsorshipID']}"
                    tree.insert("", "end", iid=iid, values=row_vals,
                                tags=(tag_sub,))

    # ─────────────────────────────────────────────────────────────────────────
    # SHARED EXCEL WRITER  (used by both export functions)
    # Writes one child block:
    #   • existing sponsor rows  (filled)
    #   • extra_blank blank rows for manual writing
    #   • child-info columns merged vertically across ALL rows
    # Returns the next available excel row number.
    # ─────────────────────────────────────────────────────────────────────────
    def _write_child_block(ws, child, start_excel_row, sno,
                           extra_blank, grp_idx,
                           fills, data_font,
                           thin, medium,
                           N_CHILD_COLS, TOTAL_COLS,
                           c_align_top, l_align_top,
                           c_align, l_align,
                           blank_donor_fill):
        from openpyxl.styles import Border

        def _border(left=thin, right=thin, top=thin, bottom=thin):
            return Border(left=left, right=right, top=top, bottom=bottom)

        fill       = fills[grp_idx % 2]
        n_existing = len(child["sponsors"])
        total_rows = n_existing + extra_blank   # total rows for this child block

        r = start_excel_row

        for row_offset in range(total_rows):
            is_first  = (row_offset == 0)
            is_last   = (row_offset == total_rows - 1)
            is_blank  = (row_offset >= n_existing)   # blank manual-write row

            top_side    = medium if is_first else thin
            bottom_side = medium if is_last  else thin

            # ── Child info columns (1-6) ──────────────────────────────────────
            child_vals = [
                sno                        if is_first else "",
                child["RegistrationNumber"] if is_first else "",
                child["ChildName"]          if is_first else "",
                child["Age"]                if is_first else "",
                child["Class"]              if is_first else "",
                child["School"]             if is_first else "",
            ]
            for col_offset, val in enumerate(child_vals):
                col_num = col_offset + 1
                c = ws.cell(row=r, column=col_num, value=val)
                c.font      = data_font
                c.fill      = fill
                c.alignment = (c_align_top if col_num in (1, 2, 4)
                               else l_align_top)
                c.border = _border(
                    left   = medium if col_num == 1            else thin,
                    right  = medium if col_num == N_CHILD_COLS else thin,
                    top    = top_side,
                    bottom = bottom_side,
                )

            # ── Donor columns (7-11) ──────────────────────────────────────────
            if is_blank:
                donor_vals = ["", "", "", "", ""]
                row_fill   = blank_donor_fill
            else:
                sp = child["sponsors"][row_offset]
                donor_vals = [
                    sp["Donor"], sp["Amount"],
                    sp["StartDate"], sp["EndDate"], sp["Notes"],
                ]
                row_fill = fill

            for col_offset, val in enumerate(donor_vals):
                col_num = N_CHILD_COLS + col_offset + 1   # 7..11
                c = ws.cell(row=r, column=col_num, value=val)
                c.font      = data_font
                c.fill      = row_fill
                c.alignment = (c_align if col_num in (8, 9, 10) else l_align)
                c.border = _border(
                    left   = medium if col_num == N_CHILD_COLS + 1 else thin,
                    right  = medium if col_num == TOTAL_COLS        else thin,
                    top    = top_side,
                    bottom = bottom_side,
                )

            ws.row_dimensions[r].height = 18
            r += 1

        # ── Merge child-info columns vertically if block > 1 row ─────────────
        if total_rows > 1:
            last_row = start_excel_row + total_rows - 1
            for col_num in range(1, N_CHILD_COLS + 1):
                ws.merge_cells(
                    start_row=start_excel_row, start_column=col_num,
                    end_row=last_row,          end_column=col_num,
                )
                mc = ws.cell(row=start_excel_row, column=col_num)
                mc.font      = data_font
                mc.fill      = fill
                mc.alignment = (c_align_top if col_num in (1, 2, 4)
                                else l_align_top)
                mc.border = _border(
                    left   = medium if col_num == 1            else thin,
                    right  = medium if col_num == N_CHILD_COLS else thin,
                    top    = medium,
                    bottom = medium,
                )

        return r   # next available row

    # ─────────────────────────────────────────────────────────────────────────
    # EXCEL EXPORT 1 — Sponsored children only  (same as before)
    # ─────────────────────────────────────────────────────────────────────────
    def export_to_excel():
        import datetime as _dt
        from tkinter import filedialog
        try:
            from openpyxl import Workbook
            from openpyxl.styles import (Font, PatternFill, Alignment,
                                         Border, Side)
            from openpyxl.utils import get_column_letter
        except ImportError:
            messagebox.showerror("Missing Library",
                                 "openpyxl is not installed.\nRun: pip install openpyxl")
            return

        pivot_rows = _build_pivot()
        if not pivot_rows:
            messagebox.showwarning("No Data", "No sponsorship records to export.")
            return

        default_name = f"Sponsorships_{_dt.date.today().strftime('%Y%m%d')}.xlsx"
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")],
            initialfile=default_name, title="Save Sponsorship Export")
        if not filepath:
            return

        wb  = Workbook()
        ws  = wb.active
        ws.title = "Sponsorships"

        thin   = Side(style="thin",   color="B0BEC5")
        medium = Side(style="medium", color="78909C")

        c_align     = Alignment(horizontal="center", vertical="center")
        l_align     = Alignment(horizontal="left",   vertical="center",  wrap_text=True)
        c_align_top = Alignment(horizontal="center", vertical="top")
        l_align_top = Alignment(horizontal="left",   vertical="top",     wrap_text=True)

        hdr_font  = Font(name="Arial", bold=True, color="FFFFFF", size=10)
        hdr_fill_child = PatternFill("solid", start_color="0F2540")
        hdr_fill_donor = PatternFill("solid", start_color="1D6FD8")
        hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)

        data_font  = Font(name="Arial", size=9)
        fills      = [PatternFill("solid", start_color="FFFFFF"),
                      PatternFill("solid", start_color="EEF2F7")]
        blank_fill = PatternFill("solid", start_color="FFF9F0")  # warm tint for blank rows

        COLS = [
            (1,  "S.No",       7,   True),
            (2,  "Child ID",   14,  True),
            (3,  "Child Name", 26,  True),
            (4,  "Age",         7,  True),
            (5,  "Class",      12,  True),
            (6,  "School",     26,  True),
            (7,  "Donor/s",    28,  False),
            (8,  "Amount",     12,  False),
            (9,  "Start Date", 13,  False),
            (10, "End Date",   13,  False),
            (11, "Notes",      32,  False),
        ]
        N_CHILD_COLS = 6
        TOTAL_COLS   = 11

        def _bdr(left=thin, right=thin, top=thin, bottom=thin):
            return Border(left=left, right=right, top=top, bottom=bottom)

        # Title
        ws.merge_cells(start_row=1, start_column=1,
                       end_row=1,   end_column=TOTAL_COLS)
        tc = ws.cell(row=1, column=1, value="Sponsorship Records")
        tc.font      = Font(name="Arial", bold=True, size=14, color="0F2540")
        tc.alignment = Alignment(horizontal="center", vertical="center")
        tc.border    = _bdr(left=medium, right=medium, top=medium, bottom=medium)
        ws.row_dimensions[1].height = 30

        # Header
        for col_num, label, width, is_child in COLS:
            fill = hdr_fill_child if is_child else hdr_fill_donor
            c = ws.cell(row=2, column=col_num, value=label)
            c.font = hdr_font; c.fill = fill
            c.alignment = hdr_align
            c.border = _bdr(left=medium, right=medium, top=medium, bottom=medium)
            ws.column_dimensions[get_column_letter(col_num)].width = width
        ws.row_dimensions[2].height = 22

        current_row = 3
        for grp_idx, child in enumerate(pivot_rows):
            current_row = _write_child_block(
                ws, child, current_row, grp_idx + 1,
                extra_blank=1, grp_idx=grp_idx,
                fills=fills, data_font=data_font,
                thin=thin, medium=medium,
                N_CHILD_COLS=N_CHILD_COLS, TOTAL_COLS=TOTAL_COLS,
                c_align_top=c_align_top, l_align_top=l_align_top,
                c_align=c_align, l_align=l_align,
                blank_donor_fill=blank_fill,
            )

        ws.freeze_panes = "A3"
        ws.auto_filter.ref = f"A2:{get_column_letter(TOTAL_COLS)}{current_row - 1}"
        wb.save(filepath)
        messagebox.showinfo("Export Successful",
                            f"Sponsorship data exported.\n\n{filepath}")

    # ─────────────────────────────────────────────────────────────────────────
    # EXCEL EXPORT 2 — ALL CHILDREN  (sponsored first, then unsponsored)
    # ─────────────────────────────────────────────────────────────────────────
    def export_all_children_excel():
        import datetime as _dt
        from tkinter import filedialog
        try:
            from openpyxl import Workbook
            from openpyxl.styles import (Font, PatternFill, Alignment,
                                         Border, Side)
            from openpyxl.utils import get_column_letter
        except ImportError:
            messagebox.showerror("Missing Library",
                                 "openpyxl is not installed.\nRun: pip install openpyxl")
            return

        sponsored, unsponsored = _build_all_children()
        if not sponsored and not unsponsored:
            messagebox.showwarning("No Data", "No children records found.")
            return

        default_name = f"AllChildren_{_dt.date.today().strftime('%Y%m%d')}.xlsx"
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")],
            initialfile=default_name, title="Save All Children Export")
        if not filepath:
            return

        wb  = Workbook()
        ws  = wb.active
        ws.title = "All Children"

        # ── Styles ────────────────────────────────────────────────────────────
        thin   = Side(style="thin",   color="B0BEC5")
        medium = Side(style="medium", color="78909C")

        c_align     = Alignment(horizontal="center", vertical="center")
        l_align     = Alignment(horizontal="left",   vertical="center",  wrap_text=True)
        c_align_top = Alignment(horizontal="center", vertical="top")
        l_align_top = Alignment(horizontal="left",   vertical="top",     wrap_text=True)

        hdr_font       = Font(name="Arial", bold=True, color="FFFFFF", size=10)
        hdr_fill_child = PatternFill("solid", start_color="0F2540")
        hdr_fill_donor = PatternFill("solid", start_color="1D6FD8")
        hdr_align      = Alignment(horizontal="center", vertical="center",
                                   wrap_text=True)

        data_font = Font(name="Arial", size=9)

        # Sponsored group: white/light-blue alternating
        fills_sp = [PatternFill("solid", start_color="FFFFFF"),
                    PatternFill("solid", start_color="EEF2F7")]
        # Unsponsored group: very light yellow alternating (visually distinct)
        fills_un = [PatternFill("solid", start_color="FFFFF0"),
                    PatternFill("solid", start_color="FEFCE8")]

        # Blank donor rows: warm peach tint (for manual writing)
        blank_fill = PatternFill("solid", start_color="FFF3E0")

        COLS = [
            (1,  "S.No",       7,   True),
            (2,  "Child ID",   14,  True),
            (3,  "Child Name", 26,  True),
            (4,  "Age",         7,  True),
            (5,  "Class",      12,  True),
            (6,  "School",     26,  True),
            (7,  "Donor/s",    28,  False),
            (8,  "Amount",     12,  False),
            (9,  "Start Date", 13,  False),
            (10, "End Date",   13,  False),
            (11, "Notes",      32,  False),
        ]
        N_CHILD_COLS = 6
        TOTAL_COLS   = 11

        def _bdr(left=thin, right=thin, top=thin, bottom=thin):
            return Border(left=left, right=right, top=top, bottom=bottom)

        # ── Title (row 1) ─────────────────────────────────────────────────────
        ws.merge_cells(start_row=1, start_column=1,
                       end_row=1,   end_column=TOTAL_COLS)
        tc = ws.cell(row=1, column=1, value="Complete Children Sponsorship List")
        tc.font      = Font(name="Arial", bold=True, size=14, color="0F2540")
        tc.alignment = Alignment(horizontal="center", vertical="center")
        tc.border    = _bdr(left=medium, right=medium, top=medium, bottom=medium)
        ws.row_dimensions[1].height = 30

        # ── Column header (row 2) ─────────────────────────────────────────────
        for col_num, label, width, is_child in COLS:
            fill = hdr_fill_child if is_child else hdr_fill_donor
            c = ws.cell(row=2, column=col_num, value=label)
            c.font = hdr_font; c.fill = fill
            c.alignment = hdr_align
            c.border = _bdr(left=medium, right=medium, top=medium, bottom=medium)
            ws.column_dimensions[get_column_letter(col_num)].width = width
        ws.row_dimensions[2].height = 22

        current_row = 3

        # ═════════════════════════════════════════════════════════════════════
        # SECTION A — SPONSORED CHILDREN
        # ═════════════════════════════════════════════════════════════════════
        if sponsored:
            for grp_idx, child in enumerate(sponsored):
                current_row = _write_child_block(
                    ws, child, current_row, grp_idx + 1,
                    extra_blank=1, grp_idx=grp_idx,
                    fills=fills_sp, data_font=data_font,
                    thin=thin, medium=medium,
                    N_CHILD_COLS=N_CHILD_COLS, TOTAL_COLS=TOTAL_COLS,
                    c_align_top=c_align_top, l_align_top=l_align_top,
                    c_align=c_align, l_align=l_align,
                    blank_donor_fill=blank_fill,
                )

        # Gap row between sections
        ws.row_dimensions[current_row].height = 10
        current_row += 1

        # ═════════════════════════════════════════════════════════════════════
        # SECTION B — UNSPONSORED CHILDREN
        # ═════════════════════════════════════════════════════════════════════
        if unsponsored:
            for grp_idx, child in enumerate(unsponsored):
                # No existing sponsors — 2 blank rows for manual writing
                current_row = _write_child_block(
                    ws, child, current_row, grp_idx + 1,
                    extra_blank=2, grp_idx=grp_idx,
                    fills=fills_un, data_font=data_font,
                    thin=thin, medium=medium,
                    N_CHILD_COLS=N_CHILD_COLS, TOTAL_COLS=TOTAL_COLS,
                    c_align_top=c_align_top, l_align_top=l_align_top,
                    c_align=c_align, l_align=l_align,
                    blank_donor_fill=blank_fill,
                )

        # ── Freeze & save ─────────────────────────────────────────────────────
        ws.freeze_panes = "A3"
        wb.save(filepath)
        messagebox.showinfo(
            "Export Successful",
            f"All children list exported.\n\n"
            f"Sponsored:   {len(sponsored)}\n"
            f"Unsponsored: {len(unsponsored)}\n\n"
            f"{filepath}")

    # ─────────────────────────────────────────────────────────────────────────
    # HELPER CALCULATIONS
    # ─────────────────────────────────────────────────────────────────────────
    def get_remaining_amount_for_child(child_id):
        cursor.execute("""
            SELECT IFNULL(SUM(CAST(SponsorshipAmount AS REAL)), 0)
            FROM tblSponsorships WHERE ChildID=? AND IsActive=1
        """, (child_id,))
        sponsored = float(cursor.fetchone()[0] or 0)
        cursor.execute("""
            SELECT CAST(ChildRequiredAmount AS REAL)
            FROM tblChildren WHERE ChildID=?
        """, (child_id,))
        required = float(cursor.fetchone()[0] or 0)
        return required - sponsored

    def update_percentage(event=None):
        try:
            required = float(required_amount_var.get())
            given    = float(amount_entry.get())
            if required > 0:
                percent = round((given / required) * 100, 2)
                percentage_entry.config(state="normal")
                percentage_entry.delete(0, tk.END)
                percentage_entry.insert(0, str(percent))
                percentage_entry.config(state="readonly")
        except Exception:
            percentage_entry.config(state="normal")
            percentage_entry.delete(0, tk.END)
            percentage_entry.config(state="readonly")

    def update_committed_amount(event=None):
        try:
            import datetime as _dt
            s = start_date.get_date(); e = end_date.get_date()
            months = (e.year - s.year) * 12 + (e.month - s.month) + 1
            round(float(amount_entry.get()) * months, 2)
        except Exception:
            pass

    def on_amount_change(event=None):
        update_percentage()
        update_committed_amount()

    # ─────────────────────────────────────────────────────────────────────────
    # EVENT HANDLERS
    # ─────────────────────────────────────────────────────────────────────────
    def on_child_select(event=None):
        selected = child_combo.get()
        if not selected:
            return
        child_id = children_dict.get(selected)
        if child_id is None:
            return
        cursor.execute(
            "SELECT ChildRequiredAmount FROM tblChildren WHERE ChildID=?",
            (child_id,))
        result = cursor.fetchone()
        if result:
            required_amount_var.set(str(result.ChildRequiredAmount))
            update_percentage()
        remaining = get_remaining_amount_for_child(child_id)
        still_required_amount_var.set(f"{remaining:.2f}")
        if remaining <= 0:
            amount_entry.config(state="disabled")
            assign_btn.config(state="disabled")
        else:
            amount_entry.config(state="normal")
            assign_btn.config(state="normal")

    def on_donor_selected(event=None):
        nonlocal remaining_monthly
        import datetime as _dt
        selected_text = donor_combo.get()
        donor_id      = donor_dict.get(selected_text)
        if not donor_id:
            return
        try:
            cursor.execute("""
                SELECT IFNULL(SUM(MonthlyAmount), 0) FROM tblCommitments
                WHERE DonorID=? AND IsActive=1
            """, (donor_id,))
            total_monthly = float(cursor.fetchone()[0] or 0)

            cursor.execute("""
                SELECT IFNULL(SUM(CAST(SponsorshipAmount AS REAL)), 0)
                FROM tblSponsorships WHERE DonorID=? AND IsActive=1
            """, (donor_id,))
            total_allocated = float(cursor.fetchone()[0] or 0)

            remaining_monthly = total_monthly - total_allocated
            remaining_committed_entry.config(state="normal")
            remaining_committed_entry.delete(0, tk.END)
            remaining_committed_entry.insert(0, f"{remaining_monthly:.0f}")
            remaining_committed_entry.config(state="readonly")

            amount_entry.delete(0, tk.END)
            amount_entry.insert(0, f"{remaining_monthly:.0f}")

            cursor.execute("""
                SELECT CommitmentStartDate, CommitmentEndDate
                FROM tblCommitments WHERE DonorID=? AND IsActive=1
                ORDER BY CommitmentStartDate DESC LIMIT 1
            """, (donor_id,))
            commitment = cursor.fetchone()
            if commitment:
                s_raw, e_raw = commitment
                try:
                    start_date.set_date(
                        _dt.datetime.strptime(str(s_raw)[:10], "%Y-%m-%d")
                        if isinstance(s_raw, str) else s_raw)
                except Exception:
                    pass
                try:
                    if e_raw:
                        end_date.set_date(
                            _dt.datetime.strptime(str(e_raw)[:10], "%Y-%m-%d")
                            if isinstance(e_raw, str) else e_raw)
                except Exception:
                    pass

            cursor.execute("""
                SELECT IFNULL(SUM(DonationAmount), 0)
                FROM tblDonations WHERE DonorID=?
            """, (donor_id,))
            total_donated = float(cursor.fetchone()[0] or 0)

            donor_summary_label.config(
                text=(
                    f"Monthly Commitment :  Rs. {total_monthly:,.0f}\n"
                    f"Allocated (Sponsorships) :  Rs. {total_allocated:,.0f}\n"
                    f"Remaining Monthly :  Rs. {remaining_monthly:,.0f}\n"
                    f"Total Donations :  Rs. {total_donated:,.0f}"
                ), fg=ACCENT)
            update_percentage()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load donor info:\n{e}")

    def on_tree_select(event):
        selected = tree.selection()
        if not selected:
            return
        iid = selected[0]
        try:
            child_id = int(iid.split("_")[1])
        except (IndexError, ValueError):
            return
        for k, v in children_dict.items():
            if v == child_id:
                child_combo.set(k)
                on_child_select()
                break
        selected_id.set("")
        donor_combo.set("")
        donor_summary_label.config(text="Select a donor to view summary.", fg=MUTED)
        remaining_committed_entry.config(state="normal")
        remaining_committed_entry.delete(0, tk.END)
        remaining_committed_entry.config(state="readonly")
        assign_btn.grid()
        update_btn.grid_remove()

    tree.bind("<<TreeviewSelect>>", on_tree_select)

    # ─────────────────────────────────────────────────────────────────────────
    # CORE ACTIONS
    # ─────────────────────────────────────────────────────────────────────────
    def assign_sponsorship():
        try:
            if not child_combo.get().strip() or not donor_combo.get().strip():
                messagebox.showwarning("Validation Error",
                                       "Please select both a Child and a Donor.")
                return
            amount_txt = amount_entry.get().strip()
            if not amount_txt:
                messagebox.showwarning("Validation Error",
                                       "Please enter a sponsorship amount.")
                return
            amount         = float(amount_txt)
            still_required = float(still_required_amount_var.get() or 0)
            if amount > still_required:
                messagebox.showwarning(
                    "Validation Error",
                    f"Sponsorship amount (Rs. {amount:,.2f}) cannot exceed "
                    f"the still-required amount (Rs. {still_required:,.2f}).")
                return
            percentage = float(percentage_entry.get() or 0)
            start      = start_date.get_date().strftime('%Y-%m-%d')
            end_val    = end_date.get_date()
            end        = end_val.strftime('%Y-%m-%d') if end_val else None
            notes      = notes_entry.get("1.0", tk.END).strip()
            child_id   = children_dict[child_combo.get()]
            donor_id   = donor_dict[donor_combo.get()]
            if end and end < start:
                messagebox.showerror("Date Error",
                                     "End date cannot be earlier than start date.")
                return
            if end:
                cursor.execute("""
                    SELECT CommitmentID FROM tblCommitments
                    WHERE DonorID=? AND IsActive=1
                      AND DATE(CommitmentStartDate) <= DATE(?)
                      AND DATE(CommitmentEndDate)   >= DATE(?)
                    ORDER BY CommitmentStartDate LIMIT 1
                """, (donor_id, start, end))
            else:
                cursor.execute("""
                    SELECT CommitmentID FROM tblCommitments
                    WHERE DonorID=? AND IsActive=1
                      AND DATE(CommitmentStartDate) <= DATE(?)
                      AND (CommitmentEndDate IS NULL
                           OR DATE(CommitmentEndDate) >= DATE(?))
                    ORDER BY CommitmentStartDate LIMIT 1
                """, (donor_id, start, start))
            row = cursor.fetchone()
            if not row:
                cursor.execute("""
                    SELECT CommitmentStartDate, CommitmentEndDate, IsActive
                    FROM tblCommitments WHERE DonorID=?
                    ORDER BY CommitmentStartDate
                """, (donor_id,))
                all_c = cursor.fetchall()
                if not all_c:
                    messagebox.showerror("No Commitments Found",
                                         "This donor has no commitments.\n"
                                         "Please create a commitment first.")
                else:
                    details = "\n".join(
                        f"  • {c[0]} → {c[1]}  (Active: {c[2]})" for c in all_c)
                    messagebox.showerror(
                        "Commitment Date Mismatch",
                        f"No commitment covers the selected sponsorship dates:\n"
                        f"  Start: {start}\n  End:   {end or 'Not specified'}\n\n"
                        f"Available commitments:\n{details}\n\n"
                        "Please adjust the sponsorship dates.")
                return
            commitment_id = row[0]
            if selected_id.get() != "":
                cursor.execute("""
                    UPDATE tblSponsorships
                    SET DonorID=?, ChildID=?, SponsorshipAmount=?, Percentage=?,
                        StartDate=?, EndDate=?, Notes=?, IsActive=1, CommitmentID=?
                    WHERE SponsorshipID=?
                """, (donor_id, child_id, str(amount), percentage,
                      start, end, notes, commitment_id, selected_id.get()))
            else:
                cursor.execute("""
                    INSERT INTO tblSponsorships
                        (DonorID, ChildID, SponsorshipAmount, Percentage,
                         StartDate, EndDate, Notes, IsActive, CommitmentID)
                    VALUES (?, ?, ?, ?, ?, ?, ?, 1, ?)
                """, (donor_id, child_id, str(amount), percentage,
                      start, end, notes, commitment_id))
            conn.commit()
            messagebox.showinfo("Success", "Sponsorship saved successfully!")
            win.grab_release()
            win.destroy()
            open_sponsorship_mgmt()
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def update_sponsorship():
        assign_sponsorship()

    def del_sponsorship():
        selected_item = tree.selection()
        if not selected_item:
            messagebox.showwarning("No Selection", "Please select a row first.")
            return
        iid = selected_item[0]
        try:
            child_id = int(iid.split("_")[1])
        except (IndexError, ValueError):
            return
        cursor.execute("""
            SELECT s.SponsorshipID, d.FullName, s.SponsorshipAmount,
                   s.StartDate, s.EndDate
            FROM tblSponsorships s
            JOIN tblDonors d ON s.DonorID = d.DonorID
            WHERE s.ChildID = ?
            ORDER BY s.SponsorshipID
        """, (child_id,))
        sponsors = cursor.fetchall()
        if not sponsors:
            messagebox.showwarning("No Sponsorships",
                                   "No sponsorships found for this child.")
            return
        if len(sponsors) == 1:
            sid, dname, amt, sd, ed = sponsors[0]
            if not messagebox.askyesno(
                    "Confirm Delete",
                    f"Delete sponsorship by {dname} (Rs. {amt}, {sd} – {ed})?"):
                return
        else:
            pick_win = tk.Toplevel(win)
            pick_win.title("Select Sponsorship to Delete")
            pick_win.grab_set(); pick_win.focus_force()
            tk.Label(pick_win,
                     text="This child has multiple sponsorships.\nSelect one to delete:",
                     font=FONT_LBL, pady=8).pack(padx=12)
            listbox = tk.Listbox(pick_win, width=65, height=len(sponsors),
                                 font=FONT_ENTRY)
            listbox.pack(padx=12, pady=6)
            sid_map = {}
            for s in sponsors:
                label = f"{s[1]}  |  Rs. {s[2]}  |  {s[3]} – {s[4]}"
                listbox.insert(tk.END, label)
                sid_map[label] = s[0]
            chosen_sid = [None]
            def _do_pick():
                sel = listbox.curselection()
                if not sel:
                    messagebox.showwarning("No Selection",
                                           "Please select a sponsorship.",
                                           parent=pick_win)
                    return
                chosen_sid[0] = sid_map[listbox.get(sel[0])]
                pick_win.destroy()
            tk.Button(pick_win, text="Delete Selected", command=_do_pick,
                      bg=BTN_RED, fg="white", font=FONT_BTN).pack(pady=8)
            pick_win.wait_window()
            if chosen_sid[0] is None:
                return
            sid = chosen_sid[0]
            if not messagebox.askyesno("Confirm Delete",
                                       "Are you sure you want to delete this sponsorship?"):
                return
        cursor.execute("DELETE FROM tblSponsorships WHERE SponsorshipID=?", (sid,))
        conn.commit()
        messagebox.showinfo("Deleted", "Sponsorship deleted successfully.")
        load_treeview()
        clear_form()

    def clear_form():
        selected_id.set("")
        child_combo.set(""); donor_combo.set("")
        amount_entry.config(state="normal"); amount_entry.delete(0, tk.END)
        percentage_entry.config(state="normal"); percentage_entry.delete(0, tk.END)
        percentage_entry.config(state="readonly")
        remaining_committed_entry.config(state="normal")
        remaining_committed_entry.delete(0, tk.END)
        remaining_committed_entry.config(state="readonly")
        notes_entry.delete("1.0", tk.END)
        required_amount_var.set(""); still_required_amount_var.set("")
        donor_summary_label.config(text="Select a donor to view summary.", fg=MUTED)
        assign_btn.config(state="disabled"); assign_btn.grid()
        update_btn.grid_remove()

    # ─────────────────────────────────────────────────────────────────────────
    # ACTION BUTTONS
    # ─────────────────────────────────────────────────────────────────────────
    btn_lf = make_section(inner, "Actions", row=2, col=0, colspan=3)

    assign_btn = tk.Button(
        btn_lf, text="✔  Assign Sponsorship",
        command=assign_sponsorship, state="disabled",
        font=FONT_BTN, bg=BTN_GREEN, fg="white",
        relief="flat", cursor="hand2", padx=16, pady=6)
    assign_btn.grid(row=0, column=0, padx=8, pady=8)

    update_btn = tk.Button(
        btn_lf, text="✎  Update Sponsorship",
        command=update_sponsorship,
        font=FONT_BTN, bg=BTN_BLUE, fg="white",
        relief="flat", cursor="hand2", padx=16, pady=6)
    update_btn.grid(row=0, column=1, padx=8, pady=8)
    update_btn.grid_remove()

    tk.Button(btn_lf, text="✖  Clear Form", command=clear_form,
              font=FONT_BTN, bg=BTN_GRAY, fg="white",
              relief="flat", cursor="hand2", padx=16, pady=6).grid(
        row=0, column=2, padx=8, pady=8)

    tk.Button(btn_lf, text="🗑  Delete Sponsorship", command=del_sponsorship,
              font=FONT_BTN, bg=BTN_RED, fg="white",
              relief="flat", cursor="hand2", padx=16, pady=6).grid(
        row=0, column=3, padx=8, pady=8)

    tk.Button(btn_lf, text="📥  Export Sponsored", command=export_to_excel,
              font=FONT_BTN, bg=BTN_ORANGE, fg="white",
              relief="flat", cursor="hand2", padx=16, pady=6).grid(
        row=0, column=4, padx=8, pady=8)

    # ── NEW BUTTON ────────────────────────────────────────────────────────────
    tk.Button(btn_lf, text="📋  Export All Children", command=export_all_children_excel,
              font=FONT_BTN, bg=BTN_PURPLE, fg="white",
              relief="flat", cursor="hand2", padx=16, pady=6).grid(
        row=0, column=5, padx=8, pady=8)

    tk.Button(footer, text="✖  Close", command=win.destroy,
              font=FONT_BTN, bg=BTN_RED, fg="white",
              relief="flat", cursor="hand2",
              padx=24, pady=6).pack(side="right", padx=20)

    # ─────────────────────────────────────────────────────────────────────────
    # BINDINGS
    # ─────────────────────────────────────────────────────────────────────────
    donor_var_str = tk.StringVar()
    donor_combo.configure(textvariable=donor_var_str)

    child_combo.bind("<<ComboboxSelected>>",  on_child_select)
    donor_combo.bind("<<ComboboxSelected>>",  on_donor_selected)
    amount_entry.bind("<KeyRelease>",         on_amount_change)
    amount_entry.bind("<FocusOut>",           on_amount_change)
    start_date.bind("<<DateEntrySelected>>",  update_committed_amount)
    end_date.bind("<<DateEntrySelected>>",    update_committed_amount)

    win.bind('<Return>', lambda e: (
        assign_btn.invoke() if assign_btn.winfo_ismapped()
        and str(assign_btn.cget("state")) != "disabled"
        else update_btn.invoke() if update_btn.winfo_ismapped() else None
    ))

    # ─────────────────────────────────────────────────────────────────────────
    # INITIALISE DATA
    # ─────────────────────────────────────────────────────────────────────────
    load_children()
    child_combo.set_completion_list(list(children_dict.keys()))
    load_collectors()
    load_donors_for_collector()
    load_treeview()



# ============================================================

# ============================================================
# Children Reports
# ============================================================


def open_child_report_form():
    report_win = tk.Toplevel()
    report_win.title("Child Admission Report" + ORG_SUFFIX)
    report_win.geometry("900x500")

    ww, wh = 900, 500
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="From Date:").grid(row=0, column=0, padx=5)
    from_date = DateEntry(filter_frame, width=12, date_pattern="dd/mm/yyyy")
    from_date.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="To Date:").grid(row=0, column=2, padx=5)
    to_date = DateEntry(filter_frame, width=12, date_pattern="dd/mm/yyyy")
    to_date.grid(row=0, column=3, padx=5)


    def show_print_dialog():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return

        try:
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
            if not os.path.exists(downloads_folder):
                os.makedirs(downloads_folder)

            # dd/mm/yyyy → yyyy-mm-dd for file name
            f_start = datetime.strptime(from_date.get(), "%d/%m/%Y").strftime("%Y-%m-%d")
            f_end = datetime.strptime(to_date.get(), "%d/%m/%Y").strftime("%Y-%m-%d")

            file_name = f"Children_Admission_Report_{f_start}_to_{f_end}.pdf"
            file_path = os.path.join(downloads_folder, file_name)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Main headings
            c.setFont("Helvetica-Bold", 16)
            c.drawString(200, height - 40, "Alkhidmat Khawateen Trust")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(213, height - 55, "Children Admission Report")

            # Show duration in dd/mm/yyyy
            c.setFont("Helvetica", 10)
            c.drawString(219, height - 70,
                         f"Duration: {from_date.get()} to {to_date.get()}")

            # Center name
            if report_rows:
                unique_centers = set(row.CenterName for row in report_rows)
                if len(unique_centers) == 1:
                    center_name = list(unique_centers)[0]
                    c.drawString(40, height - 100, f"Center: {center_name}")
                else:
                    c.drawString(40, height - 100, "Center: Multiple Centers")

            # Table headers
            y = height - 120
            c.setFont("Helvetica-Bold", 10)
            c.drawString(40, y, "Full Name")
            c.drawString(180, y, "Father Name")
            c.drawString(320, y, "Admission Date")

            # Table data
            y -= 20
            c.setFont("Helvetica", 9)
            for row in report_rows:
                adm_date = format_db_date(row.AdmissionDate, "%d/%m/%Y")

                c.drawString(40, y, str(row.FullName))
                c.drawString(180, y, str(row.FatherName))
                c.drawString(320, y, adm_date)

                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            time.sleep(0.5)
            open_file_cross_platform(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))



    def load_report():
        tree.delete(*tree.get_children())

        # Convert dateEntry format dd/mm/yyyy → yyyy-mm-dd
        start = datetime.strptime(from_date.get(), "%d/%m/%Y").strftime("%Y-%m-%d")
        end = datetime.strptime(to_date.get(), "%d/%m/%Y").strftime("%Y-%m-%d")

        try:
            conn = get_connection()
            cursor = conn.cursor()

            query = """
                SELECT c.FullName, c.FatherName, c.AdmissionDate, ce.CenterName
                FROM tblChildren c
                JOIN tblCenters ce ON c.CenterID = ce.CenterID
                WHERE c.AdmissionDate BETWEEN ? AND ?
                ORDER BY c.AdmissionDate ASC
            """
            cursor.execute(query, (start, end))
            rows = cursor.fetchall()
            report_rows.clear()
            report_rows.extend(rows)

            for row in rows:
                adm_date = format_db_date(row.AdmissionDate, "%d/%m/%Y")
                tree.insert('', 'end',
                            values=(row.FullName, row.FatherName, adm_date, row.CenterName))

            conn.close()
            print_btn.pack(pady=10)

        except Exception as e:
            messagebox.showerror("Error", f"Failed to load report:\n{e}")


    tk.Button(filter_frame, text="Generate Report", command=load_report,
              bg="blue", fg="white", width=20).grid(row=0, column=4, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", command=show_print_dialog)
    print_btn.pack_forget()

    columns = ("Child Name", "Father Name", "Admission Date", "Center Name")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150)
    tree.pack(padx=10, pady=10, fill="both", expand=True)




# ============================================================

# -------------Start of Sponsorship Pie Chart ------------
# Ensure these imports exist in your environment:
# pip install matplotlib pillow reportlab

def open_child_sponsorship_graph_report(parent=None):
    """
    Open a window that lets user search/select a child and shows a pie chart
    of sponsorship by donor for that child. Includes Print/Save PDF capability.
    """
    # --- Local helpers / fallbacks ---
    try:
        AutocompleteCombobox  # if defined elsewhere
    except NameError:
        class AutocompleteCombobox(ttk.Combobox):
            def set_completion_list(self, completion_list):
                self._completion_list = sorted(completion_list, key=str.lower)
                self.configure(values=self._completion_list)
                self.bind('<KeyRelease>', self._handle_keyrelease)
            def _handle_keyrelease(self, event):
                if event.keysym in ("BackSpace", "Left", "Right", "Down", "Up", "Return"):
                    return
                value = self.get()
                if value == '':
                    data = self._completion_list
                else:
                    data = [item for item in self._completion_list if item.lower().startswith(value.lower())]
                self.configure(values=data)
                if data:
                    self.event_generate('<Down>')

    # --- Window ---
    win = tk.Toplevel(parent) if parent else tk.Toplevel()
    win.title("Child Sponsorship Graph" + ORG_SUFFIX)
    W, H = 900, 520
    sw = win.winfo_screenwidth(); sh = win.winfo_screenheight()
    x = (sw//2) - (W//2); y = (sh//2) - (H//2)
    win.geometry(f"{W}x{H}+{x}+{y}")
    win.transient(parent)
    win.grab_set()
    win.focus_force()
    win.bind("<Escape>", lambda e: win.destroy())

    # --- Layout frames ---
    left_frame = tk.Frame(win)
    left_frame.pack(side="left", fill="both", expand=False, padx=8, pady=8)
    right_frame = tk.Frame(win)
    right_frame.pack(side="right", fill="both", expand=True, padx=8, pady=8)

    # --- Controls on left ---
    tk.Label(left_frame, text="Search Child:").grid(row=0, column=0, sticky="w")
    child_var = tk.StringVar()
    child_combo = AutocompleteCombobox(left_frame, textvariable=child_var, width=35)
    child_combo.grid(row=1, column=0, padx=2, pady=4)
    tk.Label(left_frame, text="Matching Children:").grid(row=2, column=0, sticky="w", pady=(6,0))

    # Treeview listing children (FullName, FatherName)
    columns = ("ChildID", "FullName", "FatherName")
    tree = ttk.Treeview(left_frame, columns=columns, show="headings", height=18)
    tree.heading("ChildID", text="ID")
    tree.heading("FullName", text="Full Name")
    tree.heading("FatherName", text="Father Name")
    tree.column("ChildID", width=40)
    tree.column("FullName", width=220)
    tree.column("FatherName", width=140)
    tree.grid(row=3, column=0, sticky="nsew", pady=4)

    # Buttons
    btn_frame = tk.Frame(left_frame)
    btn_frame.grid(row=4, column=0, pady=6, sticky="w")
    print_btn = tk.Button(btn_frame, text="Print / Save PDF", state="disabled")
    print_btn.pack(side="left", padx=4)
    refresh_btn = tk.Button(btn_frame, text="Refresh", command=lambda: load_children())
    refresh_btn.pack(side="left", padx=4)

    left_frame.grid_rowconfigure(3, weight=1)
    left_frame.grid_columnconfigure(0, weight=1)

    # --- Right side: chart area and details ---
    chart_title_var = tk.StringVar(value="Select a child to view chart")
    lbl_title = tk.Label(right_frame, textvariable=chart_title_var, font=("Arial", 12, "bold"))
    lbl_title.pack(anchor="n")

    chart_canvas_frame = tk.Frame(right_frame, bd=2, relief="sunken")
    chart_canvas_frame.pack(fill="both", expand=True, padx=6, pady=6)
    # placeholder image widget
    chart_img_label = tk.Label(chart_canvas_frame)
    chart_img_label.pack(expand=True)

    # Details list under chart
    details_frame = tk.Frame(right_frame)
    details_frame.pack(fill="x", padx=6, pady=(0,6))
    details_text = tk.Text(details_frame, height=8, wrap="word")
    details_text.pack(fill="both", expand=True)

    # --- DB queries & logic ---
    def load_children():
        """
        Load children list into tree and completion list.
        """
        try:
            conn = get_connection()
            if not conn:
                raise Exception("DB connection failed")
            cur = conn.cursor()
            cur.execute("""
                SELECT ChildID, FullName, FatherName
                FROM tblChildren
                WHERE Status = 1 OR Status IS NULL
                ORDER BY FullName
            """)
            rows = cur.fetchall()
            conn.close()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load children: {e}")
            rows = []

        # populate tree
        tree.delete(*tree.get_children())
        names = []
        for r in rows:
            cid = r.ChildID
            fn = r.FullName
            father = r.FatherName or ""
            tree.insert("", "end", iid=str(cid), values=(cid, fn, father))
            names.append(f"{fn} (ID:{cid})")

        child_combo.set_completion_list(names)

    def fetch_sponsorship_summary_for_child(child_id):
        """
        FIXED: Proper SQLite data type handling and NULL-safe casting
        """
        try:
            conn = get_connection()
            if not conn:
                raise Exception("DB connection failed")
            cur = conn.cursor()

            # FIXED: Safe casting with COALESCE and NULLIF for SQLite
            cur.execute("""
                SELECT CAST(COALESCE(NULLIF(ChildRequiredAmount, ''), '0') AS REAL) AS Req 
                FROM tblChildren 
                WHERE ChildID = ?
            """, (child_id,))
            row = cur.fetchone()
            
            if not row:
                conn.close()
                return 0, []
            
            total_required = float(row[0] or 0)
            
            # Debug: print to check value
            print(f"Child {child_id} required amount: {total_required}")

            # FIXED: Better date handling for SQLite
            today_str = datetime.today().strftime('%Y-%m-%d')
            
            # FIXED: Proper aggregation with NULL-safe casting
            cur.execute("""
                SELECT
                    s.DonorID,
                    d.FullName AS DonorName,
                    SUM(CAST(COALESCE(NULLIF(s.SponsorshipAmount, ''), '0') AS REAL)) AS SponsAmount
                FROM tblSponsorships s
                JOIN tblDonors d ON s.DonorID = d.DonorID
                WHERE s.ChildID = ?
                  AND s.IsActive = 1
                  AND (s.EndDate IS NULL OR DATE(s.EndDate) >= DATE(?))
                GROUP BY s.DonorID, d.FullName
                HAVING SponsAmount > 0
                ORDER BY SponsAmount DESC
            """, (child_id, today_str))
            
            donor_rows = cur.fetchall()
            
            # Debug: print donor count
            print(f"Found {len(donor_rows)} active sponsors")

            result = []
            for r in donor_rows:
                donor_id = r[0]  # Use index instead of attribute
                donor_name = r[1]
                amount = float(r[2] or 0.0)
                
                # Debug each donor
                print(f"  Donor: {donor_name}, Amount: {amount}")
                
                # FIXED: Safe date retrieval
                cur.execute("""
                    SELECT MAX(CommitmentEndDate) AS LastEnd
                    FROM tblCommitments
                    WHERE DonorID = ? AND IsActive = 1
                """, (donor_id,))
                le = cur.fetchone()
                last_commit_end = le[0] if le and le[0] else None
                
                # FIXED: Safe donation retrieval with proper casting
                cur.execute("""
                    SELECT DonationDate, 
                           CAST(COALESCE(NULLIF(DonationAmount, ''), '0') AS REAL) AS Amount
                    FROM tblDonations
                    WHERE DonorID = ?
                    ORDER BY DonationDate DESC
                    LIMIT 1
                """, (donor_id,))
                drow = cur.fetchone()
                
                if drow:
                    last_donation_date = drow[0]
                    last_donation_amount = float(drow[1] or 0.0)
                else:
                    last_donation_date = None
                    last_donation_amount = 0.0

                result.append({
                    "donor_id": donor_id,
                    "donor_name": donor_name,
                    "amount": amount,
                    "last_commit_end": last_commit_end,
                    "last_donation_date": last_donation_date,
                    "last_donation_amount": last_donation_amount
                })

            conn.close()
            
            # Debug final result
            print(f"Total required: {total_required}, Total donors: {len(result)}")
            
            return total_required, result

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to fetch sponsorship summary: {e}")
            import traceback
            traceback.print_exc()  # Print full error for debugging
            return 0, []

    # --- Chart drawing ---
    def draw_pie_chart(child_id, child_name, father_name):
        total_required, donors = fetch_sponsorship_summary_for_child(child_id)

        details_text.delete("1.0", tk.END)
        details_text.insert(
            tk.END,
            f"Child: {child_name}\nFather: {father_name}\n"
            f"Required Amount: Rs. {total_required:.2f}\n\n"
        )

        if total_required <= 0:
            chart_img_label.config(text="Required amount not defined or zero", image="", bg="lightyellow")
            print_btn.config(state="disabled")
            return

        # ---- Prepare pie data ----
        labels = []
        sizes = []
        colors = []

        cmap = matplotlib.colormaps.get_cmap("tab20")

        total_sponsored = 0.0

        for idx, d in enumerate(donors):
            amt = float(d["amount"] or 0)
            if amt <= 0:
                continue

            total_sponsored += amt
            labels.append(d["donor_name"])
            sizes.append(amt)
            colors.append(cmap(idx % 20))

        remaining = max(total_required - total_sponsored, 0)

        # ---- Add UNSPONSORED slice ----
        if remaining > 0:
            labels.append("Not Sponsored")
            sizes.append(remaining)
            colors.append("#D3D3D3")  # light gray

        if not sizes:
            chart_img_label.config(text="No sponsorship data found", image="", bg="lightyellow")
            print_btn.config(state="disabled")
            details_text.insert(tk.END, "No active sponsorships found for this child.\n")
            return

        # ---- Create Pie Chart ----
        fig = Figure(figsize=(5, 4), dpi=100)
        ax = fig.add_subplot(111)

        wedges, texts = ax.pie(
            sizes,
            startangle=90,
            colors=colors
        )

        ax.axis("equal")
        ax.set_title(f"Sponsorship Status — {child_name}")

        # ---- Legend with percentages ----
        legend_labels = []
        for lbl, val in zip(labels, sizes):
            pct = (val / total_required) * 100
            legend_labels.append(f"{lbl}: Rs.{val:.2f} ({pct:.1f}%)")

        ax.legend(
            wedges,
            legend_labels,
            loc="center left",
            bbox_to_anchor=(1, 0.5),
            fontsize=9
        )

        # ---- Save temp image ----
        tmpf = os.path.join(tempfile.gettempdir(), f"child_spons_{child_id}.png")
        fig.savefig(tmpf, bbox_inches="tight")
        fig.clf()

        # ---- Display in Tk ----
        im = Image.open(tmpf)
        cw = chart_canvas_frame.winfo_width() or 420
        ch = chart_canvas_frame.winfo_height() or 360
        im.thumbnail((cw, ch), Resampling.LANCZOS)

        tkimg = ImageTk.PhotoImage(im)
        chart_img_label.image = tkimg
        chart_img_label.config(image=tkimg, text="", bg="white")

        # ---- Details panel ----
        details_text.insert(tk.END, "Breakdown:\n")
        for d in donors:
            amt = float(d["amount"] or 0)
            if amt <= 0:
                continue
            pct = (amt / total_required) * 100
            details_text.insert(
                tk.END,
                f"- {d['donor_name']}: Rs.{amt:.2f} ({pct:.1f}%)\n"
            )

        if remaining > 0:
            pct = (remaining / total_required) * 100
            details_text.insert(
                tk.END,
                f"- Not Sponsored: Rs.{remaining:.2f} ({pct:.1f}%)\n"
            )

        print_btn.config(state="normal")


    # --- Print / Save PDF ---
    def save_pdf():
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Please select a child first.")
            return

        cid = int(sel[0])
        vals = tree.item(sel[0], "values")
        child_name = vals[1]
        father_name = vals[2]

        total_required, donors = fetch_sponsorship_summary_for_child(cid)
        if not donors:
            messagebox.showinfo("No Data", "No sponsorship data to print.")
            return

        # === Generate chart if not exists ===
        tmpf = os.path.join(tempfile.gettempdir(), f"child_spons_{cid}.png")
        if not os.path.exists(tmpf):
            draw_pie_chart(cid, child_name, father_name)

        downloads = os.path.join(os.path.expanduser("~"), "Downloads")
        if not os.path.exists(downloads):
            os.makedirs(downloads)
        filename = f"Sponsorship_{child_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(downloads, filename)

        try:
            c = canvas.Canvas(file_path, pagesize=A4)
            w, h = A4

            # === HEADER ===
            c.setFont("Helvetica-Bold", 16)
            c.drawCentredString(w/2, h - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawCentredString(w/2, h - 60, "Child Sponsorship Breakdown")
            c.setFont("Helvetica", 10)
            c.drawString(50, h - 85, f"Child: {child_name}")
            c.drawString(50, h - 100, f"Father: {father_name}")
            c.drawString(50, h - 115, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # === CHART (centered) ===
            chart_width = 260
            chart_height = 260
            chart_x = (w - chart_width) / 2
            chart_y = h - 400
            try:
                c.drawImage(tmpf, chart_x, chart_y, width=chart_width, height=chart_height, preserveAspectRatio=True)
            except Exception as e:
                print("PDF image embed failed:", e)

            # === DONOR BREAKDOWN BELOW CHART ===
            y = chart_y - 30
            c.setFont("Helvetica-Bold", 11)
            c.drawString(50, y, "Donor Breakdown:")
            y -= 18
            c.setFont("Helvetica", 10)
            for d in donors:
                ld = format_db_date(d['last_donation_date'], "%d/%m/%Y") if d['last_donation_date'] else "N/A"
                lcd = format_db_date(d['last_commit_end'], "%d/%m/%Y") if d['last_commit_end'] else "N/A"
                perc = (d['amount'] / total_required * 100) if total_required > 0 else 0
                line = (
                    f"{d['donor_name']}: Rs.{d['amount']:.2f} ({perc:.2f}%) | "
                    f"LastCommit: {lcd} | LastDonation: {ld} (Rs.{d['last_donation_amount']:.2f})"
                )
                c.drawString(50, y, line)
                y -= 14
                if y < 60:
                    c.showPage()
                    y = h - 60
                    c.setFont("Helvetica", 10)

            c.save()

            # Auto open
            try:
                os.startfile(file_path)
            except Exception:
                messagebox.showinfo("Saved", f"Report saved to: {file_path}")

        except Exception as e:
            messagebox.showerror("PDF Error", str(e))


    # --- Events ---
    def on_tree_select(event):
        sel = tree.selection()
        if not sel:
            return
        cid = int(sel[0])
        vals = tree.item(sel[0], "values")
        child_name = vals[1]
        father_name = vals[2]
        chart_title_var.set(f"{child_name} — Sponsorship")
        draw_pie_chart(cid, child_name, father_name)

    def on_child_combo_selected(event=None):
        txt = child_var.get().strip()
        if not txt:
            return
        cid = None
        if "(ID:" in txt and txt.endswith(")"):
            try:
                cid = int(txt.split("(ID:")[-1].rstrip(")"))
            except:
                cid = None
        if cid is not None and tree.exists(str(cid)):
            tree.selection_set(str(cid))
            tree.focus(str(cid))
            tree.see(str(cid))
            on_tree_select(None)
            return
        for item in tree.get_children():
            vals = tree.item(item, "values")
            nm = vals[1].lower()
            if txt.lower() in nm:
                tree.selection_set(item)
                tree.focus(item)
                tree.see(item)
                on_tree_select(None)
                return
        messagebox.showinfo("Not Found", "No matching child in the current list. Try Refresh.")

    # Bindings
    tree.bind("<<TreeviewSelect>>", on_tree_select)
    child_combo.bind("<<ComboboxSelected>>", on_child_combo_selected)
    child_combo.bind("<Return>", on_child_combo_selected)
    print_btn.config(command=save_pdf)

    # initial load
    load_children()
    child_combo.focus_set()


# ---------------End of Sponsorship Pie Chart ------------

# ============================================================
# Academic Tracking
# ============================================================


def open_academics_mgmt():

    # ─────────────────────────────────────────────────────────────────────────
    # WINDOW
    # ─────────────────────────────────────────────────────────────────────────
    win = tk.Toplevel()
    win.title("Academic Progress Report" + ORG_SUFFIX)
    win.state('zoomed')
    win.grab_set()
    win.focus_force()
    win.bind('<Escape>', lambda e: win.destroy())

    # ── Design tokens ─────────────────────────────────────────────────────────
    HEADER_BG   = "#0f2540"
    HEADER_FG   = "#f8fafc"
    SECTION_BG  = "#eef2f7"
    FRAME_BG    = "#ffffff"
    LABEL_FG    = "#374151"
    ACCENT      = "#1d6fd8"
    MUTED       = "#94a3b8"
    RO_BG       = "#f1f5f9"
    BTN_BLUE    = "#1d6fd8"
    BTN_GREEN   = "#16a34a"
    BTN_ORANGE  = "#d97706"
    BTN_RED     = "#dc2626"
    BTN_GRAY    = "#64748b"
    FONT_LBL    = ("Segoe UI", 9)
    FONT_ENTRY  = ("Segoe UI", 9)
    FONT_BOLD   = ("Segoe UI", 9,  "bold")
    FONT_HEADER = ("Segoe UI", 13, "bold")
    FONT_SEC    = ("Segoe UI", 10, "bold")
    FONT_BTN    = ("Segoe UI", 9,  "bold")
    PAD         = {"padx": 6, "pady": 4}

    win.configure(bg=SECTION_BG)

    # ─────────────────────────────────────────────────────────────────────────
    # STATE VARIABLES
    # ─────────────────────────────────────────────────────────────────────────
    selected_apr_id  = tk.IntVar(value=0)
    child_var        = tk.StringVar()
    gender_var       = tk.StringVar()
    father_var       = tk.StringVar()
    guardian_var     = tk.StringVar()
    center_var       = tk.StringVar()
    admin_var        = tk.StringVar()
    school_var       = tk.StringVar()
    class_var        = tk.StringVar()
    exam_var         = tk.StringVar()
    year_var         = tk.StringVar()
    total_var        = tk.StringVar()
    obtained_var     = tk.StringVar()
    perc_var         = tk.StringVar()
    div_var          = tk.StringVar()
    promoted_var     = tk.IntVar()
    position_var     = tk.StringVar()
    session_var      = tk.StringVar()
    img_path_var     = tk.StringVar()
    search_var       = tk.StringVar()
    child_map        = {}
    session_map      = {}
    current_image    = None   # keeps PhotoImage reference alive

    # ─────────────────────────────────────────────────────────────────────────
    # HEADER
    # ─────────────────────────────────────────────────────────────────────────
    header = tk.Frame(win, bg=HEADER_BG, height=56)
    header.pack(fill="x", side="top")
    header.pack_propagate(False)
    tk.Label(header, text="🎓  Academic Progress Reports",
             font=FONT_HEADER, bg=HEADER_BG, fg=HEADER_FG,
             pady=10).pack(side="left", padx=20)

    tk.Frame(win, bg=ACCENT, height=3).pack(fill="x")

    # ─────────────────────────────────────────────────────────────────────────
    # FOOTER  (packed before body)
    # ─────────────────────────────────────────────────────────────────────────
    footer = tk.Frame(win, bg="#e2e8f0", pady=8)
    footer.pack(fill="x", side="bottom")

    tk.Button(footer, text="✖  Close",
              command=win.destroy,
              font=FONT_BTN, bg=BTN_RED, fg="white",
              relief="flat", cursor="hand2",
              padx=24, pady=6).pack(side="right", padx=20)
    
    # ─────────────────────────────────────────────────────────────────────────
    # BODY — top: form area  |  bottom: treeview
    # ─────────────────────────────────────────────────────────────────────────
    body = tk.Frame(win, bg=SECTION_BG)
    body.pack(fill="both", expand=True)

    # ── Scrollable form canvas ────────────────────────────────────────────────
    form_canvas   = tk.Canvas(body, bg=SECTION_BG, highlightthickness=0, height=380)
    form_v_scroll = tk.Scrollbar(body, orient="vertical",   command=form_canvas.yview)
    form_h_scroll = tk.Scrollbar(body, orient="horizontal", command=form_canvas.xview)

    form_h_scroll.pack(side="bottom", fill="x")
    form_v_scroll.pack(side="right",  fill="y")
    form_canvas.pack(side="top", fill="x", expand=False)

    form_canvas.configure(yscrollcommand=form_v_scroll.set,
                          xscrollcommand=form_h_scroll.set)

    inner = tk.Frame(form_canvas, bg=SECTION_BG)
    inner_id = form_canvas.create_window((0, 0), window=inner, anchor="nw")

    def _on_inner_cfg(e):
        form_canvas.configure(scrollregion=form_canvas.bbox("all"))
    def _on_canvas_cfg(e):
        form_canvas.itemconfig(inner_id,
                               width=max(e.width, inner.winfo_reqwidth()))
    inner.bind("<Configure>",       _on_inner_cfg)
    form_canvas.bind("<Configure>", _on_canvas_cfg)

    def _mwheel(e):
        try:
            form_canvas.yview_scroll(int(-1*(e.delta/120)), "units")
        except tk.TclError:
            pass
    form_canvas.bind_all("<MouseWheel>", _mwheel)

    inner.columnconfigure((0, 1, 2), weight=1)

    # ── Section helper ────────────────────────────────────────────────────────
    def make_section(parent, title, row, col=0, colspan=1, rowspan=1):
        lf = tk.LabelFrame(parent, text=f"  {title}  ",
                           font=FONT_SEC, fg=ACCENT, bg=FRAME_BG,
                           bd=1, relief="groove", labelanchor="nw")
        lf.grid(row=row, column=col, columnspan=colspan, rowspan=rowspan,
                sticky="nsew", padx=8, pady=6)
        lf.columnconfigure((1, 3), weight=1)
        return lf

    def lbl_ro(parent, text, row, col, var, width=22):
        """Label + readonly entry."""
        tk.Label(parent, text=text, font=FONT_LBL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        e = tk.Entry(parent, textvariable=var, font=FONT_ENTRY,
                     width=width, state="readonly",
                     readonlybackground=RO_BG, relief="solid", bd=1)
        e.grid(row=row, column=col+1, sticky="ew", **PAD)
        return e

    def lbl_entry(parent, text, row, col, var, width=22):
        """Label + editable entry."""
        tk.Label(parent, text=text, font=FONT_LBL, fg=LABEL_FG,
                 bg=FRAME_BG, anchor="e").grid(
            row=row, column=col, sticky="e", **PAD)
        e = tk.Entry(parent, textvariable=var, font=FONT_ENTRY,
                     width=width, relief="solid", bd=1)
        e.grid(row=row, column=col+1, sticky="ew", **PAD)
        return e

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 0 — CHILD SELECTION  (full width, row 0)
    # ─────────────────────────────────────────────────────────────────────────
    sel_lf = make_section(inner, "Child Selection", row=0, col=0, colspan=3)
    sel_lf.columnconfigure((1, 3, 5), weight=1)

    tk.Label(sel_lf, text="Child Name:", font=FONT_LBL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=0, column=0, sticky="e", **PAD)
    child_combo = AutocompleteCombobox(sel_lf, textvariable=child_var,
                                       font=FONT_ENTRY, width=28)
    child_combo.grid(row=0, column=1, sticky="ew", **PAD)

    lbl_ro(sel_lf, "Gender:",    0, 2, gender_var,  width=14)
    lbl_ro(sel_lf, "Father:",    0, 4, father_var,  width=22)
    lbl_ro(sel_lf, "Guardian:",  1, 0, guardian_var, width=22)
    lbl_ro(sel_lf, "Center:",    1, 2, center_var,  width=22)
    lbl_ro(sel_lf, "Administrator:", 1, 4, admin_var, width=22)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 1 — ACADEMIC DETAILS  (col 0-1, row 1)
    # ─────────────────────────────────────────────────────────────────────────
    acad_lf = make_section(inner, "Academic Details", row=1, col=0, colspan=2)
    acad_lf.columnconfigure((1, 3), weight=1)

    lbl_entry(acad_lf, "School:",        0, 0, school_var)
    lbl_entry(acad_lf, "Class:",         0, 2, class_var,  width=14)
    lbl_entry(acad_lf, "Examination:",   1, 0, exam_var)
    lbl_entry(acad_lf, "Academic Year:", 1, 2, year_var,   width=14)
    lbl_entry(acad_lf, "Total Marks:",   2, 0, total_var,  width=12)
    lbl_entry(acad_lf, "Obtained Marks:",2, 2, obtained_var, width=12)

    tk.Label(acad_lf, text="Percentage:", font=FONT_LBL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=3, column=0, sticky="e", **PAD)
    perc_e = tk.Entry(acad_lf, textvariable=perc_var, font=FONT_ENTRY,
                      width=12, state="readonly",
                      readonlybackground=RO_BG, relief="solid", bd=1)
    perc_e.grid(row=3, column=1, sticky="ew", **PAD)

    lbl_entry(acad_lf, "Division:", 3, 2, div_var, width=14)

    tk.Label(acad_lf, text="Position:", font=FONT_LBL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=4, column=0, sticky="e", **PAD)
    tk.Entry(acad_lf, textvariable=position_var, font=FONT_ENTRY,
             width=12, relief="solid", bd=1).grid(
        row=4, column=1, sticky="ew", **PAD)

    tk.Label(acad_lf, text="Promoted:", font=FONT_LBL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=4, column=2, sticky="e", **PAD)
    tk.Checkbutton(acad_lf, text="Yes", variable=promoted_var,
                   font=FONT_LBL, bg=FRAME_BG,
                   activebackground=FRAME_BG).grid(
        row=4, column=3, sticky="w", **PAD)

    tk.Label(acad_lf, text="Session:", font=FONT_LBL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=5, column=0, sticky="e", **PAD)
    session_combo = ttk.Combobox(acad_lf, textvariable=session_var,
                                 font=FONT_ENTRY, width=20, state="readonly")
    session_combo.grid(row=5, column=1, sticky="ew", **PAD)

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 2 — RESULT IMAGE  (col 2, row 1)
    # ─────────────────────────────────────────────────────────────────────────
    img_lf = make_section(inner, "Result Image", row=1, col=2)
    img_lf.columnconfigure(0, weight=1)

    # Fixed-size image container
    img_container = tk.Frame(img_lf, width=200, height=200,
                             bg="#d1d5db", relief="solid", bd=1)
    img_container.grid(row=0, column=0, columnspan=2,
                       padx=10, pady=8, sticky="n")
    img_container.pack_propagate(False)

    img_label = tk.Label(img_container, text="No Image",
                         bg="#d1d5db", font=FONT_LBL)
    img_label.pack(expand=True)

    tk.Label(img_lf, text="Image Path:", font=FONT_LBL,
             fg=LABEL_FG, bg=FRAME_BG, anchor="e").grid(
        row=1, column=0, sticky="e", **PAD)
    tk.Entry(img_lf, textvariable=img_path_var, font=FONT_ENTRY,
             width=20, state="readonly",
             readonlybackground=RO_BG, relief="solid", bd=1).grid(
        row=1, column=1, sticky="ew", **PAD)

    tk.Button(img_lf, text="📁  Browse",
              command=lambda: browse_image(),
              font=FONT_LBL, bg=BTN_BLUE, fg="white",
              relief="flat", cursor="hand2",
              padx=8, pady=4).grid(
        row=2, column=0, columnspan=2, pady=(4, 8))

    # ─────────────────────────────────────────────────────────────────────────
    # SECTION 3 — ACTION BUTTONS  (full width, row 2)
    # ─────────────────────────────────────────────────────────────────────────
    btn_lf = make_section(inner, "Actions", row=2, col=0, colspan=3)

    save_btn = tk.Button(
        btn_lf, text="💾  Save Report",
        command=lambda: save_report(),
        font=FONT_BTN, bg=BTN_BLUE, fg="white",
        relief="flat", cursor="hand2", padx=16, pady=6)
    save_btn.grid(row=0, column=0, padx=8, pady=6)

    update_btn = tk.Button(
        btn_lf, text="✎  Update Report",
        command=lambda: update_report(),
        font=FONT_BTN, bg=BTN_ORANGE, fg="white",
        relief="flat", cursor="hand2", padx=16, pady=6)
    update_btn.grid(row=0, column=1, padx=8, pady=6)
    update_btn.grid_remove()   # hidden until a row is selected

    tk.Button(btn_lf, text="🖨  Print Selected",
              command=lambda: print_selected_report(),
              font=FONT_BTN, bg=BTN_GREEN, fg="white",
              relief="flat", cursor="hand2", padx=16, pady=6).grid(
        row=0, column=2, padx=8, pady=6)

    tk.Button(btn_lf, text="✖  Clear Form",
              command=lambda: clear_form(),
              font=FONT_BTN, bg=BTN_GRAY, fg="white",
              relief="flat", cursor="hand2", padx=16, pady=6).grid(
        row=0, column=3, padx=8, pady=6)

    # ─────────────────────────────────────────────────────────────────────────
    # SEARCH BAR + TREEVIEW
    # ─────────────────────────────────────────────────────────────────────────
    search_frame = tk.Frame(body, bg=SECTION_BG)
    search_frame.pack(fill="x", padx=10, pady=(6, 2))

    tk.Label(search_frame, text="🔍  Search Child / Year:",
             font=FONT_BOLD, bg=SECTION_BG, fg=LABEL_FG).pack(
        side="left", padx=(6, 4))
    search_box = AutocompleteCombobox(search_frame,
                                      textvariable=search_var,
                                      font=FONT_ENTRY, width=32)
    search_box.pack(side="left", padx=4)
    tk.Button(search_frame, text="Search",
              command=lambda: on_search(),
              font=FONT_LBL, bg=BTN_BLUE, fg="white",
              relief="flat", cursor="hand2",
              padx=10, pady=3).pack(side="left", padx=4)
    tk.Button(search_frame, text="Show All",
              command=lambda: (search_var.set(""), load_reports()),
              font=FONT_LBL, bg=BTN_GRAY, fg="white",
              relief="flat", cursor="hand2",
              padx=10, pady=3).pack(side="left", padx=4)

    # Treeview section
    tree_lf = tk.LabelFrame(body, text="  Academic Records  ",
                             font=FONT_SEC, fg=ACCENT, bg=FRAME_BG,
                             bd=1, relief="groove", labelanchor="nw")
    tree_lf.pack(fill="both", expand=True, padx=8, pady=(2, 4))

    tv_v = tk.Scrollbar(tree_lf, orient="vertical")
    tv_h = tk.Scrollbar(tree_lf, orient="horizontal")

    cols = ("ID", "Child", "Center", "Class", "Year", "Percentage", "Division")
    tree = ttk.Treeview(
        tree_lf, columns=cols, show="headings",
        yscrollcommand=tv_v.set,
        xscrollcommand=tv_h.set
    )
    tv_v.config(command=tree.yview)
    tv_h.config(command=tree.xview)

    col_w = {"ID": 55, "Child": 200, "Center": 180,
             "Class": 80, "Year": 100, "Percentage": 90, "Division": 90}
    for col in cols:
        tree.heading(col, text=col)
        tree.column(col, width=col_w.get(col, 120), minwidth=60)

    tree.tag_configure("odd",  background="#f8fafc")
    tree.tag_configure("even", background="#ffffff")

    tv_h.pack(side="bottom", fill="x")
    tv_v.pack(side="right",  fill="y")
    tree.pack(side="left", fill="both", expand=True)

    tree.bind("<<TreeviewSelect>>", lambda e: on_tree_select())

    # ─────────────────────────────────────────────────────────────────────────
    # LOGIC FUNCTIONS
    # ─────────────────────────────────────────────────────────────────────────

    def display_image(path):
        nonlocal current_image
        if not path or not os.path.exists(path):
            img_label.config(image="", text="No Image",
                             bg="#d1d5db")
            img_container.config(bg="#d1d5db")
            return
        try:
            img = Image.open(path)
            img.thumbnail((200, 200))
            current_image = ImageTk.PhotoImage(img)
            img_label.config(image=current_image, text="", bg="white")
            img_container.config(bg="white")
        except Exception as e:
            img_label.config(image="", text="Load Error", bg="#fca5a5")
            print("Image load error:", e)

    def browse_image():
        path = filedialog.askopenfilename(
            title="Select Result Image",
            filetypes=[("Image Files", "*.jpg *.jpeg *.png *.gif")])
        if path:
            img_path_var.set(path)
            display_image(path)

    def calc_percentage(*_):
        try:
            total    = float(total_var.get())
            obtained = float(obtained_var.get())
            perc     = (obtained / total * 100) if total > 0 else 0
            perc_var.set(f"{perc:.2f}")
        except Exception:
            perc_var.set("")

    total_var.trace("w",   calc_percentage)
    obtained_var.trace("w", calc_percentage)

    def load_children():
        conn   = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT c.ChildID, c.FullName, c.Gender, c.FatherName,
                   c.GuardianName, c.SchoolName, c.Class,
                   ce.CenterName, ce.AdministratorName, c.SessionID
            FROM tblChildren c
            LEFT JOIN tblCenters ce ON c.CenterID = ce.CenterID
        """)
        rows = cursor.fetchall()
        conn.close()
        child_map.clear()
        for row in rows:
            child_map[row.FullName] = {
                "id"      : row.ChildID,
                "gender"  : row.Gender,
                "father"  : row.FatherName,
                "guardian": row.GuardianName,
                "school"  : row.SchoolName,
                "class"   : row.Class,
                "center"  : row.CenterName,
                "admin"   : row.AdministratorName,
                "session" : row.SessionID,
            }
        child_combo.set_completion_list(list(child_map.keys()))
        search_box.set_completion_list(list(child_map.keys()))

    def load_sessions():
        conn   = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT SessionID, SessionName FROM tblSessions")
        rows = cursor.fetchall()
        conn.close()
        session_map.clear()
        for r in rows:
            session_map[r.SessionName] = r.SessionID
        session_combo["values"] = list(session_map.keys())

    def on_child_select(event=None):
        name = child_var.get()
        if name not in child_map:
            return
        info = child_map[name]
        gender_var.set(info["gender"]   or "")
        father_var.set(info["father"]   or "")
        guardian_var.set(info["guardian"] or "")
        center_var.set(info["center"]   or "")
        admin_var.set(info["admin"]     or "")
        school_var.set(info["school"]   or "")
        class_var.set(info["class"]     or "")
        for sess_name, sid in session_map.items():
            if sid == info["session"]:
                session_var.set(sess_name)
                break

    child_combo.bind("<<ComboboxSelected>>", on_child_select)

    def load_reports(filter_text=""):
        tree.delete(*tree.get_children())
        conn   = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT r.ReportID, c.FullName, ce.CenterName,
                   r.Class, r.AcademicYear,
                   CASE WHEN CAST(r.TotalMarks AS REAL) > 0
                        THEN ROUND(CAST(r.ObtainedMarks AS REAL) /
                             CAST(r.TotalMarks AS REAL) * 100, 2)
                        ELSE 0 END AS Percentage,
                   r.Division
            FROM tblAcademicProgressReports r
            JOIN  tblChildren c  ON r.ChildID  = c.ChildID
            LEFT JOIN tblCenters ce ON c.CenterID = ce.CenterID
            WHERE c.FullName LIKE ? OR r.AcademicYear LIKE ?
            ORDER BY r.ReportID DESC
        """, (f"%{filter_text}%", f"%{filter_text}%"))
        rows = cursor.fetchall()
        conn.close()
        for idx, row in enumerate(rows):
            tag = "odd" if idx % 2 else "even"
            tree.insert("", "end", tags=(tag,),
                        values=(
                            str(row.ReportID),
                            str(row.FullName),
                            str(row.CenterName or ""),
                            str(row.Class      or ""),
                            str(row.AcademicYear or ""),
                            str(row.Percentage or ""),
                            str(row.Division   or ""),
                        ))

    def on_search(event=None):
        load_reports(search_var.get().strip())

    def on_search_highlight(event=None):
        """Select the first tree row matching the search text."""
        txt = search_var.get().strip().lower()
        if not txt:
            return
        CHILD_COL = 1
        for item in tree.get_children():
            vals = tree.item(item, "values")
            try:
                if txt in str(vals[CHILD_COL]).lower():
                    tree.selection_set(item)
                    tree.focus(item)
                    tree.see(item)
                    on_tree_select()
                    break
            except Exception:
                continue

    search_box.bind("<<ComboboxSelected>>", on_search_highlight)
    search_box.bind("<Return>",             on_search_highlight)

    def on_tree_select():
        selected = tree.selection()
        if not selected:
            return
        item = selected[0]
        apr_id = tree.item(item, "values")[0]
        selected_apr_id.set(apr_id)
        try:
            conn   = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT ReportID, ChildID, ChildName, Gender, FatherOrGuardianName,
                       AghoshCenterName, AdministratorName, School, Class,
                       Examination, AcademicYear, TotalMarks, ObtainedMarks,
                       CASE WHEN CAST(TotalMarks AS REAL) > 0
                            THEN ROUND(CAST(ObtainedMarks AS REAL) /
                                 CAST(TotalMarks AS REAL) * 100, 2)
                            ELSE 0 END AS Percentage,
                       Division, Promoted, Position,
                       ResultImagePath, SessionID
                FROM tblAcademicProgressReports
                WHERE ReportID = ?
            """, (apr_id,))
            row = cursor.fetchone()
            conn.close()
            if not row:
                return
            child_var.set(row.ChildName    or "")
            gender_var.set(row.Gender      or "")
            father_var.set(row.FatherOrGuardianName or "")
            center_var.set(row.AghoshCenterName     or "")
            admin_var.set(row.AdministratorName     or "")
            school_var.set(row.School      or "")
            class_var.set(row.Class        or "")
            exam_var.set(row.Examination   or "")
            year_var.set(row.AcademicYear  or "")
            total_var.set(row.TotalMarks   or "")
            obtained_var.set(row.ObtainedMarks or "")
            perc_var.set(row.Percentage    or "")
            div_var.set(row.Division       or "")
            promoted_var.set(int(row.Promoted) if row.Promoted is not None else 0)
            position_var.set(row.Position  or "")
            img_path_var.set(row.ResultImagePath or "")
            display_image(row.ResultImagePath)
            # Match session name
            for sname, sid in session_map.items():
                if sid == row.SessionID:
                    session_var.set(sname)
                    break
            # Show Update, hide Save
            save_btn.grid_remove()
            update_btn.grid()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load report:\n{e}")

    def clear_form():
        selected_apr_id.set(0)
        for v in (child_var, gender_var, father_var, guardian_var,
                  center_var, admin_var, school_var, class_var,
                  exam_var, year_var, total_var, obtained_var,
                  perc_var, div_var, position_var, session_var,
                  img_path_var):
            v.set("")
        promoted_var.set(0)
        img_label.config(image="", text="No Image", bg="#d1d5db")
        img_container.config(bg="#d1d5db")
        save_btn.grid()
        update_btn.grid_remove()
        child_combo.focus_set()

    def save_report():
        if not child_var.get():
            messagebox.showwarning("Validation", "Please select a child.")
            return
        if child_var.get() not in child_map:
            messagebox.showwarning("Validation", "Invalid child selected.")
            return
        try:
            conn   = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO tblAcademicProgressReports
                    (ChildID, ChildName, Gender, FatherOrGuardianName,
                     AghoshCenterName, AdministratorName, School, Class,
                     Examination, AcademicYear, TotalMarks, ObtainedMarks,
                     Division, Promoted, Position, ResultImagePath, SessionID)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                child_map[child_var.get()]["id"],
                child_var.get(),
                gender_var.get(),
                father_var.get() or guardian_var.get(),
                center_var.get(),
                admin_var.get(),
                school_var.get(),
                class_var.get(),
                exam_var.get(),
                year_var.get(),
                total_var.get(),
                obtained_var.get(),
                div_var.get(),
                promoted_var.get(),
                position_var.get(),
                img_path_var.get(),
                session_map.get(session_var.get(), None),
            ))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Academic report saved successfully.")
            load_reports()
            clear_form()
        except Exception as e:
            messagebox.showerror("DB Error", str(e))

    def update_report():
        if selected_apr_id.get() == 0:
            messagebox.showwarning("No Selection",
                                   "Please select a report to update.")
            return
        if child_var.get() not in child_map:
            messagebox.showwarning("Validation", "Invalid child selected.")
            return
        try:
            conn   = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE tblAcademicProgressReports
                SET ChildID=?, ChildName=?, Gender=?, FatherOrGuardianName=?,
                    AghoshCenterName=?, AdministratorName=?, School=?, Class=?,
                    Examination=?, AcademicYear=?, TotalMarks=?, ObtainedMarks=?,
                    Division=?, Promoted=?, Position=?, ResultImagePath=?, SessionID=?
                WHERE ReportID=?
            """, (
                child_map[child_var.get()]["id"],
                child_var.get(),
                gender_var.get(),
                father_var.get() or guardian_var.get(),
                center_var.get(),
                admin_var.get(),
                school_var.get(),
                class_var.get(),
                exam_var.get(),
                year_var.get(),
                total_var.get(),
                obtained_var.get(),
                div_var.get(),
                promoted_var.get(),
                position_var.get(),
                img_path_var.get(),
                session_map.get(session_var.get(), None),
                selected_apr_id.get(),
            ))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Academic report updated successfully.")
            load_reports()
            clear_form()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to update report:\n{e}")

    def print_selected_report():
        item = tree.focus()
        if not item:
            messagebox.showwarning("No Selection",
                                   "Please select a record to print.")
            return
        apr_id = tree.item(item, "values")[0]
        try:
            conn   = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT ChildName, Gender, FatherOrGuardianName,
                       AghoshCenterName, AdministratorName, School, Class,
                       Examination, AcademicYear, TotalMarks, ObtainedMarks,
                       CASE WHEN CAST(TotalMarks AS REAL) > 0
                            THEN ROUND(CAST(ObtainedMarks AS REAL) /
                                 CAST(TotalMarks AS REAL) * 100, 2)
                            ELSE 0 END AS Percentage,
                       Division, Promoted, Position, ResultImagePath
                FROM tblAcademicProgressReports WHERE ReportID=?
            """, (apr_id,))
            data = cursor.fetchone()
            conn.close()
            if not data:
                messagebox.showerror("Error", "Could not find the selected report.")
                return

            (child_name, gender, father, center, admin, school, class_name,
             exam, year, total, obtained, perc, division,
             promoted, position, img_path) = data

            downloads = os.path.join(os.path.expanduser("~"), "Downloads")
            os.makedirs(downloads, exist_ok=True)
            file_name = f"AcademicReport_{child_name}_{apr_id}.pdf".replace(" ", "_")
            file_path = os.path.join(downloads, file_name)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Header
            y = height - 30
            c.setFont("Helvetica-Bold", 16)
            txt = "Al-Khidmat Khawateen Trust Pakistan"
            c.drawCentredString(width / 2, y, txt)
            y -= 20
            c.setFont("Helvetica-Bold", 13)
            c.drawCentredString(width / 2, y, str(center or ""))
            y -= 20
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width / 2, y, "Academic Progress Report")
            y -= 6
            c.setLineWidth(1)
            c.line(40, y, width - 40, y)
            y -= 20

            c.setFont("Helvetica", 11)

            def draw_line(label, value):
                nonlocal y
                if y < 80:
                    c.showPage()
                    y = height - 50
                    c.setFont("Helvetica", 11)
                c.setFont("Helvetica-Bold", 11)
                c.drawString(50, y, f"{label}:")
                c.setFont("Helvetica", 11)
                c.drawString(210, y, str(value or ""))
                y -= 18

            draw_line("Child Name",       child_name)
            draw_line("Gender",           gender)
            draw_line("Father / Guardian",father)
            draw_line("Aghosh Center",    center)
            draw_line("Administrator",    admin)
            draw_line("School",           school)
            draw_line("Class",            class_name)
            draw_line("Examination",      exam)
            draw_line("Academic Year",    year)
            draw_line("Total Marks",      total)
            draw_line("Obtained Marks",   obtained)
            draw_line("Percentage",       f"{perc}%" if perc else "")
            draw_line("Division",         division)
            draw_line("Promoted",         "Yes" if promoted else "No")
            draw_line("Position",         position)

            y -= 8
            c.setLineWidth(0.5)
            c.line(40, y, width - 40, y)
            y -= 12

            # Result image
            if img_path and os.path.exists(img_path):
                try:
                    from reportlab.lib.utils import ImageReader as _IR
                    ir   = _IR(img_path)
                    iw, ih = ir.getSize()
                    max_w, max_h = 400, 380
                    ratio = min(max_w / iw, max_h / ih)
                    dw, dh = iw * ratio, ih * ratio
                    x_off = (width - dw) / 2
                    if y - dh < 30:
                        c.showPage()
                        y = height - 50
                    c.drawImage(img_path, x_off, y - dh,
                                width=dw, height=dh,
                                preserveAspectRatio=True)
                    y -= dh + 10
                except Exception as img_e:
                    messagebox.showwarning("Image Error",
                                           f"Could not embed image:\n{img_e}")

            # Footer
            c.setFont("Helvetica-Oblique", 8)
            from datetime import datetime as _dt
            c.drawString(40, 20,
                         f"Generated: {_dt.now().strftime('%d/%m/%Y %H:%M')}")

            c.save()
            messagebox.showinfo("Success", f"PDF saved to:\n{file_path}")
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # ─────────────────────────────────────────────────────────────────────────
    # KEYBOARD SHORTCUTS
    # ─────────────────────────────────────────────────────────────────────────
    win.bind('<Return>', lambda e: (
        save_btn.invoke()   if save_btn.winfo_ismapped()   else
        update_btn.invoke() if update_btn.winfo_ismapped() else None
    ))

    # ─────────────────────────────────────────────────────────────────────────
    # INIT
    # ─────────────────────────────────────────────────────────────────────────
    load_sessions()
    load_children()
    load_reports()



    
    
# ============================================================

# ----------- Start of Helth Managemnt -----------
   
def open_health_mgmt():
    import tempfile
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    import os
    from datetime import datetime

    win = tk.Toplevel()
    win.title("Child Health Management" + ORG_SUFFIX)
    ww, wh = 1500, 720
    sw = win.winfo_screenwidth(); sh = win.winfo_screenheight()
    x = (sw // 2) - (ww // 2); y = (sh // 2) - (wh // 2)
    win.geometry(f"{ww}x{wh}+{x}+{y}")
    win.transient()
    win.grab_set()
    win.focus_force()
    win.bind('<Escape>', lambda e: win.destroy())

    # Pressing Enter -> Save (only when not editing using multi-line fields)
    # We'll bind it to the window so Enter will call save_record
    # but avoid interfering when focus is on Text widgets
    def on_enter(event=None):
        # if focus in a Text widget, ignore (allow newline)
        widget = win.focus_get()
        if isinstance(widget, tk.Text):
            return
        save_record()
    win.bind('<Return>', on_enter)

    # ---------------- variables ----------------
    child_var = tk.StringVar()
    regno_var = tk.StringVar()
    dob_var = tk.StringVar()
    gender_var = tk.StringVar()
    country_var = tk.StringVar()

    bmi_var = tk.StringVar()
    height_var = tk.StringVar()
    weight_var = tk.StringVar()
    diagnosis_var = tk.StringVar()

    vision_r_var = tk.StringVar()
    vision_l_var = tk.StringVar()
    vision_diag_var = tk.StringVar()

    hearing_r_var = tk.StringVar()
    hearing_l_var = tk.StringVar()
    hearing_diag_var = tk.StringVar()

    speech_var = tk.StringVar()
    blood_group_var = tk.StringVar()
    ecg_var = tk.StringVar()
    prescribed_var = tk.StringVar()
    special_treatment_var = tk.StringVar()
    medofficer_var = tk.StringVar()
    med_institution_var = tk.StringVar()

    # internal state
    child_map = {}   # name -> dict with id/reg/dob/gender/country
    selected_health_id = None

    # ---------- UI layout ----------
    main_frame = tk.Frame(win, padx=8, pady=8)
    main_frame.pack(fill='both', expand=True)

    left_frame = tk.Frame(main_frame)
    left_frame.pack(side='left', fill='both', expand=True, padx=(0,8))

    right_frame = tk.Frame(main_frame, width=360)
    right_frame.pack(side='right', fill='y')

    # ---- Form labelframe ----
    form = tk.LabelFrame(left_frame, text="Health Examination", padx=10, pady=10)
    form.pack(fill='x', padx=4, pady=4)

    # Row 0: Child autocomplete
    tk.Label(form, text="Child:").grid(row=0, column=0, sticky='e', padx=4, pady=4)
    child_combo = AutocompleteCombobox(form, textvariable=child_var, width=40)
    child_combo.grid(row=0, column=1, columnspan=3, sticky='w', padx=4, pady=4)

    # Row 1: Registration / DOB / Gender / Country
    tk.Label(form, text="Reg. No:").grid(row=1, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=regno_var, width=20, state='readonly').grid(row=1, column=1, sticky='w', padx=4, pady=4)

    tk.Label(form, text="DOB:").grid(row=1, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=dob_var, width=18, state='readonly').grid(row=1, column=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Gender:").grid(row=2, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=gender_var, width=20, state='readonly').grid(row=2, column=1, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Country:").grid(row=2, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=country_var, width=18).grid(row=2, column=3, sticky='w', padx=4, pady=4)

    # Separator
    ttk.Separator(form, orient='horizontal').grid(row=3, column=0, columnspan=4, sticky="ew", pady=6)

    # Physical exam
    tk.Label(form, text="BMI:").grid(row=4, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=bmi_var, width=15).grid(row=4, column=1, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Height:").grid(row=4, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=height_var, width=12).grid(row=4, column=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Weight:").grid(row=5, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=weight_var, width=15).grid(row=5, column=1, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Diagnosis:").grid(row=5, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=diagnosis_var, width=30).grid(row=5, column=3, sticky='w', padx=4, pady=4)

    # Vision
    tk.Label(form, text="Vision R:").grid(row=6, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=vision_r_var, width=12).grid(row=6, column=1, sticky='w', padx=4, pady=4)
    tk.Label(form, text="Vision L:").grid(row=6, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=vision_l_var, width=12).grid(row=6, column=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Vision Diagnosis:").grid(row=7, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=vision_diag_var, width=50).grid(row=7, column=1, columnspan=3, sticky='w', padx=4, pady=4)

    # Hearing
    tk.Label(form, text="Hearing R:").grid(row=8, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=hearing_r_var, width=12).grid(row=8, column=1, sticky='w', padx=4, pady=4)
    tk.Label(form, text="Hearing L:").grid(row=8, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=hearing_l_var, width=12).grid(row=8, column=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Hearing Diagnosis:").grid(row=9, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=hearing_diag_var, width=50).grid(row=9, column=1, columnspan=3, sticky='w', padx=4, pady=4)

    # Speech
    tk.Label(form, text="Speech:").grid(row=10, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=speech_var, width=50).grid(row=10, column=1, columnspan=3, sticky='w', padx=4, pady=4)

    # Investigations
    tk.Label(form, text="Blood Group:").grid(row=11, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=blood_group_var, width=12).grid(row=11, column=1, sticky='w', padx=4, pady=4)
    tk.Label(form, text="ECG:").grid(row=11, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=ecg_var, width=12).grid(row=11, column=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Prescribed Medicine:").grid(row=12, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=prescribed_var, width=50).grid(row=12, column=1, columnspan=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Special Treatment:").grid(row=13, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=special_treatment_var, width=50).grid(row=13, column=1, columnspan=3, sticky='w', padx=4, pady=4)

    tk.Label(form, text="Medical Officer:").grid(row=14, column=0, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=medofficer_var, width=30).grid(row=14, column=1, sticky='w', padx=4, pady=4)
    tk.Label(form, text="Medical Institution:").grid(row=14, column=2, sticky='e', padx=4, pady=4)
    tk.Entry(form, textvariable=med_institution_var, width=30).grid(row=14, column=3, sticky='w', padx=4, pady=4)

    # Buttons: Save / Update / Clear / Print
    button_frame = tk.Frame(form)
    button_frame.grid(row=15, column=0, columnspan=4, pady=(10,0))

    save_btn = tk.Button(button_frame, text="Save", bg="blue", fg="white", width=14)
    save_btn.grid(row=0, column=0, padx=6)
    update_btn = tk.Button(button_frame, text="Update", bg="orange", fg="white", width=14)
    update_btn.grid(row=0, column=1, padx=6)
    clear_btn = tk.Button(button_frame, text="Clear", bg="gray", fg="white", width=14)
    clear_btn.grid(row=0, column=2, padx=6)
    print_btn = tk.Button(button_frame, text="Print", bg="green", fg="white", width=14)
    print_btn.grid(row=0, column=3, padx=6)

    # Initially Update is hidden
    update_btn.grid_remove()

    # ----------- Right frame: Treeview of health records ----------
    tv_frame = tk.LabelFrame(right_frame, text="Health Records", padx=6, pady=6)
    tv_frame.pack(fill='both', expand=True, padx=4, pady=4)

    cols = ("ID", "RegNo", "Child", "ExamDate", "Diagnosis")
    tree = ttk.Treeview(tv_frame, columns=cols, show='headings', height=20)
    for c in cols:
        tree.heading(c, text=c)
        tree.column(c, width=120 if c != "Diagnosis" else 200)
    tree.pack(fill='both', expand=True)

    # ---------------- DB operations ----------------
    def load_children():
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT ChildID, FullName, RegistrationNumber, DateOfBirth, Gender
                FROM tblChildren
                WHERE Status = 1
                ORDER BY FullName
            """)
            rows = cursor.fetchall()
            conn.close()
            for r in rows:
                child_map[r.FullName] = {
                    "id": r.ChildID,
                    "reg": getattr(r, "RegistrationNumber", "") or "",
                    "dob": getattr(r, "DateOfBirth", None),
                    "gender": getattr(r, "Gender", "") or "",
                    "country": getattr(r, "Country", "") or ""
                }
            names = list(child_map.keys())
            child_combo.set_completion_list(names)
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load children:\n{e}")

    def on_child_selected(event=None):
        name = child_var.get()
        if not name or name not in child_map:
            return
        info = child_map[name]
        regno_var.set(info["reg"])
        dob_val = info["dob"]
        if dob_val:
            # some DB rows may be datetime objects
            try:
                dob_var.set(format_db_date(dob_val, "%d/%m/%Y"))
            except:
                dob_var.set(str(dob_val))
        else:
            dob_var.set("")
        gender_var.set(info["gender"] or "")
        country_var.set(info.get("country", "") or "")

    child_combo.bind("<<ComboboxSelected>>", on_child_selected)

    # ---------- CRUD for Health records ----------
    def clear_form():
        nonlocal selected_health_id
        selected_health_id = None
        regno_var.set("")
        dob_var.set("")
        child_var.set("")
        gender_var.set("")
        country_var.set("")

        bmi_var.set(""); height_var.set(""); weight_var.set(""); diagnosis_var.set("")
        vision_r_var.set(""); vision_l_var.set(""); vision_diag_var.set("")
        hearing_r_var.set(""); hearing_l_var.set(""); hearing_diag_var.set("")
        speech_var.set(""); blood_group_var.set(""); ecg_var.set(""); prescribed_var.set("")
        special_treatment_var.set(""); medofficer_var.set(""); med_institution_var.set("")
        update_btn.grid_remove()
        save_btn.grid()
        # reset focus
        child_combo.focus_set()

    def load_records():
        tree.delete(*tree.get_children())
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT h.HealthID, c.RegistrationNumber, c.FullName, h.CreatedOn, h.PhysicalDiagnosis
                FROM tblHealth h
                JOIN tblChildren c ON h.ChildID = c.ChildID
                ORDER BY h.CreatedOn DESC
            """)

            rows = cursor.fetchall()
            conn.close()
            for r in rows:
                # CreatedDate may be datetime
                created = getattr(r, "CreatedOn", None)
                tree.insert("", "end", iid=r.HealthID, values=(
                    r.HealthID, r.RegistrationNumber, r.FullName,
                    format_db_date(created, "%d/%m/%Y") if created else "",
                    getattr(r, "PhysicalDiagnosis", "") or ""
                ))

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load health records:\n{e}")

    def on_tree_select(event=None):
        nonlocal selected_health_id
        sel = tree.selection()
        if not sel:
            return
        hid = int(sel[0])
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT h.*, c.FullName, c.RegistrationNumber, c.DateOfBirth, c.Gender
                FROM tblHealth h
                JOIN tblChildren c ON h.ChildID = c.ChildID
                WHERE h.HealthID = ?
            """, (hid,))
            row = cursor.fetchone()
            conn.close()
            if not row:
                return
            selected_health_id = hid
            # fill form
            child_name = getattr(row, "FullName", "")
            child_var.set(child_name)
            regno_var.set(getattr(row, "RegistrationNumber", ""))
            dob_val = getattr(row, "DateOfBirth", None)
            if dob_val:
                try:
                    dob_var.set(format_db_date(dob_val, "%d/%m/%Y"))
                except:
                    dob_var.set(str(dob_val))
            gender_var.set(getattr(row, "Gender", "") or "")

            bmi_var.set(getattr(row, "BMI", "") or "")
            height_var.set(getattr(row, "Height", "") or "")
            weight_var.set(getattr(row, "Weight", "") or "")
            #diagnosis_var.set(getattr(row, "Diagnosis", "") or "")
            diagnosis_var.set(getattr(row, "PhysicalDiagnosis", "") or "")

            vision_r_var.set(getattr(row, "VisionRight", "") or "")
            vision_l_var.set(getattr(row, "VisionLeft", "") or "")
            vision_diag_var.set(getattr(row, "VisionDiagnosis", "") or "")

            hearing_r_var.set(getattr(row, "HearingRight", "") or "")
            hearing_l_var.set(getattr(row, "HearingLeft", "") or "")
            hearing_diag_var.set(getattr(row, "HearingDiagnosis", "") or "")

            speech_var.set(getattr(row, "Speech", "") or "")
            blood_group_var.set(getattr(row, "BloodGroup", "") or "")
            ecg_var.set(getattr(row, "ECG", "") or "")
            prescribed_var.set(getattr(row, "PrescribedMedicine", "") or "")
            special_treatment_var.set(getattr(row, "SpecialTreatment", "") or "")
            #medofficer_var.set(getattr(row, "MedicalOfficer", "") or "")
            medofficer_var.set(getattr(row, "MedicalOfficerName", "") or "")
            med_institution_var.set(getattr(row, "MedicalInstitutionName", "") or "")
            #med_institution_var.set(getattr(row, "MedicalInstitution", "") or "")

            # Show update button
            update_btn.grid()
            save_btn.grid_remove()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load record:\n{e}")

    tree.bind('<<TreeviewSelect>>', on_tree_select)

    # ---------- Save / Update ----------
    def save_record(event=None):
        nonlocal selected_health_id
        if not child_var.get().strip():
            messagebox.showwarning("Validation", "Please select a child before saving.")
            return
        child_info = child_map.get(child_var.get())
        if not child_info:
            messagebox.showwarning("Validation", "Selected child is invalid.")
            return
        child_id = child_info["id"]
        try:
            conn = get_connection()
            cursor = conn.cursor()
            now = datetime.now()
            cursor.execute("""
                INSERT INTO tblHealth (
                    ChildID, CreatedOn, BMI, Height, Weight, PhysicalDiagnosis,
                    VisionRight, VisionLeft, VisionDiagnosis,
                    HearingRight, HearingLeft, HearingDiagnosis,
                    Speech, BloodGroup, ECG, PrescribedMedicine,
                    SpecialTreatment, MedicalOfficerName, MedicalInstitutionName
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                child_id,
                now,
                bmi_var.get(),
                height_var.get(),
                weight_var.get(),
                diagnosis_var.get(),
                vision_r_var.get(),
                vision_l_var.get(),
                vision_diag_var.get(),
                hearing_r_var.get(),
                hearing_l_var.get(),
                hearing_diag_var.get(),
                speech_var.get(),
                blood_group_var.get(),
                ecg_var.get(),
                prescribed_var.get(),
                special_treatment_var.get(),
                medofficer_var.get(),
                med_institution_var.get()
            ))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Health record saved.")
            clear_form()
            load_records()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to save record:\n{e}")

    def update_record():
        nonlocal selected_health_id
        if not selected_health_id:
            messagebox.showwarning("Select", "Please select a record to update.")
            return
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE tblHealth SET
                    BMI=?, Height=?, Weight=?, PhysicalDiagnosis=?,
                    VisionRight=?, VisionLeft=?, VisionDiagnosis=?,
                    HearingRight=?, HearingLeft=?, HearingDiagnosis=?,
                    Speech=?, BloodGroup=?, ECG=?, PrescribedMedicine=?,
                    SpecialTreatment=?, MedicalOfficerName=?, MedicalInstitutionName=?, CreatedOn=?
                    WHERE HealthID = ?

            """, (
                bmi_var.get(),
                height_var.get(),
                weight_var.get(),
                diagnosis_var.get(),
                vision_r_var.get(),
                vision_l_var.get(),
                vision_diag_var.get(),
                hearing_r_var.get(),
                hearing_l_var.get(),
                hearing_diag_var.get(),
                speech_var.get(),
                blood_group_var.get(),
                ecg_var.get(),
                prescribed_var.get(),
                special_treatment_var.get(),
                medofficer_var.get(),
                med_institution_var.get(),
                datetime.now(),
                selected_health_id
            ))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Health record updated.")
            clear_form()
            load_records()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to update record:\n{e}")

    # ---------- Print ----------


    def print_record():
        """Generate a horizontal A5 health card for the selected child with photo and open it (temp file)."""
        child_name = child_var.get().strip()
        if not child_name:
            messagebox.showwarning("Select", "Please select a child first.")
            return

        try:
            # Get ChildID from in-memory map or fallback to DB
            child_id = None
            if 'child_map' in globals() and isinstance(child_map, dict):
                info = child_map.get(child_name)
                if info:
                    child_id = info.get("id")

            if not child_id:
                conn = get_connection()
                cursor = conn.cursor()
                cursor.execute("SELECT ChildID FROM tblChildren WHERE FullName = ?", (child_name,))
                row = cursor.fetchone()
                conn.close()
                if row:
                    child_id = row.ChildID
                else:
                    messagebox.showerror("Not Found", f"Could not find ChildID for '{child_name}'")
                    return

            # Fetch registration number, center, and photo path
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT c.RegistrationNumber, c.DateOfBirth, c.Gender, c.FatherName, c.PhotoPath, ce.CenterName
                FROM tblChildren c
                LEFT JOIN tblCenters ce ON c.CenterID = ce.CenterID
                WHERE c.ChildID = ?
            """, (child_id,))
            info_row = cursor.fetchone()
            conn.close()

            registration_no = getattr(info_row, "RegistrationNumber", "") or ""
            dob_val = getattr(info_row, "DateOfBirth", None)
            dob_str = format_db_date(dob_val, "%d/%m/%Y") if dob_val else ""
            gender_val = getattr(info_row, "Gender", "") or ""
            father_name = getattr(info_row, "FatherName", "") or ""
            center_name = getattr(info_row, "CenterName", "") or "Aghosh Center"
            photo_path = getattr(info_row, "PhotoPath", "") or ""

            # === Prepare PDF ===
            tmp_dir = tempfile.gettempdir()
            safe_name = child_name.replace(" ", "_").replace("/", "_")
            tmp_pdf = os.path.join(tmp_dir, f"HealthCard_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")

            c = canvas.Canvas(tmp_pdf, pagesize=landscape(A5))
            w, h = landscape(A5)

            left = 30
            right = w - 30
            y = h - 30

            # === Draw Photo if exists ===
            if photo_path and os.path.exists(photo_path):
                try:
                    # Resize image proportionally (thumbnail)
                    with Image.open(photo_path) as im:
                        im.thumbnail((80, 80))  # Max size
                        img_tmp = os.path.join(tmp_dir, f"tmp_photo_{os.path.basename(photo_path)}")
                        im.save(img_tmp)

                    # Place at top right corner under header
                    c.drawImage(img_tmp, right - 90, y - 60, width=80, height=80, preserveAspectRatio=True, mask='auto')
                except Exception as img_err:
                    print(f"Photo error: {img_err}")

            # === Header ===
            c.setFont("Helvetica-Bold", 16)
            c.drawString(left, y, "Al-Khidmat Khawateen Trust Pakistan")

            c.setFont("Helvetica-Bold", 13)
            c.drawString(left, y - 22, center_name)

            c.setFont("Helvetica-Bold", 14)
            c.drawString(left, y - 44, "CHILD HEALTH EXAMINATION CARD")

            y -= 68  # shift below header & photo

            # === Personal Information ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Personal Information")
            c.line(left, y - 2, right, y - 2)
            # Generated On
            c.setFont("Helvetica-Oblique", 7)
            c.drawRightString(right, y, f"Health Card Generated On: {datetime.now().strftime('%d-%m-%Y')}")
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Aghosh ID:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, registration_no)

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 200, y, "Name:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 280, y, child_name)

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 430, y, "Gender:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 490, y, gender_val)
            y -= 16

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Date of Birth:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, dob_str)

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 200, y, "Father Name:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 280, y, father_name)
            y -= 24

            # === Physical Examination ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Physical Examination")
            c.line(left, y - 2, right, y - 2)
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "BMI:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, bmi_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 180, y, "Height:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 220, y, height_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 340, y, "Weight:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 410, y, weight_var.get())
            y -= 16

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Diagnosis:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, diagnosis_var.get())
            y -= 20

            # === Vision ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Vision")
            c.line(left, y - 2, right, y - 2)
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Right:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, vision_r_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 180, y, "Left:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 220, y, vision_l_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 340, y, "Diagnosis:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 410, y, vision_diag_var.get())
            y -= 20

            # === Hearing ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Hearing")
            c.line(left, y - 2, right, y - 2)
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Right:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, hearing_r_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 180, y, "Left:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 220, y, hearing_l_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 340, y, "Diagnosis:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 410, y, hearing_diag_var.get())
            y -= 20

            # === Speech / Investigations ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Speech / Investigations")
            c.line(left, y - 2, right, y - 2)
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Speech:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, speech_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 340, y, "Blood Group:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 410, y, blood_group_var.get())
            y -= 16

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "ECG:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, ecg_var.get())
            y -= 20

            # === Treatment ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Treatment / Medication")
            c.line(left, y - 2, right, y - 2)
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Prescribed Medicine:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 140, y, prescribed_var.get()[:80])
            y -= 16

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Special Treatment:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 140, y, special_treatment_var.get()[:80])
            y -= 20

            # === Medical Officer ===
            c.setFont("Helvetica", 11)
            c.drawString(left, y, "Medical Officer")
            c.line(left, y - 2, right, y - 2)
            y -= 18

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left, y, "Name:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 70, y, medofficer_var.get())

            c.setFont("Helvetica-Bold", 10)
            c.drawString(left + 200, y, "Institution:")
            c.setFont("Helvetica", 10)
            c.drawString(left + 280, y, med_institution_var.get())
            y -= 24

            c.save()

            # Open PDF in viewer
            try:
                os.startfile(tmp_pdf)
            except Exception:
                messagebox.showinfo("Saved", f"Report created: {tmp_pdf}")

        except Exception as e:
            messagebox.showerror("Print Error", f"Failed to generate health card:\n{e}")


# ------------ End of Print Record Health Card ----------


    # Wire the buttons
    save_btn.config(command=save_record)
    update_btn.config(command=update_record)
    clear_btn.config(command=clear_form)
    print_btn.config(command=print_record)

    # load initial data and records
    load_children()
    load_records()
    # set focus on child
    child_combo.focus_set()

    
# --------------end of Health mangement -----------

# --------------Start of Extra Cariculum activities -----------
def open_activities_mgmt():
    win = tk.Toplevel()
    win.title("Manage Extracurricular Activities" + ORG_SUFFIX)
    win.geometry("900x600")
    win.grab_set()
    win.focus_force()

    # ---------- Variables ----------
    activity_id_var = tk.StringVar()
    activity_name_var = tk.StringVar()
    date_var = tk.StringVar()
    organized_by_var = tk.StringVar()
    desc_var = tk.StringVar()
    center_var = tk.StringVar()
    session_var = tk.StringVar()

    center_map = {}
    session_map = {}

    # ---------- Load Centers ----------
    def load_centers():
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT CenterID, CenterName FROM tblCenters")
        rows = cursor.fetchall()
        conn.close()
        for r in rows:
            center_map[r.CenterName] = r.CenterID
        center_combo["values"] = list(center_map.keys())
        if rows:
            first_center = list(center_map.keys())[0]
            center_combo.set(first_center)


    # ---------- Load Sessions ----------
    '''def load_sessions():
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT SessionID, SessionName FROM tblSessions WHERE IsCurrentSession = 1")
        rows = cursor.fetchall()
        conn.close()
        for r in rows:
            session_map[r.SessionName] = r.SessionID
        session_combo["values"] = list(session_map.keys())
        # Auto select current session
        if rows:
            session_var.set(rows[0].SessionName)'''
            
    def load_sessions():
        """Load all sessions into dropdown but auto-select the current one."""
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT SessionID, SessionName, IsCurrentSession FROM tblSessions")
        rows = cursor.fetchall()
        conn.close()

        session_combo["values"] = [r.SessionName for r in rows]
        session_map.clear()
        for r in rows:
            session_map[r.SessionName] = r.SessionID

        # Auto-select the current session (IsCurrentSession = 1)
        for r in rows:
            if getattr(r, "IsCurrentSession", 0) == 1:
                session_var.set(r.SessionName)
                break
        else:
            # If no current session, just set the first one
            if rows:
                session_var.set(rows[0].SessionName)

        # Load activities for selected session on startup
        load_activities_for_selected_session()


    def on_session_select(event):
        """Triggered when a session is selected in the dropdown."""
        load_activities_for_selected_session()


    def load_activities_for_selected_session():
        """Helper to filter activities by selected session."""
        selected_session = session_var.get()
        if not selected_session:
            return

        session_id = session_map.get(selected_session)
        if not session_id:
            return

        # Modify load_activities to accept an optional session_id filter
        load_activities(session_id=session_id)


    # ---------- Save / Update ----------
    def save_activity():
        """Save a new activity record into tblActivities."""
        try:
            name = activity_name_var.get().strip()
            date_text = date_var.get().strip()
            center_name = center_combo.get().strip()
            session_name = session_combo.get().strip()
            desc = desc_var.get().strip()
            organized_by = organized_by_var.get().strip()
            #participants = participants_var.get().strip()
            # --- Validate Total Participants (must be numeric) ---
            participants = participants_var.get().strip()
            if not participants.isdigit():
                messagebox.showwarning("Validation Error", "Total Participants must contain only numeric digits (0–9).")
                participants_var.set("")        # clear the textbox
                form.focus_set()                # bring focus back to form
                #form.after(100, lambda: form.winfo_children()[<index_of_participants_entry>].focus_set())  # optional: focus back to entry
                return
            #participants_val = int(participants)


            # Validation
            if not name:
                messagebox.showwarning("Validation", "Please enter activity name.")
                return
            if not date_text:
                messagebox.showwarning("Validation", "Please enter activity date (DD/MM/YYYY).")
                return

            from datetime import datetime
            try:
                date_obj = datetime.strptime(date_text, "%d/%m/%Y")
                sql_date = date_obj.strftime("%Y-%m-%d")  # SQL 2005 safe
            except Exception:
                messagebox.showerror("Date Error", "Invalid date format. Use DD/MM/YYYY.")
                return

            center_id = center_map.get(center_name)
            session_id = session_map.get(session_name)

            if not center_id or not session_id:
                messagebox.showwarning("Validation", "Please select valid Center and Session.")
                return

            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO tblActivities
                (ActivityName, ActivityDate, Description, OrganizedBy, TotalParticipants, CenterID, SessionID)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (
                name,
                sql_date,
                desc or None,
                organized_by or None,
                participants or None,
                center_id,
                session_id
            ))
            #conn.commit()
            #conn.close()

            messagebox.showinfo("Success", "Activity saved successfully.")


        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to save activity:\n{e}")

        conn.commit()
        conn.close()
        clear_form()
        #load_activities()
        selected_session_name = session_combo.get().strip()
        selected_session_id = session_map.get(selected_session_name)
        load_activities(selected_session_id)


    # ---------- Delete ----------
    
    def delete_activity():
        sel = tree.selection()
        if not sel:
            messagebox.showwarning("Select", "Please select an activity to delete.")
            return
        aid = activity_id_var.get().strip()
        #aid = sel[0]
        if not messagebox.askyesno("Confirm Delete", "Are you sure you want to delete this activity?"):
            return
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("DELETE FROM tblActivities WHERE ActivityID = ?", (aid,))
            conn.commit()
            conn.close()
            messagebox.showinfo("Deleted", "Activity deleted.")
            clear_form()
            #load_activities()
            selected_session_name = session_combo.get().strip()
            selected_session_id = session_map.get(selected_session_name)
            load_activities(selected_session_id)
            save_btn.config(state="normal")
            update_btn.config(state="normal")
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to delete activity:\n{e}")


    # ---------- Load / Display Activities ----------
    def load_activities(session_id=None):
        """Load EC activities into the Treeview, optionally filtered by SessionID."""
        try:
            # Clear existing rows first
            for item in tree.get_children():
                tree.delete(item)

            conn = get_connection()
            cursor = conn.cursor()

            if session_id:
                cursor.execute("""
                    SELECT 
                        a.ActivityID,
                        a.ActivityName,
                        a.ActivityDate,
                        a.OrganizedBy,
                        a.TotalParticipants,
                        a.Description,
                        c.CenterName,
                        s.SessionName
                    FROM tblActivities a
                    LEFT JOIN tblCenters c ON a.CenterID = c.CenterID
                    LEFT JOIN tblSessions s ON a.SessionID = s.SessionID
                    WHERE a.SessionID = ?
                    ORDER BY a.ActivityDate DESC
                """, (session_id,))
            else:
                cursor.execute("""
                    SELECT 
                        a.ActivityID,
                        a.ActivityName,
                        a.ActivityDate,
                        a.OrganizedBy,
                        a.TotalParticipants,
                        a.Description,
                        c.CenterName,
                        s.SessionName
                    FROM tblActivities a
                    LEFT JOIN tblCenters c ON a.CenterID = c.CenterID
                    LEFT JOIN tblSessions s ON a.SessionID = s.SessionID
                    ORDER BY a.ActivityDate DESC
                """)

            rows = cursor.fetchall()
            conn.close()

            for row in rows:
                # Format date safely
                activity_date = ""
                try:
                    if row.ActivityDate:
                        activity_date = format_db_date(row.ActivityDate, "%d/%m/%Y")
                except Exception:
                    activity_date = str(row.ActivityDate) or ""

                # Insert into Treeview (store ActivityID as iid for easy update/delete)
                tree.insert(
                    "",
                    "end",
                    iid=row.ActivityID,  # this is important for update/delete
                    values=(
                        row.ActivityName or "",
                        activity_date,
                        row.CenterName or "",
                        row.SessionName or "",
                        row.OrganizedBy or "",
                        row.TotalParticipants or ""
                    )
                )

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load activities:\n{e}")


    
   
    # ---------- On Select ----------
    def on_select(event):
        
        sel = tree.selection()
        if not sel:
            return
        item_id = sel[0]
        vals = tree.item(item_id, "values")
        if not vals:
            return

        activity_id_var.set(item_id)  # store ActivityID (iid)
        activity_name_var.set(vals[0])
        date_var.set(vals[1])
        center_var.set(vals[2])
        session_var.set(vals[3])
        organized_by_var.set(vals[4])
        participants_var.set(vals[5])

        # Fetch description separately
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT Description FROM tblActivities WHERE ActivityID = ?", (item_id,))
        row = cursor.fetchone()
        conn.close()
        desc_var.set(row.Description if row else "")

        # Disable Save button, enable Update/Delete
        save_btn.config(state="disabled")
        update_btn.config(state="normal")





    # ---------- Clear ----------
    def clear_form():
        activity_id_var.set("")
        activity_name_var.set("")
        #date_var.set("")
        organized_by_var.set("")
        desc_var.set("")
        #center_var.set("")
        participants_var.set("")
        # do not clear session (keep current)
        tree.selection_remove(tree.selection())
        #center_combo.focus_set()
        activity_name_entry.focus_set()

    # ---------- update activity -------------
    def update_activity():
        aid = activity_id_var.get().strip()
        if not aid:
            messagebox.showwarning("Select", "Please select an activity to update.")
            return

        # Basic validation
        name = activity_name_var.get().strip()
        if not name:
            messagebox.showwarning("Validation", "Activity name is required.")
            return

        # Convert date from dd/mm/YYYY to YYYY-MM-DD (SQL friendly)
        date_text = date_var.get().strip()
        activity_date = None
        if date_text:
            try:
                activity_date = datetime.strptime(date_text, "%d/%m/%Y").strftime("%Y-%m-%d")
            except Exception:
                messagebox.showwarning("Date Error", "Activity Date must be in dd/mm/YYYY format.")
                return

        # center and session ids
        center_name = center_combo.get().strip()
        session_name = session_combo.get().strip()
        center_id = center_map.get(center_name)
        session_id = session_map.get(session_name)

        desc = desc_var.get().strip()
        org = organized_by_var.get().strip()
        participants = participants_var.get().strip()
        try:
            participants_val = int(participants) if participants != "" else None
        except ValueError:
            messagebox.showwarning("Validation", "Total Participants must be an integer.")
            return

        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                UPDATE tblActivities
                SET ActivityName = ?, ActivityDate = ?, CenterID = ?, SessionID = ?,
                    Description = ?, OrganizedBy = ?, TotalParticipants = ?
                WHERE ActivityID = ?
            """, (
                name,
                activity_date,
                center_id,
                session_id,
                desc,
                org,
                participants_val,
                aid
            ))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Activity updated.")
            save_btn.config(state="normal")
            update_btn.config(state="disabled")
            clear_form()
            #load_activities()
            selected_session_name = session_combo.get().strip()
            selected_session_id = session_map.get(selected_session_name)
            load_activities(selected_session_id)
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to update activity:\n{e}")
            
    # ----------- prit pdf --------------
    def print_activities_report():
        """Generate a temporary A4 PDF showing all activities of the selected session."""
        try:
            session_name = session_combo.get().strip()
            if not session_name:
                messagebox.showwarning("Select Session", "Please select a session first.")
                return

            session_id = session_map.get(session_name)
            if not session_id:
                messagebox.showwarning("Error", "Invalid session selected.")
                return

            # --- Fetch all activities for this session ---
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT a.ActivityName, a.ActivityDate, a.OrganizedBy, a.TotalParticipants, a.Description, 
                    c.CenterName, s.SessionName
                FROM tblActivities a
                LEFT JOIN tblCenters c ON a.CenterID = c.CenterID
                LEFT JOIN tblSessions s ON a.SessionID = s.SessionID
                WHERE a.SessionID = ?
                ORDER BY a.ActivityDate ASC
            """, (session_id,))
            rows = cursor.fetchall()
            conn.close()

            if not rows:
                messagebox.showinfo("No Data", f"No activities found for session '{session_name}'.")
                return

            # --- Determine center name (from first record) ---
            center_name = rows[0].CenterName or "N/A"

            # --- Prepare PDF in temp memory ---
            tmp_dir = tempfile.gettempdir()
            tmp_pdf = os.path.join(
                tmp_dir,
                f"EC_Activities_{session_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )

            

            doc = SimpleDocTemplate(tmp_pdf, pagesize=A4, leftMargin=40, rightMargin=40, topMargin=40, bottomMargin=30)
            styles = getSampleStyleSheet()
            elements = []

            # --- Header ---
            title_style = styles["Title"]
            subtitle_style = styles["Heading2"]
            normal_style = styles["Normal"]

            elements.append(Paragraph("<b>Al-Khidmat Khawateen Trust Pakistan</b>", title_style))
            elements.append(Spacer(1, 8))
            elements.append(Paragraph(f"<b>{center_name}</b>", subtitle_style))
            elements.append(Spacer(1, 6))
            elements.append(Paragraph(f"<b>Extra-Curricular Activities - Session: {session_name}</b>", normal_style))
            elements.append(Spacer(1, 14))

            # --- Table Data ---
            table_data = [["#", "Activity Name", "Date", "Organized By", "Participants", "Description"]]
            for i, r in enumerate(rows, start=1):
                date_str = ""
                try:
                    if r.ActivityDate:
                        date_str = format_db_date(r.ActivityDate, "%d/%m/%Y")
                except:
                    date_str = str(r.ActivityDate) or ""

                table_data.append([
                    str(i),
                    r.ActivityName or "",
                    date_str,
                    r.OrganizedBy or "",
                    str(r.TotalParticipants or ""),
                    (r.Description or "")[:80],
                ])

            # --- Table Style ---
            tbl = Table(table_data, colWidths=[25, 120, 70, 100, 70, 160])
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("BOX", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))

            elements.append(tbl)
            elements.append(Spacer(1, 14))
            elements.append(Paragraph(
                f"<font size=8><i>Generated on {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}</i></font>",
                normal_style
            ))

            doc.build(elements)

            # --- Open PDF directly for printing (not saved permanently) ---
            try:
                os.startfile(tmp_pdf)
            except Exception:
                messagebox.showinfo("Report Created", f"PDF created temporarily at:\n{tmp_pdf}")

        except Exception as e:
            messagebox.showerror("Print Error", f"Failed to generate report:\n{e}")

    

    # ---------- Layout ----------
    form = tk.LabelFrame(win, text="Add / Edit Activity", padx=10, pady=10)
    form.pack(fill="x", padx=10, pady=5)

    tk.Label(form, text="Center:").grid(row=0, column=0, sticky="e")
    center_combo = ttk.Combobox(form, textvariable=center_var, width=30, state="readonly")
    center_combo.grid(row=0, column=1, padx=5, pady=3)

    tk.Label(form, text="Session:").grid(row=0, column=2, sticky="e")
    session_combo = ttk.Combobox(form, textvariable=session_var, width=30, state="readonly")
    session_combo.grid(row=0, column=3, padx=5, pady=3)
    session_combo.bind("<<ComboboxSelected>>", on_session_select)


    tk.Label(form, text="Activity Name:").grid(row=1, column=0, sticky="e")
    activity_name_entry = tk.Entry(form, textvariable=activity_name_var, width=33)
    activity_name_entry.grid(row=1, column=1, padx=5, pady=3)
 
    tk.Label(form, text="Activity Date:").grid(row=1, column=2, sticky="e")

    date_picker = DateEntry(
        form,
        textvariable=date_var,
        date_pattern='dd/mm/yyyy',   # Display in day/month/year format
        width=30,
        background='darkblue',
        foreground='white',
        borderwidth=2
    )
    date_picker.grid(row=1, column=3, padx=5, pady=3)

    tk.Label(form, text="Organized By:").grid(row=2, column=0, sticky="e")
    tk.Entry(form, textvariable=organized_by_var, width=22).grid(row=2, column=1, padx=5, pady=3)
    
    tk.Label(form, text="Total Participants:").grid(row=2, column=2, sticky="e")
    participants_var = tk.StringVar()
    tk.Entry(form, textvariable=participants_var, width=22).grid(row=2, column=3, padx=5, pady=3)
    
    tk.Label(form, text="One Line Descp:").grid(row=2, column=4, sticky="e")
    tk.Entry(form, textvariable=desc_var, width=12).grid(row=2, column=5, padx=5, pady=3)
    


    
    '''save_btn = tk.Button(form, text="Save / Add", command=save_activity, bg="blue", fg="white", width=12)
    save_btn.grid(row=3, column=0, pady=10)

    tk.Button(form, text="Clear", command=clear_form, width=12).grid(row=3, column=1, pady=10)
    tk.Button(form, text="Delete", command=delete_activity, bg="red", fg="white", width=12).grid(row=3, column=3, pady=10, sticky="e")
    update_btn = tk.Button(form, text="Update", command=update_activity, bg="Green", fg="white", width=12)
    update_btn.grid(row=3, column=2, pady=10)
    
    tk.Button(form, text="Print Activities", command=print_activities_report, bg="green", fg="white", width=12).grid(row=3, column=4, pady=10)'''
    # --- Buttons Row (Aligned Nicely) ---
    button_frame = tk.Frame(form)
    button_frame.grid(row=3, column=0, columnspan=5, pady=10)

    save_btn = tk.Button(button_frame, text="Save / Add", command=save_activity, bg="blue", fg="white", width=12)
    save_btn.pack(side="left", padx=5)

    update_btn = tk.Button(button_frame, text="Update", command=update_activity, bg="green", fg="white", width=12)
    update_btn.pack(side="left", padx=5)

    tk.Button(button_frame, text="Clear", command=clear_form, width=12).pack(side="left", padx=5)
    tk.Button(button_frame, text="Delete", command=delete_activity, bg="red", fg="white", width=12).pack(side="left", padx=5)
    tk.Button(button_frame, text="Print Activities", command=print_activities_report, bg="darkgreen", fg="white", width=14).pack(side="left", padx=5)



    # ---------- Treeview ----------
    # === Treeview for Activities ===
    cols = ("Activity Name", "Date", "Center", "Session", "Organized By", "Participants")
    tree = ttk.Treeview(win, columns=cols, show="headings")

    for col in cols:
        tree.heading(col, text=col)
        tree.column(col, width=130)

    tree.pack(fill="both", expand=True, padx=10, pady=10)

    # Bind selection event
    tree.bind("<<TreeviewSelect>>", on_select)
    activity_name_entry.focus_set()
    


    # ---------- Initialize ----------
    load_centers()
    load_sessions()
    load_activities()


# --------------End of Extra Cariculum activities -----------

# ------------ Start of Enroll Child in a activity -----------

def open_child_activities_mgmt():
    import tkinter as tk
    from tkinter import ttk, messagebox
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    import tempfile, os
    from datetime import datetime

    win = tk.Toplevel()
    win.title("Child Extra-Curricular Participation Management" + ORG_SUFFIX)
    win.geometry("1100x650")
    win.grab_set()

    # --- Center window ---
    win.update_idletasks()
    w, h = 1100, 650
    x = (win.winfo_screenwidth() // 2) - (w // 2)
    y = (win.winfo_screenheight() // 2) - (h // 2)
    win.geometry(f"{w}x{h}+{x}+{y}")
    win.bind("<Escape>", lambda e: win.destroy())

    # --- Variables ---
    session_var = tk.StringVar()
    activity_var = tk.StringVar()
    search_var = tk.StringVar()

    sessions_map, activities_map = {}, {}
    all_children = []  # [(ChildID, RegNo, FullName), ...] for search

    # --- Top Frame ---
    top = tk.LabelFrame(win, text="Session & Activity", padx=10, pady=10)
    top.pack(fill="x", padx=10, pady=5)

    tk.Label(top, text="Session:").grid(row=0, column=0, sticky="e")
    session_combo = ttk.Combobox(top, textvariable=session_var, width=30, state="readonly")
    session_combo.grid(row=0, column=1, padx=5)

    tk.Label(top, text="Activity:").grid(row=0, column=2, sticky="e", padx=(10,0))
    activity_combo = ttk.Combobox(top, textvariable=activity_var, width=40, state="readonly")
    activity_combo.grid(row=0, column=3, padx=5)

    # --- Search Bar ---
    tk.Label(top, text="Search Child:").grid(row=0, column=4, sticky="e", padx=(30, 5))
    search_entry = tk.Entry(top, textvariable=search_var, width=30)
    search_entry.grid(row=0, column=5, padx=5)
    # we'll call search_child on Return or when the field changes
    search_entry.bind("<KeyRelease>", lambda e: search_child())

    # --- Treeview ---
    cols = ("RegNo", "Child Name", "Remarks", "Award", "Position")
    tree = ttk.Treeview(win, columns=cols, show="headings", height=18)
    for col in cols:
        tree.heading(col, text=col)
        tree.column(col, width=200)
    tree.column("RegNo", width=100)
    tree.pack(fill="both", expand=True, padx=10, pady=8)

    # --- Inline cell edit handler ---
    def edit_cell(event):
        region = tree.identify("region", event.x, event.y)
        if region != "cell":
            return

        col = tree.identify_column(event.x)   # like "#3"
        row = tree.identify_row(event.y)
        if not row:
            return

        # disallow editing RegNo and Child Name (#1 and #2)
        if col in ("#1", "#2"):
            return

        # compute bbox and column index
        bbox = tree.bbox(row, col)
        if not bbox:
            return
        x, y, width, height = bbox
        col_index = int(col[1:]) - 1  # 0-based for cols list
        current_val = tree.item(row, "values")[col_index] if tree.item(row, "values") else ""

        entry = tk.Entry(tree, font=("Helvetica", 10))
        entry.insert(0, current_val)
        entry.place(x=x, y=y, width=width, height=height)
        entry.focus()

        def save_edit(event=None):
            new_val = entry.get().strip()
            # update tree view cell
            vals = list(tree.item(row, "values"))
            vals[col_index] = new_val
            tree.item(row, values=vals)
            entry.destroy()

        entry.bind("<Return>", save_edit)
        entry.bind("<FocusOut>", save_edit)

    tree.bind("<Double-1>", edit_cell)

    # --- Buttons ---
    btn_frame = tk.Frame(win)
    btn_frame.pack(fill="x", pady=5)
    save_btn = tk.Button(btn_frame, text="Save Participation", bg="blue", fg="white", width=18, command=lambda: save_participation())
    save_btn.pack(side="left", padx=10)
    tk.Button(btn_frame, text="Print Report", bg="green", fg="white", width=18, command=lambda: print_participation_report()).pack(side="right", padx=10)

    # -------------------------
    # Helper: get currently selected session id and activity id
    def _get_selected_session_id():
        return sessions_map.get(session_var.get())

    def _get_selected_activity_id():
        return activities_map.get(activity_var.get())

    # --- Load sessions ---
    def load_sessions():
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT SessionID, SessionName, IsCurrentSession FROM tblSessions ORDER BY SessionName")
            rows = cursor.fetchall()
            conn.close()

            sessions_map.clear()
            for r in rows:
                sessions_map[getattr(r, "SessionName", r[1])] = getattr(r, "SessionID", r[0])
            session_combo["values"] = list(sessions_map.keys())

            # select current session if any, otherwise pick first
            current = None
            for r in rows:
                is_current = getattr(r, "IsCurrentSession", None)
                if is_current:
                    current = getattr(r, "SessionName", r[1])
                    break
            if current:
                session_var.set(current)
            elif rows:
                session_var.set(getattr(rows[0], "SessionName", rows[0][1]))

            # whenever session selection changes, reload activities and children
            session_combo.bind("<<ComboboxSelected>>", lambda e: load_activities_for_session())

            # initial load
            load_activities_for_session()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load sessions:\n{e}")

    # --- Load activities for currently selected session ---
    def load_activities_for_session():
        try:
            # clear search box and activity selection visually
            activity_var.set('')
            search_var.set('')

            sid = _get_selected_session_id()
            if not sid:
                # clear activity list and children tree
                activities_map.clear()
                activity_combo["values"] = []
                load_children_for_session()  # will do nothing when no session selected
                return

            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT ActivityID, ActivityName FROM tblActivities WHERE SessionID=? ORDER BY ActivityName ASC", (sid,))
            acts = cursor.fetchall()
            conn.close()

            activities_map.clear()
            for a in acts:
                name = getattr(a, "ActivityName", a[1])
                aid = getattr(a, "ActivityID", a[0])
                activities_map[name] = aid

            activity_combo["values"] = list(activities_map.keys())

            # bind activity selection -> reload children for that activity
            activity_combo.bind("<<ComboboxSelected>>", lambda e: load_children_for_session())

            # If there's at least one activity, select the first by default
            if acts:
                first_name = getattr(acts[0], "ActivityName", acts[0][1])
                activity_var.set(first_name)

            # load children for current session and activity
            load_children_for_session()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load activities:\n{e}")

    # --- Load children for the chosen session and (optional) activity ---
    def load_children_for_session():
        """Populate tree with children of the selected session.
           If an activity is selected, also load saved participation data for that activity."""
        try:
            # clear tree
            for item in tree.get_children():
                tree.delete(item)
            all_children.clear()

            sid = _get_selected_session_id()
            if not sid:
                return

            aid = _get_selected_activity_id()  # may be None

            conn = get_connection()
            cursor = conn.cursor()

            if aid:
                # join to get participation values for the selected activity
                query = """
                    SELECT c.ChildID, c.RegistrationNumber, c.FullName,
                           IFNULL(p.ParticipationRemarks, '') AS ParticipationRemarks,
                           IFNULL(p.AwardReceived, '') AS AwardReceived,
                           IFNULL(p.Position, '') AS Position
                    FROM tblChildren c
                    LEFT JOIN tblChildActivities p
                      ON c.ChildID = p.ChildID AND p.ActivityID = ?
                    WHERE c.SessionID = ?
                    ORDER BY c.FullName
                """
                cursor.execute(query, (aid, sid))
            else:
                # no activity selected — just list children (blank participation fields)
                query = """
                    SELECT c.ChildID, c.RegistrationNumber, c.FullName,
                           '' AS ParticipationRemarks, '' AS AwardReceived, '' AS Position
                    FROM tblChildren c
                    WHERE c.SessionID = ?
                    ORDER BY c.FullName
                """
                cursor.execute(query, (sid,))

            rows = cursor.fetchall()
            conn.close()

            for r in rows:
                cid = getattr(r, "ChildID", r[0])
                reg = getattr(r, "RegistrationNumber", r[1]) or ""
                name = getattr(r, "FullName", r[2]) or ""
                remarks = getattr(r, "ParticipationRemarks", None)
                if remarks is None:
                    # depending on driver, alias may be present in tuple positions
                    remarks = r[3] if len(r) > 3 else ""
                award = getattr(r, "AwardReceived", None)
                if award is None:
                    award = r[4] if len(r) > 4 else ""
                pos = getattr(r, "Position", None)
                if pos is None:
                    pos = r[5] if len(r) > 5 else ""

                # populate all_children for search
                all_children.append((cid, reg, name))

                # insert into tree; store ChildID in tags for later save
                tree.insert("", "end",
                            values=(reg, name, remarks or "", award or "", pos or ""),
                            tags=(cid,))
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load children for session:\n{e}")

    # --- Search children within currently-loaded list ---
    def search_child():
        q = search_var.get().strip().lower()
        # clear current view and re-insert only matching children (but preserve participation if present)
        for i in tree.get_children():
            tree.delete(i)

        # We will query DB for participation values (to show saved data when searching)
        sid = _get_selected_session_id()
        aid = _get_selected_activity_id()

        # build a mapping of childID->(reg, name, remarks, award, pos) for fast lookup
        lookup = {}
        try:
            conn = get_connection()
            cursor = conn.cursor()
            if aid:
                query = """
                    SELECT c.ChildID, c.RegistrationNumber, c.FullName,
                           IFNULL(p.ParticipationRemarks, '') AS ParticipationRemarks,
                           IFNULL(p.AwardReceived, '') AS AwardReceived,
                           IFNULL(p.Position, '') AS Position
                    FROM tblChildren c
                    LEFT JOIN tblChildActivities p
                      ON c.ChildID = p.ChildID AND p.ActivityID = ?
                    WHERE c.SessionID = ?
                """
                cursor.execute(query, (aid, sid))
            else:
                query = """
                    SELECT c.ChildID, c.RegistrationNumber, c.FullName,
                           '' AS ParticipationRemarks, '' AS AwardReceived, '' AS Position
                    FROM tblChildren c
                    WHERE c.SessionID = ?
                """
                cursor.execute(query, (sid,))
            rows = cursor.fetchall()
            conn.close()
            for r in rows:
                cid = getattr(r, "ChildID", r[0])
                reg = getattr(r, "RegistrationNumber", r[1]) or ""
                name = getattr(r, "FullName", r[2]) or ""
                remarks = getattr(r, "ParticipationRemarks", None)
                if remarks is None:
                    remarks = r[3] if len(r) > 3 else ""
                award = getattr(r, "AwardReceived", None)
                if award is None:
                    award = r[4] if len(r) > 4 else ""
                pos = getattr(r, "Position", None)
                if pos is None:
                    pos = r[5] if len(r) > 5 else ""
                lookup[cid] = (reg, name, remarks or "", award or "", pos or "")
        except Exception as e:
            messagebox.showerror("DB Error", f"Search failed:\n{e}")
            return

        # now filter and insert matches
        for cid, reg, name in all_children:
            if not q or q in name.lower() or q in reg.lower():
                vals = lookup.get(cid, (reg, name, "", "", ""))
                tree.insert("", "end", values=vals, tags=(cid,))

    # --- Save participation: insert/update records from the treeview into tblChildActivities ---
    def save_participation():
        aid = _get_selected_activity_id()
        if not aid:
            messagebox.showwarning("Select", "Please select an activity.")
            return

        try:
            conn = get_connection()
            cursor = conn.cursor()
            count = 0

            for row_id in tree.get_children():
                vals = tree.item(row_id, "values")
                tags = tree.item(row_id, "tags")
                if not vals or not tags:
                    continue

                reg_no, name, remarks, award, position = vals
                child_id = tags[0]

                # Only save if user entered something (non-empty remarks/award/position)
                if (remarks and str(remarks).strip()) or (award and str(award).strip()) or (position and str(position).strip()):
                    # check if record exists
                    cursor.execute("SELECT COUNT(*) FROM tblChildActivities WHERE ChildID=? AND ActivityID=?", (child_id, aid))
                    exists = cursor.fetchone()[0]

                    if exists and exists > 0:
                        cursor.execute("""
                            UPDATE tblChildActivities
                            SET ParticipationRemarks=?, AwardReceived=?, Position=?
                            WHERE ChildID=? AND ActivityID=?
                        """, (remarks, award, position, child_id, aid))
                    else:
                        cursor.execute("""
                            INSERT INTO tblChildActivities (ChildID, ActivityID, ParticipationRemarks, AwardReceived, Position)
                            VALUES (?, ?, ?, ?, ?)
                        """, (child_id, aid, remarks, award, position))
                    count += 1

            conn.commit()
            conn.close()

            messagebox.showinfo("Saved", f"{count} participation record(s) saved successfully.")
            # reload children for the same session/activity to reflect DB state
            load_children_for_session()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to save participation:\n{e}")

    # --- Print participation report for visible rows (temporary file) ---
    def print_participation_report():
        session_name = session_var.get().strip()
        activity_name = activity_var.get().strip()
        if not session_name or not activity_name:
            messagebox.showwarning("Missing Info", "Select both session and activity first.")
            return

        rows = [tree.item(i, "values") for i in tree.get_children()]
        if not rows:
            messagebox.showwarning("No Data", "No participation data to print.")
            return

        tmp_pdf = os.path.join(tempfile.gettempdir(), f"EC_Activities_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
        try:
            c = canvas.Canvas(tmp_pdf, pagesize=A4)
            w, h = A4
            y = h - 60

            c.setFont("Helvetica-Bold", 16)
            c.drawCentredString(w/2, y, "Al-Khidmat Khawateen Trust Pakistan")
            y -= 20
            c.setFont("Helvetica-Bold", 12)
            c.drawCentredString(w/2, y, f"Extra-Curricular Activities Report ({session_name})")
            y -= 15
            c.setFont("Helvetica", 11)
            c.drawCentredString(w/2, y, f"Activity: {activity_name}")
            y -= 30

            c.setFont("Helvetica-Bold", 10)
            c.drawString(40, y, "Reg No")
            c.drawString(120, y, "Child Name")
            c.drawString(320, y, "Remarks")
            c.drawString(480, y, "Award")
            c.drawString(550, y, "Position")
            y -= 10
            c.line(40, y, w - 40, y)
            y -= 15

            c.setFont("Helvetica", 9)
            for reg, name, remarks, award, position in rows:
                if y < 60:
                    c.showPage()
                    y = h - 60
                c.drawString(40, y, str(reg))
                c.drawString(120, y, str(name))
                c.drawString(320, y, str(remarks))
                c.drawString(480, y, str(award))
                c.drawString(550, y, str(position))
                y -= 15

            c.save()
            try:
                os.startfile(tmp_pdf)
            except Exception:
                messagebox.showinfo("Report Generated", f"PDF created temporarily at:\n{tmp_pdf}")
        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # --- Initialize ---
    load_sessions()




# ---------- End of Enroll --------------
    
# ============================================================
# Child Profile Report
# ============================================================
    
def open_child_profile_report_form():
    #def open_child_profile_report_form():
    global tree
    tree = None
    report_win = tk.Toplevel()
    report_win.title("Child Profile Report" + ORG_SUFFIX)
    report_win.geometry("900x600")

    # Center on screen
    ww, wh = 900, 600
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    # === Load children from database ===
    children_data = []
    child_lookup = {}  # FullName -> (ChildID, FatherName)
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT ChildID, FullName, FatherName FROM tblChildren ORDER BY FullName")
        rows = cursor.fetchall()
        for row in rows:
            children_data.append((row.FullName, row.FatherName))
            child_lookup[row.FullName] = (row.ChildID, row.FatherName)
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load children: {e}")
        return

    # === Search Frame ===
    search_frame = tk.Frame(report_win)
    search_frame.pack(pady=10)

    tk.Label(search_frame, text="Search Child:", font=("Arial", 10)).grid(row=0, column=0, padx=5)
    search_var = tk.StringVar()
    search_combo = ttk.Combobox(search_frame, textvariable=search_var, values=list(child_lookup.keys()), width=40)
    search_combo.grid(row=0, column=1, padx=5)

    # === Tree Frame ===
    tree_frame = tk.Frame(report_win)
    tree_frame.pack(fill="both", expand=True, padx=10, pady=10)

    columns = ("Full Name", "Father Name")
    tree = ttk.Treeview(tree_frame, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=300)

    vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=vsb.set)
    vsb.pack(side='right', fill='y')
    tree.pack(side="left", fill="both", expand=True)

    # Populate
    for child in children_data:
        tree.insert('', tk.END, values=child)

    # Setup tag for highlight
    tree.tag_configure("match", background="lightblue")

    # === Functions ===
    def highlight_matching_rows(event=None):
        search_text = search_var.get().lower()
        first_match_id = None
        for item in tree.get_children():
            values = tree.item(item, "values")
            full_name = values[0].lower()
            if search_text in full_name:
                tree.item(item, tags=("match",))
                if first_match_id is None:
                    first_match_id = item
            else:
                tree.item(item, tags=())
        if first_match_id:
            tree.see(first_match_id)
            tree.selection_set(first_match_id)

    def search_child():
        highlight_matching_rows()

    # Bind after defining function
    search_combo.bind("<KeyRelease>", highlight_matching_rows)
    tk.Button(search_frame, text="Search", command=search_child).grid(row=0, column=2, padx=5)

    
    # === MAKE TREEVIEW SCROLLABLE ===
    vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=vsb.set)
    vsb.pack(side='right', fill='y')

    tree.pack(side="left", fill="both", expand=True)
    
    
    
    #highlighted the search name
    def highlight_matching_rows(event=None):
        search_text = search_var.get().lower()
        first_match_id = None
        for item in tree.get_children():
            values = tree.item(item, "values")
            full_name = values[0].lower()  # assuming Full Name is first column
            if search_text in full_name:
                tree.item(item, tags=("match",))
                if first_match_id is None:
                    first_match_id = item
            else:
                tree.item(item, tags=())
        
        # Auto-scroll to first match
        if first_match_id:
            tree.see(first_match_id)
            tree.selection_set(first_match_id)
        tree.tag_configure("match", background="lightblue")

    
    
    #Create pdf
# ============================================================
# Profile PDF Generation
# ============================================================
    def create_child_profile_pdf(child_id, child_name, father_name):
        conn = get_connection()
        cursor = conn.cursor()

        # Fetch child data
        cursor.execute("SELECT * FROM tblChildren WHERE ChildID = ?", (child_id,))
        row = cursor.fetchone()

        if not row:
            messagebox.showerror("Error", f"No record found for ID {child_id}")
            return

        columns = [col[0] for col in cursor.description]
        child_data = dict(zip(columns, row))

        # Fetch Center Name
        center_id = child_data.get("CenterID")
        center_name = ""
        if center_id:
            cursor.execute("SELECT CenterName FROM tblCenters WHERE CenterID = ?", (center_id,))
            center_row = cursor.fetchone()
            center_name = center_row[0] if center_row else ""
            
        # Fetch Session Name
        session_id = child_data.get("SessionID")
        #messagebox.showinfo("", session_id)
        session_name = ""
        if center_id:
            cursor.execute("SELECT SessionName FROM tblSessions WHERE SessionID = ?", (session_id,))
            session_row = cursor.fetchone()
            session_name = session_row[0] if session_row else ""

        # UI Window
        preview_win = Toplevel()
        preview_win.title("Child Profile Preview" + ORG_SUFFIX)
        preview_win.state("zoomed")
        preview_win.grab_set()
        preview_win.focus_force()

        # === SCROLLABLE FRAME ===
        container = Frame(preview_win)
        container.pack(fill='both', expand=True)

        canvas = Canvas(container)
        scroll_y = Scrollbar(container, orient="vertical", command=canvas.yview)
        scroll_y.pack(side=RIGHT, fill=Y)
        canvas.pack(side=LEFT, fill=BOTH, expand=True)
        canvas.configure(yscrollcommand=scroll_y.set)

        # Create a frame inside the canvas
        scrollable_frame = Frame(canvas)
        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")

        # Update scroll region
        def configure_scroll_region(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        scrollable_frame.bind("<Configure>", configure_scroll_region)

        # Bind mouse wheel
        def on_mouse_wheel(event):
            canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        canvas.bind_all("<MouseWheel>", on_mouse_wheel)

        # === CHILD PHOTO ===
        try:
            photo_frame = Frame(scrollable_frame, width=130, height=150, bd=2, relief="groove")
            photo_frame.grid(row=0, column=3, rowspan=5, padx=20, pady=10, sticky="ne")
            photo_frame.grid_propagate(False)  # Prevent the frame from resizing to image size

            photo_path = child_data.get('PhotoPath')
            if photo_path and os.path.exists(photo_path):
                ext = os.path.splitext(photo_path)[1].lower()
                if ext in ['.png', '.jpg', '.jpeg']:
                    img = Image.open(photo_path)
                    img = img.resize((120, 140))
                    photo_img = ImageTk.PhotoImage(img)
                    img_label = Label(photo_frame, image=photo_img)
                    img_label.place(relx=0.5, rely=0.5, anchor="center")
                    photo_frame.image = photo_img  # Keep a reference
                else:
                    Label(photo_frame, text="Invalid format", font=("Arial", 9)).place(relx=0.5, rely=0.5, anchor="center")
            else:
                Label(photo_frame, text="Photo not found", font=("Arial", 9)).place(relx=0.5, rely=0.5, anchor="center")
        except Exception as e:
            Label(photo_frame, text="Error loading photo", font=("Arial", 9)).place(relx=0.5, rely=0.5, anchor="center")


        # Header
        Label(scrollable_frame, text="Aghosh Program", font=("Arial", 16, "bold")).grid(
            row=0, column=0, columnspan=3, pady=10, sticky="nw")
        
        Label(scrollable_frame, text="AlKhidmat Khawateen Trust Pakistan", font=("Arial", 16, "bold")).grid(
            row=0, column=1, columnspan=3, pady=10, sticky="nw")
        

        Label(scrollable_frame, text=f"Child Name: {child_name}", font=("Arial", 12, "bold")).grid(
            row=1, column=0, sticky="w", padx=20)
       
        Label(scrollable_frame, text=f"Session: {session_name}", font=("Arial", 12, "bold")).grid(
            row=2, column=0, sticky="w", padx=20)

        Label(scrollable_frame, text=f"Center: {center_name}", font=("Arial", 12, "bold")).grid(
            row=3, column=0, sticky="w", padx=20)

        Label(scrollable_frame, text=f"Date: {datetime.now().strftime('%d-%m-%Y')}", font=("Arial", 12, "bold")).grid(
            row=4, column=0, sticky="e", padx=20)

        # === GRID SECTION DISPLAY ===
        row_counter = [5]  # mutable list to update inside nested function

        def display_section(title, fields):
            row_start = row_counter[0]
            Label(scrollable_frame, text=title, font=("Arial", 12, "bold"), fg="blue").grid(
                row=row_start, column=0, columnspan=4, sticky="w", padx=20, pady=10)
            row_counter[0] += 1

            col = 0
            #for field in fields:
            #    if field in child_data:
            #        val = child_data[field]
            #        if val:
            #            Label(scrollable_frame, text=f"{field}:", font=("Arial", 10, "bold")).grid(
            #                row=row_counter[0], column=col, sticky="w", padx=20)

            #            Label(scrollable_frame, text=str(val), wraplength=200, justify="left").grid(
            #                row=row_counter[0], column=col + 1, sticky="w", padx=20)

            #            col += 2
            #            if col >= 8:
            #                col = 0
            #                row_counter[0] += 1
                            
            for field in fields:

                # ----- SPECIAL RULE FOR MOTHER DEATH DATE -----
                if field == "MotherDeathDate":
                    mother_status = str(child_data.get("MotherStatus", "")).strip().lower()

                    if mother_status == "alive":
                        val = "None"   # Explicitly show None
                    else:
                        val = child_data.get(field)

                else:
                    val = child_data.get(field)

                # Skip completely empty fields (except "None" which we want to show)
                if val in [None, "", "null"]:
                    continue

                Label(scrollable_frame, text=f"{field}:", font=("Arial", 10, "bold")).grid(
                    row=row_counter[0], column=col, sticky="w", padx=20
                )

                Label(scrollable_frame, text=str(val), wraplength=200, justify="left").grid(
                    row=row_counter[0], column=col + 1, sticky="w", padx=20
                )

                col += 2
                if col >= 8:
                    col = 0
                    row_counter[0] += 1


            row_counter[0] += 2  # Add spacing between sections

        # Display all sections in grid format
        display_section("Basic Information", [
            "FullName", "FatherName", "Gender", "DateOfBirth", "AdmissionDate", "RegistrationNumber",
            "SchoolName", "Class", "Intelligence", "Disability", "HealthCondition", "Status", "ChildRequiredAmount"
        ])

        display_section("Father Information", [
            "ReasonFatherDeath", "FatherDeathDate", "FatherOccupation", "FatherDesignation"
        ])

        display_section("Mother Information", [
            "MotherName", "MotherStatus", "MotherDeathDate", "MotherCNIC"
        ])

        display_section("Address", [
            "PermanentAddress", "TemporaryAddress"
        ])

        # Siblings
        sibling_fields = []
        for i in range(1, 6):
            sibling_fields.extend([
                f"Sibling{i}Name", f"Sibling{i}Gender", f"Sibling{i}DOB"
            ])
        display_section("Siblings", sibling_fields)

        # Meet Persons
        meet_fields = []
        for i in range(1, 6):
            meet_fields.extend([
                f"MeetPerson{i}Name", f"MeetPerson{i}CNIC", f"MeetPerson{i}Contact"
            ])
        display_section("Authorized to Meet", meet_fields)

        # Introducer
        display_section("Introducer", [
            "IntroducerName", "Introducer CNIC", "IntroducerContact", "Introducer Address"
        ])

        # Guardian
        display_section("Guardian", [
            "GuardianName", "GuardianRelation", "GuardianCNIC", "GuardianContact", "Guardianaddress"
        ])

        # Documents
        display_section("Documents", [
            "DocSchoolCertificate", "DocBForm", "DocFather CNIC", "DocMotherCNIC", "DocFatherDeathCert"
        ])

        # Close cursor
        cursor.close()
        conn.close()

    
    
        # PDF Button
       

        def generate_pdf():
            try:
                temp_pdf = tempfile.mktemp(".pdf")
                pdf = pdfcanvas.Canvas(temp_pdf, pagesize=A4)
                width, height = A4
                x_margin = 50
                y = height - 50

                base_dir = os.path.dirname(os.path.abspath(__file__))  # Path of your script
                # Sample Center Name
                center_name = child_data.get("CenterName", "Aghosh Karachi")
                current_date = datetime.now().strftime("%d-%m-%Y")

                # ---- HEADER ----
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawString(x_margin, y, f"Center: {center_name}")

                pdf.setFont("Helvetica", 10)
                pdf.drawRightString(width - x_margin, y, f"Report Generated On: {current_date}")

                y -= 25
                pdf.setFont("Helvetica-Bold", 18)
                pdf.drawCentredString(width / 2, y, "Alkhidmat Khawateen Trust Pakistan")
                y -= 20
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawCentredString(width / 2, y, "Aghosh Program")
                y -= 20
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawCentredString(width / 2, y, "Child Profile")
                
                y -= 30
                pdf.setFont("Helvetica", 10)
                pdf.drawRightString(width / 2, y, f"Session: {session_name}")
                
            

                # ---- CHILD PHOTO ----
                image_path = child_data.get("PhotoPath", "")
                if image_path and os.path.exists(image_path):
                    try:
                        pdf.drawImage(image_path, x=width - 150, y=height - 180, width=100, height=100, preserveAspectRatio=True)
                    except Exception as e:
                        messagebox.showerror("Image Error", f"Could not add child photo: {str(e)}")

                # ---- FIELD LABELS MAPPING ----
                field_labels = {
                    "FullName": "Full Name", "FatherName": "Father's Name", "Gender": "Gender",
                    "DateOfBirth": "Date of Birth", "AdmissionDate": "Admission Date", "RegistrationNumber": "Registration Number",
                    "SchoolName": "School Name", "Class": "Class", "Intelligence": "Child Category",
                    "Disability": "Disability", "HealthCondition": "Health Condition", "Status": "Child Status",
                    "ChildRequiredAmount": "Monthly Required Amount",
                    "ReasonFatherDeath": "Reason of Father's Death", "FatherDeathDate": "Father's Death Date",
                    "FatherOccupation": "Father's Occupation", "FatherDesignation": "Father's Designation",
                    "MotherName": "Mother's Name", "MotherStatus": "Mother's Status", "MotherDeathDate": "Mother's Death Date",
                    "MotherCNIC": "Mother's CNIC",
                    "PermanentAddress": "Permanent Address", "TemporaryAddress": "Temporary Address",
                    "IntroducerName": "Introducer Name", "Introducer CNIC": "Introducer CNIC", "IntroducerContact": "Introducer Contact",
                    "Introducer Address": "Introducer Address",
                    "GuardianName": "Guardian Name", "GuardianRelation": "Guardian Relation", "GuardianCNIC": "Guardian CNIC",
                    "GuardianContact": "Guardian Contact", "Guardianaddress": "Guardian Address",
                    "DocSchoolCertificate": "School Certificate", "DocBForm": "B-Form",
                    "DocFather CNIC": "Father's CNIC", "DocMotherCNIC": "Mother's CNIC", "DocFatherDeathCert": "Father's Death Certificate"
                }

                # Add siblings and meet-person field labels dynamically
                for i in range(1, 6):
                    field_labels[f"Sibling{i}Name"] = f"Sibling {i} Name"
                    field_labels[f"Sibling{i}Gender"] = f"Sibling {i} Gender"
                    field_labels[f"Sibling{i}DOB"] = f"Sibling {i} Date of Birth"
                    field_labels[f"MeetPerson{i}Name"] = f"Meet Person {i} Name"
                    field_labels[f"MeetPerson{i}CNIC"] = f"Meet Person {i} CNIC"
                    field_labels[f"MeetPerson{i}Contact"] = f"Meet Person {i} Contact"

                def draw_section(title, fields):
                    nonlocal y
                    pdf.setFont("Helvetica-Bold", 12)
                    pdf.setFillColorRGB(0, 0, 1)
                    pdf.drawString(x_margin, y, title)
                    y -= 20
                    pdf.setFillColorRGB(0, 0, 0)
                    pdf.setFont("Helvetica", 10)

                    for field in fields:

                        # ---- SPECIAL RULE: Mother Death Date ----
                        if field == "MotherDeathDate":
                            if child_data.get("MotherStatus", "").strip().lower() == "alive":
                                continue  # Skip death date if mother is alive

                        value = str(child_data.get(field, "")).strip()

                        if value in ["", "None", "null"]:
                            continue  # Skip empty values

                        if y < 100:
                            pdf.showPage()
                            y = height - 50

                        label = field_labels.get(field, field)
                        pdf.drawString(x_margin, y, f"{label}:")
                        pdf.drawString(x_margin + 150, y, value)
                        y -= 15

                    # Dotted separator
                    y -= 5
                    for i in range(x_margin, int(width - x_margin), 5):
                        pdf.drawString(i, y, ".")
                    y -= 15

                # ---- DATA SECTIONS ----
                draw_section("Basic Information", [
                    "FullName", "FatherName", "Gender", "DateOfBirth", "AdmissionDate", "RegistrationNumber",
                    "SchoolName", "Class", "Intelligence", "Disability", "HealthCondition", "Status", "ChildRequiredAmount"
                ])

                draw_section("Father Information", [
                    "ReasonFatherDeath", "FatherDeathDate", "FatherOccupation", "FatherDesignation"
                ])

                draw_section("Mother Information", [
                    "MotherName", "MotherStatus", "MotherDeathDate", "MotherCNIC"
                ])

                draw_section("Address", [
                    "PermanentAddress", "TemporaryAddress"
                ])

                sibling_fields = [f"Sibling{i}{detail}" for i in range(1, 6) for detail in ["Name", "Gender", "DOB"]]
                draw_section("Siblings", sibling_fields)

                meet_fields = [f"MeetPerson{i}{detail}" for i in range(1, 6) for detail in ["Name", "CNIC", "Contact"]]
                draw_section("Authorized to Meet", meet_fields)

                draw_section("Introducer", [
                    "IntroducerName", "Introducer CNIC", "IntroducerContact", "Introducer Address"
                ])

                draw_section("Guardian", [
                    "GuardianName", "GuardianRelation", "GuardianCNIC", "GuardianContact", "Guardianaddress"
                ])
#no need of documents section here
                #draw_section("Documents", [
                #    "DocSchoolCertificate", "DocBForm", "DocFather CNIC", "DocMotherCNIC", "DocFatherDeathCert"
                #])
                
                # For Documents display
                            # Add a new page for documents
                pdf.showPage()
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawCentredString(width / 2, height - 50, "Attached Documents")

                doc_fields = [
                    ("Father CNIC", child_data.get("DocFatherCNIC", "")),
                    ("Mother CNIC", child_data.get("DocMotherCNIC", "")),
                    ("B-Form", child_data.get("DocBForm", "")),
                    ("Father Death Certificate", child_data.get("DocFatherDeathCert", "")),
                    ("School Certificate", child_data.get("DocSchoolCertificate", "")),
                ]

                y = height - 100
                image_width = 200
                image_height = 150
                gap = 30

                for label, doc_path in doc_fields:
                    if doc_path:
                        full_path = os.path.join(base_dir, doc_path)
                        if os.path.exists(full_path) and full_path.lower().endswith((".jpg", ".jpeg", ".png")):
                            try:
                                pdf.setFont("Helvetica", 12)
                                pdf.drawString(x_margin, y, label + ":")
                                pdf.drawImage(full_path, x_margin + 150, y - image_height + 20, width=image_width, height=image_height, preserveAspectRatio=True)
                                y -= (image_height + gap)

                                if y < 150:
                                    pdf.showPage()
                                    y = height - 100
                            except Exception as e:
                                messagebox.showwarning("Document Error", f"Could not load {label}:\n{str(e)}")                                      
                
                
                
                # Block end Documents display
                pdf.save()

                # Open the generated PDF
                try:
                    open_file_cross_platform(temp_pdf)
                except Exception as e:
                    messagebox.showerror("Open Error", str(e))
                    
                


            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate PDF:\n{e}")



        #Button(preview_win, text="Generate PDF", bg="green", fg="white", font=("Arial", 12), command=generate_pdf).place(x=650, y=600)
        btn_frame = Frame(preview_win, width=160, height=50)
        btn_frame.place(x=1050, y=30)  # Adjust y to position it at top

        # Add the button inside that frame
        Button(
            btn_frame,
            text="Generate PDF",
            bg="blue",
            fg="white",
            font=("Arial", 12, "bold"),
            command=generate_pdf
        ).pack(fill="both", expand=True)




# ============================================================
# Clean Child Profile PDF (no data no display)
# ============================================================
        def generate_clean_pdf():
            try:
                temp_pdf = tempfile.mktemp(".pdf")
                pdf = pdfcanvas.Canvas(temp_pdf, pagesize=A4)
                width, height = A4
                x_margin = 50
                y = height - 50

                base_dir = os.path.dirname(os.path.abspath(__file__))
                center_name = child_data.get("CenterName", "Aghosh Homes Karachi")
                current_date = datetime.now().strftime("%d-%m-%Y")

                # ---- HEADER ----
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawString(x_margin, y, f"Center: {center_name}")

                pdf.setFont("Helvetica", 10)
                pdf.drawRightString(width - x_margin, y, f"Report Generated On: {current_date}")

                y -= 25
                pdf.setFont("Helvetica-Bold", 18)
                pdf.drawCentredString(width / 2, y, "Alkhidmat Khawateen Trust Pakistan")

                y -= 20
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawCentredString(width / 2, y, "Aghosh Program")

                y -= 20
                pdf.setFont("Helvetica-Bold", 12)
                pdf.drawCentredString(width / 2, y, "Child Profile")

                y -= 30
                pdf.setFont("Helvetica", 10)
                pdf.drawRightString(width / 2, y, f"Session: {session_name}")

                # ---- CHILD PHOTO ----
                image_path = child_data.get("PhotoPath", "")
                if image_path and os.path.exists(image_path):
                    try:
                        pdf.drawImage(image_path, x=width - 150, y=height - 180,
                                    width=100, height=100, preserveAspectRatio=True)
                    except:
                        pass
          
                        
                # ---- FIELD LABELS ----
                field_labels = {
                    "FullName": "Full Name", "FatherName": "Father's Name", "Gender": "Gender",
                    "DateOfBirth": "Date of Birth", "AdmissionDate": "Admission Date",
                    "RegistrationNumber": "Registration Number",
                    "SchoolName": "School Name", "Class": "Class", "Intelligence": "Child Category",
                    "Disability": "Disability", "HealthCondition": "Health Condition",
                    "Status": "Child Status", "ChildRequiredAmount": "Monthly Required Amount",
                    "ReasonFatherDeath": "Reason of Father's Death",
                    "FatherDeathDate": "Father's Death Date",
                    "FatherOccupation": "Father's Occupation",
                    "FatherDesignation": "Father's Designation",
                    "MotherName": "Mother's Name", "MotherStatus": "Mother's Status",
                    "MotherDeathDate": "Mother's Death Date", "MotherCNIC": "Mother's CNIC",
                    "PermanentAddress": "Permanent Address",
                    "TemporaryAddress": "Temporary Address",
                    "IntroducerName": "Introducer Name",
                    "Introducer CNIC": "Introducer CNIC",
                    "IntroducerContact": "Introducer Contact",
                    "Introducer Address": "Introducer Address",
                    "GuardianName": "Guardian Name",
                    "GuardianRelation": "Guardian Relation",
                    "GuardianCNIC": "Guardian CNIC",
                    "GuardianContact": "Guardian Contact",
                    "Guardianaddress": "Guardian Address"
                }

                # Add dynamic sibling / meet person labels
                for i in range(1, 6):
                    field_labels[f"Sibling{i}Name"] = f"Sibling {i} Name"
                    field_labels[f"Sibling{i}Gender"] = f"Sibling {i} Gender"
                    field_labels[f"Sibling{i}DOB"] = f"Sibling {i} DOB"
                    field_labels[f"MeetPerson{i}Name"] = f"Meet Person {i} Name"
                    field_labels[f"MeetPerson{i}CNIC"] = f"Meet Person {i} CNIC"
                    field_labels[f"MeetPerson{i}Contact"] = f"Meet Person {i} Contact"

                # ---- CLEAN SECTION DRAWER ----
                def draw_clean_section(title, fields):
                    nonlocal y

                    # Filter out empty fields
                    #filtered = [
                    #    f for f in fields
                    #    if str(child_data.get(f, "")).strip() not in ["", "None", "null"]
                    #]
                    
                    filtered = []

                    for f in fields:
                        value = str(child_data.get(f, "")).strip()

                        # ---- SPECIAL RULE: Mother Death Date ----
                        if f == "MotherDeathDate":
                            if child_data.get("MotherStatus", "").strip().lower() == "alive":
                                continue  # skip death date if mother is alive

                        if value not in ["", "None", "null"]:
                            filtered.append(f)


                    if not filtered:
                        return

                    pdf.setFont("Helvetica-Bold", 12)
                    pdf.setFillColorRGB(0, 0, 1)
                    pdf.drawString(x_margin, y, title)
                    y -= 20

                    pdf.setFillColorRGB(0, 0, 0)
                    pdf.setFont("Helvetica", 10)

                    for field in filtered:
                        if y < 100:
                            pdf.showPage()
                            y = height - 50

                        label = field_labels.get(field, field)
                        value = str(child_data.get(field, ""))

                        pdf.drawString(x_margin, y, f"{label}:")
                        pdf.drawString(x_margin + 150, y, value)
                        y -= 15

                    y -= 10

                # ---- SECTIONS ----
                draw_clean_section("Basic Information", [
                    "FullName", "FatherName", "Gender", "DateOfBirth", "AdmissionDate",
                    "RegistrationNumber", "SchoolName", "Class", "Intelligence",
                    "Disability", "HealthCondition", "Status", "ChildRequiredAmount"
                ])

                draw_clean_section("Father Information", [
                    "ReasonFatherDeath", "FatherDeathDate",
                    "FatherOccupation", "FatherDesignation"
                ])

                draw_clean_section("Mother Information", [
                    "MotherName", "MotherStatus", "MotherDeathDate", "MotherCNIC"
                ])

                draw_clean_section("Address", [
                    "PermanentAddress", "TemporaryAddress"
                ])

                # ---- FIXED SIBLING CLEANING ----
                sibling_fields = []
                for i in range(1, 6):
                    name = str(child_data.get(f"Sibling{i}Name", "")).strip()
                    if name:  # ONLY include sibling if name exists
                        sibling_fields.extend([
                            f"Sibling{i}Name",
                            f"Sibling{i}Gender",
                            f"Sibling{i}DOB"
                        ])

                if sibling_fields:
                    draw_clean_section("Siblings", sibling_fields)

                # ---- Meet Persons ----
                meet_fields = []
                for i in range(1, 6):
                    name = str(child_data.get(f"MeetPerson{i}Name", "")).strip()
                    if name:
                        meet_fields.extend([
                            f"MeetPerson{i}Name",
                            f"MeetPerson{i}CNIC",
                            f"MeetPerson{i}Contact"
                        ])

                if meet_fields:
                    draw_clean_section("Authorized to Meet", meet_fields)

                # ---- Introducer ----
                draw_clean_section("Introducer", [
                    "IntroducerName", "Introducer CNIC",
                    "IntroducerContact", "Introducer Address"
                ])

                # ---- Guardian ----
                draw_clean_section("Guardian", [
                    "GuardianName", "GuardianRelation",
                    "GuardianCNIC", "GuardianContact",
                    "Guardianaddress"
                ])

                # ---- Documents Page ----
                pdf.showPage()
                pdf.setFont("Helvetica-Bold", 14)
                pdf.drawCentredString(width / 2, height - 50, "Attached Documents")

                doc_fields = [
                    ("Father CNIC", child_data.get("DocFatherCNIC", "")),
                    ("Mother CNIC", child_data.get("DocMotherCNIC", "")),
                    ("B-Form", child_data.get("DocBForm", "")),
                    ("Father Death Certificate", child_data.get("DocFatherDeathCert", "")),
                    ("School Certificate", child_data.get("DocSchoolCertificate", "")),
                    ("Other Documemt 1", child_data.get("OtherDoc1", "")),
                    ("Other Documemt 2", child_data.get("OtherDoc2", "")),
                    ("Other Documemt 3", child_data.get("OtherDoc3", "")),
                ]

                y = height - 100
                image_width = 200
                image_height = 150
                gap = 30

                for label, path in doc_fields:
                    if path:
                        full_path = os.path.join(base_dir, path)
                        if os.path.exists(full_path):
                            try:
                                pdf.setFont("Helvetica", 12)
                                pdf.drawString(x_margin, y, label + ":")
                                pdf.drawImage(full_path,
                                            x_margin + 150,
                                            y - image_height + 20,
                                            width=image_width,
                                            height=image_height,
                                            preserveAspectRatio=True)
                                y -= (image_height + gap)

                                if y < 150:
                                    pdf.showPage()
                                    y = height - 100
                            except:
                                pass

                pdf.save()

                open_file_cross_platform(temp_pdf)

            except Exception as e:
                messagebox.showerror("Error", f"Failed to generate PDF:\n{e}")




        Button(
            btn_frame,
            text="Generate Clean PDF",
            bg="darkgreen",
            fg="white",
            font=("Arial", 12, "bold"),
            command=generate_clean_pdf
        ).pack(fill="both", expand=True, pady=(55, 0))
# ============================================================



    # === Generate Report Button ===
    def generate_report():
        selected = tree.selection()
        if not selected:
            messagebox.showwarning("No Selection", "Please select a child from the list.")
            return
        child_name = tree.item(selected[0], 'values')[0]
        child_id, father_name = child_lookup.get(child_name, (None, None))
        if child_id:
            create_child_profile_pdf(child_id, child_name, father_name)

    btn_frame = tk.Frame(report_win)
    btn_frame.pack(pady=10)
    tk.Button(btn_frame, text="Generate Profile Report", command=generate_report, width=30).pack()




# ============================================================
# Child Profile for Donor Report
# ============================================================

def open_child_profile_for_donor_report():
 
    win = tk.Toplevel(root)
    win.title("Child Profile for Donors" + ORG_SUFFIX)
    win.geometry("900x550")
    win.transient(root)
    win.grab_set()
 
    win.update_idletasks()
    width  = 900
    height = 650
    x = (win.winfo_screenwidth()  // 2) - (width  // 2)
    y = (win.winfo_screenheight() // 2) - (height // 2)
    win.geometry(f"{width}x{height}+{x}+{y}")
 
    # ── SEARCH BAR ────────────────────────────────────────────────────────────
    search_var    = tk.StringVar()
    donor_name_var = tk.StringVar()
 
    top_frame = tk.Frame(win)
    top_frame.pack(fill="x", padx=10, pady=5)
 
    tk.Label(top_frame, text="Search Child:", font=("Arial", 11)).pack(side="left")
    search_entry = tk.Entry(top_frame, textvariable=search_var, width=30)
    search_entry.pack(side="left", padx=5)
    search_entry.focus_set()
 
    tk.Label(top_frame, text="Donor Name:", font=("Arial", 11)).pack(side="left", padx=(20, 5))
    donor_entry = tk.Entry(top_frame, textvariable=donor_name_var, width=30)
    donor_entry.pack(side="left")
 
    # ── TREEVIEW ──────────────────────────────────────────────────────────────
    cols = ("RegNo", "ChildName", "FatherName")
    tree = ttk.Treeview(win, columns=cols, show="headings", height=18)
    for col in cols:
        tree.heading(col, text=col)
        tree.column(col, width=250)
    tree.pack(fill="both", expand=True, padx=10, pady=10)
 
    # ── LOAD CHILDREN ─────────────────────────────────────────────────────────
    def load_children(filter_text=""):
        tree.delete(*tree.get_children())
        conn = get_connection()
        if not conn:
            messagebox.showerror("Database Error", "Database connection failed.")
            return
        cursor = conn.cursor()
        cursor.execute("""
            SELECT CenterID FROM tblCenters
            WHERE OperationalStatus = 'Active' LIMIT 1
        """)
        center_row = cursor.fetchone()
        if not center_row:
            conn.close()
            messagebox.showerror("Error", "No Active Center found.")
            return
        active_center_id = center_row.CenterID
        query  = """
            SELECT RegistrationNumber, FullName, FatherName
            FROM tblChildren WHERE CenterID = ?
        """
        params = [active_center_id]
        if filter_text:
            query += " AND FullName LIKE ?"
            params.append(f"%{filter_text}%")
        query += " ORDER BY FullName"
        cursor.execute(query, params)
        for row in cursor.fetchall():
            tree.insert("", "end", values=(
                row.RegistrationNumber, row.FullName, row.FatherName))
        conn.close()
 
    load_children()
    search_entry.bind("<KeyRelease>", lambda e: load_children(search_var.get()))
 
    # ── PDF GENERATION ────────────────────────────────────────────────────────
    def generate_pdf():
        if not tree.selection():
            messagebox.showwarning("Required", "Please select a child.")
            return
 
        donor_name = donor_name_var.get().strip()
        reg_no     = tree.item(tree.selection()[0])["values"][0]
 
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM tblChildren WHERE RegistrationNumber = ?",
                       (reg_no,))
        row = cursor.fetchone()
        if not row:
            messagebox.showerror("Error", "Child record not found.")
            conn.close()
            return
        columns = [col[0] for col in cursor.description]
        child   = dict(zip(columns, row))
 
        cursor.execute("""
            SELECT CenterID, CenterName, Region, City FROM tblCenters
            WHERE OperationalStatus = 'Active' LIMIT 1
        """)
        center_row = cursor.fetchone()
        if not center_row:
            messagebox.showerror("Error", "No Active Center found.")
            conn.close()
            return
        center_cols = [col[0] for col in cursor.description]
        center      = dict(zip(center_cols, center_row))
        conn.close()
 
        def g(key, default="-"):
            val = child.get(key)
            return str(val).strip() if val not in (None, "", " ") else default
 
        def fmt(d):
            if not d:
                return "-"
            try:
                if isinstance(d, str):
                    from datetime import datetime as _dt
                    d = _dt.strptime(d[:10], "%Y-%m-%d")
                return d.strftime("%d-%b-%Y")
            except Exception:
                return "-"
 
        guardian_name     = g("GuardianName")     if child.get("GuardianName")     else g("MotherName")
        guardian_relation = g("GuardianRelation") if child.get("GuardianRelation") else "Mother"
        guardian_contact  = g("GuardianContact")
 
        if child.get("MotherStatus", "").strip().lower() == "alive":
            mother_alive_dead  = "Alive"
            mother_death_cause = "-"
        else:
            mother_alive_dead  = f"Dead ({fmt(child.get('MotherDeathDate'))})"
            mother_death_cause = "-"
 
        from reportlab.lib.pagesizes      import A4
        from reportlab.pdfgen             import canvas
        from reportlab.lib.colors         import HexColor
        from reportlab.lib.units          import inch
        from reportlab.pdfbase.pdfmetrics import stringWidth
        import os, win32api
 
        downloads = os.path.join(os.path.expanduser("~"), "Downloads")
        os.makedirs(downloads, exist_ok=True)
        pdf_path = os.path.join(downloads, f"{reg_no} - {g('FullName')}.pdf")
 
        c = canvas.Canvas(pdf_path, pagesize=A4)
        w, h = A4
 
        BG_COLOR  = HexColor("#eff4a8")
        GREEN     = HexColor("#88c468")
        BORDER    = HexColor("#000000")
        TEXT      = HexColor("#000000")
        ROW_H     = 32
        LM        = 30
        RM        = 30
        TW        = w - LM - RM
        FONT_SIZE = 10
 
        def mid(y, row_h=ROW_H):
            return y + row_h * 0.45
 
        def wrap_text(text, max_width, font_name, font_size):
            text  = str(text)
            if stringWidth(text, font_name, font_size) <= max_width:
                return [text]
            words, lines, current_line = text.split(), [], ""
            for word in words:
                test = (current_line + " " + word).strip()
                if stringWidth(test, font_name, font_size) <= max_width:
                    current_line = test
                else:
                    if current_line:
                        lines.append(current_line)
                    current_line = word
            if current_line:
                lines.append(current_line)
            return lines[:2]
 
        # ── Header ───────────────────────────────────────────────────────────
        IMG_W   = 1 * inch
        IMG_H   = 1 * inch
        y_top   = h - 0.5 * inch
        img_y   = y_top - IMG_H
        logo_x  = LM
        photo_x = w - RM - IMG_W
 
        text_left     = logo_x  + IMG_W + 10
        text_right    = photo_x - 10
        text_center_x = (text_left + text_right) / 2
        line_gap      = 16
        text_y3 = img_y
        text_y2 = img_y + line_gap
        text_y1 = img_y + 2 * line_gap
 
        logo_path = resource_path("logo/akh_logo.png")
        if os.path.exists(logo_path):
            try:
                c.drawImage(logo_path, logo_x, img_y, width=IMG_W, height=IMG_H)
            except Exception:
                pass
 
        # ── FIX: use _resolve_photo_path instead of resource_path ────────────
        photo_path = _resolve_photo_path(child.get("PhotoPath"))
        if photo_path:
            try:
                c.drawImage(photo_path, photo_x, img_y,
                            width=IMG_W, height=IMG_H,
                            preserveAspectRatio=True, mask="auto")
            except Exception:
                pass
 
        c.setFont("Helvetica-Bold", 14)
        c.setFillColor(TEXT)
        c.drawCentredString(text_center_x, text_y1, "ALKHIDMAT KHAWATEEN TRUST PAKISTAN")
        c.drawCentredString(text_center_x, text_y2, "ORPHAN CARE PROGRAM")
        c.drawCentredString(text_center_x, text_y3, "AGHOSH ORPHAN PROFILE")
 
        y = img_y - 35
 
        COL_WIDTHS = [TW * 0.16, TW * 0.34, TW * 0.125, TW * 0.375]
 
        def section_header(y, label):
            c.setFillColor(GREEN)
            c.setStrokeColor(BORDER)
            c.setLineWidth(0.5)
            c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
            c.setFillColor(TEXT)
            c.setFont("Helvetica-Bold", 12)
            c.drawString(LM + 5, mid(y), label)
            return y - ROW_H
 
        def empty_row(y):
            return y - (ROW_H * 0.3)
 
        def draw_text_lines(x, y, lines, font_name, is_bold=False):
            font = f"{font_name}-Bold" if is_bold else font_name
            c.setFont(font, FONT_SIZE)
            if len(lines) == 1:
                c.drawString(x, y + ROW_H * 0.45, lines[0])
            else:
                c.drawString(x, y + ROW_H * 0.65, lines[0])
                c.drawString(x, y + ROW_H * 0.25, lines[1])
 
        def data_row_6col(y, field1, value1, field2, value2, field3, value3):
            _R = 30 + 50 + 30 + 50 + 30 + 50
            cw = [TW*38.4/_R, TW*81.6/_R, TW*30/_R, TW*30/_R, TW*30./_R, TW*30/_R]
            c.setFillColor(BG_COLOR); c.setStrokeColor(BORDER); c.setLineWidth(0.5)
            c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
            x_pos = LM
            for i in range(5):
                x_pos += cw[i]
                c.line(x_pos, y, x_pos, y + ROW_H)
            c.setFillColor(TEXT)
            offsets = [0] * 6
            offsets[0] = LM
            for i in range(1, 6):
                offsets[i] = offsets[i-1] + cw[i-1]
            for i in range(3):
                fi, vi = i*2, i*2+1
                draw_text_lines(offsets[fi]+3, y,
                                wrap_text([field1,field2,field3][i], cw[fi]-6, "Helvetica", FONT_SIZE),
                                "Helvetica", is_bold=True)
                draw_text_lines(offsets[vi]+3, y,
                                wrap_text([value1,value2,value3][i], cw[vi]-6, "Helvetica", FONT_SIZE),
                                "Helvetica", is_bold=False)
            return y - ROW_H
 
        def data_row_4col(y, field1, value1, field2, value2):
            c.setFillColor(BG_COLOR); c.setStrokeColor(BORDER); c.setLineWidth(0.5)
            c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
            x_pos = LM
            for i in range(3):
                x_pos += COL_WIDTHS[i]
                c.line(x_pos, y, x_pos, y + ROW_H)
            c.setFillColor(TEXT)
            pairs = [(field1, COL_WIDTHS[0], LM, True),
                     (value1, COL_WIDTHS[1], LM+COL_WIDTHS[0], False),
                     (field2, COL_WIDTHS[2], LM+COL_WIDTHS[0]+COL_WIDTHS[1], True),
                     (value2, COL_WIDTHS[3], LM+COL_WIDTHS[0]+COL_WIDTHS[1]+COL_WIDTHS[2], False)]
            for text, cw, x0, bold in pairs:
                lines = wrap_text(text, cw - 6, "Helvetica", FONT_SIZE)
                draw_text_lines(x0 + 3, y, lines, "Helvetica", is_bold=bold)
            return y - ROW_H
 
        # Center info
        col_w = TW / 3
        c.setFillColor("#88c468"); c.setStrokeColor(BORDER); c.setLineWidth(0.5)
        c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
        c.line(LM+col_w, y, LM+col_w, y+ROW_H)
        c.line(LM+col_w*2, y, LM+col_w*2, y+ROW_H)
        c.setFont("Helvetica-Bold", FONT_SIZE); c.setFillColor(TEXT)
        c.drawCentredString(LM+col_w*0.5, mid(y), "Aghosh")
        c.drawCentredString(LM+col_w*1.5, mid(y), "Region")
        c.drawCentredString(LM+col_w*2.5, mid(y), "City/Town")
        y -= ROW_H
        c.setFillColor(BG_COLOR); c.setStrokeColor(BORDER)
        c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
        c.line(LM+col_w, y, LM+col_w, y+ROW_H)
        c.line(LM+col_w*2, y, LM+col_w*2, y+ROW_H)
        c.setFont("Helvetica", FONT_SIZE); c.setFillColor(TEXT)
        c.drawCentredString(LM+col_w*0.5, mid(y), center.get("CenterName") or "-")
        c.drawCentredString(LM+col_w*1.5, mid(y), center.get("Region")     or "-")
        c.drawCentredString(LM+col_w*2.5, mid(y), center.get("City")       or "-")
        y -= ROW_H
 
        y = empty_row(y)
        y = section_header(y, "Donor Name")
        c.setFillColor(BG_COLOR); c.setStrokeColor(BORDER)
        c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
        c.setFont("Helvetica", FONT_SIZE); c.setFillColor(TEXT)
        donor_display = f"Respected {donor_name}" if donor_name else "Respected ______________________"
        c.drawCentredString(w/2, mid(y), donor_display)
        y -= ROW_H
 
        y = empty_row(y)
        y = section_header(y, "Orphan Details")
        y = data_row_4col(y, "Student Name",  g("FullName"),            "Date Of Birth",  fmt(child.get("DateOfBirth")))
        y = data_row_4col(y, "ID",            g("RegistrationNumber"),  "Admission Date", fmt(child.get("AdmissionDate")))
        y = data_row_4col(y, "Class",         g("Class"),              "School Name",    g("SchoolName"))
        y = data_row_6col(y, "Father's Name", g("FatherName"),
                             "Date Of Death", fmt(child.get("FatherDeathDate")),
                             "Cause Of Death", g("ReasonFatherDeath"))
        y = data_row_6col(y, "Mother's Name", g("MotherName"),
                             "Alive/Dead (Date)", mother_alive_dead,
                             "Cause Of Death", mother_death_cause)
        y = data_row_4col(y, "Address", g("PermanentAddress"), "Contact No.", guardian_contact)
 
        y = empty_row(y)
        y = section_header(y, "Guardian Information")
        y = data_row_4col(y, "Guardian Name",  guardian_name,            "Relation With Child", guardian_relation)
        y = data_row_4col(y, "Qualification",  "-",                      "Profession",          "-")
        y = data_row_4col(y, "Monthly Income", "-",                      "Financial Status",    "Mustahiq")
        y = data_row_4col(y, "Address",        g("PermanentAddress"),    "Contact No.",         guardian_contact)
 
        y = empty_row(y)
        y = section_header(y, "Sponsorship Duration")
        c.setFillColor(BG_COLOR); c.setStrokeColor(BORDER)
        c.rect(LM, y, TW, ROW_H, fill=1, stroke=1)
        c.line(LM+TW/2, y, LM+TW/2, y+ROW_H)
        c.setFont("Helvetica", FONT_SIZE); c.setFillColor(TEXT)
        c.drawString(LM+3,       mid(y), "Start Month _________________ Year _______________")
        c.drawString(LM+TW/2+3, mid(y), "Expiry Month _________________ Year ______________")
 
        footer_y = 50
        c.setFont("Helvetica-Bold", 12); c.setFillColor(TEXT)
        c.drawCentredString(w/2, footer_y+40, "Head Office: Alkhidmat 76-E, PECHS Block 6, Karachi, Pakistan.")
        c.drawCentredString(w/2, footer_y+20, "Phone: 0800 77 77 8")
        c.drawCentredString(w/2, footer_y,    "Email: info@alkhidmatkhawateen.org, www.alkhidmatkhawateen.org")
 
        c.save()
        win32api.ShellExecute(0, "open", pdf_path, None, ".", 1)
 
    # ── DOCX GENERATION ───────────────────────────────────────────────────────
    def generate_docx():
        if not tree.selection():
            messagebox.showwarning("Required", "Please select a child.")
            return
 
        donor_name = donor_name_var.get().strip()
        reg_no     = tree.item(tree.selection()[0])["values"][0]
 
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM tblChildren WHERE RegistrationNumber = ?",
                       (reg_no,))
        row = cursor.fetchone()
        if not row:
            messagebox.showerror("Error", "Child record not found.")
            conn.close()
            return
        columns = [col[0] for col in cursor.description]
        child   = dict(zip(columns, row))
 
        cursor.execute("""
            SELECT CenterID, CenterName, Region, City FROM tblCenters
            WHERE OperationalStatus = 'Active' LIMIT 1
        """)
        center_row = cursor.fetchone()
        if not center_row:
            messagebox.showerror("Error", "No Active Center found.")
            conn.close()
            return
        center_cols = [col[0] for col in cursor.description]
        center      = dict(zip(center_cols, center_row))
        conn.close()
 
        from docx              import Document
        from docx.shared       import Pt, Inches, RGBColor, Twips
        from docx.enum.text    import WD_ALIGN_PARAGRAPH
        from docx.enum.table   import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.oxml.ns      import qn
        from docx.oxml         import OxmlElement
        from datetime          import datetime as _dt
        import os, win32api
 
        def g(key, default="-"):
            val = child.get(key)
            return str(val).strip() if val not in (None, "", " ") else default
 
        def fmt(d):
            if not d:
                return "-"
            try:
                if isinstance(d, str):
                    d = _dt.strptime(d[:10], "%Y-%m-%d")
                return d.strftime("%d-%b-%Y")
            except Exception:
                return "-"
 
        guardian_name     = g("GuardianName")     if child.get("GuardianName")     else g("MotherName")
        guardian_relation = g("GuardianRelation") if child.get("GuardianRelation") else "Mother"
        guardian_contact  = g("GuardianContact")
 
        if child.get("MotherStatus", "").strip().lower() == "alive":
            mother_alive_dead  = "Alive"
            mother_death_cause = "-"
        else:
            mother_alive_dead  = f"Dead ({fmt(child.get('MotherDeathDate'))})"
            mother_death_cause = "-"
 
        COLOR_BG    = "EFF4A8"
        COLOR_GREEN = "88C468"
        COLOR_WHITE = "FFFFFF"
        COLOR_BLACK = "000000"
 
        PAGE_W   = 11906
        MARGIN   = 720
        TW       = PAGE_W - 2 * MARGIN
        ROW_H_PT = 22
        FONT_SZ  = 10
 
        C4 = [int(TW*0.16), int(TW*0.34), int(TW*0.125),
              TW - int(TW*0.16) - int(TW*0.34) - int(TW*0.125)]
        C6 = [int(TW*0.16), int(TW*0.34), int(TW*0.125), int(TW*0.125),
              int(TW*0.125),
              TW - int(TW*0.16) - int(TW*0.34) - 3*int(TW*0.125)]
        C3 = [TW//3, TW//3, TW - 2*(TW//3)]
 
        def set_shading(cell, fill_hex):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill_hex)
            for old in tcPr.findall(qn("w:shd")):
                tcPr.remove(old)
            tcPr.append(shd)
 
        def set_cell_borders(cell, color="000000", sz=4):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tbl_brd = OxmlElement("w:tcBorders")
            for side in ["top","bottom","left","right"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "single"); el.set(qn("w:sz"),    str(sz))
                el.set(qn("w:space"), "0");      el.set(qn("w:color"), color)
                tbl_brd.append(el)
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcPr.append(tbl_brd)
 
        def no_borders(cell):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tbl_brd = OxmlElement("w:tcBorders")
            for side in ["top","bottom","left","right"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "nil")
                tbl_brd.append(el)
            for old in tcPr.findall(qn("w:tcBorders")):
                tcPr.remove(old)
            tcPr.append(tbl_brd)
 
        def set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            mar = OxmlElement("w:tcMar")
            for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:w"), str(val)); el.set(qn("w:type"), "dxa")
                mar.append(el)
            for old in tcPr.findall(qn("w:tcMar")):
                tcPr.remove(old)
            tcPr.append(mar)
 
        def set_cell_valign(cell, align="center"):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            va = OxmlElement("w:vAlign")
            va.set(qn("w:val"), align)
            for old in tcPr.findall(qn("w:vAlign")):
                tcPr.remove(old)
            tcPr.append(va)
 
        def set_row_height(row, pt=ROW_H_PT):
            tr = row._tr; trPr = tr.get_or_add_trPr()
            trH = OxmlElement("w:trHeight")
            trH.set(qn("w:val"),   str(int(pt * 20)))
            trH.set(qn("w:hRule"), "atLeast")
            for old in trPr.findall(qn("w:trHeight")):
                trPr.remove(old)
            trPr.append(trH)
 
        def set_cell_width(cell, dxa):
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(dxa)); tcW.set(qn("w:type"), "dxa")
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcPr.append(tcW)
 
        def cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT,
                      size=FONT_SZ, color=COLOR_BLACK):
            cell.paragraphs[0].clear()
            p = cell.paragraphs[0]; p.alignment = align
            run = p.add_run(str(text))
            run.bold = bold; run.font.size = Pt(size)
            run.font.color.rgb = RGBColor.from_string(color)
            run.font.name = "Calibri"
 
        def add_spacer(doc, space_pt=3):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.line_spacing = Pt(space_pt)
 
        def _init_tbl(tbl):
            tblPr = tbl._tbl.find(qn("w:tblPr"))
            tblW  = OxmlElement("w:tblW")
            tblW.set(qn("w:w"), str(TW)); tblW.set(qn("w:type"), "dxa")
            tblPr.append(tblW)
            tblBrd = OxmlElement("w:tblBorders")
            for side in ["top","left","bottom","right","insideH","insideV"]:
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"), "none"); el.set(qn("w:sz"), "0")
                el.set(qn("w:space"), "0"); el.set(qn("w:color"), "auto")
                tblBrd.append(el)
            tblPr.append(tblBrd)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
 
        doc  = Document()
        sect = doc.sections[0]
        sect.page_width    = Twips(PAGE_W)
        sect.page_height   = Twips(16838)
        sect.top_margin    = Inches(0.5)
        sect.bottom_margin = Inches(0.5)
        sect.left_margin   = Inches(0.5)
        sect.right_margin  = Inches(0.5)
 
        style = doc.styles["Normal"]
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after  = Pt(0)
        style.font.name = "Calibri"
        style.font.size = Pt(FONT_SZ)
 
        # Header table
        LOGO_W  = 1260; PHOTO_W = 1260
        TITLE_W = TW - LOGO_W - PHOTO_W
 
        hdr_tbl = doc.add_table(rows=1, cols=3)
        _init_tbl(hdr_tbl)
        hdr_row = hdr_tbl.rows[0]
        set_row_height(hdr_row, 100)
 
        logo_cell  = hdr_row.cells[0]
        title_cell = hdr_row.cells[1]
        photo_cell = hdr_row.cells[2]
 
        for cell in (logo_cell, title_cell, photo_cell):
            no_borders(cell)
            set_cell_valign(cell, "center")
            set_cell_margins(cell, top=0, bottom=0, left=0, right=0)
 
        set_cell_width(logo_cell,  LOGO_W)
        set_cell_width(title_cell, TITLE_W)
        set_cell_width(photo_cell, PHOTO_W)
 
        logo_path = resource_path("logo/akh_logo.png")
        if os.path.exists(logo_path):
            logo_cell.paragraphs[0].clear()
            logo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = logo_cell.paragraphs[0].add_run()
            run.add_picture(logo_path, width=Pt(85))
 
        title_cell.paragraphs[0].clear()
        for line in ["ALKHIDMAT KHAWATEEN TRUST PAKISTAN",
                     "ORPHAN CARE PROGRAM",
                     "AGHOSH ORPHAN PROFILE"]:
            p = (title_cell.paragraphs[0]
                 if line == "ALKHIDMAT KHAWATEEN TRUST PAKISTAN"
                 else title_cell.add_paragraph())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.bold = True; run.font.size = Pt(14)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
 
        # ── FIX: use _resolve_photo_path instead of resource_path ────────────
        photo_path = _resolve_photo_path(child.get("PhotoPath"))
        if photo_path:
            try:
                photo_cell.paragraphs[0].clear()
                photo_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = photo_cell.paragraphs[0].add_run()
                run.add_picture(photo_path, height=Pt(85))
            except Exception:
                pass
 
        add_spacer(doc, 4)
 
        def std_table(col_widths, rows):
            n_cols = len(col_widths)
            tbl = doc.add_table(rows=len(rows), cols=n_cols)
            _init_tbl(tbl)
            for r_idx, cells in enumerate(rows):
                row = tbl.rows[r_idx]
                set_row_height(row, ROW_H_PT)
                for c_idx, (txt, bold, align, bg) in enumerate(cells):
                    cell = row.cells[c_idx]
                    set_shading(cell, bg)
                    set_cell_borders(cell)
                    set_cell_valign(cell, "center")
                    set_cell_margins(cell, top=40, bottom=40, left=80, right=40)
                    set_cell_width(cell, col_widths[c_idx])
                    cell_text(cell, txt, bold=bold, align=align)
            return tbl
 
        def section_header(label):
            tbl = doc.add_table(rows=1, cols=1)
            _init_tbl(tbl)
            row  = tbl.rows[0]
            set_row_height(row, ROW_H_PT)
            cell = row.cells[0]
            set_shading(cell, COLOR_GREEN)
            set_cell_borders(cell)
            set_cell_valign(cell, "center")
            set_cell_margins(cell, top=40, bottom=40, left=80, right=40)
            set_cell_width(cell, TW)
            cell_text(cell, label, bold=True, size=11,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
 
        L = WD_ALIGN_PARAGRAPH.LEFT
        C = WD_ALIGN_PARAGRAPH.CENTER
 
        def row4(f1, v1, f2, v2):
            return [(f1,True,L,COLOR_BG),(v1,False,L,COLOR_BG),
                    (f2,True,L,COLOR_BG),(v2,False,L,COLOR_BG)]
 
        def row6(f1,v1,f2,v2,f3,v3):
            return [(f1,True,L,COLOR_BG),(v1,False,L,COLOR_BG),
                    (f2,True,L,COLOR_BG),(v2,False,L,COLOR_BG),
                    (f3,True,L,COLOR_BG),(v3,False,L,COLOR_BG)]
 
        # Center info
        ci_rows = [
            [(t,True,C,COLOR_GREEN) for t in ["Aghosh","Region","City/Town"]],
            [(center.get("CenterName") or "-", False, C, COLOR_BG),
             (center.get("Region")     or "-", False, C, COLOR_BG),
             (center.get("City")       or "-", False, C, COLOR_BG)],
        ]
        std_table(C3, ci_rows)
        add_spacer(doc, 4)
 
        section_header("Donor Name")
        donor_display = f"Respected {donor_name}" if donor_name else "Respected ______________________"
        std_table([TW], [[(donor_display, False, C, COLOR_BG)]])
        add_spacer(doc, 4)
 
        section_header("Orphan Details")
        std_table(C4, [
            row4("Student Name",  g("FullName"),           "Date Of Birth",  fmt(child.get("DateOfBirth"))),
            row4("ID",            g("RegistrationNumber"), "Admission Date", fmt(child.get("AdmissionDate"))),
            row4("Class",         g("Class"),              "School Name",    g("SchoolName")),
        ])
        std_table(C6, [
            row6("Father's Name", g("FatherName"),
                 "Date Of Death", fmt(child.get("FatherDeathDate")),
                 "Cause Of Death", g("ReasonFatherDeath")),
            row6("Mother's Name", g("MotherName"),
                 "Alive/Dead (Date)", mother_alive_dead,
                 "Cause Of Death", mother_death_cause),
        ])
        std_table(C4, [row4("Address", g("PermanentAddress"), "Contact No.", guardian_contact)])
        add_spacer(doc, 4)
 
        section_header("Guardian Information")
        std_table(C4, [
            row4("Guardian Name",  guardian_name,            "Relation With Child", guardian_relation),
            row4("Qualification",  "-",                      "Profession",          "-"),
            row4("Monthly Income", "-",                      "Financial Status",    "Mustahiq"),
            row4("Address",        g("PermanentAddress"),    "Contact No.",         guardian_contact),
        ])
        add_spacer(doc, 4)
 
        section_header("Sponsorship Duration")
        std_table([TW//2, TW - TW//2], [[
            ("Start Month _________________ Year _______________",  False, L, COLOR_BG),
            ("Expiry Month _________________ Year ______________", False, L, COLOR_BG),
        ]])
        add_spacer(doc, 6)
 
        for line in [
            "Head Office: Alkhidmat 76-E, PECHS Block 6, Karachi, Pakistan.",
            "Phone: 0800 77 77 8",
            "Email: info@alkhidmatkhawateen.org, www.alkhidmatkhawateen.org",
        ]:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            run = p.add_run(line)
            run.bold = True; run.font.size = Pt(12)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0, 0, 0)
 
        downloads = os.path.join(os.path.expanduser("~"), "Downloads")
        os.makedirs(downloads, exist_ok=True)
        docx_path = os.path.join(downloads, f"{reg_no} - {g('FullName')}.docx")
        doc.save(docx_path)
        win32api.ShellExecute(0, "open", docx_path, None, ".", 1)
 
    # ── PRINT ALL CHILDREN ────────────────────────────────────────────────────
    def print_all_children():
        if not messagebox.askyesno(
                "Confirm Print",
                "This will print profiles of ALL children to the default printer.\n\n"
                "Do you want to continue?"):
            return
 
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM tblChildren ORDER BY FullName")
        children = cursor.fetchall()
        if not children:
            messagebox.showwarning("No Data", "No children found.")
            return
        cursor.execute("""
            SELECT CenterName, Region, City FROM tblCenters
            WHERE OperationalStatus = 'Active' LIMIT 1
        """)
        center = cursor.fetchone()
        conn.close()
        if not center:
            messagebox.showerror("Error", "No Active Center found.")
            return
 
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen        import canvas as rl_canvas
        from reportlab.lib.colors    import HexColor
        from tempfile                import mkdtemp
        import shutil, time
 
        temp_dir = mkdtemp(prefix="aghosh_print_")
        try:
            for child in children:
                reg_no   = child.RegistrationNumber
                pdf_path = os.path.join(temp_dir, f"Child_Profile_{reg_no}.pdf")
 
                c = rl_canvas.Canvas(pdf_path, pagesize=A4)
                w, h = A4
                green = HexColor("#8DBB2F")
 
                logo_path = resource_path("logo/akh_logo.png")
                if os.path.exists(logo_path):
                    c.drawImage(logo_path, 40, h-90, width=80, height=60)
 
                # ── FIX: use _resolve_photo_path instead of resource_path ────
                photo_path = _resolve_photo_path(child.PhotoPath)
                if photo_path:
                    try:
                        c.drawImage(photo_path, w-120, h-95, width=65, height=80)
                    except Exception:
                        pass
 
                c.setFont("Helvetica-Bold", 12)
                c.drawCentredString(w/2, h-50,  "ALKHIDMAT KHAWATEEN TRUST PAKISTAN")
                c.drawCentredString(w/2, h-65,  "ORPHAN CARE PROGRAM")
                c.drawCentredString(w/2, h-80,  "AGHOSH ORPHAN PROFILE")
 
                c.setFillColor(green)
                c.rect(40, h-120, w-80, 20, fill=1)
                c.setFillColor("black"); c.setFont("Helvetica", 10)
                c.drawString(50,    h-115, f"Aghosh: {center.CenterName}")
                c.drawCentredString(w/2, h-115, f"Region: {center.Region}")
                c.drawRightString(w-50, h-115,  f"City: {center.City}")
 
                y = h - 160
 
                def fmt(d):
                    return format_db_date(d, "%d/%m/%Y") if d else ""
 
                def row(label, value):
                    nonlocal y
                    c.rect(40, y, 200, 20)
                    c.rect(240, y, w-280, 20)
                    c.drawString(45,  y+6, label)
                    c.drawString(245, y+6, str(value))
                    y -= 20
 
                guardian_name     = child.GuardianName if child.GuardianName else child.MotherName
                guardian_relation = child.GuardianRelation if child.GuardianName else "Mother"
 
                c.setFillColor(green)
                c.rect(40, y, w-80, 20, fill=1)
                c.setFillColor("black")
                c.drawString(45, y+6, "Orphan Details")
                y -= 20
 
                row("Student Name",          child.FullName)
                row("ID",                    child.RegistrationNumber)
                row("Date of Birth",         fmt(child.DateOfBirth))
                row("Admission Date",        fmt(child.AdmissionDate))
                row("Class",                 child.Class)
                row("School Name",           child.SchoolName)
                row("Father Name",           child.FatherName)
 
                mother_status = ("Alive" if child.MotherStatus == "Alive"
                                 else f"Deceased ({fmt(child.MotherDeathDate)})")
                row("Mother Name", f"{child.MotherName} — {mother_status}")
                row("Cause of Father Death", child.ReasonFatherDeath or "")
                row("Address",               child.PermanentAddress)
 
                c.setFillColor(green)
                c.rect(40, y, w-80, 20, fill=1)
                c.setFillColor("black")
                c.drawString(45, y+6, "Guardian Information")
                y -= 20
 
                row("Guardian Name",     guardian_name)
                row("Relation",          guardian_relation)
                row("Financial Status",  "Poor")
                row("Monthly Income",    "Very Low")
                row("Required Monthly",  child.ChildRequiredAmount)
 
                c.setFillColor(green)
                c.rect(40, y, w-80, 20, fill=1)
                c.setFillColor("black")
                c.drawString(45, y+6, "Sponsorship Duration")
                y -= 20
 
                row("Start Month / Year",  "________________________")
                row("Expiry Month / Year", "________________________")
 
                y -= 30
                c.setFont("Helvetica", 9)
                c.drawCentredString(w/2, y,    "Head Office: 76-E Block 6 PECHS, Karachi, Sindh, Pakistan")
                c.drawCentredString(w/2, y-15, "Phone: 0800-77778 | 021-34304985-88")
                c.drawCentredString(w/2, y-30, "Email: info@alkhidmatkhawateen.org | www.alkhidmatkhawateen.org")
 
                c.save()
                os.startfile(pdf_path, "print")
                time.sleep(1.5)
        finally:
            time.sleep(5)
            shutil.rmtree(temp_dir, ignore_errors=True)
 
    # ── BUTTONS ───────────────────────────────────────────────────────────────
    btn_frame = tk.Frame(win)
    btn_frame.pack(pady=10)
 
    tk.Button(btn_frame, text="Generate PDF",
              font=("Arial",11,"bold"), bg="#2E8B57", fg="white", width=18,
              command=generate_pdf).pack(side="left", padx=10)
 
    tk.Button(btn_frame, text="Generate Docx",
              font=("Arial",11,"bold"), bg="#2E8B57", fg="white", width=18,
              command=generate_docx).pack(side="left", padx=10)
 
    tk.Button(btn_frame, text="Print All Profiles",
              font=("Arial",11,"bold"), bg="#2E8B47", fg="white", width=18,
              command=print_all_children).pack(side="left", padx=10)


    #tk.Button(win, text="Generate PDF", command=generate_pdf, font=("Arial", 12), bg="#2E8B57", fg="white", width=15).grid(row=0, column=0, padx=5)
    #tk.Button(win, text="Print All Profiles", command=print_all_children, font=("Arial", 12), bg="#2E8B57", fg="white", width=10).grid(row=0, column=1, padx=5)    
    
    
# ============================================================

# ============================================================
# Donations Management
# ============================================================

def open_donation_mgmt():
    donation_win = tk.Toplevel(root)
    donation_win.title("Donations Management" + ORG_SUFFIX)
    #donation_win.geometry("950x600")
 
    # Get screen dimension
    screen_width = donation_win.winfo_screenwidth()
    screen_height = donation_win.winfo_screenheight()
    # Window size
    win_width = 950
    win_height = 600
    # Calculate position
    x = (screen_width // 2) - (win_width // 2)
    y = (screen_height // 2) - (win_height // 2)

    # Set geometry to center
    donation_win.geometry(f"{win_width}x{win_height}+{x}+{y}")
    
    donation_win.grab_set()
    donation_win.focus_force()

    form_frame = tk.Frame(donation_win)
    form_frame.pack(pady=10)

    # === DONATION TYPE ===
    tk.Label(form_frame, text="Donation Type:").grid(row=0, column=0, padx=5, pady=5, sticky="e")
    donation_type_var = tk.StringVar()
    donation_type_combo = ttk.Combobox(form_frame, textvariable=donation_type_var, state="readonly", width=30)
    donation_type_combo['values'] = ("Cash", "Check", "Online")
    donation_type_combo.grid(row=0, column=1, padx=5, pady=5)
    donation_type_combo.set("Cash")
    donation_type_combo.focus_set()
    # === CENTER ===
    tk.Label(form_frame, text="Center:").grid(row=1, column=0, padx=5, pady=5, sticky="e")
    center_var = tk.StringVar()
    center_combo = ttk.Combobox(form_frame, textvariable=center_var, state="readonly", width=30)
    center_combo.grid(row=1, column=1, padx=5, pady=5)

    # === COLLECTOR ===
    tk.Label(form_frame, text="Collector:").grid(row=2, column=0, padx=5, pady=5, sticky="e")
    collector_var = tk.StringVar()
    collector_combo = ttk.Combobox(form_frame, textvariable=collector_var, state="readonly", width=30)
    collector_combo.grid(row=2, column=1, padx=5, pady=5)

    # === DONOR (filtered by collector) ===
    tk.Label(form_frame, text="Donor:").grid(row=3, column=0, padx=5, pady=5, sticky="e")
    donor_var = tk.StringVar()
    donor_combo = ttk.Combobox(form_frame, textvariable=donor_var, state="readonly", width=30)
    donor_combo.grid(row=3, column=1, padx=5, pady=5)

    # === AMOUNT ===
    tk.Label(form_frame, text="Donation Amount:").grid(row=4, column=0, padx=5, pady=5, sticky="e")
    amount_entry = tk.Entry(form_frame, width=33)
    amount_entry.grid(row=4, column=1, padx=5, pady=5)

    # === DATE ===
    tk.Label(form_frame, text="Donation Date:").grid(row=5, column=0, padx=5, pady=5, sticky="e")
    donation_date = DateEntry(form_frame, width=30)
    donation_date.grid(row=5, column=1, padx=5, pady=5)

    # ===== LOAD DROPDOWNS FROM DATABASE =====
    center_dict = {}
    collector_dict_by_center = {}
    donor_dict_by_collector = {}

    conn = get_connection()
    if conn:
        try:
            cursor = conn.cursor()

            # === Load Centers ===
            cursor.execute("SELECT CenterID, CenterName FROM tblCenters")
            centers = cursor.fetchall()
            center_combo['values'] = [c.CenterName for c in centers]
            center_dict = {c.CenterName: c.CenterID for c in centers}

            # === Load Collectors by Center ===
            for center in centers:
                cursor.execute("SELECT CollectorID, FullName FROM tblCollectors WHERE CenterID=?", (center.CenterID,))
                collectors = cursor.fetchall()
                collector_dict_by_center[center.CenterName] = {c.FullName: c.CollectorID for c in collectors}

                # Also for each collector, prepare donor mapping
                for c in collectors:
                    cursor.execute("SELECT DonorID, FullName FROM tblDonors WHERE CollectorID=?", (c.CollectorID,))
                    donors = cursor.fetchall()
                    donor_dict_by_collector[c.FullName] = {d.FullName: d.DonorID for d in donors}

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load data: {e}")
        finally:
            conn.close()

    # ==== WHEN CENTER SELECTED, UPDATE COLLECTOR COMBO ====
    def update_collector_combo(event):
        selected_center = center_var.get()
        collectors = collector_dict_by_center.get(selected_center, {})
        collector_combo['values'] = list(collectors.keys())
        collector_combo.set('')
        # Also reset donor list
        donor_combo['values'] = []
        donor_combo.set('')

    center_combo.bind("<<ComboboxSelected>>", update_collector_combo)
    
    # ==== Set default selected center ====
    if centers:
        first_center_name = centers[0].CenterName
        center_combo.set(first_center_name)
       # Also load collectors for that center immediately
        collectors = collector_dict_by_center.get(first_center_name, {})
        collector_combo['values'] = list(collectors.keys())
        collector_combo.set('')
        donor_combo['values'] = []
        donor_combo.set('')

    # ==== WHEN COLLECTOR SELECTED, UPDATE DONOR COMBO ====
    def update_donor_combo(event):
        selected_collector = collector_var.get()
        donors = donor_dict_by_collector.get(selected_collector, {})
        donor_combo['values'] = list(donors.keys())
        donor_combo.set('')

    collector_combo.bind("<<ComboboxSelected>>", update_donor_combo)
    
    
   


    # ===== BUTTON ACTIONS =====
    def clear_form():
        donation_type_combo.set("Cash")
        #center_combo.set('')
        collector_combo.set('')
        donor_combo.set('')
        amount_entry.delete(0, tk.END)
        donation_date.set_date(datetime.now())
        donor_combo['values'] = []

    # ===== ADD DONATION =====
    def add_donation():
        try:
            donation_type = donation_type_var.get()
            center_name = center_var.get()
            collector_name = collector_var.get()
            donor_name = donor_var.get()
            amount = amount_entry.get().strip()
            date = donation_date.get_date()
            date_str = date.strftime('%Y-%m-%d')

            if not center_name or not collector_name or not donor_name or not amount:
                messagebox.showwarning("Input Error", "All fields are required.")
                return

            center_id = center_dict.get(center_name)
            collector_id = collector_dict_by_center[center_name][collector_name]
            donor_id = donor_dict_by_collector[collector_name][donor_name]

            conn = get_connection()
            if conn:
                cursor = conn.cursor()

                # ---- VALIDATE DONATION DATE AGAINST COMMITMENT ----
                cursor.execute("""
                    SELECT MIN(CommitmentStartDate)
                    FROM tblCommitments
                    WHERE DonorID = ? AND IsActive = 1
                """, (donor_id,))

                row = cursor.fetchone()
                if row and row[0]:
                    commitment_start = parse_db_date(row[0])
                    if commitment_start:
                        commitment_start = commitment_start.date()

                    if commitment_start and date < commitment_start:
                        messagebox.showerror(
                            "Invalid Donation Date",
                            f"Donation date cannot be earlier than Active commitment start date "
                            f"({format_db_date(commitment_start, '%d-%m-%Y')})."
                        )
                        donation_date.focus_set()
                        conn.close()
                        return


                # ---- INSERT DONATION ----
                cursor.execute("""
                    INSERT INTO tblDonations
                        (DonationType, DonationAmount, DonationDate, CenterID, CollectorID, DonorID)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (donation_type, amount, date_str, center_id, collector_id, donor_id))

                conn.commit()
                conn.close()

                messagebox.showinfo("Success", "Donation added successfully.")
                load_donations()
                clear_form()

        except Exception as e:
            messagebox.showerror("Error", f"Failed to add donation: {e}")


    # ===== LOAD DONATIONS INTO TREEVIEW =====
    tree_frame = tk.Frame(donation_win)
    tree_frame.pack(padx=20, pady=10, fill="both", expand=True)

    columns = ("Type", "Amount", "Date", "Center", "Collector", "Donor")
    tree = ttk.Treeview(tree_frame, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150)
    tree.pack(fill="both", expand=True)
    
    #tree.bind("<<TreeviewSelect>>", on_tree_select)

    def load_donations():
        for row in tree.get_children():
            tree.delete(row)
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT 
                        d.DonationType, 
                        d.DonationAmount, 
                        d.DonationDate,
                        c.CenterName, 
                        col.FullName, 
                        do.FullName
                    FROM tblDonations d
                    LEFT JOIN tblCenters c ON d.CenterID = c.CenterID
                    LEFT JOIN tblCollectors col ON d.CollectorID = col.CollectorID
                    LEFT JOIN tblDonors do ON d.DonorID = do.DonorID
                    ORDER BY d.DonationDate DESC
                """)
                for row in cursor.fetchall():
                    tree.insert('', 'end', values=[str(x) for x in row])
                conn.close()
            except Exception as e:
                messagebox.showerror("DB Error", f"Failed to load donations: {e}")

    load_donations()
    #tree.bind("<<TreeviewSelect>>", on_tree_select)

    def load_donations_by_donor(donor_id):
        for row in tree.get_children():
            tree.delete(row)
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT 
                        d.DonationType, 
                        d.DonationAmount, 
                        d.DonationDate,
                        c.CenterName, 
                        col.FullName, 
                        do.FullName
                    FROM tblDonations d
                    LEFT JOIN tblCenters c ON d.CenterID = c.CenterID
                    LEFT JOIN tblCollectors col ON d.CollectorID = col.CollectorID
                    LEFT JOIN tblDonors do ON d.DonorID = do.DonorID
                    WHERE d.DonorID = ?
                    ORDER BY d.DonationDate DESC
                """, (donor_id,))
                for row in cursor.fetchall():
                    tree.insert('', 'end', values=[str(x) for x in row])
                conn.close()
            except Exception as e:
                messagebox.showerror("DB Error", f"Failed to load donations: {e}")

    def on_donor_selected(event):
        #selected_donor_name = donor_var.get()
        #donor_id = donor_dict_by_collector.get(selected_donor_name)
        #if donor_id:
        #    load_donations_by_donor(donor_id)
            
        selected_collector = collector_var.get()
        selected_donor = donor_var.get()
        donor_id = None
        if selected_collector in donor_dict_by_collector:
            donor_id = donor_dict_by_collector[selected_collector].get(selected_donor)
        if donor_id:
            load_donations_by_donor(donor_id)
            
    donor_combo.bind("<<ComboboxSelected>>", on_donor_selected)
    
    #--------show update button
    def show_update_button(tree_item):
    # If already exists, destroy and recreate to reset
        for widget in btn_frame.winfo_children():
            if str(widget).endswith("update_button"):
                widget.destroy()

        global update_button
        update_button = tk.Button(btn_frame, text="Update", bg="blue", fg="white", width=15,
                                name="update_button",
                                command=lambda: update_donation(tree_item))
        update_button.grid(row=0, column=3, padx=5)


    
    #print receipt
    def show_print_button(tree_item, donation_data):
        # Remove old button if exists
        for widget in btn_frame.winfo_children():
            if str(widget).endswith("print_button"):
                widget.destroy()

        global print_button
        print_button = tk.Button(
            btn_frame, text="Print Receipt", bg="purple", fg="white", width=15,
            name="print_button",
            command=lambda: print_receipt(donation_data)
        )
        print_button.grid(row=0, column=4, padx=5)


        

    def print_receipt(data):
        try:
            donation_type, amount, donation_date_str, center_name, collector_name, donor_name = data

            # Temporary PDF file
            tmpfile = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            c = canvas.Canvas(tmpfile.name, pagesize=A5)

            width, height = A5

            # --- HEADER ---
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width/2, height-40, "Alkhidmat Khawateen Trust Pakistan")

            c.setFont("Helvetica-Bold", 12)
            c.drawCentredString(width/2, height-60, "Aghosh Program")

            c.setFont("Helvetica", 11)
            c.drawCentredString(60, height-80, f"{center_name}")

            c.setFont("Helvetica", 10)
            c.drawRightString(width-40, height-80, f"Donation Date: {donation_date_str}")

            # Receipt date = current system date
            receipt_date = datetime.now().strftime("%Y-%m-%d")
            c.setFont("Helvetica", 10)
            c.drawString(40, height-100, f"Receipt Date: {receipt_date}")

            # Horizontal line
            c.line(40, height-110, width-40, height-110)

            # --- BODY DETAILS ---
            y = height - 140
            line_gap = 20
            c.setFont("Helvetica", 11)
            c.drawString(50, y, f"Donor Name: {donor_name}")
            y -= line_gap
            c.drawString(50, y, f"Collector: {collector_name}")
            y -= line_gap
            c.drawString(50, y, f"Donation Type: {donation_type}")
            y -= line_gap
            c.drawString(50, y, f"Donation Amount: Rs. {amount}")
            y -= line_gap

            # Footer
            c.line(40, 80, width-40, 80)
            c.setFont("Helvetica-Oblique", 10)
            c.drawCentredString(width/2, 65, "Thank you for your generous support!")

            # Finalize PDF
            c.showPage()
            c.save()

            # Auto open PDF
            os.startfile(tmpfile.name)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

   

    
    
# on donation selection to edit
    def on_tree_select(event):
        selected_item = tree.focus()
        if not selected_item:
            return
        data = tree.item(selected_item, 'values')
        if not data:
            return
        
        # Unpack values based on your columns order
        donation_type, amount, donation_date_str, center_name, collector_name, donor_name = data
        
        # Load into form
        donation_type_combo.set(donation_type)
        center_combo.set(center_name)
        update_collector_combo(None)  # load collectors for this center
        collector_combo.set(collector_name)
        update_donor_combo(None)      # load donors for this collector
        donor_combo.set(donor_name)
        center_combo['state'] = 'disabled'
        collector_combo['state'] = 'disabled'
        donor_combo['state'] = 'disabled'
        add_button.config(state="disabled")

        amount_entry.delete(0, tk.END)
        amount_entry.insert(0, amount)
        donation_date.set_date(datetime.strptime(donation_date_str, "%Y-%m-%d"))
        

        # Enable update mode
        #add_button.config(text="Update Donation", command=lambda: update_donation(selected_item))
        #tree.bind("<<TreeviewSelect>>", on_tree_select)
        # Show Update Button dynamically
        show_update_button(selected_item)
        show_print_button(selected_item, data)
    
    
        
    #def update_donation(tree_item):
    def update_donation(tree_item):
        try:
            donation_type = donation_type_var.get()
            center_name = center_var.get()
            collector_name = collector_var.get()
            donor_name = donor_var.get()
            amount = amount_entry.get().strip()
            date_str = donation_date.get_date().strftime('%Y-%m-%d')

            center_id = center_dict.get(center_name)
            collector_id = collector_dict_by_center.get(center_name, {}).get(collector_name)
            donor_id = donor_dict_by_collector.get(collector_name, {}).get(donor_name)

            old_amount, old_date = tree.item(tree_item, 'values')[1:3]

            conn = get_connection()
            if conn:
                cursor = conn.cursor()
                cursor.execute("""
                    UPDATE tblDonations
                    SET DonationType=?, DonationAmount=?, DonationDate=?, CenterID=?, CollectorID=?, DonorID=?
                    WHERE DonationAmount=? AND DonationDate=?
                """, (donation_type, amount, date_str, center_id, collector_id, donor_id, old_amount, old_date))
                conn.commit()
                conn.close()

            messagebox.showinfo("Success", "Donation updated successfully.")
            load_donations()
            clear_form()

            # Hide update button after done
            update_button.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Failed to update donation: {e}")
    


    
    
    
    

    # ===== BUTTONS =====
    btn_frame = tk.Frame(donation_win)
    btn_frame.pack(pady=10)

    add_button = tk.Button(btn_frame, text="Add Donation", command=add_donation, width=15, bg="green", fg="white")
    add_button.grid(row=0, column=0, padx=5)
    #tk.Button(btn_frame, text="Add Donation", command=add_donation, width=15, bg="green", fg="white").grid(row=0, column=0, padx=5)
    tk.Button(btn_frame, text="Clear", command=clear_form, width=10).grid(row=0, column=1, padx=5)
    tk.Button(btn_frame, text="Close", command=donation_win.destroy, width=10).grid(row=0, column=2, padx=5)

    donation_win.bind("<Escape>", lambda e: donation_win.destroy())
    donation_win.bind("<Return>", lambda e: add_donation())
    tree.bind("<<TreeviewSelect>>", on_tree_select)



    

# ============================================================

    

# ============================================================
# Sessions Management
# ============================================================



def open_session_mgmt():
    conn = get_connection()
    cursor = conn.cursor()
    def load_sessions():
        session_tree.delete(*session_tree.get_children())
        cursor.execute("SELECT SessionID, SessionName, StartDate, EndDate, IsCurrentSession FROM tblSessions")
        for row in cursor.fetchall():
            session_tree.insert("", "end", values=(row.SessionID, row.SessionName, str(row.StartDate)[:10] if row.StartDate else '', row.EndDate.date(), "Yes" if row.IsCurrentSession else "No"))

    def clear_form():
        session_id_var.set("")
        name_entry.delete(0, tk.END)
        save_btn.config(text="Add Session")
        #start_date.set_date("")
        #end_date.set_date("")
        is_current_var.set(0)
        

    def save_session():
        name = name_entry.get().strip()
        is_current = is_current_var.get()
        session_id = session_id_var.get()

        try:
            # Get and validate dates
            start = start_date.get()
            end = end_date.get()

            if not start or not end:
                messagebox.showerror("Error", "Both Start Date and End Date are required.")
                return
            if not name:
                messagebox.showerror("Error", "Session name is required.")
                return

            # Convert to proper format
            start_dt = datetime.strptime(start, "%m/%d/%y")
            end_dt = datetime.strptime(end, "%m/%d/%y")

            # Format for SQL Server
            start_str = start_dt.strftime("%Y-%m-%d")
            end_str = end_dt.strftime("%Y-%m-%d")

            if start_dt > end_dt:
                messagebox.showerror("Error", "Start date must be before End date.")
                return

            # ✅ Check for duplicates (same Start and End dates, and not same SessionID if editing)
            if session_id:
                cursor.execute("""
                    SELECT COUNT(*) FROM tblSessions 
                    WHERE StartDate = ? AND EndDate = ? AND SessionID != ?""",
                    (start_str, end_str, session_id))
            else:
                cursor.execute("""
                    SELECT COUNT(*) FROM tblSessions 
                    WHERE StartDate = ? AND EndDate = ?""",
                    (start_str, end_str))
            
            if cursor.fetchone()[0] > 0:
                messagebox.showwarning("Duplicate", "Session already saved.")
                return

            # Save to DB
            if is_current:
                cursor.execute("UPDATE tblSessions SET IsCurrentSession = 0 WHERE IsCurrentSession = 1")

            if session_id:
                cursor.execute("""
                    UPDATE tblSessions 
                    SET SessionName=?, StartDate=?, EndDate=?, IsCurrentSession=? 
                    WHERE SessionID=?""",
                    (name, start_str, end_str, is_current, session_id))
            else:
                cursor.execute("""
                    INSERT INTO tblSessions (SessionName, StartDate, EndDate, IsCurrentSession) 
                    VALUES (?, ?, ?, ?)""",
                    (name, start_str, end_str, is_current))

            conn.commit()
            save_btn.config(text="Add Session")
            name_entry.focus_set()
            messagebox.showinfo("Success", "Session saved successfully.")
            load_sessions()
            clear_form()

        except Exception as e:
            conn.rollback()
            #messagebox.showerror("Database Error", str(e))




    def select_session(event):
        selected = session_tree.focus()
        if not selected:
            return
        values = session_tree.item(selected, "values")
        session_id_var.set(values[0])
        name_entry.delete(0, tk.END)
        name_entry.insert(0, values[1])
        start_date.set_date(datetime.strptime(values[2], "%Y-%m-%d").date())
        end_date.set_date(datetime.strptime(values[3], "%Y-%m-%d").date())
        
        #start_date.set_date(values[2])
        #end_date.set_date(values[3])
        is_current_var.set(1 if values[4] == "Yes" else 0)
        save_btn.config(text="Update Session")
        name_entry.focus_set()

    # UI
    window = tk.Toplevel()
    window.title("Session Management" + ORG_SUFFIX)
    window.geometry("700x500")
    window.grab_set()
    window.focus_force()
    window.bind("<Escape>", lambda e: window.destroy())
    window.bind('<Return>', lambda e: save_session())

    # Variables
    session_id_var = tk.StringVar()
    is_current_var = tk.IntVar()

    # Form
    tk.Label(window, text="Add New Session", font=("Arial", 16, "bold")).pack(pady=10)
    form_frame = tk.Frame(window)
    form_frame.pack(pady=10, fill="x", padx=20)
    
    
    
    # === FORM FIELDS ===
    #form_frame = tk.Frame(window)
    #form_frame.pack(padx=30, pady=10)
    

    # Session Name
    tk.Label(form_frame, text="Session Name:").grid(row=2, column=0, padx=(10, 2), pady=5, sticky="e")
    name_entry = tk.Entry(form_frame, width=20)
    name_entry.grid(row=2, column=1, padx=(2, 10), pady=5, sticky="w")

    # Start Date
    tk.Label(form_frame, text="Start Date:").grid(row=3, column=0, padx=(10, 2), pady=5, sticky="e")
    start_date = DateEntry(form_frame, width=18)
    start_date.grid(row=3, column=1, padx=(2, 10), pady=5, sticky="w")

    # End Date
    tk.Label(form_frame, text="End Date:").grid(row=4, column=0, padx=(10, 2), pady=5, sticky="e")
    end_date = DateEntry(form_frame, width=18)
    end_date.grid(row=4, column=1, padx=(2, 10), pady=5, sticky="w")


    tk.Checkbutton(form_frame, text="Is Current Session", variable=is_current_var).grid(row=5, column=1, sticky="w", pady=5)

    save_btn = tk.Button(form_frame, text="Add Session", command=save_session, width=15, bg="lightgreen")
    save_btn.grid(row=6, column=1, pady=10, sticky="w")

    tk.Button(form_frame, text="Clear", command=clear_form).grid(row=6, column=1, padx=120, pady=10, sticky="w")

    # Treeview
    tree_frame = tk.Frame(window)
    tree_frame.pack(pady=10, fill="both", expand=True)

    columns = ("SessionID", "SessionName", "StartDate", "EndDate", "IsCurrent")
    session_tree = ttk.Treeview(tree_frame, columns=columns, show="headings", height=10)
    for col in columns:
        session_tree.heading(col, text=col)
        session_tree.column(col, width=100)
    session_tree.pack(fill="both", expand=True)
    session_tree.bind("<<TreeviewSelect>>", select_session)
    name_entry.focus_set()
    # Load existing
    load_sessions()



    
# ============================================================
# Commitment Management
import datetime as dt
# ============================================================

def _calc_monthly(total_amount: float,
                  start: dt.date,
                  end: dt.date) -> float:
    """
    Divide the total commitment amount evenly across the months covered
    by [start, end] (inclusive of both boundary months).
    Returns 0 if the range is invalid.
    """
    months = (end.year - start.year) * 12 + (end.month - start.month) + 1
    if months <= 0:
        return 0.0
    return round(total_amount / months, 2)
 
 
# ─────────────────────────────────────────────────────────────────────────────
# open_donor_commitments  (FIXED)
# ─────────────────────────────────────────────────────────────────────────────
 
def open_donor_commitments():
 
    commit_win = tk.Toplevel()
    commit_win.title("Manage Donor Commitments" + ORG_SUFFIX)
    ww, wh = 820, 580
    sw = commit_win.winfo_screenwidth()
    sh = commit_win.winfo_screenheight()
    commit_win.geometry(f"{ww}x{wh}+{(sw-ww)//2}+{(sh-wh)//2}")
    commit_win.transient()
    commit_win.grab_set()
    commit_win.focus_force()
    commit_win.bind("<Escape>", lambda e: commit_win.destroy())
 
    # ── TOP FILTER BAR  — Collector → Donor ──────────────────────────────────
    top_frame = tk.Frame(commit_win)
    top_frame.pack(pady=8, padx=10, fill="x")
 
    tk.Label(top_frame, text="Collector:").grid(row=0, column=0, padx=5, sticky="e")
    collector_var   = tk.StringVar()
    collector_combo = ttk.Combobox(top_frame, textvariable=collector_var,
                                   width=28, state="readonly")
    collector_combo.grid(row=0, column=1, padx=5)
 
    tk.Label(top_frame, text="Select Donor:").grid(row=0, column=2, padx=5, sticky="e")
    donor_var   = tk.StringVar()
    donor_combo = AutocompleteCombobox(top_frame, textvariable=donor_var, width=34)
    donor_combo.grid(row=0, column=3, padx=5)
 
    # ── LOAD COLLECTORS ───────────────────────────────────────────────────────
    collector_lookup = {"-- All Collectors --": None}
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT CollectorID, FullName FROM tblCollectors ORDER BY FullName")
        for row in cursor.fetchall():
            collector_lookup[row.FullName] = row.CollectorID
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load collectors:\n{e}")
 
    collector_combo["values"] = list(collector_lookup.keys())
    collector_combo.current(0)
 
    # ── LOAD DONORS filtered by collector ────────────────────────────────────
    donor_dict = {}   # {DonorName: DonorID}
 
    def load_donors_for_collector(event=None):
        cid = collector_lookup.get(collector_var.get())
        try:
            conn = get_connection()
            cursor = conn.cursor()
            if cid is None:
                cursor.execute(
                    "SELECT DonorID, FullName FROM tblDonors "
                    "WHERE IsActive=1 ORDER BY FullName")
            else:
                cursor.execute(
                    "SELECT DonorID, FullName FROM tblDonors "
                    "WHERE IsActive=1 AND CollectorID=? ORDER BY FullName",
                    (cid,))
            rows = cursor.fetchall()
            conn.close()
            donor_dict.clear()
            for row in rows:
                donor_dict[row.FullName] = row.DonorID
            donor_combo.set_completion_list(list(donor_dict.keys()))
            donor_var.set("")
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")
 
    collector_combo.bind("<<ComboboxSelected>>", load_donors_for_collector)
    load_donors_for_collector()   # initial load
 
    # ── TREEVIEW ──────────────────────────────────────────────────────────────
    columns = ("Start Date", "End Date", "Commitment Amount",
               "Monthly Amount", "Notes", "Active")
 
    tree_frame = tk.Frame(commit_win)
    tree_frame.pack(padx=10, pady=6, fill="both", expand=True)
 
    v_scroll = tk.Scrollbar(tree_frame, orient="vertical")
    h_scroll = tk.Scrollbar(tree_frame, orient="horizontal")
 
    tree = ttk.Treeview(
        tree_frame, columns=columns, show="headings", height=8,
        yscrollcommand=v_scroll.set,
        xscrollcommand=h_scroll.set
    )
    v_scroll.config(command=tree.yview)
    h_scroll.config(command=tree.xview)
 
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=120, minwidth=80)
 
    h_scroll.pack(side="bottom", fill="x")
    v_scroll.pack(side="right",  fill="y")
    tree.pack(side="left", fill="both", expand=True)
 
    # ── COMMITMENT FORM ───────────────────────────────────────────────────────
    form_frame = tk.Frame(commit_win)
    form_frame.pack(pady=6, padx=10)
 
    tk.Label(form_frame, text="Start Date:").grid(row=0, column=0, padx=5, pady=4, sticky="e")
    start_date = DateEntry(form_frame, width=15)
    start_date.grid(row=0, column=1, padx=5)
 
    tk.Label(form_frame, text="End Date:").grid(row=0, column=2, padx=5, pady=4, sticky="e")
    end_date = DateEntry(form_frame, width=15)
    end_date.grid(row=0, column=3, padx=5)
 
    tk.Label(form_frame, text="Commitment Amount:").grid(row=1, column=0, padx=5, pady=4, sticky="e")
    amount_entry = tk.Entry(form_frame, width=20)
    amount_entry.grid(row=1, column=1, padx=5)
 
    # ── FIX: display the calculated monthly amount (read-only) ───────────────
    tk.Label(form_frame, text="Monthly Amount:").grid(row=1, column=2, padx=5, pady=4, sticky="e")
    monthly_var   = tk.StringVar()
    monthly_entry = tk.Entry(form_frame, width=20, textvariable=monthly_var,
                             state="readonly", readonlybackground="#f1f5f9")
    monthly_entry.grid(row=1, column=3, padx=5)
 
    tk.Label(form_frame, text="Notes:").grid(row=2, column=0, padx=5, pady=4, sticky="e")
    notes_entry = tk.Text(form_frame, height=2, width=40)
    notes_entry.grid(row=2, column=1, columnspan=3, padx=5, pady=4)
 
    is_active_var = tk.IntVar(value=1)
    tk.Checkbutton(form_frame, text="Active",
                   variable=is_active_var).grid(row=3, column=1, sticky="w", padx=5)
 
    # ── live-update Monthly Amount whenever amount/dates change ──────────────
    def refresh_monthly(*_):
        try:
            amt   = float(amount_entry.get().strip())
            s     = start_date.get_date()
            e     = end_date.get_date()
            mthly = _calc_monthly(amt, s, e)
            monthly_var.set(f"{mthly:.2f}")
        except Exception:
            monthly_var.set("")
 
    amount_entry.bind("<KeyRelease>", refresh_monthly)
    start_date.bind("<<DateEntrySelected>>", refresh_monthly)
    end_date.bind("<<DateEntrySelected>>",   refresh_monthly)
 
    # ── BUTTONS ───────────────────────────────────────────────────────────────
    button_frame = tk.Frame(commit_win)
    button_frame.pack(pady=8)
 
    add_btn = tk.Button(button_frame, text="Add Commitment",
                        width=20, bg="blue", fg="white")
    add_btn.grid(row=0, column=0, padx=10)
 
    update_btn = tk.Button(button_frame, text="Update Commitment",
                           width=20, bg="green", fg="white")
    update_btn.grid(row=0, column=1, padx=10)
    update_btn.grid_remove()
 
    tk.Button(button_frame, text="Close",
              width=15, command=commit_win.destroy).grid(row=0, column=2, padx=10)
 
    # ── FUNCTIONS ─────────────────────────────────────────────────────────────
    selected_commitment_id = None
 
    def load_commitments(event=None):
        donor_name = donor_var.get()
        donor_id   = donor_dict.get(donor_name)
        if not donor_id:
            return
        tree.delete(*tree.get_children())
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT CommitmentID, CommitmentStartDate, CommitmentEndDate,
                       CommitmentAmount, MonthlyAmount, Notes, IsActive
                FROM tblCommitments
                WHERE DonorID=?
                ORDER BY CommitmentStartDate DESC
            """, (donor_id,))
            rows = cursor.fetchall()
            conn.close()
            for row in rows:
                cid, s, e, amt, mamt, notes, active = row
                tree.insert("", "end", iid=cid, values=(
                    format_db_date(s, "%d/%m/%Y") if s else "",
                    format_db_date(e, "%d/%m/%Y") if e else "",
                    f"Rs.{float(amt):.2f}"  if amt  else "0.00",
                    f"Rs.{float(mamt):.2f}" if mamt else "0.00",
                    notes or "",
                    "Yes" if active else "No"
                ))
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load commitments:\n{e}")
 
    # ── FIX #3: donor selection drives button visibility ─────────────────────
    def on_donor_change(event=None):
        """
        Called whenever the donor combobox selection changes.
        - Loads that donor's commitments into the tree.
        - If the donor has NO commitments  → show Add, hide Update.
        - If the donor already has commitments → hide Add, show Update
          and pre-populate the most-recent commitment in the form.
        """
        reset_form()          # clear form & restore default button state
        load_commitments()    # populate tree
 
        donor_name = donor_var.get()
        donor_id   = donor_dict.get(donor_name)
        if not donor_id:
            return
 
        # Check whether the treeview has any rows after load_commitments()
        children = tree.get_children()
        if children:
            # Pre-select the first (most-recent) commitment and show Update
            tree.selection_set(children[0])
            tree.focus(children[0])
            on_tree_select(None)   # will flip to Update mode
        # else: reset_form() already shows Add / hides Update
 
    donor_combo.bind("<<ComboboxSelected>>", on_donor_change)
 
    # ── Reset form ────────────────────────────────────────────────────────────
    def reset_form():
        nonlocal selected_commitment_id
        selected_commitment_id = None
        add_btn.grid()
        update_btn.grid_remove()
        try:
            start_date.set_date(dt.datetime.now())
            end_date.set_date(dt.datetime.now())
        except Exception:
            pass
        amount_entry.delete(0, tk.END)
        monthly_var.set("")
        notes_entry.delete("1.0", tk.END)
        is_active_var.set(1)
 
    # ── Add commitment ────────────────────────────────────────────────────────
    def add_commitment():
        donor_name = donor_combo.get().strip()
        if not donor_name:
            messagebox.showwarning("Validation Error", "Please select a donor.")
            return
        donor_id = donor_dict.get(donor_name)
        if not donor_id:
            messagebox.showerror("Error", "Invalid donor selected.")
            return
 
        try:
            new_start = start_date.get_date()
            new_end   = end_date.get_date()
        except Exception as e:
            messagebox.showwarning("Validation Error", f"Invalid dates: {e}")
            return
 
        if new_end <= new_start:
            messagebox.showwarning("Validation Error",
                                   "End date must be after Start date.")
            return
 
        amount_txt = amount_entry.get().strip()
        if not amount_txt:
            messagebox.showwarning("Validation Error",
                                   "Please enter commitment amount.")
            return
        try:
            amount_val = float(amount_txt)
        except Exception:
            messagebox.showwarning("Validation Error",
                                   "Commitment amount must be a valid number.")
            return
 
        # FIX #1: calculate MonthlyAmount
        monthly_val = _calc_monthly(amount_val, new_start, new_end)
 
        try:
            notes_text = notes_entry.get("1.0", tk.END).strip()
        except TypeError:
            notes_text = notes_entry.get().strip()
 
        is_active_val = 1 if is_active_var.get() == 1 else 0
 
        try:
            conn   = get_connection()
            cursor = conn.cursor()
 
            # Check max existing end date
            cursor.execute(
                "SELECT MAX(CommitmentEndDate) FROM tblCommitments WHERE DonorID=?",
                (donor_id,))
            row = cursor.fetchone()
            max_end_raw  = row[0] if row else None
            max_end_date = None
            if max_end_raw:
                if isinstance(max_end_raw, dt.datetime):
                    max_end_date = max_end_raw.date()
                elif isinstance(max_end_raw, dt.date):
                    max_end_date = max_end_raw
                else:
                    try:
                        max_end_date = dt.datetime.strptime(
                            str(max_end_raw)[:10], "%Y-%m-%d").date()
                    except Exception:
                        max_end_date = None
 
            if max_end_date:
                if new_start <= max_end_date or new_end <= max_end_date:
                    conn.close()
                    messagebox.showwarning(
                        "Validation Error",
                        f"The donor already has a commitment ending on "
                        f"{max_end_date.strftime('%d-%m-%Y')}.\n"
                        "New commitment must start and end after that date."
                    )
                    return
 
            if is_active_val == 1:
                cursor.execute(
                    "SELECT COUNT(*) FROM tblCommitments "
                    "WHERE DonorID=? AND IsActive=1", (donor_id,))
                active_count = cursor.fetchone()[0] or 0
                if active_count > 0:
                    confirm = messagebox.askyesno(
                        "Active Commitment Exists",
                        "The donor already has an active commitment.\n\n"
                        "Do you want to deactivate the old commitment(s) "
                        "and add this new one as active?"
                    )
                    if confirm:
                        cursor.execute(
                            "UPDATE tblCommitments SET IsActive=0 "
                            "WHERE DonorID=? AND IsActive=1", (donor_id,))
                    else:
                        conn.close()
                        return
 
            # FIX #1: INSERT now includes MonthlyAmount
            cursor.execute("""
                INSERT INTO tblCommitments
                    (DonorID, CommitmentStartDate, CommitmentEndDate,
                     CommitmentAmount, MonthlyAmount, Notes, IsActive)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (donor_id,
                  new_start.strftime("%Y-%m-%d"),
                  new_end.strftime("%Y-%m-%d"),
                  amount_val,
                  monthly_val,       # ← was missing
                  notes_text,
                  is_active_val))
 
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Commitment saved successfully.")
            load_commitments()
            reset_form()
 
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to add commitment:\n{e}")
 
    add_btn.config(command=add_commitment)
 
    # ── Tree select → populate form ───────────────────────────────────────────
    def on_tree_select(event=None):
        nonlocal selected_commitment_id
        item = tree.focus()
        if not item:
            return
        selected_commitment_id = int(item)
        vals = tree.item(item, "values")
 
        # vals: (Start Date, End Date, CommitmentAmount, MonthlyAmount, Notes, Active)
        try:
            start_date.set_date(dt.datetime.strptime(vals[0], "%d/%m/%Y"))
        except Exception:
            pass
        try:
            end_date.set_date(dt.datetime.strptime(vals[1], "%d/%m/%Y"))
        except Exception:
            pass
 
        # FIX #5: strip "Rs." prefix properly before inserting into entry
        amount_entry.delete(0, tk.END)
        amount_entry.insert(0, vals[2].replace("Rs.", "").strip())
 
        monthly_var.set(vals[3].replace("Rs.", "").strip())
 
        notes_entry.delete("1.0", tk.END)
        notes_entry.insert("1.0", vals[4])
        is_active_var.set(1 if vals[5] == "Yes" else 0)
 
        # FIX #3: switch to Update mode
        update_btn.grid()
        add_btn.grid_remove()
 
    tree.bind("<<TreeviewSelect>>", on_tree_select)
 
    # ── Update commitment ─────────────────────────────────────────────────────
    def update_commitment():
        if not selected_commitment_id:
            return
 
        try:
            new_start = start_date.get_date()
            new_end   = end_date.get_date()
        except Exception as e:
            messagebox.showwarning("Validation Error", f"Invalid dates: {e}")
            return
 
        if new_end <= new_start:
            messagebox.showwarning("Validation Error",
                                   "End date must be after Start date.")
            return
 
        amount_txt = amount_entry.get().strip()
        if not amount_txt:
            messagebox.showwarning("Validation Error", "Please enter commitment amount.")
            return
        try:
            amount_val = float(amount_txt)
        except Exception:
            messagebox.showwarning("Validation Error",
                                   "Commitment amount must be a valid number.")
            return
 
        # FIX #2: recalculate MonthlyAmount on update
        monthly_val = _calc_monthly(amount_val, new_start, new_end)
 
        try:
            conn   = get_connection()
            cursor = conn.cursor()
            # FIX #2: UPDATE now writes MonthlyAmount
            cursor.execute("""
                UPDATE tblCommitments
                SET CommitmentStartDate=?,
                    CommitmentEndDate=?,
                    CommitmentAmount=?,
                    MonthlyAmount=?,
                    Notes=?,
                    IsActive=?
                WHERE CommitmentID=?
            """, (
                new_start.strftime("%Y-%m-%d"),
                new_end.strftime("%Y-%m-%d"),
                amount_val,
                monthly_val,       # ← was missing
                notes_entry.get("1.0", "end").strip(),
                is_active_var.get(),
                selected_commitment_id
            ))
            conn.commit()
            conn.close()
            messagebox.showinfo("Success", "Commitment updated successfully.")
            load_commitments()
            reset_form()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to update commitment:\n{e}")
 
    update_btn.config(command=update_commitment)
 
    # Enter triggers Add or Update depending on which button is visible
    commit_win.bind("<Return>",
                    lambda e: add_btn.invoke()
                    if add_btn.winfo_ismapped() else update_btn.invoke())

# ============================================================



# ============================================================

def open_user_mgmt():
    open_placeholder("User Access Control")



def open_Exel_For_AKF():
    #messagebox.showinfo("Report", "Opening Performance Report...")
    
    """
    Export all child data from the database into the master sheet Excel format.
    Matches the Format_Master_sheet.xlsx column layout exactly.
    Saves to the user's Downloads folder and opens the file.
    """
    import os
    import datetime
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    try:
        # ── 1. Fetch all data from DB ─────────────────────────────────────────
        conn = get_connection()
        cursor = conn.cursor()

        # Children + center info
        cursor.execute("""
            SELECT
                c.ChildID,
                c.FullName,
                c.FatherName,
                c.FatherDeathDate,
                c.MotherName,
                c.MotherDeathDate,
                c.Intelligence,
                c.RegistrationNumber,
                c.Status,
                c.AdmissionDate,
                c.DateOfBirth,
                c.Class,
                c.SchoolName,
                c.GuardianName,
                c.GuardianRelation,
                c.GuardianCNIC,
                c.GuardianContact,
                c.PhotoPath,
                c.DocSchoolCertificate,
                ce.CenterName,
                ce.Region,
                ce.AdministratorName,
                ce.AdminContactNumber
            FROM tblChildren c
            LEFT JOIN tblCenters ce ON c.CenterID = ce.CenterID
            ORDER BY c.ChildID
        """)
        children = cursor.fetchall()

        # Academic reports — latest per child
        cursor.execute("""
            SELECT ChildID, AcademicYear, Class
            FROM tblAcademicProgressReports
            ORDER BY ChildID, ReportID DESC
        """)
        acad_rows = cursor.fetchall()
        # Build dict: ChildID -> list of (AcademicYear, Class)
        acad_map = {}
        for row in acad_rows:
            acad_map.setdefault(row.ChildID, []).append(
                (row.AcademicYear, row.Class))

        # Sponsorships — active per child → donor info
        cursor.execute("""
            SELECT
                s.ChildID,
                d.OfficeDonorID,
                d.FullName     AS DonorName,
                d.DonationType,
                s.SponsorshipAmount
            FROM tblSponsorships s
            JOIN tblDonors d ON s.DonorID = d.DonorID
            WHERE s.IsActive = 1
            ORDER BY s.ChildID, s.SponsorshipID
        """)
        spon_rows = cursor.fetchall()
        spon_map = {}   # ChildID -> first active sponsorship row
        for row in spon_rows:
            if row.ChildID not in spon_map:
                spon_map[row.ChildID] = row

        conn.close()

        # ── 2. Create workbook ────────────────────────────────────────────────
        wb = Workbook()
        ws = wb.active
        ws.title = "Sheet1"

        # ── 3. Define columns (exactly matching the template) ─────────────────
        HEADERS = [
            "Sr. No",                           # A  1
            "Region",                           # B  2
            "Aghosh",                           # C  3
            "Administrator Name",               # D  4
            "Contact Number",                   # E  5
            "Child Category",                   # F  6
            "ERP ID",                           # G  7
            "Old ID",                           # H  8
            "Status",                           # I  9
            "Aghosh ID",                        # J  10
            "Child Name",                       # K  11
            "Father Name",                      # L  12
            "Father Death Date",                # M  13
            "Mother Name",                      # N  14
            "Mother Death Date (if Died)",      # O  15
            "H.O Donor Code",                   # P  16
            "INGO Code",                        # Q  17
            "Donor Stake",                      # R  18
            "Donor Category",                   # S  19
            "HO Organization Name",             # T  20
            "Organization Donor Name",          # U  21
            "Corporate Donor Name",             # V  22
            "HO Individual Donor Name",         # W  23
            "Regional Donor Name",              # X  24
            "Aghosh Local Donor Name",          # Y  25
            "Care of (C/O)",                    # Z  26
            "Refused By:",                      # AA 27
            "Admission Date",                   # AB 28
            "DOB",                              # AC 29
            "Age",                              # AD 30
            "Age Round",                        # AE 31
            "Class 2025-2026",                  # AF 32
            "",                                 # AG 33  (blank)
            "",                                 # AH 34  (blank)
            "",                                 # AI 35  (blank)
            "",                                 # AJ 36  (blank)
            "",                                 # AK 37  (blank)
            "",                                 # AL 38  (blank)
            "",                                 # AM 39  (blank)
            "School Name",                      # AN 40
            "Guardian Name",                    # AO 41
            "Relation",                         # AP 42
            "CNIC #2",                          # AQ 43
            "Contact No",                       # AR 44
            "Profile (Yes/No)",                 # AS 45
            "Sponsorship Certificate (Yes/No)", # AT 46
            "Academic Progress Reports (Yes/No)",# AU 47
            "",                                 # AV 48  (blank trailing col)
        ]

        # Column widths matching the template
        COL_WIDTHS = {
            1:  10.33,  # A
            2:  13.11,  # B
            3:  26.00,  # C
            4:  22.55,  # D
            5:  19.22,  # E
            6:  17.55,  # F
            7:  10.78,  # G
            8:  10.55,  # H
            9:  19.78,  # I
            10: 13.78,  # J
            11: 28.22,  # K
            12: 22.78,  # L
            13: 20.89,  # M
            14: 27.55,  # N
            15: 29.11,  # O
            16: 19.11,  # P
            17: 14.55,  # Q
            18: 15.78,  # R
            19: 18.66,  # S
            20: 24.78,  # T
            21: 27.44,  # U
            22: 25.22,  # V
            23: 28.00,  # W
            24: 40.55,  # X
            25: 27.55,  # Y
            26: 16.33,  # Z
            27: 15.11,  # AA
            28: 18.55,  # AB
            29: 16.00,  # AC
            30: 23.89,  # AD
            31: 14.33,  # AE
            32: 16.66,  # AF
            33: 18.22,  # AG
            34: 16.66,  # AH
            35: 16.66,  # AI
            36: 16.66,  # AJ
            37: 16.66,  # AK
            38: 16.66,  # AL
            39: 16.66,  # AM
            40: 42.00,  # AN
            41: 27.55,  # AO
            42: 12.78,  # AP
            43: 15.78,  # AQ
            44: 40.66,  # AR
            45: 18.66,  # AS
            46: 32.78,  # AT
            47: 36.22,  # AU
            48:  6.22,  # AV
        }

        # ── 4. Style helpers ──────────────────────────────────────────────────
        HEADER_FILL = PatternFill("solid", fgColor="92D050")
        HEADER_FONT = Font(name="Arial", bold=True, size=10)
        HEADER_ALIGN = Alignment(horizontal="center", vertical="center",
                                 wrap_text=True)

        DATA_FONT  = Font(name="Arial", size=10)
        DATA_ALIGN = Alignment(vertical="center", wrap_text=False)

        THIN = Side(style="thin", color="000000")
        BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)

        DATE_FMT = "DD-MMM-YYYY"   # e.g. 18-Dec-2014

        def safe_date(val):
            """Parse a date string from DB (YYYY-MM-DD) → datetime.date or None."""
            if not val:
                return None
            if isinstance(val, (datetime.date, datetime.datetime)):
                return val
            try:
                return datetime.datetime.strptime(str(val)[:10], "%Y-%m-%d").date()
            except Exception:
                return None

        # ── 5. Write header row ───────────────────────────────────────────────
        ws.row_dimensions[1].height = 16.95
        for col_idx, header in enumerate(HEADERS, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font   = HEADER_FONT
            cell.fill   = HEADER_FILL
            cell.alignment = HEADER_ALIGN
            cell.border = BORDER
            ws.column_dimensions[get_column_letter(col_idx)].width = \
                COL_WIDTHS.get(col_idx, 15)

        # Academic year columns in order (cols 32-39)
        ACAD_YEARS = [
            "2019-20", "2020-2021", "2021-2022", "2022-2023",
            "2023-24",  "2024-25",  "2025-26",   "2026-27"
        ]

        # ── 6. Write data rows ────────────────────────────────────────────────
        for sr_no, child in enumerate(children, 1):
            row_idx = sr_no + 1   # data starts at row 2

            child_id   = child.RegistrationNumber
            spon       = spon_map.get(child_id)       # may be None
            acad_list  = acad_map.get(child_id, [])   # list of (year, class)



            # Dates
            dob_date    = safe_date(child.DateOfBirth)
            adm_date    = safe_date(child.AdmissionDate)
            f_death     = safe_date(child.FatherDeathDate)
            m_death     = safe_date(child.MotherDeathDate)

            # Age formula (references DOB cell = AC{row})
            ac_cell = f"AC{row_idx}"
            age_formula  = (f'=DATEDIF({ac_cell},TODAY(),"Y")&"Years,"'
                            f'&DATEDIF({ac_cell},TODAY(),"YM")&" Months,"'
                            f'&DATEDIF({ac_cell},TODAY(),"MD")&" Days"')
            age_round    = (f'=DATEDIF({ac_cell},TODAY(),"Y")&"."'
                            f'&DATEDIF({ac_cell},TODAY(),"YM")')

            # Build the 48 values in column order
            row_values = [
                sr_no,                              # A  Sr. No
                child.Region or "",                 # B  Region
                child.CenterName or "",             # C  Aghosh
                child.AdministratorName or "",      # D  Administrator Name
                child.AdminContactNumber or "",     # E  Contact Number
                child.Intelligence or "",           # F  Child Category
                "",                                 # G  ERP ID (not in DB)
                "",                                 # H  Old ID (not in DB)
                #child.Status or "",                 # I  Status
                "Active" if str(child.Status or "").strip() == "1" else "Inactive",  # I  Status
                child_id,                           # J  Aghosh ID (ChildID)
                child.FullName or "",               # K  Child Name
                child.FatherName or "",             # L  Father Name
                f_death,                            # M  Father Death Date
                child.MotherName or "",             # N  Mother Name
                m_death,                            # O  Mother Death Date
                spon.OfficeDonorID if spon else "", # P  H.O Donor Code
                "",                                 # Q  INGO Code
                spon.SponsorshipAmount if spon else "", # R  Donor Stake
                spon.DonationType if spon else "",  # S  Donor Category
                "",                                 # T  HO Organization Name
                "",                                 # U  Organization Donor Name
                "",                                 # V  Corporate Donor Name
                "",                                 # W  HO Individual Donor Name
                "",                                 # X  Regional Donor Name
                spon.DonorName if spon else "",     # Y  Aghosh Local Donor Name
                "",                                 # Z  Care of (C/O)
                "",                                 # AA Refused By
                adm_date,                           # AB Admission Date
                dob_date,                           # AC DOB
                age_formula,                        # AD Age
                age_round,                          # AE Age Round
                child.Class or "",                  # AF Class 2025-2026 (from tblChildren)
                "",                                 # AG blank
                "",                                 # AH blank
                "",                                 # AI blank
                "",                                 # AJ blank
                "",                                 # AK blank
                "",                                 # AL blank
                "",                                 # AM blank
                child.SchoolName or "",             # AN School Name
                child.GuardianName or "",           # AO Guardian Name
                child.GuardianRelation or "",       # AP Relation
                child.GuardianCNIC or "",           # AQ CNIC #2
                child.GuardianContact or "",        # AR Contact No
                "Yes" if child.PhotoPath and child.PhotoPath.strip() else "No",  # AS Profile
                "",                                 # AT Sponsorship Certificate
                "",                                 # AU Academic Progress Reports
                "",                                 # AV blank
            ]

            for col_idx, value in enumerate(row_values, 1):
                cell = ws.cell(row=row_idx, column=col_idx)
                cell.font      = DATA_FONT
                cell.alignment = DATA_ALIGN
                cell.border    = BORDER

                if isinstance(value, (datetime.date, datetime.datetime)):
                    cell.value          = value
                    cell.number_format  = DATE_FMT
                elif isinstance(value, str) and value.startswith("="):
                    cell.value = value   # formula
                else:
                    cell.value = value

            ws.row_dimensions[row_idx].height = 15

        # ── 7. Freeze header row ──────────────────────────────────────────────
        ws.freeze_panes = "A2"

        # ── 8. Auto-filter on header row ──────────────────────────────────────
        ws.auto_filter.ref = f"A1:{get_column_letter(len(HEADERS))}1"

        # ── 9. Save ───────────────────────────────────────────────────────────
        downloads = os.path.join(os.path.expanduser("~"), "Downloads")
        os.makedirs(downloads, exist_ok=True)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        file_path = os.path.join(downloads, f"Children_Master_Export_{timestamp}.xlsx")

        wb.save(file_path)
        messagebox.showinfo(
            "Export Complete",
            f"✅  Exported {len(children)} children.\n\nFile saved to:\n{file_path}"
        )
        os.startfile(file_path)

    except Exception as e:
        messagebox.showerror("Export Error", f"Failed to export:\n{e}")





def open_Aghosh_summary_report_form():
    messagebox.showinfo("Report", "Opening Aghosh Summary Report...")
    
# ============================================================
# Donation Report
# ============================================================
    
def open_donation_report_form():
    report_win = tk.Toplevel()
    report_win.title("Donation Report by Date Range" + ORG_SUFFIX)
    report_win.geometry("900x500")

    # Center window
    ww, wh = 900, 500
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    # Filter section
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="From Date:").grid(row=0, column=0, padx=5)
    from_date = DateEntry(filter_frame, width=12)
    from_date.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="To Date:").grid(row=0, column=2, padx=5)
    to_date = DateEntry(filter_frame, width=12)
    to_date.grid(row=0, column=3, padx=5)

    report_rows = []
    grand_total = tk.DoubleVar(value=0.0)
    date_range_str = tk.StringVar(value="")

    # Treeview
    columns = ("Donor", "Collector", "Center", "Date", "Donation Amount")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150)

    tree.pack_forget()

    # Grand Total label
    total_label = tk.Label(report_win, text="", font=("Arial", 12, "bold"))
    total_label.pack_forget()

    # Print button
    print_btn = tk.Button(report_win, text="Print Report", width=20, bg="green", fg="white")
    print_btn.pack_forget()

    # Load report data
    def load_report():
        tree.delete(*tree.get_children())
        report_rows.clear()
        grand_total.set(0.0)

        start = from_date.get_date().strftime('%Y-%m-%d')
        end = to_date.get_date().strftime('%Y-%m-%d')
        date_range_str.set(f"From: {start} To: {end}")

        try:
            conn = get_connection()
            cursor = conn.cursor()
            # Cast DonationAmount to float to avoid SUM error on nvarchar
            cursor.execute("""
                SELECT d.FullName AS DonorName, col.FullName AS CollectorName, ce.CenterName, 
                    t.DonationDate, CAST(t.DonationAmount AS REAL) AS Amount
                FROM tblDonations t
                LEFT JOIN tblDonors d ON t.DonorID = d.DonorID
                LEFT JOIN tblCollectors col ON t.CollectorID = col.CollectorID
                LEFT JOIN tblCenters ce ON t.CenterID = ce.CenterID
                WHERE t.DonationDate BETWEEN ? AND ?
                ORDER BY t.DonationDate, d.FullName
            """, (start, end))
            rows = cursor.fetchall()

            total = 0.0
            for row in rows:
                donor, collector, center, donation_date, amount = row
                donation_date_str = format_db_date(donation_date, "%d/%m/%Y") if donation_date else ""
                amount = amount or 0.0
                tree.insert('', 'end', values=(donor, collector, center, donation_date_str, f"Rs.{amount:.2f}"))
                report_rows.append({
                    "Donor": donor,
                    "Collector": collector,
                    "Center": center,
                    "Date": donation_date_str,
                    "Amount": amount
                })
                total += amount  # accumulate total here
                #grand_total.set(grand_total.get() + amount)
            grand_total.set(total)


            if rows:
                tree.pack(padx=10, pady=10, fill="both", expand=True)
                total_label.config(text=f"Grand Total Donations: Rs.{total:.2f}")
                total_label.pack()
                print_btn.pack(pady=10)
            else:
                tree.pack_forget()
                total_label.pack_forget()
                print_btn.pack_forget()
            conn.close()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load donations:\n{e}")

    # Print report


    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No data to print.")
            return

        try:
            # Get Downloads folder
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
            if not os.path.exists(downloads_folder):
                os.makedirs(downloads_folder)

            # Create file path
            file_name = f"Donation_Report_{date_range_str.get().replace(' ', '_').replace(':', '-')}.pdf"
            file_path = os.path.join(downloads_folder, file_name)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(200, height - 60, "Donation Report by Date")
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, date_range_str.get())

            # Table headers
            y = height - 100
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(150, y, "Collector")
            c.drawString(270, y, "Center")
            c.drawString(400, y, "Date")
            c.drawString(480, y, "Amount")

            y -= 20
            c.setFont("Helvetica", 9)

            # Table rows
            for r in report_rows:
                c.drawString(50, y, str(r["Donor"] or ""))
                c.drawString(150, y, str(r["Collector"] or ""))
                c.drawString(270, y, str(r["Center"] or ""))
                c.drawString(400, y, r["Date"])
                c.drawString(480, y, f"Rs.{r['Amount']:.2f}")
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            # Grand total
            y -= 10
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, f"Grand Total Donations: Rs.{grand_total.get():.2f}")

            c.save()

            # Auto open the PDF
            open_file_cross_platform(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))


    # Buttons
    tk.Button(filter_frame, text="Generate Report", command=load_report, bg="blue", fg="white", width=20).grid(row=0, column=4, padx=10)
    print_btn.config(command=print_report)


    
# ============================================================
    
# ============================================================
# Child Sponsorship Report
# ============================================================
    
def open_child_sponsorship_report_form():
    report_win = tk.Toplevel()
    report_win.title("Child Sponsorship Report" + ORG_SUFFIX)
    report_win.geometry("900x500")

    # Center window
    ww, wh = 900, 500
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    # Fetch all children for search box
    child_dict = {}
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT ChildID, FullName FROM tblChildren ORDER BY FullName")
        children = cursor.fetchall()
        child_dict = {row.FullName: row.ChildID for row in children}
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load children:\n{e}")
        return

    # Filter section
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Select Child:").grid(row=0, column=0, padx=5)

    child_var = tk.StringVar()
    child_combo = AutocompleteCombobox(filter_frame, textvariable=child_var, width=40)
    child_combo.set_completion_list(list(child_dict.keys()))
    child_combo.grid(row=0, column=1, padx=5)

    report_rows = []

    # Treeview setup
    columns = ("Center", "Collector", "Donor", "Monthly", "Percentage", "Amount")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=130)
    tree.pack_forget()  # hidden initially

    # Print button setup
    print_btn = tk.Button(report_win, text="Print Report", width=20, bg="green", fg="white")
    print_btn.pack_forget()

    # Report loading function
    def load_report(event=None):
        selected_child = child_var.get()
        child_id = child_dict.get(selected_child)
        if not child_id:
            return
        
        tree.delete(*tree.get_children())
        report_rows.clear()

        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT ce.CenterName,
                    col.FullName AS CollectorName,
                    d.FullName AS DonorName,
                    CAST(s.SponsorshipAmount AS REAL) AS SponsorshipAmount,
                    s.Percentage
                FROM tblSponsorships s
                JOIN tblDonors d ON s.DonorID = d.DonorID
                JOIN tblCollectors col ON d.CollectorID = col.CollectorID
                JOIN tblCenters ce ON col.CenterID = ce.CenterID
                WHERE s.ChildID = ?
            """, (child_id,))
            rows = cursor.fetchall()

            for row in rows:
                center, collector, donor, monthly, percent = row
                allocated = (monthly * percent) / 100
                tree.insert('', 'end', values=(center, collector, donor,
                                               f"Rs.{monthly:.2f}",
                                               f"{percent}%", f"Rs.{allocated:.2f}"))
                report_rows.append({
                    "Center": center,
                    "Collector": collector,
                    "Donor": donor,
                    "Monthly": monthly,
                    "Percentage": percent,
                    "Allocated": allocated
                })

            if rows:
                tree.pack(padx=10, pady=10, fill="both", expand=True)
                print_btn.pack(pady=10)
            else:
                tree.pack_forget()
                print_btn.pack_forget()
            conn.close()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load sponsorship data:\n{e}")

    # Print report function

    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No data to print.")
            return

        try:
            #  Get user's Downloads folder
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")

            #  Generate unique filename
            file_name = f"Child_Sponsorship_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_folder, file_name)

            # === Create PDF ===
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(180, height - 60, "Child Sponsorship Report")
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, f"Child: {child_var.get()}")

            y = height - 100
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Center")
            c.drawString(150, y, "Collector")
            c.drawString(250, y, "Donor")
            c.drawString(350, y, "Monthly")
            c.drawString(420, y, "%")
            #c.drawString(460, y, "Allocated")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r['Center'])
                c.drawString(150, y, r['Collector'])
                c.drawString(250, y, r['Donor'])
                c.drawString(350, y, f"Rs.{r['Monthly']:.2f}")
                c.drawString(420, y, f"{r['Percentage']}%")
                #c.drawString(460, y, f"Rs.{r['Allocated']:.2f}")
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()

            # ✅ Small delay for file write
            time.sleep(1)

            messagebox.showinfo("Success", f"Report saved in Downloads:\n{file_path}")

            # 📂 Auto open PDF
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))


    # Bind events
    child_combo.bind("<<ComboboxSelected>>", load_report)
    print_btn.config(command=print_report)



    
    
# ============================================================
    
# ============================================================
# Sponsorship Report
# ============================================================

    
def open_donor_child_map_report():
    report_win = tk.Toplevel()
    report_win.title("Donor Sponsorship Allocation Report" + ORG_SUFFIX)
    report_win.geometry("900x500")

    ww, wh = 900, 500
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    # Center Combo
    tk.Label(filter_frame, text="Center:").grid(row=0, column=0, padx=5)
    center_var = tk.StringVar()
    center_combo = ttk.Combobox(filter_frame, textvariable=center_var, state="readonly", width=25)
    center_combo.grid(row=0, column=1, padx=5)

    # Collector Combo
    tk.Label(filter_frame, text="Collector:").grid(row=0, column=2, padx=5)
    collector_var = tk.StringVar()
    collector_combo = ttk.Combobox(filter_frame, textvariable=collector_var, state="readonly", width=25)
    collector_combo.grid(row=0, column=3, padx=5)

    # Donor Combo
    tk.Label(filter_frame, text="Donor:").grid(row=0, column=4, padx=5)
    donor_var = tk.StringVar()
    donor_combo = ttk.Combobox(filter_frame, textvariable=donor_var, state="readonly", width=25)
    donor_combo.grid(row=0, column=5, padx=5)

    # Dictionaries
    center_dict = {}
    collector_dict_by_center = {}
    donor_dict_by_collector = {}

    report_rows = []
    monthly_commitment = tk.StringVar(value="Monthly Commitment: Rs. 0.00")

    # Treeview
    columns = ("Child Name", "Percentage", "Sponsored Amount")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=200)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    total_frame = tk.Frame(report_win)
    total_frame.pack()
    tk.Label(total_frame, textvariable=monthly_commitment, font=("Helvetica", 12, "bold")).pack()

    print_btn = tk.Button(report_win, text="Print Report", command=lambda: print_report(report_rows, donor_var.get(), monthly_commitment.get()))
    print_btn.pack(pady=10)
    print_btn.pack_forget()

    # Load centers, collectors, donors
    def load_centers_collectors_donors():
        try:
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("SELECT CenterID, CenterName FROM tblCenters")
            centers = cursor.fetchall()
            center_dict.update({row.CenterName: row.CenterID for row in centers})
            center_combo['values'] = list(center_dict.keys())

            # Collectors per center
            for center in centers:
                cursor.execute("SELECT CollectorID, FullName FROM tblCollectors WHERE CenterID=?", (center.CenterID,))
                collectors = cursor.fetchall()
                collector_dict_by_center[center.CenterName] = {c.FullName: c.CollectorID for c in collectors}

                # Donors per collector
                for c in collectors:
                    cursor.execute("SELECT DonorID, FullName FROM tblDonors WHERE CollectorID=?", (c.CollectorID,))
                    donors = cursor.fetchall()
                    donor_dict_by_collector.setdefault(c.FullName, {})
                    donor_dict_by_collector[c.FullName].update({d.FullName: d.DonorID for d in donors})

            # Defaults
            if centers:
                first_center = centers[0].CenterName
                center_combo.set(first_center)
                load_collectors()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load data:\n{e}")
        finally:
            conn.close()

    # Handlers
    def load_collectors(event=None):
        selected_center = center_var.get()
        collectors = collector_dict_by_center.get(selected_center, {})
        collector_combo['values'] = list(collectors.keys())
        if collectors:
            collector_combo.set(next(iter(collectors)))
            load_donors()

    def load_donors(event=None):
        selected_collector = collector_var.get()
        donors = donor_dict_by_collector.get(selected_collector, {})
        donor_combo['values'] = list(donors.keys())
        if donors:
            donor_combo.set(next(iter(donors)))
            load_report()

    # Report loader
    def load_report(event=None):
        tree.delete(*tree.get_children())
        report_rows.clear()
        monthly_commitment.set("Monthly Commitment: Rs. 0.00")
        print_btn.pack_forget()

        selected_collector = collector_var.get()
        selected_donor = donor_var.get()

        if not selected_collector or not selected_donor:
            return

        donor_id = donor_dict_by_collector.get(selected_collector, {}).get(selected_donor)
        if not donor_id:
            return

        try:
            conn = get_connection()
            cursor = conn.cursor()

            # Get total monthly commitment from active commitments
            cursor.execute("""
                SELECT IFNULL(SUM(MonthlyAmount), 0)
                FROM tblCommitments
                WHERE DonorID = ? AND IsActive = 1
            """, (donor_id,))
            monthly_amount = float(cursor.fetchone()[0] or 0)

            # Get all active sponsorships for donor
            cursor.execute("""
                SELECT c.FullName AS ChildName,
                    CAST(s.SponsorshipAmount AS REAL) AS SponsorshipAmount
                FROM tblSponsorships s
                JOIN tblChildren c ON s.ChildID = c.ChildID
                WHERE s.DonorID = ? AND s.IsActive = 1
            """, (donor_id,))
            rows = cursor.fetchall()

            if rows:
                monthly_commitment.set(f"Monthly Commitment: Rs. {monthly_amount:.2f}")
                for row in rows:
                    child_name = row.ChildName
                    allocated = float(row.SponsorshipAmount or 0)
                    percentage = (allocated / monthly_amount * 100) if monthly_amount > 0 else 0
                    tree.insert('', 'end', values=(
                        child_name,
                        f"{percentage:.2f}%",
                        f"Rs. {allocated:.2f}"
                    ))
                    report_rows.append((child_name, percentage, allocated))

                print_btn.pack(pady=10)

            conn.close()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load sponsorship data:\n{e}")




    def print_report(data, donor_name, total_commitment):
        if not data:
            messagebox.showinfo("No Data", "No data to print.")
            return

        try:
            # Save in temporary folder
            temp_pdf = os.path.join(tempfile.gettempdir(), f"DonorChildMapping_{donor_name.replace(' ', '_')}.pdf")

            c = canvas.Canvas(temp_pdf, pagesize=A4)
            width, height = A4

            # Header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(180, height - 60, "Donor Sponsorship Allocation Report")
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 90, f"Donor: {donor_name}")
            c.drawString(50, height - 105, total_commitment)

            # Table headers
            y = height - 130
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Child Name")
            c.drawString(250, y, "Percentage")
            c.drawString(350, y, "Allocated Amount")

            # Table rows
            y -= 20
            c.setFont("Helvetica", 9)
            for row in data:
                c.drawString(50, y, row[0])
                c.drawString(250, y, f"{row[1]}%")
                c.drawString(350, y, f"Rs. {row[2]:.2f}")
                y -= 15
                if y < 40:  # New page if space is low
                    c.showPage()
                    y = height - 40

            c.save()

            # Auto open PDF
            open_file_cross_platform(temp_pdf)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # Bind
    center_combo.bind("<<ComboboxSelected>>", load_collectors)
    collector_combo.bind("<<ComboboxSelected>>", load_donors)
    donor_combo.bind("<<ComboboxSelected>>", load_report)

    # Initial load
    load_centers_collectors_donors()
    

# ============================================================

    
# ============================================================
# Donor Report
# ============================================================

def open_donor_report_form():
    report_win = tk.Toplevel()
    report_win.title("Donor Report" + ORG_SUFFIX)
    report_win.geometry("950x500")

    # Center window on screen
    ww, wh = 950, 500
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    # === Filter Frame ===
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)
   
    # Storage for donors to be printed
    donor_rows = []
    
    collector_dict = {}   # maps CollectorName -> CollectorID
    donor_dict = {}       # maps DonorName -> DonorID

    center_var = tk.StringVar()
    collector_var = tk.StringVar()
    donor_var = tk.StringVar()
    
    tk.Label(filter_frame, text="Center:").grid(row=0, column=0, padx=5)
    center_var = tk.StringVar()
    
    tk.Label(filter_frame, text="Collector:").grid(row=0, column=2, padx=5)
    collector_var = tk.StringVar()
    
    # Create the three comboboxes (define BEFORE handlers that use them)
    center_combo = ttk.Combobox(filter_frame, textvariable=center_var, state="readonly", width=28)
    center_combo.grid(row=0, column=1, padx=6, pady=6)

    collector_combo = ttk.Combobox(filter_frame, textvariable=collector_var, state="readonly", width=28)
    collector_combo.grid(row=0, column=3, padx=6, pady=6)

    # Print Button
    print_btn = tk.Button(
        report_win,
        text="Print Report",
        bg="green",
        fg="white",
        width=20,
        command=lambda: show_print_dialog()
    )
    print_btn.pack_forget()  # Hide initially

    # === Treeview for results ===
    columns = ("Donor Name", "Office Donor ID", "Collector Name")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")

    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=200 if col != "Collector Name" else 250)

    tree.pack(padx=10, pady=10, fill="both", expand=True)

    # Dictionaries for dropdown data
    center_dict = {}
    collector_dict_by_center = {}
    # Populate collector_dict on form load so it's not empty
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT CollectorID, FullName
            FROM tblCollectors
            WHERE 1=1 -- add filtering if needed
            ORDER BY FullName
        """)
        collectors = cursor.fetchall()
        conn.close()

        collector_dict.clear()
        for row in collectors:
            collector_dict[row.FullName] = row.CollectorID

        # Load into collector combo if you have one
        collector_combo['values'] = list(collector_dict.keys())

    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load collectors:\n{e}")


    # === LOAD DONORS ===
    def load_donors_for_report(event=None):
        tree.delete(*tree.get_children())
        donor_rows.clear()
        print_btn.pack_forget()

        selected_center = center_var.get()
        selected_collector = collector_var.get()

        if not selected_center or not selected_collector:
            return

        collector_id = collector_dict_by_center.get(selected_center, {}).get(selected_collector)
        if not collector_id:
            return

        try:
            conn = get_connection()
            cursor = conn.cursor()

            cursor.execute("""
                SELECT 
                    d.FullName AS DonorName,
                    d.OfficeDonorID,
                    col.FullName AS CollectorName,
                    cmt.CommitmentStartDate,
                    cmt.CommitmentEndDate,
                    cmt.MonthlyAmount
                FROM tblDonors d
                INNER JOIN tblCollectors col ON d.CollectorID = col.CollectorID
                INNER JOIN tblCommitments cmt ON d.DonorID = cmt.DonorID
                WHERE col.CollectorID = ?
                ORDER BY d.FullName, cmt.CommitmentStartDate
            """, (collector_id,))

            rows = cursor.fetchall()
            conn.close()

            if not rows:
                messagebox.showinfo("No Data", "No donors found for the selected collector.")
                return

            for row in rows:
                donor_name = row.DonorName
                office_id = row.OfficeDonorID
                collector_name = row.CollectorName
                commitment_period = f"{format_db_date(row.CommitmentStartDate, '%Y-%m-%d')} to {format_db_date(row.CommitmentEndDate, '%Y-%m-%d')}"
                monthly_amt = f"Rs. {row.MonthlyAmount:.2f}"

                # Insert into Treeview
                tree.insert('', 'end', values=(donor_name, office_id, collector_name))

                # Store for printing
                donor_rows.append({
                    "DonorName": donor_name,
                    "OfficeDonorID": office_id,
                    "CollectorName": collector_name,
                    "CommitmentPeriod": commitment_period,
                    "MonthlyAmount": monthly_amt
                })

            print_btn.pack(pady=10)

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load donor data:\n{e}")


    

    # === EVENTS ===
   # --- helpers (place inside open_donor_report_form) ---
    def _safe_fmt_date(dt):
        try:
            return format_db_date(dt, '%Y-%m-%d') if dt else ""
        except Exception:
            return ""

    def rebuild_tree_from_donor_rows():
        """Treeview shows one row per donor (grouping happens in PDF)."""
        tree.delete(*tree.get_children())
        seen = set()
        for r in donor_rows:
            key = (r["DonorName"], r.get("OfficeDonorID") or "", r.get("CollectorName") or "")
            if key in seen:
                continue
            seen.add(key)
            tree.insert("", "end", values=key)
        # Toggle print button
        if seen:
            print_btn.pack(pady=10)
        else:
            print_btn.pack_forget()


    # --- events (replace your existing ones) ---
    def on_center_change(event=None):
        """When center changes: populate collector list, clear donor_rows/tree.
        Optionally auto-select if only 1 collector."""
        selected_center = center_var.get()
        donor_rows.clear()
        tree.delete(*tree.get_children())
        print_btn.pack_forget()

        # reset collector combo
        collector_combo.set('')
        collector_combo['values'] = []

        if not selected_center:
            return

        # refresh collectors for this center (uses collector_dict_by_center built on form load)
        collectors_map = collector_dict_by_center.get(selected_center, {})
        collector_combo['values'] = list(collectors_map.keys())

        # If exactly one collector, auto-select and load
        if len(collectors_map) == 1:
            only_name = next(iter(collectors_map.keys()))
            collector_combo.set(only_name)
            on_collector_change()  # will fill donor_rows + tree


    def on_collector_change(event=None):
        """Load donors + commitments for the selected collector INTO donor_rows,
        then rebuild the Treeview from donor_rows. PDF reads donor_rows too."""
        selected_collector = collector_var.get()
        donor_rows.clear()
        tree.delete(*tree.get_children())
        print_btn.pack_forget()

        if not selected_collector:
            return

        # Resolve collector_id from either mapping (by center) or the flat dict
        collector_id = (
            collector_dict_by_center.get(center_var.get(), {}).get(selected_collector)
            or collector_dict.get(selected_collector)
        )
        if not collector_id:
            return

        try:
            conn = get_connection()
            cursor = conn.cursor()

            # Pull donors under this collector + any active commitments (left join so donors with no commitments still show)
            cursor.execute("""
                SELECT 
                    d.DonorID,
                    d.FullName            AS DonorName,
                    d.OfficeDonorID,
                    col.FullName          AS CollectorName,
                    c.CommitmentID,
                    c.MonthlyAmount,
                    c.CommitmentStartDate,
                    c.CommitmentEndDate
                FROM tblDonors d
                JOIN tblCollectors col ON d.CollectorID = col.CollectorID
                LEFT JOIN tblCommitments c 
                    ON c.DonorID = d.DonorID AND c.IsActive = 1
                WHERE d.CollectorID = ? AND d.IsActive = 1
                ORDER BY d.FullName, c.CommitmentStartDate
            """, (collector_id,))
            rows = cursor.fetchall()
            conn.close()

            # Build donor_rows as the single source of truth for both Treeview & PDF
            # Ensure each donor appears at least once, even if no commitments.
            current_donor_id = None
            had_commitment_for_donor = False

            # We'll accumulate per-donor; when donor changes and had no commitments, push a placeholder row.
            buffer_rows = []
            def _flush_placeholder_if_needed():
                nonlocal had_commitment_for_donor
                if current_donor_id is not None and not had_commitment_for_donor and buffer_rows:
                    # Add one placeholder row for "no commitments"
                    d = buffer_rows[-1]  # last seen donor info
                    donor_rows.append({
                        "DonorName":      d["DonorName"],
                        "OfficeDonorID":  d["OfficeDonorID"],
                        "CollectorName":  d["CollectorName"],
                        "CommitmentPeriod": "-",
                        "MonthlyAmount":    "-"
                    })
                had_commitment_for_donor = False

            for row in rows:
                donor_id = row.DonorID
                # donor boundary?
                if donor_id != current_donor_id:
                    _flush_placeholder_if_needed()
                    current_donor_id = donor_id
                    had_commitment_for_donor = False

                donor_info = {
                    "DonorName":     row.DonorName,
                    "OfficeDonorID": row.OfficeDonorID or "",
                    "CollectorName": row.CollectorName or ""
                }
                buffer_rows.append(donor_info)

                if row.CommitmentID:
                    had_commitment_for_donor = True
                    period = f"{_safe_fmt_date(row.CommitmentStartDate)} to {_safe_fmt_date(row.CommitmentEndDate)}".strip()
                    if period == " to ":
                        period = "-"
                    monthly = f"Rs. {float(row.MonthlyAmount):.2f}" if row.MonthlyAmount is not None else "-"
                    donor_rows.append({
                        **donor_info,
                        "CommitmentPeriod": period,
                        "MonthlyAmount": monthly
                    })

            # flush last donor if it had no commitments
            _flush_placeholder_if_needed()

            # Now rebuild the Treeview strictly from donor_rows
            rebuild_tree_from_donor_rows()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load donors & commitments:\n{e}")







    center_combo.bind("<<ComboboxSelected>>", on_center_change)
    collector_combo.bind("<<ComboboxSelected>>", on_collector_change)

    
    
    # === LOAD CENTERS & COLLECTORS ===
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT CenterID, CenterName FROM tblCenters")
        centers = cursor.fetchall()
        center_dict = {row.CenterName: row.CenterID for row in centers}
        center_combo['values'] = list(center_dict.keys())

        for center in centers:
            cursor.execute("SELECT CollectorID, FullName FROM tblCollectors WHERE CenterID=?", (center.CenterID,))
            collectors = cursor.fetchall()
            collector_dict_by_center[center.CenterName] = {c.FullName: c.CollectorID for c in collectors}
        
        # Default first center and collector selected
        if centers:
            first_center = centers[0].CenterName
            center_combo.set(first_center)
            collectors = collector_dict_by_center.get(first_center, {})
            collector_combo['values'] = list(collectors.keys())

            if collectors:
                first_collector = next(iter(collectors))
                collector_combo.set(first_collector)
                
                load_donors_for_report()

                # Now explicitly check to show print button if there are donors
                if donor_rows:
                    print_btn.pack(pady=10)
                else:
                    print_btn.pack_forget()
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load centers and collectors:\n{e}")
        
        # === PRINT REPORT ===

    def show_print_dialog():
        if not donor_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return

        try:
            import tempfile
            import threading

            # Create a temporary PDF file in system temp folder
            tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            file_path = tmp_file.name
            tmp_file.close()  # Close so ReportLab can write into it

            # Create PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Collector Name (same for all rows)
            collector_name = donor_rows[0].get("CollectorName", "") if donor_rows else ""

            # Title
            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(200, height - 60, "Donor Report with Commitments")

            # Current date/time in dd/mm/yyyy
            c.setFont("Helvetica", 8)
            now = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
            c.drawString(450, height - 80, f"Generated: {now}")

            y = height - 110

            # Draw Collector Name
            c.setFont("Helvetica", 12)
            if collector_name:
                c.drawString(50, height - 95, f"Collector: {collector_name}")

            # Group donors
            current_donor = None

            for row in donor_rows:
                donor_name = row["DonorName"]
                office_id = row["OfficeDonorID"]
                period = row["CommitmentPeriod"]
                monthly_amount = row["MonthlyAmount"]

                # Convert CommitmentPeriod date(s) to dd/mm/yyyy if needed
                try:
                    # If formatted like YYYY-MM-DD to YYYY-MM-DD
                    if " to " in period:
                        p1, p2 = period.split(" to ")
                        p1 = datetime.strptime(p1.strip(), "%Y-%m-%d").strftime("%d/%m/%Y")
                        p2 = datetime.strptime(p2.strip(), "%Y-%m-%d").strftime("%d/%m/%Y")
                        period = f"{p1} to {p2}"
                    else:
                        # Single date
                        period = datetime.strptime(period.strip(), "%Y-%m-%d").strftime("%d/%m/%Y")
                except:
                    pass  # If not a date format, leave as is

                if donor_name != current_donor:
                    current_donor = donor_name

                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(50, y, f"Donor: {donor_name} ({office_id})")
                    y -= 15

                    c.setFont("Helvetica-Bold", 9)
                    c.drawString(90, y, "Commitment Period")
                    c.drawString(250, y, "Monthly Amount")
                    y -= 12

                # Commitment detail
                c.setFont("Helvetica", 9)
                c.drawString(90, y, period)
                c.drawString(250, y, str(monthly_amount))
                y -= 12

                if y < 40:  # Page break
                    c.showPage()
                    y = height - 40

            c.save()

            # Auto open the temporary PDF
            os.startfile(file_path)

            # Auto delete file after 15 seconds (enough time for PDF viewer to open)
            def delete_temp():
                try:
                    time.sleep(15)
                    if os.path.exists(file_path):
                        os.remove(file_path)
                except:
                    pass

            threading.Thread(target=delete_temp, daemon=True).start()

            #messagebox.showinfo("Success", "PDF generated in temporary memory.\nIt will auto-delete shortly.")

        except Exception as e:
            messagebox.showerror("Print Error", str(e))


    
# ============================================================

# ============================================================
# Collector List
# ============================================================
def open_collector_list_form():
    report_win = tk.Toplevel()
    report_win.title("Collector Report" + ORG_SUFFIX)
    report_win.geometry("900x500")

    ww, wh = 900, 500
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Select Center:").grid(row=0, column=0, padx=5)
    center_var = tk.StringVar()
    center_combo = ttk.Combobox(filter_frame, textvariable=center_var, width=40, state="readonly")
    center_combo.grid(row=0, column=1, padx=5)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20, command=lambda: show_print_dialog())
    print_btn.pack_forget()  # Hide initially

    columns = ("Collector Name", "Contact Number", "Address", "Center Name")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=200)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    collector_rows = []



    def show_print_dialog():
        if not collector_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return

        try:
            # Get user's Downloads folder
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
            if not os.path.exists(downloads_folder):
                os.makedirs(downloads_folder)

            # Create a unique filename with timestamp
            file_name = f"Collector_List_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_folder, file_name)

            # Generate PDF
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(200, height - 40, "Alkhidmat Khawateen Trust")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(240, height - 55, "Collector List")

            c.setFont("Helvetica", 10)
            selected_center = center_var.get()
            if selected_center == "All Centers":
                c.drawString(40, height - 80, "Center: All Centers")
            else:
                c.drawString(40, height - 80, f"Center: {selected_center}")

            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(40, y, "Name")
            c.drawString(200, y, "Contact Number")
            c.drawString(350, y, "Address")

            y -= 20
            c.setFont("Helvetica", 9)
            for row in collector_rows:
                c.drawString(40, y, str(row.FullName))
                c.drawString(200, y, str(row.ContactNumber))
                c.drawString(350, y, str(row.Address))
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()

            # Confirm and open the PDF
            messagebox.showinfo("Success", f"PDF saved successfully:\n{file_path}")
            os.startfile(file_path)  # Auto open

        except Exception as e:
            messagebox.showerror("Print Error", str(e))
            


    def load_collectors_for_center(event):
        tree.delete(*tree.get_children())
        collector_rows.clear()

        selected_center = center_var.get()
        if not selected_center:
            return

        try:
            conn = get_connection()
            cursor = conn.cursor()

            if selected_center == "All Centers":
                query = """
                    SELECT col.FullName, col.ContactNumber, col.Address, ce.CenterName
                    FROM tblCollectors col
                    JOIN tblCenters ce ON col.CenterID = ce.CenterID
                    ORDER BY ce.CenterName, col.FullName
                """
                cursor.execute(query)
            else:
                query = """
                    SELECT col.FullName, col.ContactNumber, col.Address, ce.CenterName
                    FROM tblCollectors col
                    JOIN tblCenters ce ON col.CenterID = ce.CenterID
                    WHERE ce.CenterName = ?
                    ORDER BY col.FullName
                """
                cursor.execute(query, (selected_center,))

            rows = cursor.fetchall()
            collector_rows.extend(rows)

            for row in rows:
                tree.insert('', 'end', values=(row.FullName, row.ContactNumber, row.Address, row.CenterName))

            if rows:
                print_btn.pack(pady=10)
            else:
                print_btn.pack_forget()

            conn.close()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load collectors:\n{e}")

    # Load centers into ComboBox
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT CenterName FROM tblCenters")
        centers = [row.CenterName for row in cursor.fetchall()]
        centers.insert(0, "All Centers")  # Add All Centers at top
        center_combo['values'] = centers
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load centers:\n{e}")

    center_combo.bind("<<ComboboxSelected>>", load_collectors_for_center)

# ============================================================

# ============================================================
# Collector Report
# ============================================================

def open_collector_report_form():
    report_win = tk.Toplevel()
    report_win.title("Collector Report" + ORG_SUFFIX)
    report_win.geometry("950x550")

    ww, wh = 950, 550
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient()
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())
    
    collector_rows = []
    grand_total_collected = 0

    
    def collectors_for_collected_amount():
        nonlocal collector_rows, grand_total_collected  # keep if still inside outer func
        #nonlocal grand_total_collected
        tree.delete(*tree.get_children())
        collector_rows.clear()
        grand_total_collected = 0

        selected_center = center_var.get()
        start = from_date.get_date().strftime("%Y-%m-%d")
        end = to_date.get_date().strftime("%Y-%m-%d")
        
        if not selected_center:
            return
        
        try:
            conn = get_connection()
            cursor = conn.cursor()

            if selected_center == "All Centers":
                query = """
                    SELECT col.FullName, col.ContactNumber, col.Address, 
                           IFNULL(SUM(CAST(don.DonationAmount AS REAL)), 0) AS TotalCollected
                    FROM tblDonations don
                    JOIN tblCollectors col ON don.CollectorID = col.CollectorID
                    WHERE don.DonationDate BETWEEN ? AND ?
                    GROUP BY col.FullName, col.ContactNumber, col.Address
                    ORDER BY col.FullName
                """
                cursor.execute(query, (start, end))
                
            else:
                query = """
                    SELECT col.FullName, col.ContactNumber, col.Address, 
                           IFNULL(SUM(CAST(don.DonationAmount AS REAL)), 0) AS TotalCollected
                    FROM tblDonations don
                    JOIN tblCollectors col ON don.CollectorID = col.CollectorID
                    JOIN tblCenters ce ON col.CenterID = ce.CenterID
                    WHERE ce.CenterName = ? AND don.DonationDate BETWEEN ? AND ?
                    GROUP BY col.FullName, col.ContactNumber, col.Address
                    ORDER BY col.FullName
                """
                cursor.execute(query, (selected_center, start, end))
                
                
            rows = cursor.fetchall()
            
            for row in rows:
                total_collected = row.TotalCollected or 0
                grand_total_collected += total_collected
                collector_rows.append({
                    "Name": row.FullName,
                    "Contact": row.ContactNumber,
                    "Address": row.Address,
                    "Collected": total_collected
                })
                tree.insert('', 'end', values=(
                    row.FullName, row.ContactNumber, row.Address,
                    f"Rs.{total_collected:.2f}"
                ))
                
            if rows:
                print_btn.pack(pady=10)
            else:
                print_btn.pack_forget()

            conn.close()
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load collectors:\n{e}")    
    
    
    

    def show_print_dialog():
        messagebox.showinfo("","L")
        if not collector_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return

        downloads_path = str(Path.home() / "Downloads")
        filename = f"DonorWiseCommitment_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(downloads_path, filename)

        #file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])
        if not file_path:
            return

        try:
            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(200, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(230, height - 55, "Collector Report")
            c.setFont("Helvetica", 10)
            c.drawString(40, height - 70, f"From: {from_date.get_date()} To: {to_date.get_date()}")
            c.drawString(40, height - 85, f"Center: {center_var.get()}")

            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(40, y, "Name")
            c.drawString(200, y, "Contact")
            c.drawString(330, y, "Address")
            c.drawString(500, y, "Collected")

            y -= 20
            c.setFont("Helvetica", 9)
            for row in collector_rows:
                c.drawString(40, y, str(row["Name"]))
                c.drawString(200, y, str(row["Contact"]))
                c.drawString(330, y, str(row["Address"]))
                c.drawString(500, y, f"Rs.{row['Collected']:.2f}")
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            y -= 10
            c.setFont("Helvetica-Bold", 11)
            c.drawString(40, y, f"Grand Total Collected: Rs.{grand_total_collected:.2f}")

            c.save()
            time.sleep(1)
            messagebox.showinfo("Success", f"PDF saved:\n{file_path}")
            os.startfile(file_path)


        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    


# Filter frame
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="From Date:").grid(row=0, column=0, padx=5)
    from_date = DateEntry(filter_frame, width=12)
    from_date.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="To Date:").grid(row=0, column=2, padx=5)
    to_date = DateEntry(filter_frame, width=12)
    to_date.grid(row=0, column=3, padx=5)

    tk.Label(filter_frame, text="Select Center:").grid(row=0, column=4, padx=5)
    center_var = tk.StringVar()
    center_combo = ttk.Combobox(filter_frame, textvariable=center_var, width=25, state="readonly")
    center_combo.grid(row=0, column=5, padx=5)

    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white", command=lambda: collectors_for_collected_amount())
    generate_btn.grid(row=0, column=6, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20, command=lambda: show_print_dialog())
    print_btn.pack_forget()

    # Treeview
    columns = ("Collector Name", "Contact Number", "Address", "Collected Amount")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=220)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

        # Load centers into ComboBox
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT CenterName FROM tblCenters")
        centers = [row.CenterName for row in cursor.fetchall()]
        centers.insert(0, "All Centers")
        center_combo['values'] = centers
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load centers:\n{e}")
        
    
# ============================================================
# Donor-wise Commitments
# ============================================================
def open_donor_wise_commitment_report():
    report_win = tk.Toplevel(root)
    report_win.title("Donor Wise Commitment Report" + ORG_SUFFIX)

    # --- Center Window ---
    ww, wh = 950, 550
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    # Variables
    status_var = tk.StringVar(value="Active")  # default
    report_rows = []

    # --- Filters Frame ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Donor Name:").grid(row=0, column=0, padx=5)

    donor_var = tk.StringVar()
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=0, column=1, padx=5)

    # Radio Buttons for status
    tk.Label(filter_frame, text="Commitment Status:").grid(row=0, column=2, padx=5)
    rb_active = tk.Radiobutton(filter_frame, text="Active", variable=status_var, value="Active")
    rb_inactive = tk.Radiobutton(filter_frame, text="Inactive", variable=status_var, value="Inactive")
    rb_both = tk.Radiobutton(filter_frame, text="Both", variable=status_var, value="Both")
    rb_active.grid(row=0, column=3, padx=2)
    rb_inactive.grid(row=0, column=4, padx=2)
    rb_both.grid(row=0, column=5, padx=2)

    # Buttons
    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white")
    generate_btn.grid(row=0, column=6, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20)
    print_btn.pack_forget()

    # --- Treeview ---
    columns = ("Donor Name", "Office Donor ID", "Start Date", "End Date", "Monthly Amount", "Notes")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    # --- Load all donors for Autocomplete ---
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FullName FROM tblDonors ORDER BY FullName")
        donor_names = [row.FullName for row in cursor.fetchall()]
        conn.close()
        donor_combo.set_completion_list(donor_names)
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    # --- Fetch Data ---
    def collectors_for_commitments():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows = []

        donor_filter = donor_var.get().strip()
        status_filter = status_var.get()

        try:
            conn = get_connection()
            cursor = conn.cursor()

            # Build WHERE conditions
            where_clauses = []
            if status_filter == "Active":
                where_clauses.append("c.IsActive = 1")
            elif status_filter == "Inactive":
                where_clauses.append("c.IsActive = 0")

            if donor_filter:
                where_clauses.append("d.FullName = ?")  # exact match from combobox

            where_sql = " AND ".join(where_clauses)
            if where_sql:
                where_sql = "WHERE " + where_sql

            query = f"""
                SELECT d.FullName, d.OfficeDonorID,
                       c.CommitmentStartDate, c.CommitmentEndDate,
                       c.MonthlyAmount, c.Notes
                FROM tblCommitments c
                JOIN tblDonors d ON c.DonorID = d.DonorID
                {where_sql}
                ORDER BY d.FullName, c.CommitmentStartDate
            """

            params = []
            if donor_filter:
                params.append(donor_filter)

            cursor.execute(query, tuple(params))
            rows = cursor.fetchall()
            conn.close()

            for row in rows:
                start_date = format_db_date(row.CommitmentStartDate, "%d/%m/%Y")
                end_date = format_db_date(row.CommitmentEndDate, "%d/%m/%Y")
                report_rows.append({
                    "DonorName": row.FullName,
                    "OfficeDonorID": row.OfficeDonorID,
                    "StartDate": start_date,
                    "EndDate": end_date,
                    "MonthlyAmount": float(row.MonthlyAmount or 0),
                    "Notes": row.Notes or ""
                })
                tree.insert("", "end", values=(
                    row.FullName,
                    row.OfficeDonorID,
                    start_date,
                    end_date,
                    f"Rs.{float(row.MonthlyAmount or 0):.2f}",
                    row.Notes or ""
                ))

            if rows:
                print_btn.pack(pady=10)
            else:
                print_btn.pack_forget()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load commitments:\n{e}")

    # --- Print PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return
        try:
            downloads_path = str(Path.home() / "Downloads")
            filename = f"DonorWiseCommitment_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_path, filename)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(200, height - 40, "Alkhidmat Khawateen Trust")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(180, height - 60, "Donor Wise Commitment Report")

            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(200, y, "Office ID")
            c.drawString(300, y, "Start Date")
            c.drawString(380, y, "End Date")
            c.drawString(460, y, "Monthly")
            c.drawString(540, y, "Notes")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r["DonorName"])
                c.drawString(200, y, str(r["OfficeDonorID"]))
                c.drawString(300, y, r["StartDate"])
                c.drawString(380, y, r["EndDate"])
                c.drawString(460, y, f"Rs.{r['MonthlyAmount']:.2f}")
                c.drawString(540, y, r["Notes"])
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            time.sleep(1)
            messagebox.showinfo("Success", f"PDF saved:\n{file_path}")
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # Bind buttons
    generate_btn.config(command=collectors_for_commitments)
    print_btn.config(command=print_report)

    # --- Bind Enter key ---
    report_win.bind("<Return>", lambda e: collectors_for_commitments())

    # Focus on donor combo
    donor_combo.focus_set()


# ============================================================
# Commitment vs Actual Donations Report
# ============================================================

def open_commitment_vs_actual_donations_report():
    report_win = tk.Toplevel(root)
    report_win.title("Commitment vs Actual Donations Report" + ORG_SUFFIX)

    # --- Center Window ---
    ww, wh = 1050, 550
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    # Variables
    status_var = tk.StringVar(value="Active")  # default
    report_rows = []

    # --- Filters Frame ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Donor Name:").grid(row=0, column=0, padx=5)

    donor_var = tk.StringVar()
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=0, column=1, padx=5)

    # Radio Buttons for status
    tk.Label(filter_frame, text="Commitment Status:").grid(row=0, column=2, padx=5)
    rb_active = tk.Radiobutton(filter_frame, text="Active", variable=status_var, value="Active")
    rb_inactive = tk.Radiobutton(filter_frame, text="Inactive", variable=status_var, value="Inactive")
    rb_both = tk.Radiobutton(filter_frame, text="Both", variable=status_var, value="Both")
    rb_active.grid(row=0, column=3, padx=2)
    rb_inactive.grid(row=0, column=4, padx=2)
    rb_both.grid(row=0, column=5, padx=2)

    # Buttons
    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white")
    generate_btn.grid(row=0, column=6, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20)
    print_btn.pack_forget()

    # --- Treeview ---
    columns = ("Donor", "Office ID", "Start Date", "End Date",
               "Monthly Commitment", "Total Donations", "Status", "Notes")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=130)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    # --- Load all donors for Autocomplete ---
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FullName FROM tblDonors ORDER BY FullName")
        donor_names = [row.FullName for row in cursor.fetchall()]
        conn.close()
        donor_combo.set_completion_list(donor_names)
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    # --- Fetch Data ---
    def load_commitment_vs_donations():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows = []

        donor_filter = donor_var.get().strip()
        status_filter = status_var.get()

        try:
            conn = get_connection()
            cursor = conn.cursor()

            # ---- Build WHERE clause ----
            where_clauses = []
            params = []

            if status_filter == "Active":
                where_clauses.append("c.IsActive = 1")
            elif status_filter == "Inactive":
                where_clauses.append("c.IsActive = 0")

            if donor_filter:
                where_clauses.append("d.FullName = ?")
                params.append(donor_filter)

            where_sql = "WHERE " + " AND ".join(where_clauses) if where_clauses else ""

            # ---- Fetch commitments ----
            cursor.execute(f"""
                SELECT 
                    d.DonorID, d.FullName, d.OfficeDonorID,
                    c.CommitmentID,
                    c.CommitmentStartDate, c.CommitmentEndDate,
                    c.MonthlyAmount, c.Notes
                FROM tblCommitments c
                JOIN tblDonors d ON c.DonorID = d.DonorID
                {where_sql}
                ORDER BY d.FullName, c.CommitmentStartDate
            """, params)

            commitments = cursor.fetchall()

            '''for row in commitments:
                donor_id = row.DonorID
                start_date = row.CommitmentStartDate
                end_date = row.CommitmentEndDate
                monthly_amount = float(row.MonthlyAmount or 0)

                # ---- Calculate months ----
                months = (
                    (end_date.year - start_date.year) * 12 +
                    (end_date.month - start_date.month) + 1
                )'''
            
            for row in commitments:
                donor_id = row.DonorID
                
                # ---- FIXED: Parse date strings to datetime objects ----
                try:
                    if isinstance(row.CommitmentStartDate, str):
                        start_date = datetime.strptime(row.CommitmentStartDate[:10], "%Y-%m-%d")
                    else:
                        start_date = row.CommitmentStartDate
                        
                    if isinstance(row.CommitmentEndDate, str):
                        end_date = datetime.strptime(row.CommitmentEndDate[:10], "%Y-%m-%d")
                    else:
                        end_date = row.CommitmentEndDate
                except Exception as e:
                    print(f"Date parsing error for donor {row.FullName}: {e}")
                    continue  # Skip this row if dates are invalid
                
                monthly_amount = float(row.MonthlyAmount or 0)

                # ---- Calculate months ----
                months = (
                    (end_date.year - start_date.year) * 12 +
                    (end_date.month - start_date.month) + 1
                )

                expected_total = monthly_amount * months

                # ---- Sum donations MONTHLY ----
                cursor.execute("""
                    SELECT 
                        CAST(strftime('%Y', DonationDate) AS INTEGER) AS Yr,
                        CAST(strftime('%m', DonationDate) AS INTEGER) AS Mn,
                        SUM(CAST(DonationAmount AS REAL)) AS MonthTotal
                    FROM tblDonations
                    WHERE DonorID = ?
                    AND DonationDate BETWEEN ? AND ?
                    GROUP BY CAST(strftime('%Y', DonationDate) AS INTEGER), CAST(strftime('%m', DonationDate) AS INTEGER)
                """, (donor_id, start_date, end_date))

                monthly_donations = cursor.fetchall()
                total_donations = sum(float(r.MonthTotal) for r in monthly_donations)

                # ---- Determine status ----
                if total_donations >= expected_total:
                    status = "Met"
                elif total_donations > 0:
                    status = "Partial"
                else:
                    status = "Not Met"

                start_fmt = start_date.strftime("%d/%m/%Y")
                end_fmt = end_date.strftime("%d/%m/%Y")

                report_rows.append({
                    "DonorName": row.FullName,
                    "OfficeDonorID": row.OfficeDonorID,
                    "StartDate": start_fmt,
                    "EndDate": end_fmt,
                    "MonthlyAmount": monthly_amount,
                    "ExpectedTotal": expected_total,
                    "TotalDonations": total_donations,
                    "Status": status,
                    "Notes": row.Notes or ""
                })

                tree.insert("", "end", values=(
                    row.FullName,
                    row.OfficeDonorID,
                    start_fmt,
                    end_fmt,
                    f"Rs.{monthly_amount:.2f}",
                    f"Rs.{total_donations:.2f}",
                    status,
                    row.Notes or ""
                ))

            print_btn.pack(pady=10) if commitments else print_btn.pack_forget()
            conn.close()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load data:\n{e}")


    # --- Print PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return
        try:
            downloads_path = str(Path.home() / "Downloads")
            filename = f"CommitmentVsDonations_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_path, filename)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(160, height - 60, "Commitment vs Actual Donations Report")

            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(200, y, "Office ID")
            c.drawString(280, y, "Start Date")
            c.drawString(360, y, "End Date")
            c.drawString(440, y, "Monthly")
            c.drawString(520, y, "Total Donated")
            c.drawString(620, y, "Status")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r["DonorName"])
                c.drawString(200, y, str(r["OfficeDonorID"]))
                c.drawString(280, y, r["StartDate"])
                c.drawString(360, y, r["EndDate"])
                c.drawString(440, y, f"Rs.{r['MonthlyAmount']:.2f}")
                c.drawString(520, y, f"Rs.{r['TotalDonations']:.2f}")
                c.drawString(620, y, r["Status"])
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            time.sleep(1)
            messagebox.showinfo("Success", f"PDF saved:\n{file_path}")
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # Bind buttons
    generate_btn.config(command=load_commitment_vs_donations)
    print_btn.config(command=print_report)

    # --- Bind Enter key ---
    report_win.bind("<Return>", lambda e: load_commitment_vs_donations())

    # Focus on donor combo
    donor_combo.focus_set()

    
# ============================================================

def open_unsponsored_Partially_sponsored_report():
    report_win = tk.Toplevel(root)
    report_win.title("Unsponsored / Partially Sponsored Children Report" + ORG_SUFFIX)

    # --- Center Window ---
    ww, wh = 1050, 550
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    filter_var = tk.StringVar(value="All")
    report_rows = []

    # --- Filter Frame ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Filter:").grid(row=0, column=0, padx=5)
    tk.Radiobutton(filter_frame, text="All", variable=filter_var, value="All").grid(row=0, column=1, padx=2)
    tk.Radiobutton(filter_frame, text="Unsponsored", variable=filter_var, value="Unsponsored").grid(row=0, column=2, padx=2)
    tk.Radiobutton(filter_frame, text="Partially Sponsored", variable=filter_var, value="Partial").grid(row=0, column=3, padx=2)

    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white")
    generate_btn.grid(row=0, column=4, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20)
    print_btn.pack_forget()

    # --- Treeview ---
    columns = ("Child", "Required", "Sponsored", "% Sponsored", "Still Need", "Sponsored Till")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=160)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    # --- SAFE CONVERTERS ---
    def to_float(val):
        try:
            return float(val)
        except:
            return 0.0

    def to_int(val):
        try:
            return int(float(val))
        except:
            return 0

    # --- Load Data ---
    def load_children_status():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows = []

        try:
            conn = get_connection()
            cursor = conn.cursor()

            cursor.execute("""
                SELECT ChildID, FullName, 
                       CAST(ChildRequiredAmount AS REAL) AS Required
                FROM tblChildren
                WHERE Status = 1
            """)
            children = cursor.fetchall()

            for child in children:
                child_id = child.ChildID
                child_name = child.FullName
                required = to_float(child.Required)

                cursor.execute("""
                    SELECT SponsorshipAmount, Percentage, CommitmentID
                    FROM tblSponsorships
                    WHERE ChildID = ? AND IsActive = 1
                """, (child_id,))
                sponsorships = cursor.fetchall()

                sponsored_amt = 0.0
                total_percent = 0
                till_date = None

                for s in sponsorships:
                    amt = to_float(s.SponsorshipAmount)
                    pct = to_int(s.Percentage)   # ✅ FIXED HERE

                    sponsored_amt += amt
                    total_percent += pct

                    if s.CommitmentID:
                        cursor.execute("""
                            SELECT CommitmentEndDate 
                            FROM tblCommitments 
                            WHERE CommitmentID = ?
                        """, (s.CommitmentID,))
                        commit = cursor.fetchone()

                        if commit and commit.CommitmentEndDate:
                            if not till_date or commit.CommitmentEndDate > till_date:
                                till_date = commit.CommitmentEndDate

                still_need = max(required - sponsored_amt, 0)

                # --- STATUS ---
                if total_percent == 0:
                    status = "Unsponsored"
                elif total_percent < 100:
                    status = "Partial"
                else:
                    status = "Fully"

                if filter_var.get() == "Unsponsored" and status != "Unsponsored":
                    continue
                if filter_var.get() == "Partial" and status != "Partial":
                    continue

                till_fmt = format_db_date(till_date, "%d/%m/%Y") if till_date else "No Commitment"

                report_rows.append({
                    "Child": child_name,
                    "Required": required,
                    "Sponsored": sponsored_amt,
                    "Percent": total_percent,
                    "StillNeed": still_need,
                    "TillDate": till_fmt
                })

                tree.insert("", "end", values=(
                    child_name,
                    f"Rs.{required:.2f}",
                    f"Rs.{sponsored_amt:.2f}",
                    f"{total_percent}%",
                    f"Rs.{still_need:.2f}",
                    till_fmt
                ))

            if report_rows:
                print_btn.pack(pady=10)
            else:
                print_btn.pack_forget()

            conn.close()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load data:\n{e}")

    # --- Print PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import A4
            from datetime import datetime
            from pathlib import Path
            import os, time

            downloads_path = str(Path.home() / "Downloads")
            filename = f"UnsponsoredChildren_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_path, filename)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            # Header
            c.setFont("Helvetica-Bold", 16)
            c.drawString(150, height - 40, "Alkhidmat Khawateen Trust")

            c.setFont("Helvetica-Bold", 12)
            c.drawString(140, height - 60, "Unsponsored / Partially Sponsored Children Report")

            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

            # Column Headers
            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "S.No")          # NEW: Serial number column
            c.drawString(90, y, "Child")         # Shifted right
            c.drawString(230, y, "Required")     # Shifted right
            c.drawString(320, y, "Sponsored")    # Shifted right
            c.drawString(410, y, "% Sponsored")  # Shifted right
            c.drawString(500, y, "Still Need")   # Shifted right

            # Data Rows
            y -= 20
            c.setFont("Helvetica", 9)
            
            serial_number = 1  # Initialize serial number counter

            for r in report_rows:
                c.drawString(50, y, str(serial_number))                    # NEW: Serial number
                c.drawString(90, y, r["Child"])                            # Shifted right
                c.drawString(230, y, f"Rs.{r['Required']:.2f}")            # Shifted right
                c.drawString(320, y, f"Rs.{r['Sponsored']:.2f}")           # Shifted right
                c.drawString(410, y, f"{r['Percent']}%")                   # Shifted right
                c.drawString(500, y, f"Rs.{r['StillNeed']:.2f}")           # Shifted right

                serial_number += 1  # Increment counter
                y -= 15
                
                if y < 40:
                    c.showPage()
                    y = height - 40
                    # Redraw headers on new page
                    c.setFont("Helvetica-Bold", 10)
                    c.drawString(50, y, "S.No")
                    c.drawString(90, y, "Child")
                    c.drawString(230, y, "Required")
                    c.drawString(320, y, "Sponsored")
                    c.drawString(410, y, "% Sponsored")
                    c.drawString(500, y, "Still Need")
                    y -= 20
                    c.setFont("Helvetica", 9)

            c.save()
            time.sleep(1)

            messagebox.showinfo("Success", f"PDF saved:\n{file_path}")
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # Bind buttons
    generate_btn.config(command=load_children_status)
    print_btn.config(command=print_report)

    report_win.bind("<Return>", lambda e: load_children_status())
    generate_btn.focus_set()

    
def open_expired_commitments_report():
    report_win = tk.Toplevel(root)
    report_win.title("Expired Commitments Report" + ORG_SUFFIX)

    # --- Center Window ---
    ww, wh = 1000, 550
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    filter_var = tk.StringVar(value="All")
    report_rows = []

    # --- Filter Frame ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Donor:").grid(row=0, column=0, padx=5)
    donor_var = tk.StringVar()
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="Filter:").grid(row=0, column=2, padx=5)
    tk.Radiobutton(filter_frame, text="All", variable=filter_var, value="All").grid(row=0, column=3, padx=2)
    tk.Radiobutton(filter_frame, text="Active Donors Only", variable=filter_var, value="Active").grid(row=0, column=4, padx=2)
    tk.Radiobutton(filter_frame, text="Inactive Donors Only", variable=filter_var, value="Inactive").grid(row=0, column=5, padx=2)

    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white")
    generate_btn.grid(row=0, column=6, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20)
    print_btn.pack_forget()

    # --- Treeview ---
    columns = ("Donor", "OfficeDonorID", "Start Date", "End Date", "Monthly Amount", "Notes", "Donor Status")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=140)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    # --- Load donors into auto-search ---
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FullName FROM tblDonors ORDER BY FullName")
        donors = [row.FullName for row in cursor.fetchall()]
        donor_combo.set_completion_list(donors)
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    # --- Load Expired Commitments ---
    def load_expired_commitments():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows = []

        selected_donor = donor_var.get().strip()

        try:
            conn = get_connection()
            cursor = conn.cursor()

            query = """
                SELECT d.FullName, d.OfficeDonorID, c.CommitmentStartDate, 
                       c.CommitmentEndDate, c.MonthlyAmount, c.Notes, d.IsActive
                FROM tblCommitments c
                JOIN tblDonors d ON c.DonorID = d.DonorID
                WHERE c.IsActive = 1
                  AND c.CommitmentEndDate < date('now')
            """

            params = []
            if selected_donor:
                query += " AND d.FullName = ?"
                params.append(selected_donor)

            query += " ORDER BY d.FullName, c.CommitmentEndDate"

            cursor.execute(query, tuple(params))
            rows = cursor.fetchall()
            conn.close()

            for row in rows:
                start = format_db_date(row.CommitmentStartDate, "%d/%m/%Y")
                end = format_db_date(row.CommitmentEndDate, "%d/%m/%Y")
                donor_status = "Active" if row.IsActive else "Inactive"

                # Radio filter
                if filter_var.get() == "Active" and not row.IsActive:
                    continue
                if filter_var.get() == "Inactive" and row.IsActive:
                    continue

                report_rows.append({
                    "Donor": row.FullName,
                    "OfficeDonorID": row.OfficeDonorID,
                    "Start": start,
                    "End": end,
                    "Monthly": row.MonthlyAmount,
                    "Notes": row.Notes or "",
                    "Status": donor_status
                })

                tree.insert("", "end", values=(
                    row.FullName, row.OfficeDonorID, start, end,
                    f"Rs.{row.MonthlyAmount:.2f}" if row.MonthlyAmount else "0.00",
                    row.Notes or "", donor_status
                ))

            if report_rows:
                print_btn.pack(pady=10)
            else:
                print_btn.pack_forget()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load commitments:\n{e}")

    # --- Print PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return
        try:
            downloads_path = str(Path.home() / "Downloads")
            filename = f"ExpiredCommitments_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_path, filename)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(200, height - 60, "Expired Commitments Report")
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(150, y, "OfficeID")
            c.drawString(230, y, "Start")
            c.drawString(310, y, "End")
            c.drawString(390, y, "Monthly")
            c.drawString(470, y, "Notes")
            c.drawString(550, y, "Status")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r["Donor"])
                c.drawString(150, y, str(r["OfficeDonorID"]))
                c.drawString(230, y, r["Start"])
                c.drawString(310, y, r["End"])
                c.drawString(390, y, f"Rs.{r['Monthly']:.2f}" if r["Monthly"] else "0.00")
                c.drawString(470, y, r["Notes"][:30])  # short notes
                c.drawString(550, y, r["Status"])
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            time.sleep(1)
            messagebox.showinfo("Success", f"PDF saved:\n{file_path}")
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # --- Bind Buttons ---
    generate_btn.config(command=load_expired_commitments)
    print_btn.config(command=print_report)

    # Enter key triggers Generate
    report_win.bind("<Return>", lambda e: load_expired_commitments())
    generate_btn.focus_set()

    
def open_upcoming_expiry_report():
    
    report_win = tk.Toplevel(root)
    report_win.title("Upcoming Expiry Commitments Report" + ORG_SUFFIX)

    # --- Center window ---
    ww, wh = 1000, 550
    sw = report_win.winfo_screenwidth()
    sh = report_win.winfo_screenheight()
    x = (sw // 2) - (ww // 2)
    y = (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")

    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    months_var = tk.IntVar(value=1)
    donor_var = tk.StringVar()
    report_rows = []

    # --- Filter Frame ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="Donor:").grid(row=0, column=0, padx=5)
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="Expiring in:").grid(row=0, column=2, padx=5)
    tk.Radiobutton(filter_frame, text="1 Month", variable=months_var, value=1).grid(row=0, column=3, padx=2)
    tk.Radiobutton(filter_frame, text="2 Months", variable=months_var, value=2).grid(row=0, column=4, padx=2)
    tk.Radiobutton(filter_frame, text="3 Months", variable=months_var, value=3).grid(row=0, column=5, padx=2)

    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white")
    generate_btn.grid(row=0, column=6, padx=10)

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20)
    print_btn.pack_forget()

    # --- Treeview ---
    columns = ("Donor", "OfficeDonorID", "Start Date", "End Date", "Monthly Amount", "Notes")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150)
    tree.pack(padx=10, pady=10, fill="both", expand=True)

    # --- Load donors into auto-search ---
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FullName FROM tblDonors ORDER BY FullName")
        donors = [row.FullName for row in cursor.fetchall()]
        donor_combo.set_completion_list(donors)
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    # --- Load upcoming expiring commitments ---
    def load_upcoming_commitments():
        
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows = []

        selected_donor = donor_var.get().strip()
        months = months_var.get()

        try:
            conn = get_connection()
            cursor = conn.cursor()

            query = """
                SELECT d.FullName, d.OfficeDonorID, c.CommitmentStartDate,
                    c.CommitmentEndDate, c.MonthlyAmount, c.Notes
                FROM tblCommitments c
                JOIN tblDonors d ON c.DonorID = d.DonorID
                WHERE c.IsActive = 1
                AND c.CommitmentEndDate BETWEEN date('now') AND date('now', '+' || ? || ' months')
            """
            params = [months]

            # Only add donor filter if something is selected
            if selected_donor:
                query += " AND d.FullName = ?"
                params.append(selected_donor)

            query += " ORDER BY c.CommitmentEndDate"

            cursor.execute(query, tuple(params))
            rows = cursor.fetchall()
            conn.close()

            for row in rows:
                start = format_db_date(row.CommitmentStartDate, "%d/%m/%Y")
                end = format_db_date(row.CommitmentEndDate, "%d/%m/%Y")

                monthly = float(row.MonthlyAmount or 0)

                report_rows.append({
                    "Donor": row.FullName,
                    "OfficeDonorID": row.OfficeDonorID,
                    "Start": start,
                    "End": end,
                    "Monthly": monthly,
                    "Notes": row.Notes or ""
                })

                tree.insert("", "end", values=(
                    row.FullName, row.OfficeDonorID, start, end,
                    f"Rs.{monthly:.2f}",
                    row.Notes or ""
                ))

            if report_rows:
                print_btn.pack(pady=10)
            else:
                print_btn.pack_forget()

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load commitments:\n{e}")


    # --- Print PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return
        try:
            downloads_path = str(Path.home() / "Downloads")
            filename = f"UpcomingExpiry_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_path, filename)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(200, height - 60, "Upcoming Expiry Commitments Report")
            c.setFont("Helvetica", 10)
            c.drawString(50, height - 80, f"Generated: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

            y = height - 110
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(150, y, "OfficeID")
            c.drawString(230, y, "Start")
            c.drawString(310, y, "End")
            c.drawString(390, y, "Monthly")
            c.drawString(470, y, "Notes")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r["Donor"])
                c.drawString(150, y, str(r["OfficeDonorID"]))
                c.drawString(230, y, r["Start"])
                c.drawString(310, y, r["End"])
                c.drawString(390, y, f"Rs.{r['Monthly']:.2f}" if r["Monthly"] else "0.00")
                c.drawString(470, y, r["Notes"][:30])
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            time.sleep(1)
            messagebox.showinfo("Success", f"PDF saved:\n{file_path}")
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", str(e))

    # --- Bind buttons ---
    generate_btn.config(command=load_upcoming_commitments)
    print_btn.config(command=print_report)
    report_win.bind("<Return>", lambda e: load_upcoming_commitments())
    generate_btn.focus_set()


def open_donor_giving_history_report():
    report_win = tk.Toplevel(root)
    report_win.title("Donor Giving History Report" + ORG_SUFFIX)
    ww, wh = 950, 550
    sw, sh = report_win.winfo_screenwidth(), report_win.winfo_screenheight()
    x, y = (sw // 2) - (ww // 2), (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")
    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())
    report_win.bind("<Return>", lambda e: generate_report())

    # --- Filters ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    tk.Label(filter_frame, text="From Date:").grid(row=0, column=0, padx=5)
    from_date = DateEntry(filter_frame, width=12)
    from_date.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="To Date:").grid(row=0, column=2, padx=5)
    to_date = DateEntry(filter_frame, width=12)
    to_date.grid(row=0, column=3, padx=5)

    donor_var = tk.StringVar()
    tk.Label(filter_frame, text="Search Donor:").grid(row=0, column=4, padx=5)
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=0, column=5, padx=5)

    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white", command=lambda: generate_report())
    generate_btn.grid(row=0, column=6, padx=10)

    # --- Treeview ---
    columns = ("Donor", "Donation Date", "Amount", "Consistency")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=200 if col == "Donor" else 120)
    tree.pack(fill="both", expand=True, padx=10, pady=10)

    # Color tags
    tree.tag_configure("consistent", background="#d4fcd4")  # light green
    tree.tag_configure("irregular", background="#fcd4d4")  # light red

    report_rows = []

    # --- Load donors into search combobox ---
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FullName FROM tblDonors WHERE IsActive = 1 ORDER BY FullName")
        donor_names = [row.FullName for row in cursor.fetchall()]
        donor_combo.set_completion_list(donor_names)
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    # --- Generate Report ---
    def generate_report():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows.clear()

        start = from_date.get_date().strftime("%Y-%m-%d")
        end = to_date.get_date().strftime("%Y-%m-%d")
        donor_filter = donor_var.get().strip()

        try:
            conn = get_connection()
            cursor = conn.cursor()

            if donor_filter:
                cursor.execute("""
                    SELECT d.FullName, don.DonationDate, don.DonationAmount
                    FROM tblDonations don
                    JOIN tblDonors d ON don.DonorID = d.DonorID
                    WHERE d.FullName = ? AND don.DonationDate BETWEEN ? AND ?
                    ORDER BY d.FullName, don.DonationDate
                """, (donor_filter, start, end))
            else:
                cursor.execute("""
                    SELECT d.FullName, don.DonationDate, don.DonationAmount
                    FROM tblDonations don
                    JOIN tblDonors d ON don.DonorID = d.DonorID
                    WHERE don.DonationDate BETWEEN ? AND ?
                    ORDER BY d.FullName, don.DonationDate
                """, (start, end))

            rows = cursor.fetchall()
            conn.close()

            donor_donations = {}
            for r in rows:
                donor, date, amount = r
                if donor not in donor_donations:
                    donor_donations[donor] = []
                donor_donations[donor].append((date, float(amount)))

            # Process consistency
            for donor, donations in donor_donations.items():
                donations.sort(key=lambda x: x[0])  # sort by date
                months = {format_db_date(d[0], "%Y-%m") for d in donations}
                total_months = (to_date.get_date().year - from_date.get_date().year) * 12 + (to_date.get_date().month - from_date.get_date().month + 1)
                consistency = "Consistent" if len(months) >= total_months * 0.8 else "Irregular"

                for date, amount in donations:
                    row_data = {
                        "Donor": donor,
                        "DonationDate": format_db_date(date, "%d-%m-%Y"),
                        "Amount": amount,
                        "Consistency": consistency
                    }
                    report_rows.append(row_data)
                    tree.insert("", "end", values=(donor, row_data["DonationDate"], f"Rs.{amount:.2f}", consistency),
                                tags=("consistent" if consistency == "Consistent" else "irregular",))

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load donor history:\n{e}")

    # --- Print to PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return

        try:
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
            file_name = f"Donor_Giving_History_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_folder, file_name)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(180, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(200, height - 60, "Donor Giving History Report")

            y = height - 90
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(200, y, "Date")
            c.drawString(300, y, "Amount")
            c.drawString(400, y, "Consistency")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r["Donor"])
                c.drawString(200, y, r["DonationDate"])
                c.drawString(300, y, f"Rs.{r['Amount']:.2f}")
                c.drawString(400, y, r["Consistency"])
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", f"Failed to save PDF:\n{e}")

    # Print button
    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20, command=print_report)
    print_btn.pack(pady=10)

    
def open_donor_defaulters_report():
    report_win = tk.Toplevel(root)
    report_win.title("Donor Defaulters Report" + ORG_SUFFIX)
    ww, wh = 1100, 620
    sw, sh = report_win.winfo_screenwidth(), report_win.winfo_screenheight()
    x, y = (sw // 2) - (ww // 2), (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")
    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())
    report_win.bind("<Return>", lambda e: generate_report())

    # ─────────────────────────────────────────────────────────────────────────
    # FILTER BAR  (two rows)
    # ─────────────────────────────────────────────────────────────────────────
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=8, padx=10, fill="x")

    # Row 0 — dates + collector
    tk.Label(filter_frame, text="From Date:").grid(row=0, column=0, padx=5, sticky="e")
    from_date = DateEntry(filter_frame, width=12)
    from_date.grid(row=0, column=1, padx=5)

    tk.Label(filter_frame, text="To Date:").grid(row=0, column=2, padx=5, sticky="e")
    to_date = DateEntry(filter_frame, width=12)
    to_date.grid(row=0, column=3, padx=5)

    tk.Label(filter_frame, text="Collector:").grid(row=0, column=4, padx=5, sticky="e")
    collector_var  = tk.StringVar()
    collector_combo = ttk.Combobox(filter_frame, textvariable=collector_var,
                                   width=28, state="readonly")
    collector_combo.grid(row=0, column=5, padx=5)

    # Row 1 — donor search + generate button
    donor_var = tk.StringVar()
    tk.Label(filter_frame, text="Search Donor:").grid(row=1, column=0, padx=5,
                                                       pady=(6, 0), sticky="e")
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=1, column=1, columnspan=3, padx=5, pady=(6, 0), sticky="w")

    generate_btn = tk.Button(
        filter_frame, text="Generate Report", width=20,
        bg="blue", fg="white", command=lambda: generate_report()
    )
    generate_btn.grid(row=1, column=5, padx=10, pady=(6, 0))

    # ─────────────────────────────────────────────────────────────────────────
    # TREEVIEW + SCROLLBARS
    # ─────────────────────────────────────────────────────────────────────────
    columns = ("Donor", "OfficeDonorID", "Start", "End",
               "Monthly", "Paid", "Status", "Notes")

    tree_frame = tk.Frame(report_win)
    tree_frame.pack(fill="both", expand=True, padx=10, pady=(4, 0))

    v_scroll = tk.Scrollbar(tree_frame, orient="vertical")
    h_scroll = tk.Scrollbar(tree_frame, orient="horizontal")

    tree = ttk.Treeview(
        tree_frame, columns=columns, show="headings",
        yscrollcommand=v_scroll.set,
        xscrollcommand=h_scroll.set
    )

    v_scroll.config(command=tree.yview)
    h_scroll.config(command=tree.xview)

    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=150, minwidth=100)

    tree.tag_configure("full",    background="#ffe4e4", foreground="#b91c1c")
    tree.tag_configure("partial", background="#fff7e4", foreground="#92400e")

    h_scroll.pack(side="bottom", fill="x")
    v_scroll.pack(side="right",  fill="y")
    tree.pack(side="left", fill="both", expand=True)

    report_rows = []

    # ─────────────────────────────────────────────────────────────────────────
    # LOAD COLLECTORS
    # ─────────────────────────────────────────────────────────────────────────
    # collector_lookup: CollectorName -> CollectorID
    # "" entry means "All Collectors"
    collector_lookup = {"-- All Collectors --": None}

    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute(
            "SELECT CollectorID, FullName FROM tblCollectors ORDER BY FullName")
        for row in cursor.fetchall():
            collector_lookup[row.FullName] = row.CollectorID
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load collectors:\n{e}")

    collector_combo["values"] = list(collector_lookup.keys())
    collector_combo.current(0)          # default = All Collectors

    # ─────────────────────────────────────────────────────────────────────────
    # LOAD / RELOAD DONORS based on selected collector
    # ─────────────────────────────────────────────────────────────────────────
    def load_donors_for_collector(event=None):
        selected_collector = collector_var.get()
        cid = collector_lookup.get(selected_collector)   # None = all

        try:
            conn = get_connection()
            cursor = conn.cursor()
            if cid is None:
                cursor.execute(
                    "SELECT FullName FROM tblDonors WHERE IsActive = 1 ORDER BY FullName")
            else:
                cursor.execute(
                    "SELECT FullName FROM tblDonors "
                    "WHERE IsActive = 1 AND CollectorID = ? ORDER BY FullName",
                    (cid,))
            donor_names = [r.FullName for r in cursor.fetchall()]
            conn.close()
            donor_combo.set_completion_list(donor_names)
            donor_var.set("")           # clear previous selection
        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    collector_combo.bind("<<ComboboxSelected>>", load_donors_for_collector)
    load_donors_for_collector()         # initial load (all donors)

    # ─────────────────────────────────────────────────────────────────────────
    # SAFE HELPERS
    # ─────────────────────────────────────────────────────────────────────────
    def safe_float(val):
        try:
            return float(val) if val is not None else 0.0
        except (TypeError, ValueError):
            return 0.0

    def safe_date_fmt(val):
        if val is None:
            return ""
        try:
            if isinstance(val, str):
                return datetime.strptime(val[:10], "%Y-%m-%d").strftime("%d-%m-%Y")
            return val.strftime("%d-%m-%Y")
        except Exception:
            return str(val)[:10]

    # ─────────────────────────────────────────────────────────────────────────
    # GENERATE REPORT
    # ─────────────────────────────────────────────────────────────────────────
    def generate_report():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows.clear()

        start        = from_date.get_date().strftime("%Y-%m-%d")
        end          = to_date.get_date().strftime("%Y-%m-%d")
        donor_filter = donor_var.get().strip()

        selected_collector = collector_var.get()
        collector_id = collector_lookup.get(selected_collector)  # None = all

        try:
            conn   = get_connection()
            cursor = conn.cursor()

            query = """
            WITH DonationSummary AS (
                SELECT
                    DonorID,
                    SUM(CAST(DonationAmount AS REAL)) AS TotalPaid
                FROM tblDonations
                WHERE DonationDate BETWEEN ? AND ?
                GROUP BY DonorID
            )
            SELECT
                d.FullName,
                d.OfficeDonorID,
                c.CommitmentStartDate,
                c.CommitmentEndDate,
                c.MonthlyAmount,
                IFNULL(ds.TotalPaid, 0)  AS PaidAmount,
                CASE
                    WHEN ds.TotalPaid IS NULL THEN 'Full Defaulter'
                    ELSE 'Partial Defaulter'
                END AS DefaulterType,
                c.Notes
            FROM tblCommitments c
            JOIN tblDonors d ON c.DonorID = d.DonorID
            LEFT JOIN DonationSummary ds ON d.DonorID = ds.DonorID
            WHERE c.IsActive = 1
              AND c.CommitmentStartDate <= ?
              AND c.CommitmentEndDate   >= ?
              AND (
                    ds.TotalPaid IS NULL
                    OR CAST(ds.TotalPaid AS REAL) < CAST(IFNULL(c.MonthlyAmount, 0) AS REAL)
              )
            """

            params = [start, end, end, start]

            if collector_id is not None:
                query += " AND d.CollectorID = ?"
                params.append(collector_id)

            if donor_filter:
                query += " AND d.FullName = ?"
                params.append(donor_filter)

            query += " ORDER BY d.FullName"

            cursor.execute(query, params)
            rows = cursor.fetchall()
            conn.close()

            for r in rows:
                donor, officeid, start_date, end_date, monthly, paid, dtype, notes = r

                monthly_f = safe_float(monthly)
                paid_f    = safe_float(paid)
                start_fmt = safe_date_fmt(start_date)
                end_fmt   = safe_date_fmt(end_date)

                row_data = {
                    "Donor"        : donor     or "",
                    "OfficeDonorID": officeid  or "",
                    "Start"        : start_fmt,
                    "End"          : end_fmt,
                    "Monthly"      : monthly_f,
                    "Paid"         : paid_f,
                    "Type"         : dtype     or "",
                    "Notes"        : notes     or "",
                }
                report_rows.append(row_data)

                tag = "full" if dtype == "Full Defaulter" else "partial"
                tree.insert(
                    "", "end",
                    values=(
                        row_data["Donor"],
                        row_data["OfficeDonorID"],
                        row_data["Start"],
                        row_data["End"],
                        f"Rs.{monthly_f:,.2f}",
                        f"Rs.{paid_f:,.2f}",
                        row_data["Type"],
                        row_data["Notes"],
                    ),
                    tags=(tag,)
                )

            if not rows:
                messagebox.showinfo("No Results",
                                    "No defaulters found for the selected criteria.")

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to load defaulters:\n{e}")

    # ─────────────────────────────────────────────────────────────────────────
    # PRINT / PDF
    # ─────────────────────────────────────────────────────────────────────────
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No defaulter data to print.")
            return

        try:
            import tempfile, os, time
            from reportlab.lib.pagesizes import A4
            from reportlab.pdfgen import canvas
            from datetime import datetime as dt

            temp_pdf  = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_path = temp_pdf.name
            temp_pdf.close()

            c = canvas.Canvas(temp_path, pagesize=A4)
            width, height = A4
            y = height - 50

            # ── Header ────────────────────────────────────────────────────
            c.setFont("Helvetica-Bold", 16)
            c.drawCentredString(width / 2, y, "Alkhidmat Khawateen Trust Pakistan")
            y -= 22

            c.setFont("Helvetica-Bold", 12)
            c.drawCentredString(width / 2, y, "Donor Defaulters Report")
            y -= 18

            c.setFont("Helvetica", 9)
            c.drawCentredString(
                width / 2, y,
                f"Period: {from_date.get_date().strftime('%d/%m/%Y')} "
                f"to {to_date.get_date().strftime('%d/%m/%Y')}"
            )
            y -= 14

            # Collector line (only if one is selected)
            selected_collector = collector_var.get()
            if selected_collector and selected_collector != "-- All Collectors --":
                c.setFont("Helvetica", 9)
                c.drawCentredString(
                    width / 2, y,
                    f"Collector: {selected_collector}"
                )
                y -= 14

            y -= 10

            # ── Table header ──────────────────────────────────────────────
            c.setFont("Helvetica-Bold", 9)
            headers = [
                ("Donor",     40),
                ("Office ID", 150),
                ("Start",     215),
                ("End",       275),
                ("Monthly",   335),
                ("Paid",      400),
                ("Status",    460),
            ]
            for text, xpos in headers:
                c.drawString(xpos, y, text)

            y -= 10
            c.line(40, y, width - 40, y)
            y -= 15

            # ── Table rows ────────────────────────────────────────────────
            c.setFont("Helvetica", 9)
            for row in report_rows:
                if y < 60:
                    c.showPage()
                    y = height - 50
                    c.setFont("Helvetica", 9)

                c.drawString(40,  y, str(row["Donor"])[:28])
                c.drawString(150, y, str(row["OfficeDonorID"] or ""))
                c.drawString(215, y, row["Start"])
                c.drawString(275, y, row["End"])
                c.drawRightString(385, y, f"Rs.{row['Monthly']:,.2f}")
                c.drawRightString(445, y, f"Rs.{row['Paid']:,.2f}")
                c.drawString(460, y, str(row["Type"]))
                y -= 14

            # ── Footer ────────────────────────────────────────────────────
            y -= 10
            c.setFont("Helvetica-Oblique", 8)
            c.drawString(40, y,
                         f"Generated on {dt.now().strftime('%d/%m/%Y %H:%M')}")

            c.save()

            open_file_cross_platform(temp_path)

            def cleanup():
                time.sleep(15)
                try:
                    os.remove(temp_path)
                except Exception:
                    pass

            import threading
            threading.Thread(target=cleanup, daemon=True).start()

        except Exception as e:
            messagebox.showerror("Print Error", f"Failed to generate PDF:\n{e}")

    # ─────────────────────────────────────────────────────────────────────────
    # PRINT BUTTON
    # ─────────────────────────────────────────────────────────────────────────
    print_btn = tk.Button(
        report_win, text="🖨  Print Report",
        bg="green", fg="white", width=20,
        command=print_report
    )
    print_btn.pack(pady=8)

    
def open_sponsorship_vs_commitment_validation():
    report_win = tk.Toplevel(root)
    report_win.title("Sponsorship vs Commitment Validation Report" + ORG_SUFFIX)
    ww, wh = 950, 550
    sw, sh = report_win.winfo_screenwidth(), report_win.winfo_screenheight()
    x, y = (sw // 2) - (ww // 2), (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")
    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())
    report_win.bind("<Return>", lambda e: generate_report())

    # --- Filters ---
    filter_frame = tk.Frame(report_win)
    filter_frame.pack(pady=10)

    donor_var = tk.StringVar()
    tk.Label(filter_frame, text="Search Donor:").grid(row=0, column=0, padx=5)
    donor_combo = AutocompleteCombobox(filter_frame, textvariable=donor_var, width=30)
    donor_combo.grid(row=0, column=1, padx=5)

    generate_btn = tk.Button(filter_frame, text="Generate Report", width=20, bg="blue", fg="white",
                             command=lambda: generate_report())
    generate_btn.grid(row=0, column=2, padx=10)

    # --- Treeview ---
    columns = ("Donor", "OfficeDonorID", "Commitment Monthly", "Sponsorship Total", "Difference", "Notes")
    tree = ttk.Treeview(report_win, columns=columns, show="headings")
    for col in columns:
        tree.heading(col, text=col)
        tree.column(col, width=160)
    tree.pack(fill="both", expand=True, padx=10, pady=10)

    report_rows = []

    # --- Load donors for autocomplete ---
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT FullName FROM tblDonors WHERE IsActive = 1 ORDER BY FullName")
        donor_names = [row.FullName for row in cursor.fetchall()]
        donor_combo.set_completion_list(donor_names)
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donors:\n{e}")

    # --- Generate Report ---
    def generate_report():
        nonlocal report_rows
        tree.delete(*tree.get_children())
        report_rows.clear()

        donor_filter = donor_var.get().strip()

        try:
            conn = get_connection()
            cursor = conn.cursor()

            if donor_filter:
                query = """
                    SELECT d.FullName, d.OfficeDonorID,
                           IFNULL(c.MonthlyAmount, 0) AS CommitmentMonthly,
                           IFNULL(SUM(CAST(s.SponsorshipAmount AS REAL)), 0) AS SponsorshipTotal,
                           IFNULL(c.Notes, '') AS Notes
                    FROM tblDonors d
                    LEFT JOIN tblCommitments c ON d.DonorID = c.DonorID AND c.IsActive = 1
                    LEFT JOIN tblSponsorships s ON d.DonorID = s.DonorID AND s.IsActive = 1
                    WHERE d.FullName = ?
                    GROUP BY d.FullName, d.OfficeDonorID, c.MonthlyAmount, c.Notes
                    ORDER BY d.FullName
                """
                cursor.execute(query, (donor_filter,))
            else:
                query = """
                    SELECT d.FullName, d.OfficeDonorID,
                           IFNULL(c.MonthlyAmount, 0) AS CommitmentMonthly,
                           IFNULL(SUM(CAST(s.SponsorshipAmount AS REAL)), 0) AS SponsorshipTotal,
                           IFNULL(c.Notes, '') AS Notes
                    FROM tblDonors d
                    LEFT JOIN tblCommitments c ON d.DonorID = c.DonorID AND c.IsActive = 1
                    LEFT JOIN tblSponsorships s ON d.DonorID = s.DonorID AND s.IsActive = 1
                    GROUP BY d.FullName, d.OfficeDonorID, c.MonthlyAmount, c.Notes
                    ORDER BY d.FullName
                """
                cursor.execute(query)

            rows = cursor.fetchall()
            conn.close()

            for r in rows:
                donor, officeid, monthly, sponsored, notes = r
                monthly = float(monthly or 0)
                sponsored = float(sponsored or 0)
                diff = monthly - sponsored

                row_data = {
                    "Donor": donor,
                    "OfficeDonorID": officeid,
                    "Monthly": monthly,
                    "Sponsored": sponsored,
                    "Difference": diff,
                    "Notes": notes or ""
                }
                report_rows.append(row_data)

                # highlight over-allocation
                tag = "normal"
                if sponsored > monthly:
                    tag = "over"
                elif sponsored < monthly:
                    tag = "under"

                tree.insert("", "end", values=(
                    donor, officeid,
                    f"Rs.{monthly:.2f}",
                    f"Rs.{sponsored:.2f}",
                    f"Rs.{diff:.2f}",
                    notes
                ), tags=(tag,))

            # color coding
            tree.tag_configure("over", background="#ffcccc")   # red for over-allocated
            tree.tag_configure("under", background="#fff3cd")  # yellow for under-allocated
            tree.tag_configure("normal", background="#ccffcc") # green for exact match

        except Exception as e:
            messagebox.showerror("DB Error", f"Failed to validate:\n{e}")

    # --- Print to PDF ---
    def print_report():
        if not report_rows:
            messagebox.showinfo("No Data", "No report data to print.")
            return

        try:
            downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
            file_name = f"Sponsorship_vs_Commitment_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            file_path = os.path.join(downloads_folder, file_name)

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont("Helvetica-Bold", 16)
            c.drawString(150, height - 40, "Alkhidmat Khawateen Trust")
            c.setFont("Helvetica-Bold", 12)
            c.drawString(150, height - 60, "Sponsorship vs Commitment Validation Report")

            y = height - 90
            c.setFont("Helvetica-Bold", 10)
            c.drawString(50, y, "Donor")
            c.drawString(180, y, "OfficeID")
            c.drawString(270, y, "Monthly")
            c.drawString(360, y, "Sponsored")
            c.drawString(450, y, "Difference")
            c.drawString(530, y, "Notes")

            y -= 20
            c.setFont("Helvetica", 9)
            for r in report_rows:
                c.drawString(50, y, r["Donor"])
                c.drawString(180, y, str(r["OfficeDonorID"]))
                c.drawString(270, y, f"Rs.{r['Monthly']:.2f}")
                c.drawString(360, y, f"Rs.{r['Sponsored']:.2f}")
                c.drawString(450, y, f"Rs.{r['Difference']:.2f}")
                c.drawString(530, y, r["Notes"])
                y -= 15
                if y < 40:
                    c.showPage()
                    y = height - 40

            c.save()
            os.startfile(file_path)

        except Exception as e:
            messagebox.showerror("Print Error", f"Failed to save PDF:\n{e}")

    print_btn = tk.Button(report_win, text="Print Report", bg="green", fg="white", width=20, command=print_report)
    print_btn.pack(pady=10)

    
def open_overall_impact_report():

    report_win = tk.Toplevel(root)
    report_win.title("Overall Impact Report" + ORG_SUFFIX)
    ww, wh = 700, 400
    sw, sh = report_win.winfo_screenwidth(), report_win.winfo_screenheight()
    x, y = (sw // 2) - (ww // 2), (sh // 2) - (wh // 2)
    report_win.geometry(f"{ww}x{wh}+{x}+{y}")
    report_win.transient(root)
    report_win.grab_set()
    report_win.focus_force()
    report_win.bind("<Escape>", lambda e: report_win.destroy())

    summary_frame = tk.Frame(report_win)
    summary_frame.pack(pady=30, fill="both", expand=True)

    lbl_total_children = tk.Label(summary_frame, text="Total Sponsored Children: ", font=("Arial", 12))
    lbl_total_children.grid(row=0, column=0, sticky="w", pady=10, padx=20)

    lbl_total_commitments = tk.Label(summary_frame, text="Total Funds Committed: ", font=("Arial", 12))
    lbl_total_commitments.grid(row=1, column=0, sticky="w", pady=10, padx=20)

    lbl_total_donations = tk.Label(summary_frame, text="Total Funds Collected: ", font=("Arial", 12))
    lbl_total_donations.grid(row=2, column=0, sticky="w", pady=10, padx=20)

    lbl_percentage = tk.Label(summary_frame, text="Fulfillment %: ", font=("Arial", 12))
    lbl_percentage.grid(row=3, column=0, sticky="w", pady=10, padx=20)

    # Values
    val_children = tk.Label(summary_frame, text="", font=("Arial", 12, "bold"), fg="blue")
    val_children.grid(row=0, column=1, sticky="w")

    val_commitments = tk.Label(summary_frame, text="", font=("Arial", 12, "bold"), fg="blue")
    val_commitments.grid(row=1, column=1, sticky="w")

    val_donations = tk.Label(summary_frame, text="", font=("Arial", 12, "bold"), fg="blue")
    val_donations.grid(row=2, column=1, sticky="w")

    val_percentage = tk.Label(summary_frame, text="", font=("Arial", 12, "bold"), fg="blue")
    val_percentage.grid(row=3, column=1, sticky="w")

    # ------------------ LOAD DATA ------------------
    try:
        conn = get_connection()
        cursor = conn.cursor()

        cursor.execute("SELECT COUNT(DISTINCT ChildID) AS TotalChildren FROM tblSponsorships WHERE IsActive = 1")
        total_children = cursor.fetchone().TotalChildren or 0

        cursor.execute("SELECT IFNULL(SUM(CommitmentAmount),0) AS TotalCommitments FROM tblCommitments WHERE IsActive = 1")
        total_commitments = float(cursor.fetchone().TotalCommitments or 0)

        cursor.execute("SELECT IFNULL(SUM(CAST(DonationAmount AS REAL)),0) AS TotalDonations FROM tblDonations")
        total_donations = float(cursor.fetchone().TotalDonations or 0)

        conn.close()

        fulfillment_pct = (total_donations / total_commitments * 100) if total_commitments > 0 else 0

        val_children.config(text=f"{total_children}")
        val_commitments.config(text=f"Rs.{total_commitments:,.2f}")
        val_donations.config(text=f"Rs.{total_donations:,.2f}")
        val_percentage.config(text=f"{fulfillment_pct:.2f}%")

    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load impact report:\n{e}")

    # ------------------ PRINT PDF BUTTON ------------------
    def print_pdf():
        try:
            # Temp file
            tmp_pdf = os.path.join(
                tempfile.gettempdir(),
                f"OverallImpact_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            )

            c = canvas.Canvas(tmp_pdf, pagesize=landscape(A4))
            w, h = landscape(A4)
            y = h - 50

            # Header
            c.setFont("Helvetica-Bold", 20)
            c.drawCentredString(w/2, y, "Al-Khidmat Khawateen Trust Pakistan")
            y -= 40
            
            c.setFont("Helvetica-Bold", 18)
            c.drawCentredString(w/2, y, "Aghosh Program")
            y -= 45

            c.setFont("Helvetica-Bold", 16)
            c.drawCentredString(w/2, y, "Overall Impact Report")
            y -= 50

            c.setFont("Helvetica-Bold", 12)
            c.drawString(60, y, "Total Sponsored Children:")
            c.setFont("Helvetica", 12)
            c.drawString(300, y, str(val_children.cget("text")))
            y -= 25

            c.setFont("Helvetica-Bold", 12)
            c.drawString(60, y, "Total Funds Committed:")
            c.setFont("Helvetica", 12)
            c.drawString(300, y, val_commitments.cget("text"))
            y -= 25

            c.setFont("Helvetica-Bold", 12)
            c.drawString(60, y, "Total Funds Collected:")
            c.setFont("Helvetica", 12)
            c.drawString(300, y, val_donations.cget("text"))
            y -= 25

            c.setFont("Helvetica-Bold", 12)
            c.drawString(60, y, "Fulfillment Percentage:")
            c.setFont("Helvetica", 12)
            c.drawString(300, y, val_percentage.cget("text"))
            y -= 40

            c.setFont("Helvetica-Oblique", 10)
            c.drawRightString(w - 30, 30, f"Generated: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

            c.save()

            # Open PDF automatically
            try:
                os.startfile(tmp_pdf)
            except:
                messagebox.showinfo("PDF Generated", tmp_pdf)

        except Exception as e:
            messagebox.showerror("PDF Error", f"Failed to generate PDF:\n{e}")

    # Button in window
    btn_print = tk.Button(report_win, text="Print Report (PDF)", bg="green", fg="white",
                          font=("Arial", 12), width=20, command=print_pdf)
    btn_print.pack(pady=10)



# ============================================================
def open_reports():
    '''report_win = tk.Toplevel(root)
    report_win.title("Reports" + ORG_SUFFIX)
    report_win.geometry("700x400")
    #report_win.grab_set()
    report_win.focus_force()
    report_win.grab_set()'''
    report_win = tk.Toplevel(root)
    report_win.title("Reports" + ORG_SUFFIX)

    # Window size
    win_width = 700
    win_height = 400

    # Get screen dimension
    screen_width = report_win.winfo_screenwidth()
    screen_height = report_win.winfo_screenheight()

    # Calculate position
    x = (screen_width // 2) - (win_width // 2)
    y = (screen_height // 2) - (win_height // 2)

    # Set geometry to center
    report_win.geometry(f"{win_width}x{win_height}+{x}+{y}")

    # Focus and grab
    report_win.focus_force()
    report_win.grab_set()


    tk.Label(report_win, text="Select Report", font=("Arial", 14)).pack(pady=15)

    button_frame = tk.Frame(report_win)
    button_frame.pack(pady=10)

    # Row 1
    tk.Button(button_frame, text="Children Report", command=open_child_report_form,
              width=20, height=2, bg="blue", fg="white").grid(row=0, column=0, padx=10, pady=10)

    tk.Button(button_frame, text="Donations Report", command=open_donation_report_form,
              width=20, height=2, bg="blue", fg="white").grid(row=0, column=1, padx=10, pady=10)

    tk.Button(button_frame, text="Donor Child Map", command=open_donor_child_map_report,
              width=20, height=2, bg="blue", fg="white").grid(row=0, column=2, padx=10, pady=10)

    # Row 2
    tk.Button(button_frame, text="Exel For AKF", command=open_Exel_For_AKF,
              width=20, height=2, bg="blue", fg="white").grid(row=1, column=0, padx=10, pady=10)

    tk.Button(button_frame, text="Child Profile", command=open_child_profile_report_form,
              width=20, height=2, bg="blue", fg="white").grid(row=1, column=1, padx=10, pady=10)

    tk.Button(button_frame, text="Donor Report", command=open_donor_report_form,
              width=20, height=2, bg="blue", fg="white").grid(row=1, column=2, padx=10, pady=10)

    # Row 3
    tk.Button(button_frame, text="Collector's Collections", command=open_collector_report_form,
              width=20, height=2, bg="blue", fg="white").grid(row=2, column=0, padx=10, pady=10)
    
    tk.Button(button_frame, text="Collector List", command=open_collector_list_form,
              width=20, height=2, bg="blue", fg="white").grid(row=2, column=1, padx=10, pady=10)
    
    tk.Button(button_frame, text="Child Sponsorship", command=open_child_sponsorship_report_form,
              width=20, height=2, bg="blue", fg="white").grid(row=2, column=2, padx=10, pady=10)

    # Close Button
    tk.Button(report_win, text="Close", command=report_win.destroy,
              width=12).pack(pady=20)

    report_win.bind("<Escape>", lambda e: report_win.destroy())


# ============================================================
# Popup for Expiry Commitments
# ============================================================
import calendar
from datetime import datetime

def show_expiry_popup(root):
    try:
        today = datetime.today()

        # Find last day of next month
        next_month = today.month + 1 if today.month < 12 else 1
        next_year = today.year if today.month < 12 else today.year + 1
        last_day = calendar.monthrange(next_year, next_month)[1]
        expiry_till = datetime(next_year, next_month, last_day).strftime("%d-%b-%Y")

        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT COUNT(*) AS ExpiringSoon
            FROM tblCommitments
            WHERE IsActive = 1
              AND CommitmentEndDate BETWEEN date('now') AND date('now', '+1 months')
        """)
        count = cursor.fetchone().ExpiringSoon
        conn.close()

        if count > 0:
            # Create popup
            popup = tk.Toplevel(root)
            popup.title("Upcoming Expiry Reminder" + ORG_SUFFIX)
            popup.geometry("420x160")
            popup.transient(root)
            popup.grab_set()
            popup.focus_force()

            # Center on screen
            sw = popup.winfo_screenwidth()
            sh = popup.winfo_screenheight()
            x = (sw // 2) - (420 // 2)
            y = (sh // 2) - (160 // 2)
            popup.geometry(f"420x160+{x}+{y}")

            # Label with message
            msg = tk.Label(
                popup,
                text=f"{count} commitments are about to expire till {expiry_till}.",
                font=("Arial", 11),
                wraplength=380,
                justify="center"
            )
            msg.pack(expand=True, pady=20)

            # OK button
            ok_btn = tk.Button(
                popup,
                text="OK",
                width=12,
                command=lambda: [popup.destroy(), open_upcoming_expiry_report()]
            )
            ok_btn.pack(pady=10)

            # Bind click anywhere → open window
            popup.bind("<Button-1>", lambda e: [popup.destroy(), open_upcoming_expiry_report()])

    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to check expiring commitments:\n{e}")

#---------------_


def check_expiring_commitments():
    """Check commitments ending in current month and show popup if any."""
    try:
        conn = get_connection()
        cursor = conn.cursor()

        today = datetime.today()
        start_month = today.replace(day=1)
        last_day = calendar.monthrange(today.year, today.month)[1]
        end_month = today.replace(day=last_day)

        cursor.execute("""
            SELECT c.CommitmentID, d.FullName AS DonorName, col.FullName AS CollectorName,
                   c.CommitmentStartDate, c.CommitmentEndDate, c.CommitmentAmount, d.DonorID
            FROM tblCommitments c
            JOIN tblDonors d ON c.DonorID = d.DonorID
            JOIN tblCollectors col ON d.CollectorID = col.CollectorID
            WHERE c.CommitmentEndDate BETWEEN ? AND ?
        """, (start_month, end_month))
        rows = cursor.fetchall()
        conn.close()

        if not rows:
            return  # No expiring commitments → do nothing

        # === Create popup window ===
        popup = tk.Toplevel(root)
        popup.title("Expiring Commitments" + ORG_SUFFIX)
        popup.geometry("800x400")
        popup.transient(root)
        popup.grab_set()
        popup.focus_force()

        tk.Label(popup, text=f"Commitments expiring till {end_month.strftime('%d-%b-%Y')}",
                 font=("Arial", 12, "bold")).pack(pady=10)

        frame = tk.Frame(popup)
        frame.pack(fill="both", expand=True, padx=10, pady=5)

        tree = ttk.Treeview(frame, columns=("Donor", "Collector", "Start", "End", "Amount", "Action"), show="headings")
        tree.heading("Donor", text="Donor")
        tree.heading("Collector", text="Collector")
        tree.heading("Start", text="Start Date")
        tree.heading("End", text="End Date")
        tree.heading("Amount", text="Commitment Amount")
        tree.heading("Action", text="")

        tree.column("Donor", width=150)
        tree.column("Collector", width=150)
        tree.column("Start", width=100)
        tree.column("End", width=100)
        tree.column("Amount", width=120)
        tree.column("Action", width=80)

        tree.pack(fill="both", expand=True)

        # Fill treeview with commitments
        for row in rows:
            cid, donor, collector, start, end, amount, donor_id = row
            tree.insert("", "end", iid=cid, values=(
                donor, collector, format_db_date(start, "%d/%m/%Y"),
                format_db_date(end, "%d/%m/%Y"), f"Rs.{amount:.2f}", "Detail.."
            ))

        # On click "Detail.."
        def on_tree_click(event):
            item = tree.identify_row(event.y)
            if not item:
                return
            col = tree.identify_column(event.x)
            if col == "#6":  # "Action" column
                cid = int(item)
                for r in rows:
                    if r[0] == cid:
                        open_commitment_details(root, r)

        tree.bind("<Button-1>", on_tree_click)

    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to check expiring commitments:\n{e}")


def open_commitment_details(root, commitment_row):
    """Open detail window showing donor's donations for this commitment."""
    cid, donor, collector, start, end, amount, donor_id = commitment_row

    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("""
            SELECT DonationDate, DonationAmount, DonationType
            FROM tblDonations
            WHERE DonorID = ? AND DonationDate >= ?
            ORDER BY DonationDate
        """, (donor_id, start))
        donations = cursor.fetchall()
        conn.close()
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to load donations:\n{e}")
        return

    # === Create detail window ===
    detail_win = tk.Toplevel(root)
    detail_win.title(f"Donations for {donor}" + ORG_SUFFIX)
    detail_win.geometry("700x400")
    detail_win.transient(root)
    detail_win.grab_set()
    detail_win.focus_force()

    tk.Label(detail_win, text=f"Donor: {donor} | Collector: {collector}", font=("Arial", 11, "bold")).pack(pady=5)
    tk.Label(detail_win, text=f"Commitment: {format_db_date(start, '%d/%m/%Y')} - {format_db_date(end, '%d/%m/%Y')} "
                              f"| Amount: Rs.{amount:.2f}", font=("Arial", 10)).pack(pady=5)

    frame = tk.Frame(detail_win)
    frame.pack(fill="both", expand=True, padx=10, pady=5)

    tree = ttk.Treeview(frame, columns=("Date", "Amount", "Type"), show="headings")
    tree.heading("Date", text="Date")
    tree.heading("Amount", text="Amount")
    tree.heading("Type", text="Type")

    tree.column("Date", width=120)
    tree.column("Amount", width=120)
    tree.column("Type", width=150)
    tree.pack(fill="both", expand=True)

    # Fill donations
    for d in donations:
        ddate, damount, dtype = d
        tree.insert("", "end", values=(format_db_date(ddate, "%d/%m/%Y"), f"Rs.{float(damount):.2f}", dtype))

    # Print PDF
    def print_pdf():
        downloads_folder = os.path.join(os.path.expanduser("~"), "Downloads")
        if not os.path.exists(downloads_folder):
            os.makedirs(downloads_folder)

        file_name = f"Commitment_Detail_{donor}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        file_path = os.path.join(downloads_folder, file_name)

        c = canvas.Canvas(file_path, pagesize=A4)
        width, height = A4

        c.setFont("Helvetica-Bold", 14)
        c.drawString(200, height - 40, "Commitment Detail Report")

        c.setFont("Helvetica", 10)
        c.drawString(50, height - 70, f"Donor: {donor}")
        c.drawString(50, height - 85, f"Collector: {collector}")
        c.drawString(50, height - 100, f"Commitment: {format_db_date(start, '%d/%m/%Y')} - {format_db_date(end, '%d/%m/%Y')}")
        c.drawString(50, height - 115, f"Commitment Amount: Rs.{amount:.2f}")

        y = height - 140
        c.setFont("Helvetica-Bold", 10)
        c.drawString(50, y, "Donation Date")
        c.drawString(150, y, "Donated Amount")
        c.drawString(300, y, "Type")

        y -= 20
        c.setFont("Helvetica", 9)
        for d in donations:
            ddate, damount, dtype = d
            c.drawString(50, y, format_db_date(ddate, "%d/%m/%Y"))
            c.drawString(150, y, f"Rs.{float(damount):.2f}")
            c.drawString(300, y, dtype)
            y -= 15
            if y < 40:
                c.showPage()
                y = height - 40

        c.save()
        os.startfile(file_path)

    tk.Button(detail_win, text="Print Report", bg="green", fg="white", command=print_pdf).pack(pady=10)


# ============================================================
# Check Expiring Commitments
# ============================================================
def check_expiring_commitments_this_month():
    return
    try:
        conn = get_connection()
        cursor = conn.cursor()

        query = """
            SELECT COUNT(*) AS CountExpiring
            FROM tblCommitments
            WHERE IsActive = 1
              AND CAST(strftime('%m', CommitmentEndDate) AS INTEGER) = CAST(strftime('%m', 'now') AS INTEGER)
              AND CAST(strftime('%Y', CommitmentEndDate) AS INTEGER) = CAST(strftime('%Y', 'now') AS INTEGER)
        """
        cursor.execute(query)
        result = cursor.fetchone()
        conn.close()

        count_expiring = result.CountExpiring if result else 0
        if count_expiring > 0:
            today = datetime.now()
            last_day = calendar.monthrange(today.year, today.month)[1]
            messagebox.showinfo("Upcoming Expiries",
                                f"{count_expiring} commitment /s are about to expire this month "
                                f"till {last_day} {today.strftime('%B %Y')}.")
            check_expiring_commitments()
            #-------------
            
            
            
            #-------------
            
            
            
    except Exception as e:
        messagebox.showerror("DB Error", f"Failed to check expiring commitments:\n{e}")

def open_Sponsorship_Certificate_mgmt():
    """
    Sponsorship Certificate Generator
    Allows selecting a child and donor, then generates a PDF certificate
    """
    
    # ─────────────────────────────────────────────────────────────────────────
    # WINDOW SETUP
    # ─────────────────────────────────────────────────────────────────────────
    win = tk.Toplevel()
    win.title("Sponsorship Certificate Generator" + ORG_SUFFIX)
    win.geometry("600x450")
    win.resizable(False, False)
    win.grab_set()
    win.focus_force()
    
    # Center window
    win.update_idletasks()
    w = win.winfo_width()
    h = win.winfo_height()
    sw = win.winfo_screenwidth()
    sh = win.winfo_screenheight()
    x = (sw - w) // 2
    y = (sh - h) // 2
    win.geometry(f"{w}x{h}+{x}+{y}")
    
    # Design tokens
    HEADER_BG = "#0f2540"
    HEADER_FG = "#f8fafc"
    SECTION_BG = "#eef2f7"
    FRAME_BG = "#ffffff"
    LABEL_FG = "#374151"
    ACCENT = "#1d6fd8"
    BTN_GREEN = "#16a34a"
    BTN_GRAY = "#64748b"
    FONT_LBL = ("Segoe UI", 10)
    FONT_ENTRY = ("Segoe UI", 10)
    FONT_HEADER = ("Segoe UI", 13, "bold")
    FONT_BTN = ("Segoe UI", 10, "bold")
    
    win.configure(bg=SECTION_BG)
    
    # ─────────────────────────────────────────────────────────────────────────
    # HEADER
    # ─────────────────────────────────────────────────────────────────────────
    header = tk.Frame(win, bg=HEADER_BG, height=56)
    header.pack(fill="x", side="top")
    header.pack_propagate(False)
    tk.Label(header, text="📄  Sponsorship Certificate Generator",
             font=FONT_HEADER, bg=HEADER_BG, fg=HEADER_FG,
             pady=10).pack(side="left", padx=20)
    
    tk.Frame(win, bg=ACCENT, height=3).pack(fill="x")
    
    # ─────────────────────────────────────────────────────────────────────────
    # BODY
    # ─────────────────────────────────────────────────────────────────────────
    body = tk.Frame(win, bg=FRAME_BG, padx=30, pady=20)
    body.pack(fill="both", expand=True)
    
    # State variables
    selected_child_id = tk.IntVar(value=0)
    selected_donor_id = tk.IntVar(value=0)
    sponsorship_data = {}
    child_dict = {}
    donor_dict = {}
    
    # ─────────────────────────────────────────────────────────────────────────
    # CHILD SELECTION
    # ─────────────────────────────────────────────────────────────────────────
    tk.Label(body, text="Select Child:", font=FONT_LBL, fg=LABEL_FG,
             bg=FRAME_BG, anchor="w").grid(row=0, column=0, sticky="w", pady=(0, 5))
    
    child_var = tk.StringVar()
    child_combo = AutocompleteCombobox(body, textvariable=child_var, 
                                       font=FONT_ENTRY, width=50)
    child_combo.grid(row=1, column=0, sticky="ew", pady=(0, 15))
    
    # ─────────────────────────────────────────────────────────────────────────
    # DONOR SELECTION
    # ─────────────────────────────────────────────────────────────────────────
    tk.Label(body, text="Select Donor/Sponsor:", font=FONT_LBL, fg=LABEL_FG,
             bg=FRAME_BG, anchor="w").grid(row=2, column=0, sticky="w", pady=(0, 5))
    
    donor_var = tk.StringVar()
    donor_combo = ttk.Combobox(body, textvariable=donor_var, 
                               font=FONT_ENTRY, width=50, state="readonly")
    donor_combo.grid(row=3, column=0, sticky="ew", pady=(0, 15))
    
    # ─────────────────────────────────────────────────────────────────────────
    # SPONSORSHIP INFO DISPLAY
    # ─────────────────────────────────────────────────────────────────────────
    info_frame = tk.LabelFrame(body, text="  Sponsorship Information  ",
                               font=("Segoe UI", 9, "bold"), fg=ACCENT,
                               bg=FRAME_BG, bd=1, relief="groove")
    info_frame.grid(row=4, column=0, sticky="ew", pady=(0, 20))
    
    info_text = tk.Text(info_frame, height=6, width=60, font=FONT_LBL,
                       relief="flat", bg="#f8fafc", state="disabled")
    info_text.pack(padx=10, pady=10)
    
    # ─────────────────────────────────────────────────────────────────────────
    # BUTTONS
    # ─────────────────────────────────────────────────────────────────────────
    btn_frame = tk.Frame(body, bg=FRAME_BG)
    btn_frame.grid(row=5, column=0, pady=(10, 0))
    
    generate_btn = tk.Button(btn_frame, text="📄  Generate Certificate",
                            font=FONT_BTN, bg=BTN_GREEN, fg="white",
                            relief="flat", cursor="hand2",
                            padx=20, pady=8, state="disabled")
    generate_btn.pack(side="left", padx=5)
    
    tk.Button(btn_frame, text="✖  Close",
              command=win.destroy,
              font=FONT_BTN, bg=BTN_GRAY, fg="white",
              relief="flat", cursor="hand2",
              padx=20, pady=8).pack(side="left", padx=5)
    
    body.columnconfigure(0, weight=1)
    
    # ─────────────────────────────────────────────────────────────────────────
    # DATA LOADING FUNCTIONS
    # ─────────────────────────────────────────────────────────────────────────
    def load_children():
        """Load all sponsored children"""
        try:
            conn = get_connection()
            if not conn:
                raise Exception("DB connection failed")
            
            cursor = conn.cursor()
            cursor.execute("""
                SELECT DISTINCT c.ChildID, c.FullName, c.RegistrationNumber
                FROM tblChildren c
                INNER JOIN tblSponsorships s ON c.ChildID = s.ChildID
                WHERE s.IsActive = 1
                ORDER BY c.FullName
            """)
            
            rows = cursor.fetchall()
            conn.close()
            
            nonlocal child_dict
            child_dict = {
                f"{r.FullName} ({r.RegistrationNumber})": r.ChildID
                for r in rows
            }
            
            child_combo.set_completion_list(list(child_dict.keys()))
            
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load children:\n{e}")
    
    def load_donors_for_child(child_id):
        """Load donors who sponsor the selected child"""
        try:
            conn = get_connection()
            if not conn:
                raise Exception("DB connection failed")
            
            cursor = conn.cursor()
            cursor.execute("""
                SELECT DISTINCT d.DonorID, d.FullName, s.SponsorshipID,
                       s.StartDate, s.EndDate, s.SponsorshipAmount
                FROM tblDonors d
                INNER JOIN tblSponsorships s ON d.DonorID = s.DonorID
                WHERE s.ChildID = ? AND s.IsActive = 1
                ORDER BY d.FullName
            """, (child_id,))
            
            rows = cursor.fetchall()
            conn.close()
            
            nonlocal donor_dict, sponsorship_data
            donor_dict = {}
            sponsorship_data = {}
            
            for r in rows:
                key = f"{r.FullName} (ID:{r.DonorID})"
                donor_dict[key] = r.DonorID
                sponsorship_data[r.DonorID] = {
                    'sponsorship_id': r.SponsorshipID,
                    'start_date': r.StartDate,
                    'end_date': r.EndDate,
                    'amount': r.SponsorshipAmount
                }
            
            donor_combo['values'] = list(donor_dict.keys())
            
            if donor_dict:
                donor_combo.set("")
                update_info_display()
            else:
                messagebox.showinfo("No Donors", 
                                   "No active sponsors found for this child.")
                
        except Exception as e:
            messagebox.showerror("Error", f"Failed to load donors:\n{e}")
    
    def update_info_display():
        """Update the information display panel"""
        info_text.config(state="normal")
        info_text.delete("1.0", tk.END)
        
        child_text = child_var.get()
        donor_text = donor_var.get()
        
        if child_text and child_text in child_dict:
            child_id = child_dict[child_text]
            selected_child_id.set(child_id)
            
            # Get child details
            conn = get_connection()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT FullName, RegistrationNumber, Class, SchoolName
                FROM tblChildren WHERE ChildID = ?
            """, (child_id,))
            
            child_data = cursor.fetchone()
            conn.close()
            
            if child_data:
                info_text.insert(tk.END, 
                    f"Child: {child_data.FullName}\n"
                    f"Registration: {child_data.RegistrationNumber}\n"
                    f"Class: {child_data.Class or 'N/A'}\n"
                    f"School: {child_data.SchoolName or 'N/A'}\n"
                )
        
        if donor_text and donor_text in donor_dict:
            donor_id = donor_dict[donor_text]
            selected_donor_id.set(donor_id)
            
            if donor_id in sponsorship_data:
                sp_data = sponsorship_data[donor_id]
                info_text.insert(tk.END,
                    f"\nSponsorship Start: {format_db_date(sp_data['start_date'], '%d/%m/%Y') if sp_data['start_date'] else 'N/A'}\n"
                    f"Sponsorship End: {format_db_date(sp_data['end_date'], '%d/%m/%Y') if sp_data['end_date'] else 'N/A'}\n"
                )
                
                generate_btn.config(state="normal")
            else:
                generate_btn.config(state="disabled")
        else:
            generate_btn.config(state="disabled")
        
        info_text.config(state="disabled")
    
    # ─────────────────────────────────────────────────────────────────────────
    # EVENT HANDLERS
    # ─────────────────────────────────────────────────────────────────────────
    def on_child_select(event=None):
        child_text = child_var.get()
        if child_text and child_text in child_dict:
            child_id = child_dict[child_text]
            load_donors_for_child(child_id)
            update_info_display()
    
    def on_donor_select(event=None):
        update_info_display()
    
    child_combo.bind("<<ComboboxSelected>>", on_child_select)
    child_combo.bind("<Return>", on_child_select)
    donor_combo.bind("<<ComboboxSelected>>", on_donor_select)
    
    # ─────────────────────────────────────────────────────────────────────────
    # PDF GENERATION FUNCTION
    # ─────────────────────────────────────────────────────────────────────────
    def generate_certificate():
        try:
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.pdfgen import canvas
            from reportlab.lib.colors import HexColor
            import os
            from datetime import datetime

            # ───── DATA ─────
            child_id = selected_child_id.get()
            donor_id = selected_donor_id.get()

            if not child_id or not donor_id:
                messagebox.showwarning("Select", "Select child & donor")
                return

            conn = get_connection()
            cursor = conn.cursor()

            cursor.execute("""
                SELECT FullName, RegistrationNumber, Class, SchoolName,
                    PhotoPath, CenterID
                FROM tblChildren WHERE ChildID=?
            """, (child_id,))
            child = cursor.fetchone()

            cursor.execute("SELECT FullName FROM tblDonors WHERE DonorID=?", (donor_id,))
            donor = cursor.fetchone()

            cursor.execute("SELECT CenterName, City FROM tblCenters WHERE CenterID=?", (child.CenterID,))
            center = cursor.fetchone()

            conn.close()

            sp = sponsorship_data[donor_id]
            start_date = format_db_date(sp['start_date'], '%d/%m/%Y') if sp['start_date'] else ''
            end_date = format_db_date(sp['end_date'], '%d/%m/%Y') if sp['end_date'] else ''

            # ───── FILE ─────
            path = os.path.join(os.path.expanduser("~"), "Downloads")
            os.makedirs(path, exist_ok=True)
            #file = os.path.join(path, f"Cert_{child.RegistrationNumber}.pdf")
            # Clean child name (remove invalid filename characters)
            safe_name = "".join(c for c in child.FullName if c.isalnum() or c in (" ", "-")).strip()

            file_name = f"{safe_name} {child.RegistrationNumber}.pdf"

            file = os.path.join(path, file_name)

            c = canvas.Canvas(file, pagesize=landscape(A4))
            width, height = landscape(A4)

            border_color = HexColor("#1E8B2B")

            # ───── BORDER ─────
            c.setStrokeColor(border_color)
            c.setLineWidth(8)
            c.rect(20, 20, width - 40, height - 40)

            c.setLineWidth(2)
            c.rect(30, 30, width - 60, height - 60)

            # ───── HEADER ─────
            header = "logo/header_full.png"
            if os.path.exists(header):
                c.drawImage(header, 30, height - 130, width=width - 60, height=100)

            # ───── PHOTO ─────
            photo_x = width - 180
            photo_y = height - 300
            photo_w = 100
            photo_h = 120

            c.setStrokeColor(border_color)
            c.setLineWidth(2)
            c.rect(photo_x, photo_y, photo_w, photo_h)

            if child.PhotoPath:
                photo_candidates = [
                    child.PhotoPath,
                    os.path.join("Photo", os.path.basename(child.PhotoPath)),
                    os.path.join("photo", os.path.basename(child.PhotoPath))
                ]
                for photo_path in photo_candidates:
                    if os.path.exists(photo_path):
                        try:
                            c.drawImage(photo_path,
                                        photo_x + 3, photo_y + 3,
                                        width=photo_w - 6, height=photo_h - 6,
                                        preserveAspectRatio=True)
                            break
                        except:
                            pass

            # ───── TITLE ─────
            c.setFont("Helvetica-Bold", 20)
            c.setFillColor(HexColor("#1E4D8B"))
            c.drawCentredString(width / 2, height - 160, "CHILD SPONSORSHIP CERTIFICATE")

            # ───── BODY START ─────
            y = height - 200
            line_gap = 26

            c.setFont("Helvetica", 12)
            c.setFillColor(HexColor("#000"))

            c.drawString(60, y, "This is to certify that Mr./Mrs.")

            # ───── DONOR NAME ─────
            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(width / 2, y, donor.FullName)
            c.setStrokeColor(HexColor("#000000"))
            c.setLineWidth(1)
            c.line(200, y - 5, width - 200, y - 5)

            y -= line_gap
            c.setFont("Helvetica", 12)
            c.drawString(60, y, "is sponsoring the following child:")

            y -= line_gap

            # ───── SAFE LINE END ─────
            line_end_x = photo_x - 20

            def draw_line(label, value):
                nonlocal y

                c.setFont("Helvetica-Bold", 11)
                c.setFillColor(HexColor("#000000"))
                c.drawString(60, y, label)

                c.setFont("Helvetica", 12)
                c.drawString(200, y, str(value))

                c.setStrokeColor(HexColor("#CCCCCC"))
                c.setLineWidth(0.5)
                c.line(200, y - 3, line_end_x, y - 3)

                y -= line_gap

            # ───── DETAILS ─────
            draw_line("Child Name:", child.FullName)
            draw_line("Child ID:", child.RegistrationNumber)
            draw_line("Class:", child.Class or "N/A")
            draw_line("School:", child.SchoolName or "N/A")
            draw_line("Center:", center.CenterName if center else "N/A")
            draw_line("Region:", center.City if center else "N/A")
            draw_line("Start Date:", start_date)
            draw_line("Renewal Date:", end_date)

            # ───── SIGNATURES ─────
            sig_y = 110
            c.setStrokeColor(HexColor("#000000"))
            c.setLineWidth(1)
            c.line(120, sig_y, 300, sig_y)
            c.line(width - 300, sig_y, width - 120, sig_y)

            c.setFont("Helvetica", 10)
            c.setFillColor(HexColor("#000000"))
            c.drawCentredString(210, sig_y - 15, "Administrator")
            c.drawCentredString(width - 210, sig_y - 15, "Program Manager")

            # ───── FOOTER IMAGE ─────
            footer_img = "logo/urdu_text.png"
            if os.path.exists(footer_img):
                img_width = 300
                img_height = 40

                c.drawImage(
                    footer_img,
                    (width - img_width) / 2,
                    65,
                    width=img_width,
                    height=img_height,
                    preserveAspectRatio=True,
                    mask='auto'
                )

            # ───── BOTTOM INFO BAR ─────
            bar_height = 18

            # background bar
            c.setFillColor(border_color)
            c.rect(30, 30, width - 60, bar_height, stroke=0, fill=1)

            # white text
            c.setFillColor(HexColor("#FFFFFF"))
            c.setFont("Helvetica", 8)

            footer_text = "76-E, Block-6, PECHS, Karachi - Pakistan | Ph: 0800 77778 | info@alkhidmatkhawateen.org | www.alkhidmatkhawateen.org"

            c.drawCentredString(width / 2, 30 + 5, footer_text)

            # ───── SAVE ─────
            c.save()

            messagebox.showinfo("Success", f"Certificate Generated!\n\n{file}")

            try:
                os.startfile(file)
            except:
                pass

        except Exception as e:
            messagebox.showerror("Error", f"Failed to generate certificate:\n{e}")
            import traceback
            traceback.print_exc()
    
    generate_btn.config(command=generate_certificate)
    
    # ─────────────────────────────────────────────────────────────────────────
    # INITIALIZE
    # ─────────────────────────────────────────────────────────────────────────
    load_children()
    win.bind('<Escape>', lambda e: win.destroy())

def show_about():
    """Show about information."""
    messagebox.showinfo(
        "About Aghosh CareSys",
        "Aghosh Homes - Sponsorship Management System\n"
        "Developed for Alkhidmat Khawateen Trust Pakistan\n"
        "Version 5.5\n"
        "Developed by Kashif Haider Ali\n"
        "kashif.36@gmail.com"
    )


def logout(window):
    """Handle logout functionality."""
    confirm = messagebox.askyesno("Logout", "Are you sure you want to log out?")
    if confirm:
        conn = get_connection()
        if conn:
            try:
                cursor = conn.cursor()
                cursor.execute("UPDATE tblUsers SET IsLoggedIn = 0 WHERE IsLoggedIn = 1")
                conn.commit()
                conn.close()
            except Exception as e:
                messagebox.showerror("Database Error", f"Failed to update login status: {e}")
        window.destroy()
        import main
        main.open_login_window()  # Make sure main has this function


def create_menu(role):
    """Create the menu bar based on user role."""
    menubar = Menu(root)
    root.config(menu=menubar)

    # ---------------- File Menu ----------------
    file_menu = Menu(menubar, tearoff=0)
    file_menu.add_command(label="Logout", command=lambda: logout(root))
    file_menu.add_separator()
    file_menu.add_command(label="Exit", command=root.destroy)
    menubar.add_cascade(label="File", menu=file_menu)

    # ---------------- Modules Menu ----------------
    module_menu = Menu(menubar, tearoff=0)

    module_menu.add_command(label="Child Management", command=open_child_mgmt)
    module_menu.add_command(label="Donor Management", command=open_donor_mgmt)
    module_menu.add_command(label="Donor Commitments", command=open_donor_commitments)
    module_menu.add_command(label="Sponsorship Management", command=open_sponsorship_mgmt)
    module_menu.add_command(label="Collector Management", command=open_collector_mgmt)
    module_menu.add_command(label="Academics Tracking", command=open_academics_mgmt)
    module_menu.add_command(label="Health Tracking", command=open_health_mgmt)
    module_menu.add_command(label="Extracurricular Activities", command=open_activities_mgmt)
    module_menu.add_command(label="Enroll Child in Activity", command=open_child_activities_mgmt)

    module_menu.add_command(label="Donation Tracking", command=open_donation_mgmt)
    module_menu.add_command(label="Center Information", command=open_center_mgmt)
    module_menu.add_command(label="Sessions", command=open_session_mgmt)
    module_menu.add_command(label="User Access Control", command=open_user_mgmt)

    # Disable certain modules for ReadOnly users
    if role.lower() == "readonly":
        for item in [
            "Child Management", "Donor Management", "Sponsorship Management",
            "Collector Management", "User Access Control"
        ]:
            module_menu.entryconfig(item, state="disabled")

    # Add MODULES to menu bar
    menubar.add_cascade(label="Modules", menu=module_menu)

    # ---------------- Reports Menu (TOP LEVEL) ----------------
    reports_menu = Menu(menubar, tearoff=0)
    reports_menu.add_command(label="Children Report", command=open_child_report_form)
    reports_menu.add_command(label="Child Profile", command=open_child_profile_report_form)
    reports_menu.add_command(label="Generate Child Profile For a Donor", command=open_child_profile_for_donor_report)
    reports_menu.add_command(label="Donor Report", command=open_donor_report_form)
    reports_menu.add_command(label="Donations Report", command=open_donation_report_form)
    reports_menu.add_command(label="Donor Child Sponsorship Map", command=open_donor_child_map_report)
    reports_menu.add_command(label="Donor-wise Commitment Report", command=open_donor_wise_commitment_report)
    reports_menu.add_command(label="Donor Giving History Report", command=open_donor_giving_history_report)
    reports_menu.add_command(label="Donor Defaulters Report", command=open_donor_defaulters_report)
    reports_menu.add_command(label="Collector's Collections", command=open_collector_report_form)
    reports_menu.add_command(label="Collector List", command=open_collector_list_form)
    reports_menu.add_command(label="Unsponsored / Partially Sponsored Children", command=open_unsponsored_Partially_sponsored_report)
    reports_menu.add_command(label="Sponsorship vs Commitment Validation", command=open_sponsorship_vs_commitment_validation)
    reports_menu.add_command(label="Commitment vs. Actual Donations", command=open_commitment_vs_actual_donations_report)
    reports_menu.add_command(label="Expired Commitments Report", command=open_expired_commitments_report)
    reports_menu.add_command(label="Upcoming Expiry Report", command=open_upcoming_expiry_report)
    reports_menu.add_command(label="Exel For AKF", command=open_Exel_For_AKF)
    reports_menu.add_command(label="Overall Impact Report", command=open_overall_impact_report)

    # Add REPORTS as its own menu
    menubar.add_cascade(label="Reports", menu=reports_menu)

    # ---------------- Help Menu ----------------
    help_menu = Menu(menubar, tearoff=0)
    help_menu.add_command(label="About", command=show_about)
    menubar.add_cascade(label="Help", menu=help_menu)

    # Popup for near-to-expiry commitments
    root.after(1000, check_expiring_commitments_this_month)



def open_dashboard(username, role, center_id=None):
    global logged_in_role
    logged_in_role = role

    global root
    root = tk.Tk()
    root.title("AghoshCareSys - Dashboard" + ORG_SUFFIX)
    set_window_icon(root)
    root.state('zoomed')
    root.focus_force()

    # ═════════════════════════════════════════════════════════════════════════
    # DESIGN TOKENS
    # ═════════════════════════════════════════════════════════════════════════
    C = {
        "header_top"   : "#0f2540",   # deep navy
        "header_bot"   : "#1a3c5e",   # slightly lighter navy
        "sidebar"      : "#162033",   # dark sidebar
        "body"         : "#eef2f7",   # light grey page bg
        "card_bg"      : "#ffffff",
        "card_hover"   : "#f0f6ff",
        "card_border"  : "#dde5f0",
        "accent"       : "#1d6fd8",   # primary blue
        "accent_dark"  : "#154fa3",
        "success"      : "#16a34a",
        "text_dark"    : "#1e293b",
        "text_mid"     : "#475569",
        "text_light"   : "#f8fafc",
        "text_muted"   : "#94a3b8",
        "footer_bg"    : "#1e293b",
        "footer_fg"    : "#94a3b8",
        "divider"      : "#cbd5e1",
    }

    F = {
        "org"      : ("Segoe UI", 11),
        "title"    : ("Segoe UI", 20, "bold"),
        "subtitle" : ("Segoe UI", 10),
        "card_icon": ("Segoe UI", 20),
        "card_lbl" : ("Segoe UI", 9,  "bold"),
        "footer"   : ("Segoe UI", 8),
        "user_tag" : ("Segoe UI", 9,  "bold"),
        "section"  : ("Segoe UI", 8),
    }

    root.configure(bg=C["body"])

    # ═════════════════════════════════════════════════════════════════════════
    # FETCH LOGGED-IN USER
    # ═════════════════════════════════════════════════════════════════════════
    conn = get_connection()
    current_user_data = None
    if conn:
        try:
            cursor = conn.cursor()
            cursor.execute(
                "SELECT Username, Role, CenterID FROM tblUsers WHERE IsLoggedIn = 1")
            result = cursor.fetchone()
            if result:
                cursor.execute(
                    "SELECT CenterName FROM tblCenters WHERE CenterID = ?",
                    (result.CenterID,))
                cr = cursor.fetchone()
                current_user_data = {
                    "username"   : result.Username,
                    "role"       : result.Role,
                    "center_id"  : result.CenterID,
                    "center_name": cr.CenterName if cr else "Unknown",
                }
            conn.close()
        except Exception as e:
            messagebox.showerror("Database Error",
                                 f"Failed to fetch logged-in user: {e}")

    if not current_user_data:
        messagebox.showerror("Session Expired",
                             "No active session found. Please login again.")
        root.destroy()
        import main
        main.open_login_window()
        return

    logged_in_username    = current_user_data["username"]
    logged_in_role        = current_user_data["role"]
    logged_in_center_name = current_user_data["center_name"]

    # ═════════════════════════════════════════════════════════════════════════
    # HEADER
    # ═════════════════════════════════════════════════════════════════════════
    header = tk.Frame(root, bg=C["header_top"], height=110)
    header.pack(fill="x", side="top")
    header.pack_propagate(False)

    base_dir = os.path.dirname(os.path.abspath(__file__))

    # Left logo
    try:
        alkhidmat_img = Image.open(
            os.path.join(base_dir, "logo", "akh_logo.png")).resize((80, 80), Image.LANCZOS)
        alkhidmat_logo = ImageTk.PhotoImage(alkhidmat_img)
        lbl_ll = tk.Label(header, image=alkhidmat_logo, bg=C["header_top"])
        lbl_ll.image = alkhidmat_logo
        lbl_ll.place(relx=0.015, rely=0.5, anchor="w")
    except Exception:
        pass

    # Right logo
    try:
        aghosh_img = Image.open(
            os.path.join(base_dir, "logo", "aghosh_logo.png")).resize((80, 80), Image.LANCZOS)
        aghosh_logo = ImageTk.PhotoImage(aghosh_img)
        lbl_rl = tk.Label(header, image=aghosh_logo, bg=C["header_top"])
        lbl_rl.image = aghosh_logo
        lbl_rl.place(relx=0.985, rely=0.5, anchor="e")
    except Exception:
        pass

    # Centre text block
    tk.Label(header,
             text="AGHOSH CareSys",
             font=("Segoe UI", 22, "bold"),
             bg=C["header_top"], fg=C["text_light"]).place(
        relx=0.5, rely=0.30, anchor="center")

    tk.Label(header,
             text="Aghosh Homes  ·  Alkhidmat Khawateen Trust Pakistan",
             font=("Segoe UI", 10),
             bg=C["header_top"], fg=C["text_muted"]).place(
        relx=0.5, rely=0.62, anchor="center")

    # Thin accent line below header
    tk.Frame(root, bg=C["accent"], height=3).pack(fill="x")

    # ═════════════════════════════════════════════════════════════════════════
    # FOOTER (pack before body so it stays at bottom)
    # ═════════════════════════════════════════════════════════════════════════
    footer = tk.Frame(root, bg=C["footer_bg"], height=28)
    footer.pack(fill="x", side="bottom")
    footer.pack_propagate(False)

    tk.Label(footer,
             text=f"  Logged in as:  {logged_in_username}  |  Role: {logged_in_role}"
                  f"  |  Center: {logged_in_center_name}",
             font=F["footer"], bg=C["footer_bg"], fg=C["footer_fg"]).pack(
        side="left", padx=8, pady=4)

    import datetime
    tk.Label(footer,
             text=f"AghoshCareSys  ·  {datetime.date.today().strftime('%d %B %Y')}  ",
             font=F["footer"], bg=C["footer_bg"], fg=C["footer_fg"]).pack(
        side="right", padx=8, pady=4)

    # ═════════════════════════════════════════════════════════════════════════
    # MENU BAR
    # ═════════════════════════════════════════════════════════════════════════
    create_menu(logged_in_role)

    # ═════════════════════════════════════════════════════════════════════════
    # BODY — scrollable module grid
    # ═════════════════════════════════════════════════════════════════════════
    body = tk.Frame(root, bg=C["body"])
    body.pack(fill="both", expand=True)

    # Section label above buttons
    tk.Label(body,
             text="MODULES",
             font=("Segoe UI", 8, "bold"),
             bg=C["body"], fg=C["text_muted"]).pack(
        anchor="w", padx=40, pady=(18, 4))

    # ── Card helper ───────────────────────────────────────────────────────────
    CARD_W  = 175
    CARD_H  = 90
    ICON_SZ = 22

    def make_card(parent, icon, label, command, row, col):
        """A flat white card with icon + label that highlights on hover."""
        outer = tk.Frame(parent, bg=C["card_border"],
                         padx=1, pady=1)
        outer.grid(row=row, column=col, padx=10, pady=8)

        card = tk.Frame(outer, bg=C["card_bg"],
                        width=CARD_W, height=CARD_H,
                        cursor="hand2")
        card.pack()
        card.pack_propagate(False)

        ico_lbl = tk.Label(card, text=icon,
                           font=("Segoe UI", ICON_SZ),
                           bg=C["card_bg"], fg=C["accent"])
        ico_lbl.pack(pady=(14, 2))

        txt_lbl = tk.Label(card, text=label,
                           font=F["card_lbl"],
                           bg=C["card_bg"], fg=C["text_dark"],
                           wraplength=CARD_W - 16, justify="center")
        txt_lbl.pack()

        def on_enter(e):
            card.configure(bg=C["card_hover"])
            ico_lbl.configure(bg=C["card_hover"])
            txt_lbl.configure(bg=C["card_hover"])

        def on_leave(e):
            card.configure(bg=C["card_bg"])
            ico_lbl.configure(bg=C["card_bg"])
            txt_lbl.configure(bg=C["card_bg"])

        def on_click(e):
            command()

        for w in (card, ico_lbl, txt_lbl):
            w.bind("<Enter>",   on_enter)
            w.bind("<Leave>",   on_leave)
            w.bind("<Button-1>", on_click)

    # ── Module definitions — (icon, label, command, row, col) ─────────────────
    modules = [
        # row 0
        ("👶", "Child Management",          open_child_mgmt,                    0, 0),
        ("💝", "Donor Management",           open_donor_mgmt,                    0, 1),
        ("🤝", "Sponsorship",                open_sponsorship_mgmt,              0, 2),
        ("📋", "Collectors",                 open_collector_mgmt,                0, 3),
        # row 1
        ("📜", "Donor Commitments",          open_donor_commitments,             1, 0),
        ("💰", "Donations",                  open_donation_mgmt,                 1, 1),
        ("🏠", "Center Info",                open_center_mgmt,                   1, 2),
        ("🔐", "User Access",                open_user_mgmt,                     1, 3),
        # row 2
        ("📊", "Reports",                    open_reports,                       2, 0),
        ("📅", "Sessions",                   open_session_mgmt,                  2, 1),
        ("🎓", "Academics",                  open_academics_mgmt,                2, 2),
        ("🏥", "Health",                     open_health_mgmt,                   2, 3),
        # row 3
        ("🥧", "Sponsorship Pie Chart",      open_child_sponsorship_graph_report,3, 0),
        ("🎨", "Extracurricular Activities", open_activities_mgmt,               3, 1),
        ("✅", "Enroll Child in Activity",   open_child_activities_mgmt,         3, 2),
        ("💝", "Sponsorship Certificate",   open_Sponsorship_Certificate_mgmt,         3, 3),
    ]

    grid_frame = tk.Frame(body, bg=C["body"])
    grid_frame.pack(pady=4)

    for icon, label, command, row, col in modules:
        make_card(grid_frame, icon, label, command, row, col)

    root.mainloop()


# ── Test entry point ──────────────────────────────────────────────────────────
if __name__ == "__main__":
    open_dashboard("Test User", "Admin", center_id=1)

    
